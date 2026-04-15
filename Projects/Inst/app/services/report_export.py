from __future__ import annotations

from pathlib import Path

from docx import Document

from app.config import get_settings
from app.schemas import AnalysisResponse


def _build_report_text(report: AnalysisResponse) -> str:
    insights = report.insights
    engagement = insights.engagement
    lines = [
        "INSTAGRAM ANALYSIS REPORT",
        "",
        f"ID: {report.id}",
        f"URL: {report.url}",
        f"Shortcode: {report.shortcode}",
        f"Author: {report.author_username}",
        f"Media type: {report.media_type}",
        f"Created at: {report.created_at.isoformat()}",
        "",
        "METRICS",
        f"- Likes: {report.likes_count}",
        f"- Comments: {report.comments_count}",
        f"- Views: {report.view_count}",
        f"- Engagement rate: {report.engagement_rate}%",
        f"- Engagement basis: {engagement.get('basis', 'unknown')}",
        f"- ER by views: {engagement.get('by_views', 0.0):.2f}%",
        f"- ER proxy: {engagement.get('proxy', 0.0):.2f}%",
        "",
        "MODEL: CAPTURE -> RETENTION -> TRANSFER",
        f"- Capture: {insights.hook.score}/10 | {insights.hook.rationale}",
        f"- Retention: {insights.retention.score}/10 | {insights.retention.rationale}",
        f"- Transfer: {insights.transfer.score}/10 | {insights.transfer.rationale}",
        f"- Confidence: {insights.confidence_score}/100",
        "",
        "CTA",
        f"- Detected: {'yes' if insights.cta_detected else 'no'}",
        f"- Type: {insights.cta_type}",
        "",
        "AUDIENCE REACTION",
        f"- Positive comments: {insights.audience_reaction.get('positive_comments', 0)}",
        f"- Negative comments: {insights.audience_reaction.get('negative_comments', 0)}",
        f"- Questions comments: {insights.audience_reaction.get('questions_comments', 0)}",
        f"- Comments analyzed: {insights.audience_reaction.get('comments_analyzed', 0)}",
        "",
        "RECOMMENDATIONS",
    ]
    lines.extend([f"- {item}" for item in insights.actionable_recommendations])
    lines.append("")
    return "\n".join(lines)


def export_report_to_txt(report: AnalysisResponse) -> Path:
    settings = get_settings()
    reports_dir = Path(settings.reports_dir)
    reports_dir.mkdir(parents=True, exist_ok=True)

    path = reports_dir / f"analysis_{report.id}_{report.shortcode}.txt"
    path.write_text(_build_report_text(report), encoding="utf-8")
    return path


def export_report_to_docx(report: AnalysisResponse) -> Path:
    settings = get_settings()
    reports_dir = Path(settings.reports_dir)
    reports_dir.mkdir(parents=True, exist_ok=True)

    path = reports_dir / f"analysis_{report.id}_{report.shortcode}.docx"
    doc = Document()
    doc.add_heading("Instagram Analysis Report", level=1)
    doc.add_paragraph(f"ID: {report.id}")
    doc.add_paragraph(f"URL: {report.url}")
    doc.add_paragraph(f"Author: {report.author_username}")
    doc.add_paragraph(f"Media type: {report.media_type}")
    doc.add_paragraph(f"Created at: {report.created_at.isoformat()}")

    doc.add_heading("Metrics", level=2)
    engagement = report.insights.engagement
    doc.add_paragraph(f"Likes: {report.likes_count}")
    doc.add_paragraph(f"Comments: {report.comments_count}")
    doc.add_paragraph(f"Views: {report.view_count}")
    doc.add_paragraph(f"Engagement rate: {report.engagement_rate}%")
    doc.add_paragraph(f"Engagement basis: {engagement.get('basis', 'unknown')}")
    doc.add_paragraph(f"ER by views: {engagement.get('by_views', 0.0):.2f}%")
    doc.add_paragraph(f"ER proxy: {engagement.get('proxy', 0.0):.2f}%")

    doc.add_heading("Capture -> Retention -> Transfer", level=2)
    doc.add_paragraph(f"Capture: {report.insights.hook.score}/10 - {report.insights.hook.rationale}")
    doc.add_paragraph(
        f"Retention: {report.insights.retention.score}/10 - {report.insights.retention.rationale}"
    )
    doc.add_paragraph(f"Transfer: {report.insights.transfer.score}/10 - {report.insights.transfer.rationale}")
    doc.add_paragraph(f"Confidence: {report.insights.confidence_score}/100")

    doc.add_heading("CTA", level=2)
    doc.add_paragraph(f"Detected: {'yes' if report.insights.cta_detected else 'no'}")
    doc.add_paragraph(f"Type: {report.insights.cta_type}")

    doc.add_heading("Audience reaction", level=2)
    doc.add_paragraph(f"Positive comments: {report.insights.audience_reaction.get('positive_comments', 0)}")
    doc.add_paragraph(f"Negative comments: {report.insights.audience_reaction.get('negative_comments', 0)}")
    doc.add_paragraph(f"Questions comments: {report.insights.audience_reaction.get('questions_comments', 0)}")
    doc.add_paragraph(f"Comments analyzed: {report.insights.audience_reaction.get('comments_analyzed', 0)}")

    doc.add_heading("Recommendations", level=2)
    for recommendation in report.insights.actionable_recommendations:
        doc.add_paragraph(recommendation, style="List Bullet")

    doc.save(path)
    return path
