# -*- coding: utf-8 -*-
"""
Скрипт для анализа документа v7 и проверки источников
"""
import re
import shutil
from pathlib import Path
from docx import Document

def setup_files():
    """Копирует необходимые файлы если их нет"""
    base = Path(__file__).parent.parent
    version4 = base / "Version4"
    version5 = base / "Version5"
    
    version5.mkdir(exist_ok=True)
    
    # Копируем документ v7
    v7_source = version4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v7_dest = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if v7_source.exists() and not v7_dest.exists():
        shutil.copy2(v7_source, v7_dest)
        print(f"✓ Скопирован: {v7_dest.name}")
    elif v7_dest.exists():
        print(f"✓ Уже есть: {v7_dest.name}")
    else:
        print(f"✗ Не найден: {v7_source}")
        return False
    
    # Копируем список литературы
    lit_source = version4 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    lit_dest = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if lit_source.exists() and not lit_dest.exists():
        shutil.copy2(lit_source, lit_dest)
        print(f"✓ Скопирован: {lit_dest.name}")
    elif lit_dest.exists():
        print(f"✓ Уже есть: {lit_dest.name}")
    else:
        print(f"✗ Не найден: {lit_source}")
        return False
    
    return True

def extract_all_text_with_paragraphs(doc):
    """Извлекает весь текст с сохранением информации о параграфах"""
    paragraphs_data = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs_data.append({
                'text': text,
                'para': para,
                'para_idx': len(paragraphs_data)
            })
    
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs_data.append({
                        'text': cell_text,
                        'para': None,
                        'para_idx': len(paragraphs_data),
                        'is_table': True
                    })
    
    return paragraphs_data

def find_all_citations(text):
    """Находит все ссылки на источники в тексте"""
    citations = []
    
    # Квадратные скобки с номерами: [5], [1, 2], [1; 2], [1-5]
    pattern1 = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
    matches1 = re.finditer(pattern1, text)
    for match in matches1:
        numbers_str = match.group(1)
        numbers = []
        for num_str in re.split(r'[,\s;]+', numbers_str):
            if '-' in num_str:
                # Диапазон: 1-5
                start, end = map(int, num_str.split('-'))
                numbers.extend(range(start, end + 1))
            else:
                numbers.append(int(num_str.strip()))
        citations.append({
            'type': 'square_brackets',
            'full_match': match.group(0),
            'numbers': numbers,
            'position': match.span()
        })
    
    # Круглые скобки с номерами: (5), (1, 2)
    pattern2 = r'\((\d+(?:[,\s;]\s*\d+)*)\)'
    matches2 = re.finditer(pattern2, text)
    for match in matches2:
        numbers_str = match.group(1)
        numbers = []
        for num_str in re.split(r'[,\s;]+', numbers_str):
            if '-' in num_str:
                start, end = map(int, num_str.split('-'))
                numbers.extend(range(start, end + 1))
            else:
                numbers.append(int(num_str.strip()))
        citations.append({
            'type': 'round_brackets_numbers',
            'full_match': match.group(0),
            'numbers': numbers,
            'position': match.span()
        })
    
    # Круглые скобки с фамилиями и номерами: (Андреева 5), (Эдельмин, Юскис 64)
    pattern3 = r'\(([А-ЯЁа-яё\s,\-]+?)\s+(\d+)\)'
    matches3 = re.finditer(pattern3, text)
    for match in matches3:
        names_str = match.group(1).strip()
        names = [n.strip() for n in names_str.split(',')]
        citations.append({
            'type': 'round_brackets_name_number',
            'full_match': match.group(0),
            'names': names,
            'number': int(match.group(2)),
            'position': match.span()
        })
    
    return citations

def load_literature_list(literature_file):
    """Загружает список литературы из файла"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Парсим список литературы (формат: номер. Автор...)
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        # Извлекаем фамилию автора (первое слово до точки или запятой)
        author_match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
        if author_match:
            author = author_match.group(1)
        else:
            # Пробуем английские имена
            author_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', entry)
            if author_match:
                author = author_match.group(1)
            else:
                author = None
        literature[num] = {
            'number': num,
            'entry': entry,
            'author': author
        }
    
    return literature

def find_page_4_content(paragraphs_data):
    """Находит содержимое страницы 4 (примерно первые параграфы с упоминанием системного подхода)"""
    page_4_paragraphs = []
    
    keywords = ['системный семейный подход', 'андреева', 'эдельмин', 'юскис', 'адлер', 
                'гинзбург', 'климов', 'прихожан', 'толстых']
    
    for para_data in paragraphs_data[:100]:  # Первые 100 параграфов
        text_lower = para_data['text'].lower()
        if any(keyword in text_lower for keyword in keywords):
            # Проверяем, есть ли проблемные ссылки
            citations = find_all_citations(para_data['text'])
            has_problem = any(c['type'] == 'round_brackets_name_number' for c in citations)
            # Проверяем наличие ключевых слов для страницы 4
            if has_problem or 'системный' in text_lower or 'андреева' in text_lower:
                page_4_paragraphs.append(para_data)
    
    return page_4_paragraphs

def analyze_document(v7_path, literature_path):
    """Анализирует документ и находит все проблемы"""
    print("=" * 70)
    print("АНАЛИЗ ДОКУМЕНТА v7")
    print("=" * 70)
    
    # Загружаем документ
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен: {v7_path.name}")
    
    # Загружаем список литературы
    literature = load_literature_list(literature_path)
    print(f"✓ Список литературы загружен: {len(literature)} источников")
    print(f"   Источники: {min(literature.keys())} - {max(literature.keys())}")
    
    # Извлекаем текст
    paragraphs_data = extract_all_text_with_paragraphs(doc)
    print(f"✓ Извлечено параграфов: {len(paragraphs_data)}")
    
    # Анализируем каждый параграф
    all_citations = []
    problems = []
    used_sources = set()
    
    for para_data in paragraphs_data:
        text = para_data['text']
        citations = find_all_citations(text)
        
        for citation in citations:
            all_citations.append({
                'para_idx': para_data['para_idx'],
                'text_snippet': text[:150] + '...' if len(text) > 150 else text,
                'citation': citation
            })
            
            # Проверяем проблемы
            if citation['type'] == 'round_brackets_name_number':
                # Проблема: фамилия в скобках
                problems.append({
                    'type': 'name_in_brackets',
                    'para_idx': para_data['para_idx'],
                    'text': text,
                    'citation': citation,
                    'description': f"Найдена фамилия в скобках: {citation['full_match']}"
                })
            
            # Проверяем номера источников
            if 'numbers' in citation:
                for num in citation['numbers']:
                    used_sources.add(num)
                    if num > max(literature.keys()) or num < min(literature.keys()):
                        problems.append({
                            'type': 'source_not_found',
                            'para_idx': para_data['para_idx'],
                            'text': text,
                            'citation': citation,
                            'description': f"Источник {num} не найден в списке литературы (диапазон: {min(literature.keys())}-{max(literature.keys())})"
                        })
            elif 'number' in citation:
                num = citation['number']
                used_sources.add(num)
                if num > max(literature.keys()) or num < min(literature.keys()):
                    problems.append({
                        'type': 'source_not_found',
                        'para_idx': para_data['para_idx'],
                        'text': text,
                        'citation': citation,
                        'description': f"Источник {num} не найден в списке литературы (диапазон: {min(literature.keys())}-{max(literature.keys())})"
                    })
    
    # Проверяем неиспользуемые источники
    all_source_numbers = set(literature.keys())
    unused_sources = all_source_numbers - used_sources
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ АНАЛИЗА")
    print(f"{'='*70}")
    print(f"Всего найдено ссылок: {len(all_citations)}")
    print(f"Используемых источников: {len(used_sources)}")
    print(f"Неиспользуемых источников: {len(unused_sources)}")
    print(f"Найдено проблем: {len(problems)}")
    
    # Группируем проблемы по типам
    problems_by_type = {}
    for problem in problems:
        ptype = problem['type']
        if ptype not in problems_by_type:
            problems_by_type[ptype] = []
        problems_by_type[ptype].append(problem)
    
    print(f"\nПроблемы по типам:")
    for ptype, plist in problems_by_type.items():
        print(f"  {ptype}: {len(plist)}")
    
    if problems:
        print(f"\n{'='*70}")
        print("НАЙДЕННЫЕ ПРОБЛЕМЫ")
        print(f"{'='*70}")
        for i, problem in enumerate(problems[:20], 1):  # Показываем первые 20
            print(f"\n{i}. {problem['type']}: {problem['description']}")
            print(f"   Параграф {problem['para_idx'] + 1}")
            print(f"   Текст: {problem['text'][:200]}...")
        if len(problems) > 20:
            print(f"\n... и ещё {len(problems) - 20} проблем")
    
    if unused_sources:
        print(f"\n{'='*70}")
        print("НЕИСПОЛЬЗУЕМЫЕ ИСТОЧНИКИ")
        print(f"{'='*70}")
        for num in sorted(unused_sources):
            print(f"{num}. {literature[num]['entry'][:80]}...")
    
    # Ищем содержимое страницы 4
    print(f"\n{'='*70}")
    print("ПОИСК СОДЕРЖИМОГО СТРАНИЦЫ 4")
    print(f"{'='*70}")
    page_4_content = find_page_4_content(paragraphs_data)
    print(f"Найдено параграфов на странице 4: {len(page_4_content)}")
    for i, para_data in enumerate(page_4_content[:10], 1):
        print(f"\n{i}. Параграф {para_data['para_idx'] + 1}:")
        print(f"   {para_data['text'][:300]}...")
        citations = find_all_citations(para_data['text'])
        if citations:
            print(f"   Найдено ссылок: {len(citations)}")
            for cit in citations:
                print(f"     - {cit['full_match']} ({cit['type']})")
    
    return {
        'doc': doc,
        'paragraphs_data': paragraphs_data,
        'literature': literature,
        'all_citations': all_citations,
        'problems': problems,
        'used_sources': used_sources,
        'unused_sources': unused_sources,
        'page_4_content': page_4_content
    }

def main():
    version5 = Path(__file__).parent
    
    # Копируем файлы если нужно
    if not setup_files():
        print("ERROR: Не удалось скопировать файлы")
        return
    
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    if not literature_path.exists():
        print(f"ERROR: {literature_path} не найден")
        return
    
    # Анализируем документ
    analysis = analyze_document(v7_path, literature_path)
    
    # Сохраняем результаты анализа
    results_file = version5 / "ANALYSIS_RESULTS.txt"
    with open(results_file, 'w', encoding='utf-8') as f:
        f.write("РЕЗУЛЬТАТЫ АНАЛИЗА ДОКУМЕНТА v7\n")
        f.write("=" * 70 + "\n\n")
        f.write(f"Всего найдено ссылок: {len(analysis['all_citations'])}\n")
        f.write(f"Используемых источников: {len(analysis['used_sources'])}\n")
        f.write(f"Неиспользуемых источников: {len(analysis['unused_sources'])}\n")
        f.write(f"Найдено проблем: {len(analysis['problems'])}\n\n")
        
        f.write("ПРОБЛЕМЫ:\n")
        f.write("-" * 70 + "\n")
        for i, problem in enumerate(analysis['problems'], 1):
            f.write(f"\n{i}. {problem['type']}: {problem['description']}\n")
            f.write(f"   Параграф {problem['para_idx'] + 1}\n")
            f.write(f"   Текст: {problem['text']}\n")
    
    print(f"\n{'='*70}")
    print(f"✓ Результаты анализа сохранены в: {results_file.name}")

if __name__ == '__main__':
    main()

