# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v8 -> v9
Исправляет неверные номера источников и перефразирует проблемные места
"""
import re
from pathlib import Path
from docx import Document

def load_literature_list(literature_file):
    """Загружает список литературы"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
    
    return literature

def main():
    version5 = Path(__file__).parent
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v8_path.exists():
        print(f"ERROR: {v8_path} не найден")
        return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v8 -> v9")
    print("=" * 70)
    
    doc = Document(str(v8_path))
    print(f"✓ Документ загружен")
    
    literature = load_literature_list(literature_path)
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    print(f"✓ Список литературы: {len(literature)} источников (диапазон: {min_source}-{max_source})\n")
    
    fixed_count = 0
    
    # Словарь исправлений неверных номеров
    # формат: неправильный_номер -> правильный_номер
    number_fixes = {
        '[60]': '[51]',  # Франкл: 60 -> 51
        '[55]': '[47]',  # Слободчиков, Исаев: 55 -> 47
        '[59]': '[50]',  # Фонарев: 59 -> 50
        '[37]': '[18]',  # Гинзбург: 37 -> 18
        '[39]': '[33]',  # Маркова: 39 -> 33
        '[28]': '[23]',  # Карпова, Артемьева: 28 -> 23
        '[24]': '[20]',  # Журавлев, Купрейченко: 24 -> 20
        '[46]': '[38]',  # Осорина: 46 -> 38
        '[65]': '[13; 15]',  # Ядов: 65 -> 13; 15
        '[64]': '[53]',  # Эйдемиллер, Юстицкис: 64 -> 53
    }
    
    # Список замен для перефразирования (страница 4 уже исправлена в v8)
    rephrase_replacements = [
        # Параграф 77 (страница 7)
        {
            'search': '[Ананьев, 3; Кон, 31; Дружинин, 22]',
            'replace': 'Ананьев [3], Кон [31], Дружинин [22]',
            'description': 'Параграф 77: Ананьев, Кон, Дружинин'
        },
        # Параграф 105 (страница 13)
        {
            'search': '[Ананьев, 3; Божович, 12; Кон, 31]',
            'replace': 'Ананьев [3], Божович [12], Кон [31]',
            'description': 'Параграф 105: Ананьев, Божович, Кон'
        },
        # Параграф 127 (страница 16)
        {
            'search': '[Выготский, 16; Кле, 29; Кривцова, 32; Кулагина, Колюцкий, 33; Мухина, 44; Ремшмидт, 51]',
            'replace': 'Выготский [16], Кле [29], Кривцова [32], Кулагина, Колюцкий [33], Мухина [44], Ремшмидт [51]',
            'description': 'Параграф 127: Выготский, Кле, Кривцова, Кулагина, Колюцкий, Мухина, Ремшмидт'
        },
        # Параграф 195 (страница 21) - первая замена
        {
            'search': '[Деркач, Зазыкин, 20; Дружилов, 21]',
            'replace': 'Деркач, Зазыкин [20], Дружилов [21]',
            'description': 'Параграф 195: Деркач, Зазыкин, Дружилов'
        },
        # Параграф 195 (страница 21) - вторая замена
        {
            'search': '[Иванников, 26; Мерлин, 42]',
            'replace': 'Иванников [26], Мерлин [42]',
            'description': 'Параграф 195: Иванников, Мерлин'
        },
        # Параграф 236 (страница 26)
        {
            'search': '[Деркач, Зазыкин, 20; Дружилов, 21; Иванников, 26]',
            'replace': 'Деркач, Зазыкин [20], Дружилов [21], Иванников [26]',
            'description': 'Параграф 236: Деркач, Зазыкин, Дружилов, Иванников'
        },
        # Параграф 237 (страница 26) - Роджерс, Франкл
        {
            'search': '[Роджерс, 52; Франкл, 60]',
            'replace': 'Роджерс [52], Франкл [51]',
            'description': 'Параграф 237: Роджерс, Франкл'
        },
    ]
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        
        # 1. Исправляем неверные номера источников
        for wrong_num, correct_num in number_fixes.items():
            if wrong_num in text:
                text = text.replace(wrong_num, correct_num)
                if fixed_count < 20:
                    print(f"✓ Исправлен номер: {wrong_num} -> {correct_num} (параграф {para_idx + 1})")
        
        # 2. Перефразируем проблемные места
        for replacement in rephrase_replacements:
            if replacement['search'] in text:
                text = text.replace(replacement['search'], replacement['replace'])
                print(f"✓ {replacement['description']}: перефразировано")
        
        # 3. [51], [52], [53] оставляем как есть - эти источники будут добавлены в список литературы
        # Ремшмидт [51], Роджерс [52], Сапогова [53] - оставляем без изменений
        
        # 6. Исправляем неверные номера в обычных ссылках
        pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
        
        def fix_numbers_in_brackets(match):
            numbers_str = match.group(1)
            valid_numbers = []
            
            for num_str in re.split(r'[,\s;]+', numbers_str):
                num_str = num_str.strip()
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                    else:
                        # Исправляем известные неверные номера
                        if num == 60:
                            valid_numbers.append(51)  # Франкл -> 51 (будет добавлен)
                        elif num == 64:
                            valid_numbers.append(53)  # Эйдемиллер, Юстицкис -> 53 (будет добавлен)
                        elif num == 55:
                            valid_numbers.append(47)  # Слободчиков, Исаев: 55 -> 47
                        elif num == 59:
                            valid_numbers.append(50)  # Фонарев: 59 -> 50
                        elif num == 37:
                            valid_numbers.append(18)  # Гинзбург: 37 -> 18
                        elif num == 39:
                            valid_numbers.append(33)  # Маркова: 39 -> 33
                        elif num == 28:
                            valid_numbers.append(23)  # Карпова, Артемьева: 28 -> 23
                        elif num == 24:
                            valid_numbers.append(20)  # Журавлев, Купрейченко: 24 -> 20
                        elif num == 46:
                            valid_numbers.append(38)  # Осорина: 46 -> 38
                        elif num == 65:
                            valid_numbers.extend([13, 15])  # Ядов -> 13; 15
                        # 51, 52, 53 оставляем как есть - эти источники будут добавлены
                except ValueError:
                    pass
            
            if valid_numbers:
                unique_numbers = sorted(set(valid_numbers))
                return f"[{', '.join(map(str, unique_numbers))}]"
            return ""
        
        text = re.sub(pattern_numbers, fix_numbers_in_brackets, text)
        
        # Применяем изменения
        if text != original_text:
            para.text = text
            fixed_count += 1
            
            if fixed_count <= 30:
                print(f"\nПараграф {para_idx + 1}:")
                print(f"  Было: {original_text[:100]}...")
                print(f"  Стало: {text[:100]}...")
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Исправляем неверные номера
                for wrong_num, correct_num in number_fixes.items():
                    if wrong_num in text:
                        text = text.replace(wrong_num, correct_num)
                
                # Исправляем номера в скобках
                pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
                def fix_numbers(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v9_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v9_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")

if __name__ == '__main__':
    main()

