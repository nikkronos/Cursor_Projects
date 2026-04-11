# -*- coding: utf-8 -*-
"""
Тестовый скрипт для проверки сохранения изменений
"""
from pathlib import Path
from docx import Document

def test_save():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    test_path = version5 / "ТЕСТ_ИЗМЕНЕНИЙ.docx"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("Загружаю документ...")
    doc = Document(str(v7_path))
    
    # Находим параграф 46 (страница 4)
    print("\nИщу параграф 46 (страница 4)...")
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx == 45:  # Параграф 46 (индекс 45)
            original = para.text
            print(f"Исходный текст: {original[:200]}...")
            
            # Пробуем разные способы обновления
            # Способ 1: clear + add_run
            para.clear()
            para.add_run("ТЕСТ ИЗМЕНЕНИЯ: " + original)
            print(f"✓ Способ 1 применен")
            break
    
    # Сохраняем
    doc.save(str(test_path))
    print(f"\n✓ Сохранено в: {test_path.name}")
    
    # Проверяем, что изменения сохранились
    print("\nПроверяю сохраненные изменения...")
    doc2 = Document(str(test_path))
    for para_idx, para in enumerate(doc2.paragraphs):
        if para_idx == 45:
            saved_text = para.text
            print(f"Сохраненный текст: {saved_text[:200]}...")
            if "ТЕСТ ИЗМЕНЕНИЯ" in saved_text:
                print("✓ Изменения сохранились!")
            else:
                print("✗ Изменения НЕ сохранились!")

if __name__ == '__main__':
    test_save()




