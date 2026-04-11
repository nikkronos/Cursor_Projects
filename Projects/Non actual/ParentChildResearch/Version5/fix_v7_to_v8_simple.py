# -*- coding: utf-8 -*-
"""
Простой и надежный скрипт для исправления документа v7 -> v8
Использует подход из Version4: para.text = new_text
"""
import re
from pathlib import Path
from docx import Document

def load_literature_list(literature_file):
    """Загружает список литературы"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
    
    return literature

def main():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("=" * 70)
    print("ПРОСТОЕ ИСПРАВЛЕНИЕ ДОКУМЕНТА v7 -> v8")
    print("=" * 70)
    
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен")
    
    literature = load_literature_list(literature_path)
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    print(f"✓ Список литературы: {len(literature)} источников (диапазон: {min_source}-{max_source})\n")
    
    keywords_page_4 = ['системный семейный подход', 'андреева', 'эдельмин', 'юскис', 
                      'адлер', 'гинзбург', 'климов', 'прихожан', 'толстых']
    
    fixed_count = 0
    page_4_count = 0
    
    # Обрабатываем параграфы - как в Version4
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        text_lower = text.lower()
        is_page_4 = any(keyword in text_lower for keyword in keywords_page_4)
        
        # 1. Исправляем фамилии в квадратных скобках
        # Паттерн: [Фамилия, номер] или [Фамилия; номер] или [Фамилия1, номер1; Фамилия2, номер2]
        pattern = r'\[([А-ЯЁ][А-ЯЁа-яё\s,\-;]+?)\]'
        
        def replace_with_names(match):
            content = match.group(1)
            original = match.group(0)
            
            # Проверяем, есть ли фамилии (заглавные буквы)
            if not re.search(r'[А-ЯЁ][а-яё]+', content):
                return original
            
            # Парсим содержимое
            parts = [p.strip() for p in content.split(';')]
            authors_with_numbers = []
            
            for part in parts:
                # Ищем последнее число
                number_match = re.search(r'(\d+)\s*$', part)
                if number_match:
                    number = int(number_match.group(1))
                    # Извлекаем фамилии (всё до числа)
                    authors_part = part[:number_match.start()].strip().rstrip(',;')
                    authors = [a.strip() for a in authors_part.split(',')]
                    authors = [a for a in authors if a and re.match(r'^[А-ЯЁ][а-яё]+', a)]
                    
                    if authors:
                        authors_with_numbers.append({
                            'authors': authors,
                            'number': number
                        })
            
            if not authors_with_numbers:
                return original
            
            # Для страницы 4 - перефразируем
            if is_page_4:
                if len(authors_with_numbers) == 1:
                    item = authors_with_numbers[0]
                    if len(item['authors']) == 1:
                        return f"{item['authors'][0]} [{item['number']}]"
                    else:
                        return f"{', '.join(item['authors'])} [{item['number']}]"
                elif len(authors_with_numbers) == 2:
                    first = authors_with_numbers[0]
                    second = authors_with_numbers[1]
                    
                    first_str = f"{first['authors'][0]} [{first['number']}]" if len(first['authors']) == 1 else f"{', '.join(first['authors'])} [{first['number']}]"
                    
                    if len(second['authors']) == 2:
                        # "Эйдемиллер, Юстицкис, 64" -> "Эйдемиллер. Также изучал Юскис [64]"
                        second_str = f"{second['authors'][0]}. Также изучал {second['authors'][1]} [{second['number']}]"
                    elif len(second['authors']) == 1:
                        second_str = f"{second['authors'][0]} [{second['number']}]"
                    else:
                        authors_list = ', '.join(second['authors'][:-1]) + f". Также изучал {second['authors'][-1]}"
                        second_str = f"{authors_list} [{second['number']}]"
                    
                    return f"{first_str}, {second_str}"
                else:
                    # Больше двух элементов
                    result_parts = []
                    for item in authors_with_numbers:
                        if len(item['authors']) == 1:
                            result_parts.append(f"{item['authors'][0]} [{item['number']}]")
                        else:
                            result_parts.append(f"{', '.join(item['authors'])} [{item['number']}]")
                    return ', '.join(result_parts)
            else:
                # Для остальных - убираем фамилии
                numbers = [item['number'] for item in authors_with_numbers]
                unique_numbers = sorted(set(numbers))
                return f"[{', '.join(map(str, unique_numbers))}]"
        
        text = re.sub(pattern, replace_with_names, text)
        
        # 2. Исправляем неверные номера (только для не-страницы 4)
        if not is_page_4:
            pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
            
            def fix_numbers(match):
                numbers_str = match.group(1)
                valid_numbers = []
                
                for num_str in re.split(r'[,\s;]+', numbers_str):
                    num_str = num_str.strip()
                    try:
                        num = int(num_str)
                        if min_source <= num <= max_source:
                            valid_numbers.append(num)
                    except ValueError:
                        pass
                
                if valid_numbers:
                    unique_numbers = sorted(set(valid_numbers))
                    return f"[{', '.join(map(str, unique_numbers))}]"
                else:
                    return ""  # Удаляем ссылку
            
            text = re.sub(pattern_numbers, fix_numbers, text)
        
        # 3. Перефразируем текст на странице 4
        if is_page_4:
            # "Системный семейный подход Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
            # -> "Системный семейный подход, который изучали яркие представители Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
            if 'системный семейный подход' in text_lower and 'который изучали яркие представители' not in text_lower:
                # Ищем паттерн "системный семейный подход Автор [номер]"
                system_pattern = r'(системный семейный подход)\s+([А-ЯЁа-яё]+(?:\s+[А-ЯЁа-яё]+)*)\s+\[(\d+)\](?:\s*,\s*([А-ЯЁа-яё]+))?(?:\.\s*Также\s+изучал\s+([А-ЯЁа-яё]+)\s+\[(\d+)\])?'
                def rephrase_system(match):
                    approach = match.group(1)
                    author1 = match.group(2)
                    num1 = match.group(3)
                    author2 = match.group(4)
                    author3 = match.group(5)
                    num3 = match.group(6)
                    
                    if author3:
                        return f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}. Также изучал {author3} [{num3}]"
                    elif author2:
                        return f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}"
                    else:
                        return f"{approach}, который изучали яркие представители {author1} [{num1}]"
                
                text = re.sub(system_pattern, rephrase_system, text, flags=re.IGNORECASE)
            
            # "Психологические теории развития и сценариев Адлер [2]"
            # -> "Психологические теории развития и сценариев, которые изучал Адлер [2]"
            if 'психологические теории' in text_lower and 'который изучал' not in text_lower and 'который изучали' not in text_lower:
                theory_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,3})\s+теории\s+([А-ЯЁа-яё]+)\s+\[(\d+)\]'
                def rephrase_theory(match):
                    theory_words = match.group(1)
                    author = match.group(2)
                    num = match.group(3)
                    return f"{theory_words} теории, которые изучал {author} [{num}]"
                
                text = re.sub(theory_pattern, rephrase_theory, text, flags=re.IGNORECASE)
            
            # "Теории профессионального самоопределения Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
            # -> "Теории профессионального самоопределения, которые изучали Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
            if 'теории профессионального самоопределения' in text_lower and 'который изучал' not in text_lower and 'который изучали' not in text_lower:
                theory_multi_pattern = r'(теории профессионального самоопределения)\s+([А-ЯЁа-яё\s,\[\]\d]+)'
                def rephrase_theory_multi(match):
                    theory = match.group(1)
                    authors_part = match.group(2)
                    return f"{theory}, которые изучали {authors_part}"
                
                text = re.sub(theory_multi_pattern, rephrase_theory_multi, text, flags=re.IGNORECASE)
        
        # Применяем изменения - как в Version4
        if text != original_text:
            para.text = text  # Простой способ, как в Version4
            fixed_count += 1
            
            if fixed_count <= 15:
                print(f"\nПараграф {para_idx + 1} (страница 4: {is_page_4}):")
                print(f"  Было: {original_text[:100]}...")
                print(f"  Стало: {text[:100]}...")
        
        if is_page_4:
            page_4_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                # Убираем фамилии из скобок
                pattern = r'\[([А-ЯЁ][А-ЯЁа-яё\s,\-;]+?)\]'
                def replace_names(match):
                    content = match.group(1)
                    if re.search(r'[А-ЯЁ][а-яё]+', content):
                        # Парсим и извлекаем номера
                        parts = [p.strip() for p in content.split(';')]
                        numbers = []
                        for part in parts:
                            number_match = re.search(r'(\d+)\s*$', part)
                            if number_match:
                                numbers.append(int(number_match.group(1)))
                        if numbers:
                            unique_numbers = sorted(set(numbers))
                            return f"[{', '.join(map(str, unique_numbers))}]"
                    return match.group(0)
                
                text = re.sub(pattern, replace_names, text)
                
                # Исправляем неверные номера
                pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
                def fix_numbers(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v8_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v8_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    print(f"✓ Обработано параграфов страницы 4: {page_4_count}")

if __name__ == '__main__':
    main()




