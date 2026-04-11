# -*- coding: utf-8 -*-
"""
Рабочая версия скрипта - использует более надежный способ обновления параграфов
"""
import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load_literature_list(literature_file):
    """Загружает список литературы из файла"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
    
    return literature

# Функция не нужна - используем para.text напрямую

def fix_all_citations(text, literature, is_page_4=False):
    """Исправляет все проблемные ссылки в тексте"""
    fixed_text = text
    changes = []
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    
    # 1. Исправляем фамилии в квадратных скобках
    # Паттерн: [Фамилия, номер] или [Фамилия1, номер1; Фамилия2, номер2]
    pattern_names = r'\[([А-ЯЁ][А-ЯЁа-яё\s,\-;]+?)\]'
    
    def replace_names(match):
        content = match.group(1)
        original = match.group(0)
        
        # Проверяем, есть ли фамилии
        if not re.search(r'[А-ЯЁ][а-яё]+', content):
            return original
        
        # Парсим
        parts = [p.strip() for p in content.split(';')]
        authors_with_numbers = []
        
        for part in parts:
            number_match = re.search(r'(\d+)\s*$', part)
            if number_match:
                number = int(number_match.group(1))
                authors_part = part[:number_match.start()].strip().rstrip(',;')
                authors = [a.strip() for a in authors_part.split(',')]
                authors = [a for a in authors if a and re.match(r'^[А-ЯЁ][а-яё]+', a)]
                
                if authors:
                    authors_with_numbers.append({
                        'authors': authors,
                        'number': number
                    })
        
        if not authors_with_numbers:
            return original
        
        # Для страницы 4 - перефразируем
        if is_page_4:
            if len(authors_with_numbers) == 1:
                item = authors_with_numbers[0]
                if len(item['authors']) == 1:
                    result = f"{item['authors'][0]} [{item['number']}]"
                else:
                    result = f"{', '.join(item['authors'])} [{item['number']}]"
            elif len(authors_with_numbers) == 2:
                first = authors_with_numbers[0]
                second = authors_with_numbers[1]
                
                first_str = f"{first['authors'][0]} [{first['number']}]" if len(first['authors']) == 1 else f"{', '.join(first['authors'])} [{first['number']}]"
                
                if len(second['authors']) == 2:
                    second_str = f"{second['authors'][0]}. Также изучал {second['authors'][1]} [{second['number']}]"
                elif len(second['authors']) == 1:
                    second_str = f"{second['authors'][0]} [{second['number']}]"
                else:
                    authors_list = ', '.join(second['authors'][:-1]) + f". Также изучал {second['authors'][-1]}"
                    second_str = f"{authors_list} [{second['number']}]"
                
                result = f"{first_str}, {second_str}"
            else:
                result_parts = []
                for item in authors_with_numbers:
                    if len(item['authors']) == 1:
                        result_parts.append(f"{item['authors'][0]} [{item['number']}]")
                    else:
                        result_parts.append(f"{', '.join(item['authors'])} [{item['number']}]")
                result = ', '.join(result_parts)
            
            changes.append(f"Страница 4: {original} -> {result}")
            return result
        else:
            # Для остальных - убираем фамилии
            numbers = [item['number'] for item in authors_with_numbers]
            unique_numbers = sorted(set(numbers))
            result = f"[{', '.join(map(str, unique_numbers))}]"
            changes.append(f"Убраны фамилии: {original} -> {result}")
            return result
    
    fixed_text = re.sub(pattern_names, replace_names, fixed_text)
    
    # 2. Исправляем неверные номера (только для не-страницы 4)
    if not is_page_4:
        pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
        
        def fix_numbers(match):
            numbers_str = match.group(1)
            valid_numbers = []
            invalid_numbers = []
            
            for num_str in re.split(r'[,\s;]+', numbers_str):
                num_str = num_str.strip()
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                    else:
                        invalid_numbers.append(num)
                except ValueError:
                    pass
            
            if invalid_numbers:
                changes.append(f"Удалены неверные номера: {invalid_numbers}")
            
            if valid_numbers:
                unique_numbers = sorted(set(valid_numbers))
                result = f"[{', '.join(map(str, unique_numbers))}]"
                if result != match.group(0):
                    changes.append(f"Исправлены номера: {match.group(0)} -> {result}")
                return result
            else:
                changes.append(f"Удалена ссылка: {match.group(0)}")
                return ""
        
        fixed_text = re.sub(pattern_numbers, fix_numbers, fixed_text)
    
    # 3. Перефразируем текст на странице 4
    if is_page_4:
        # "Системный семейный подход Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
        # -> "Системный семейный подход, который изучали яркие представители Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
        system_pattern = r'(системный семейный подход)\s+([А-ЯЁа-яё]+(?:\s+[А-ЯЁа-яё]+)*)\s+\[(\d+)\](?:\s*,\s*([А-ЯЁа-яё]+))?(?:\.\s*Также\s+изучал\s+([А-ЯЁа-яё]+)\s+\[(\d+)\])?'
        def rephrase_system(match):
            approach = match.group(1)
            author1 = match.group(2)
            num1 = match.group(3)
            author2 = match.group(4)
            author3 = match.group(5)
            num3 = match.group(6)
            
            if author3:
                result = f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}. Также изучал {author3} [{num3}]"
            elif author2:
                result = f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}"
            else:
                result = f"{approach}, который изучали яркие представители {author1} [{num1}]"
            
            changes.append(f"Перефразировано: {match.group(0)} -> {result}")
            return result
        
        if 'который изучали яркие представители' not in fixed_text:
            fixed_text = re.sub(system_pattern, rephrase_system, fixed_text, flags=re.IGNORECASE)
        
        # "Психологические теории развития и сценариев Адлер [2]"
        # -> "Психологические теории развития и сценариев, которые изучал Адлер [2]"
        theory_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,3})\s+теории\s+([А-ЯЁа-яё]+)\s+\[(\d+)\]'
        def rephrase_theory(match):
            theory_words = match.group(1)
            author = match.group(2)
            num = match.group(3)
            result = f"{theory_words} теории, которые изучал {author} [{num}]"
            changes.append(f"Перефразировано: {match.group(0)} -> {result}")
            return result
        
        fixed_text = re.sub(theory_pattern, rephrase_theory, fixed_text, flags=re.IGNORECASE)
        
        # "Теории профессионального самоопределения Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
        # -> "Теории профессионального самоопределения, которые изучали Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
        theory_multi_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,3})\s+теории\s+([А-ЯЁа-яё\s,\[\]\d]+)'
        def rephrase_theory_multi(match):
            theory_words = match.group(1)
            authors_part = match.group(2)
            if 'который изучал' not in authors_part and 'который изучали' not in authors_part:
                result = f"{theory_words} теории, которые изучали {authors_part}"
                changes.append(f"Перефразировано: {match.group(0)} -> {result}")
                return result
            return match.group(0)
        
        fixed_text = re.sub(theory_multi_pattern, rephrase_theory_multi, fixed_text, flags=re.IGNORECASE)
    
    return fixed_text, changes

def main():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("=" * 70)
    print("РАБОЧЕЕ ИСПРАВЛЕНИЕ ДОКУМЕНТА v7 -> v8")
    print("=" * 70)
    
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен")
    
    literature = load_literature_list(literature_path)
    print(f"✓ Список литературы: {len(literature)} источников\n")
    
    keywords_page_4 = ['системный семейный подход', 'андреева', 'эдельмин', 'юскис', 
                      'адлер', 'гинзбург', 'климов', 'прихожан', 'толстых']
    
    fixed_count = 0
    page_4_count = 0
    all_changes = []
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text_lower = original_text.lower()
        is_page_4 = any(keyword in text_lower for keyword in keywords_page_4)
        
        fixed_text, changes = fix_all_citations(original_text, literature, is_page_4)
        
        if fixed_text != original_text:
            # Обновляем параграф - простой способ как в Version4
            para.text = fixed_text
            fixed_count += 1
            all_changes.extend(changes)
            
            if fixed_count <= 15:
                print(f"\nПараграф {para_idx + 1} (страница 4: {is_page_4}):")
                print(f"  Было: {original_text[:100]}...")
                print(f"  Стало: {fixed_text[:100]}...")
                for change in changes:
                    print(f"    - {change}")
        
        if is_page_4:
            page_4_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                fixed_text, _ = fix_all_citations(original_text, literature, False)
                
                if fixed_text != original_text:
                    cell.text = fixed_text
                    fixed_count += 1
    
    # Сохраняем
    doc.save(str(v8_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v8_path.name}")
    print(f"✓ Исправлено: {fixed_count}")
    print(f"✓ Страница 4: {page_4_count} параграфов")
    print(f"✓ Всего изменений: {len(all_changes)}")

if __name__ == '__main__':
    main()

