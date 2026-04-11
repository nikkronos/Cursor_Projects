# -*- coding: utf-8 -*-
"""
Тестовый скрипт - прямая замена конкретного текста
"""
from pathlib import Path
from docx import Document

def test_direct_replace():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    test_path = version5 / "ТЕСТ_ПРЯМАЯ_ЗАМЕНА.docx"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("Загружаю документ...")
    doc = Document(str(v7_path))
    
    # Ищем параграф 46 с проблемным текстом
    print("\nИщу параграф 46...")
    for para_idx, para in enumerate(doc.paragraphs):
        if para_idx == 45:  # Параграф 46 (индекс 45)
            original = para.text
            print(f"Исходный текст: {original[:150]}...")
            
            # Прямая замена конкретного текста
            if "[Андреева, 5; Эйдемиллер, Юстицкис, 64]" in original:
                new_text = original.replace(
                    "[Андреева, 5; Эйдемиллер, Юстицкис, 64]",
                    "Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
                )
                para.text = new_text
                print(f"✓ Замена выполнена")
                print(f"Новый текст: {new_text[:150]}...")
            else:
                print("✗ Текст не найден для замены")
            break
    
    # Сохраняем
    doc.save(str(test_path))
    print(f"\n✓ Сохранено в: {test_path.name}")
    
    # Проверяем сохранение
    print("\nПроверяю сохраненные изменения...")
    doc2 = Document(str(test_path))
    for para_idx, para in enumerate(doc2.paragraphs):
        if para_idx == 45:
            saved_text = para.text
            print(f"Сохраненный текст: {saved_text[:150]}...")
            if "Андреева [5], Эйдемиллер. Также изучал Юскис [64]" in saved_text:
                print("✓✓✓ ИЗМЕНЕНИЯ СОХРАНИЛИСЬ!")
            else:
                print("✗✗✗ ИЗМЕНЕНИЯ НЕ СОХРАНИЛИСЬ!")

if __name__ == '__main__':
    test_direct_replace()




