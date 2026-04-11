# -*- coding: utf-8 -*-
"""Быстрый тест - копирование файлов и базовый анализ"""
import sys
from pathlib import Path
import shutil

v5 = Path(__file__).parent
v4 = v5.parent / "Version4"

print("=" * 70)
print("БЫСТРЫЙ ТЕСТ - КОПИРОВАНИЕ ФАЙЛОВ")
print("=" * 70)

# Копируем документ v7
v7_v4 = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
v7_v5 = v5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"

if v7_v4.exists():
    if not v7_v5.exists():
        shutil.copy2(v7_v4, v7_v5)
        print(f"✓ Скопирован: {v7_v5.name}")
    else:
        print(f"✓ Уже есть: {v7_v5.name}")
else:
    print(f"✗ Не найден: {v7_v4}")

# Копируем список литературы
lit_v4 = v4 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
lit_v5 = v5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"

if lit_v4.exists():
    if not lit_v5.exists():
        shutil.copy2(lit_v4, lit_v5)
        print(f"✓ Скопирован: {lit_v5.name}")
    else:
        print(f"✓ Уже есть: {lit_v5.name}")
else:
    print(f"✗ Не найден: {lit_v4}")

print("\n" + "=" * 70)
print("ПРОВЕРКА УСТАНОВКИ БИБЛИОТЕК")
print("=" * 70)

try:
    from docx import Document
    print("✓ python-docx установлен")
    
    # Пробуем загрузить документ
    if v7_v5.exists():
        print(f"\nПробуем загрузить документ...")
        doc = Document(str(v7_v5))
        print(f"✓ Документ загружен успешно")
        print(f"  Параграфов: {len(doc.paragraphs)}")
        print(f"  Таблиц: {len(doc.tables)}")
        
        # Показываем первые несколько параграфов
        print(f"\nПервые 5 параграфов:")
        for i, para in enumerate(doc.paragraphs[:5], 1):
            text = para.text.strip()
            if text:
                print(f"  {i}. {text[:100]}...")
    
except ImportError:
    print("✗ python-docx не установлен")
    print("  Установите: pip install python-docx")
    sys.exit(1)
except Exception as e:
    print(f"✗ Ошибка: {e}")

print("\n" + "=" * 70)
print("ГОТОВО К АНАЛИЗУ")
print("=" * 70)
print("Запустите: python analyze_v7.py")




