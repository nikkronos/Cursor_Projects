# -*- coding: utf-8 -*-
"""
Улучшенная версия скрипта для исправления документа v7 -> v8
С более детальной обработкой и отладочным выводом
"""
import re
from pathlib import Path
from docx import Document

def load_literature_list(literature_file):
    """Загружает список литературы из файла"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {
            'number': num,
            'entry': entry
        }
    
    return literature

def fix_text_citations(text, literature, is_page_4=False):
    """Исправляет все проблемные ссылки в тексте"""
    original_text = text
    fixed_text = text
    
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    
    changes_made = []
    
    # 1. Исправляем фамилии ВНУТРИ квадратных скобок: [Андреева, 5] -> Андреева [5]
    # Паттерн: [Фамилия, номер] или [Фамилия1, номер1; Фамилия2, номер2]
    pattern_names_in_square = r'\[([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?(?:\s*,\s*\d+)?(?:\s*;\s*[А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?(?:\s*,\s*\d+)?)*)\]'
    
    def replace_names_in_square(match):
        content = match.group(1)
        
        # Парсим содержимое: "Андреева, 5; Эйдемиллер, Юстицкис, 64"
        # Разделяем по точкам с запятой
        parts = [p.strip() for p in content.split(';')]
        
        authors_with_numbers = []
        all_numbers = []
        
        for part in parts:
            # Парсим каждую часть: "Андреева, 5" или "Эйдемиллер, Юстицкис, 64" или "Адлер; 2"
            # Ищем последнее число в части
            number_match = re.search(r'(\d+)\s*$', part)
            if number_match:
                number = int(number_match.group(1))
                # Извлекаем фамилии (всё до последнего числа)
                authors_part = part[:number_match.start()].strip().rstrip(',;')
                authors = [a.strip() for a in authors_part.split(',')]
                # Убираем пустые и проверяем, что это фамилии
                authors = [a for a in authors if a and re.match(r'^[А-ЯЁ][а-яё]+', a)]
                
                if authors:
                    # Сохраняем все номера, даже если они вне диапазона (для страницы 4)
                    authors_with_numbers.append({
                        'authors': authors,
                        'number': number,
                        'is_valid': min_source <= number <= max_source
                    })
                    all_numbers.append(number)
        
        # Если нашли авторов с номерами, формируем правильное цитирование
        if authors_with_numbers:
            # Для страницы 4 - перефразируем согласно требованиям
            if is_page_4:
                # Для страницы 4 оставляем все номера (даже если вне диапазона), просто убираем фамилии из скобок
                result_parts = []
                
                for item in authors_with_numbers:
                    if len(item['authors']) == 1:
                        result_parts.append(f"{item['authors'][0]} [{item['number']}]")
                    else:
                        # Несколько авторов с одним номером: "Эйдемиллер, Юстицкис, 64" -> "Эйдемиллер, Юстицкис [64]"
                        authors_str = ', '.join(item['authors'])
                        result_parts.append(f"{authors_str} [{item['number']}]")
                
                # Формируем итоговый результат согласно примеру пользователя
                if len(result_parts) == 1:
                    result = result_parts[0]
                elif len(result_parts) == 2:
                    # "[Андреева, 5; Эйдемиллер, Юстицкис, 64]" -> "Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
                    first = result_parts[0]  # "Андреева [5]"
                    second = result_parts[1]  # "Эйдемиллер, Юстицкис [64]"
                    
                    # Извлекаем авторов и номер из второго элемента
                    second_match = re.match(r'^([А-ЯЁа-яё\s,]+)\s+\[(\d+)\]', second)
                    if second_match:
                        second_authors_str = second_match.group(1).strip()
                        second_number = second_match.group(2)
                        second_authors = [a.strip() for a in second_authors_str.split(',')]
                        
                        if len(second_authors) == 2:
                            # Два автора: "Эйдемиллер, Юстицкис" -> "Эйдемиллер. Также изучал Юскис [64]"
                            result = f"{first}, {second_authors[0]}. Также изучал {second_authors[1]} [{second_number}]"
                        else:
                            # Один автор
                            result = f"{first}, {second_authors[0]}. Также изучал {second[second.find('['):]}"
                    else:
                        result = f"{first}, {second}"
                else:
                    # Больше двух элементов - просто соединяем через запятую
                    result = ', '.join(result_parts)
                
                changes_made.append(f"Фамилии в квадратных скобках (страница 4): {match.group(0)} -> {result}")
                return result
            else:
                # Для остальных страниц - просто убираем фамилии, оставляем только валидные номера
                valid_numbers = [item['number'] for item in authors_with_numbers if item['is_valid']]
                if valid_numbers:
                    unique_numbers = sorted(set(valid_numbers))
                    result = f"[{', '.join(map(str, unique_numbers))}]"
                    changes_made.append(f"Фамилии в квадратных скобках удалены: {match.group(0)} -> {result}")
                    return result
                else:
                    # Все номера невалидные - удаляем ссылку
                    changes_made.append(f"Ссылка удалена (все номера неверные): {match.group(0)}")
                    return ""
        
        # Если не нашли авторов, возвращаем как есть (возможно, это просто номера)
        return match.group(0)
    
    fixed_text = re.sub(pattern_names_in_square, replace_names_in_square, fixed_text)
    
    # 2. Также исправляем фамилии в круглых скобках: (Андреева 5) -> Андреева [5]
    pattern_names_round = r'\(([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?(?:\s*,\s*[А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)?)\s+(\d+)\)'
    
    def replace_name_round(match):
        names_str = match.group(1).strip()
        source_num = match.group(2)
        
        # Исключаем ложные срабатывания
        exclude_words = ['вопрос', 'его', 'нее', 'него', 'ней']
        if any(word in names_str.lower() for word in exclude_words):
            return match.group(0)
        
        # Проверяем, что это похоже на фамилию
        if not re.match(r'^[А-ЯЁ][а-яё]+', names_str):
            return match.group(0)
        
        names = [n.strip() for n in names_str.split(',')]
        if len(names) == 1:
            result = f"{names[0]} [{source_num}]"
        else:
            result = f"{', '.join(names)} [{source_num}]"
        
        changes_made.append(f"Фамилия в круглых скобках: {match.group(0)} -> {result}")
        return result
    
    fixed_text = re.sub(pattern_names_round, replace_name_round, fixed_text)
    
    # 3. Исправляем неверные номера источников в квадратных скобках (только номера, без фамилий)
    pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
    
    def fix_numbers(match):
        numbers_str = match.group(1)
        valid_numbers = []
        invalid_numbers = []
        
        for num_str in re.split(r'[,\s;]+', numbers_str):
            num_str = num_str.strip()
            if '-' in num_str:
                try:
                    start, end = map(int, num_str.split('-'))
                    for num in range(start, end + 1):
                        if min_source <= num <= max_source:
                            valid_numbers.append(num)
                        else:
                            invalid_numbers.append(num)
                except ValueError:
                    pass
            else:
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                    else:
                        invalid_numbers.append(num)
                except ValueError:
                    pass
        
        if invalid_numbers:
            changes_made.append(f"Неверные номера удалены: {invalid_numbers}")
        
        if valid_numbers:
            unique_numbers = sorted(set(valid_numbers))
            result = f"[{', '.join(map(str, unique_numbers))}]"
            if result != match.group(0):
                changes_made.append(f"Номера исправлены: {match.group(0)} -> {result}")
            return result
        else:
            changes_made.append(f"Ссылка удалена (все номера неверные): {match.group(0)}")
            return ""
    
    fixed_text = re.sub(pattern_numbers, fix_numbers, fixed_text)
    
    # 4. Обработка страницы 4 - перефразирование текста с авторами
    if is_page_4:
        # Перефразируем "системный семейный подход Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
        # или "системный семейный подход [Андреева, 5; Эйдемиллер, Юстицкис, 64]"
        # Должно стать: "системный семейный подход, который изучали яркие представители Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
        
        # Паттерн для "подход Автор [номер], Автор2. Также изучал Автор3 [номер]"
        system_pattern1 = r'(системный семейный подход)\s+([А-ЯЁа-яё]+)\s+\[(\d+)\](?:\s*,\s*([А-ЯЁа-яё]+))?(?:\.\s*Также\s+изучал\s+([А-ЯЁа-яё]+)\s+\[(\d+)\])?'
        def rephrase_system1(match):
            approach = match.group(1)
            author1 = match.group(2)
            num1 = match.group(3)
            author2 = match.group(4)
            author3 = match.group(5)
            num3 = match.group(6)
            
            if author3:
                # Есть третий автор
                result = f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}. Также изучал {author3} [{num3}]"
            elif author2:
                # Есть второй автор
                result = f"{approach}, который изучали яркие представители {author1} [{num1}], {author2}"
            else:
                # Только один автор
                result = f"{approach}, который изучали яркие представители {author1} [{num1}]"
            
            changes_made.append(f"Перефразировано (системный подход): {match.group(0)} -> {result}")
            return result
        
        fixed_text = re.sub(system_pattern1, rephrase_system1, fixed_text, flags=re.IGNORECASE)
        
        # Перефразируем "подход Автор [номер]" -> "подход, который изучал Автор [номер]"
        approach_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,2})\s+подход\s+([А-ЯЁа-яё]+)\s+\[(\d+)\]'
        def rephrase_approach(match):
            approach_words = match.group(1)
            author = match.group(2)
            num = match.group(3)
            # Проверяем, не перефразировано ли уже
            if 'который изучал' not in fixed_text and 'который изучали' not in fixed_text:
                result = f"{approach_words} подход, который изучал {author} [{num}]"
                changes_made.append(f"Перефразировано (подход): {match.group(0)} -> {result}")
                return result
            return match.group(0)
        
        fixed_text = re.sub(approach_pattern, rephrase_approach, fixed_text, flags=re.IGNORECASE)
        
        # Перефразируем "теории Автор [номер]" -> "теории, которые изучал Автор [номер]"
        theory_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,2})\s+теории\s+([А-ЯЁа-яё]+)\s+\[(\d+)\]'
        def rephrase_theory(match):
            theory_words = match.group(1)
            author = match.group(2)
            num = match.group(3)
            if 'который изучал' not in fixed_text and 'который изучали' not in fixed_text:
                result = f"{theory_words} теории, которые изучал {author} [{num}]"
                changes_made.append(f"Перефразировано (теории): {match.group(0)} -> {result}")
                return result
            return match.group(0)
        
        fixed_text = re.sub(theory_pattern, rephrase_theory, fixed_text, flags=re.IGNORECASE)
    
    # Очищаем множественные пробелы
    fixed_text = re.sub(r'\s+', ' ', fixed_text)
    fixed_text = re.sub(r'\s*,\s*,', ',', fixed_text)
    
    return fixed_text, changes_made

def main():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    if not literature_path.exists():
        print(f"ERROR: {literature_path} не найден")
        return
    
    print("=" * 70)
    print("УЛУЧШЕННОЕ ИСПРАВЛЕНИЕ ДОКУМЕНТА v7 -> v8")
    print("=" * 70)
    
    # Загружаем документ
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен: {v7_path.name}")
    
    # Загружаем список литературы
    literature = load_literature_list(literature_path)
    print(f"✓ Список литературы загружен: {len(literature)} источников")
    print(f"   Диапазон: {min(literature.keys())}-{max(literature.keys())}\n")
    
    # Ключевые слова для страницы 4
    keywords_page_4 = ['системный семейный подход', 'андреева', 'эдельмин', 'юскис', 
                      'адлер', 'гинзбург', 'климов', 'прихожан', 'толстых']
    
    fixed_count = 0
    page_4_count = 0
    total_changes = []
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text_lower = original_text.lower()
        is_page_4 = any(keyword in text_lower for keyword in keywords_page_4)
        
        fixed_text, changes = fix_text_citations(original_text, literature, is_page_4)
        
        if changes:
            # Обновляем параграф
            para.clear()
            para.add_run(fixed_text)
            fixed_count += 1
            total_changes.extend(changes)
            
            if fixed_count <= 10:  # Показываем первые 10 для отладки
                print(f"\nПараграф {para_idx + 1} (страница 4: {is_page_4}):")
                print(f"  Исходный: {original_text[:100]}...")
                print(f"  Исправленный: {fixed_text[:100]}...")
                for change in changes:
                    print(f"    - {change}")
        
        if is_page_4:
            page_4_count += 1
    
    # Обрабатываем таблицы
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                fixed_text, changes = fix_text_citations(original_text, literature, False)
                
                if changes:
                    cell.text = fixed_text
                    fixed_count += 1
                    total_changes.extend(changes)
    
    # Сохраняем документ
    doc.save(str(v8_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v8_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    print(f"✓ Обработано параграфов страницы 4: {page_4_count}")
    print(f"✓ Всего изменений: {len(total_changes)}")
    
    if total_changes:
        print(f"\nТипы изменений:")
        change_types = {}
        for change in total_changes:
            change_type = change.split(':')[0]
            change_types[change_type] = change_types.get(change_type, 0) + 1
        for change_type, count in change_types.items():
            print(f"  {change_type}: {count}")

if __name__ == '__main__':
    main()

