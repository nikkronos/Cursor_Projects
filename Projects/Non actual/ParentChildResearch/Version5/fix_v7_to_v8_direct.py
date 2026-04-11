# -*- coding: utf-8 -*-
"""
Прямая замена проблемных мест - как в Version4
Использует конкретные замены для каждого проблемного места
"""
import re
from pathlib import Path
from docx import Document

def load_literature_list(literature_file):
    """Загружает список литературы"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
    
    return literature

def main():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("=" * 70)
    print("ПРЯМАЯ ЗАМЕНА ПРОБЛЕМНЫХ МЕСТ")
    print("=" * 70)
    
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен")
    
    literature = load_literature_list(literature_path)
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    print(f"✓ Список литературы: {len(literature)} источников\n")
    
    fixed_count = 0
    page_4_count = 0
    
    # Список замен для страницы 4
    page_4_replacements = [
        # Параграф 46
        {
            'search': '[Андреева, 5; Эйдемиллер, Юстицкис, 64]',
            'replace': 'Андреева [5], Эйдемиллер. Также изучал Юскис [64]',
            'description': 'Параграф 46: Системный семейный подход'
        },
        # Параграф 47
        {
            'search': '[Адлер; 2]',
            'replace': 'Адлер [2]',
            'description': 'Параграф 47: Психологические теории'
        },
        # Параграф 48
        {
            'search': '[Гинзбург, 18; Климов, 30; Прихожан, Толстых, 49]',
            'replace': 'Гинзбург [18], Климов [30], Прихожан, Толстых [49]',
            'description': 'Параграф 48: Теории профессионального самоопределения'
        },
    ]
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        text_lower = text.lower()
        
        # Проверяем, относится ли к странице 4
        is_page_4 = any(keyword in text_lower for keyword in [
            'системный семейный подход', 'андреева', 'эдельмин', 'юскис', 
            'адлер', 'гинзбург', 'климов', 'прихожан', 'толстых'
        ])
        
        # Замены для страницы 4
        if is_page_4:
            for replacement in page_4_replacements:
                if replacement['search'] in text:
                    text = text.replace(replacement['search'], replacement['replace'])
                    print(f"✓ {replacement['description']}: {replacement['search']} -> {replacement['replace']}")
            
            # Перефразируем текст
            # "Системный семейный подход Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
            # -> "Системный семейный подход, который изучали яркие представители Андреева [5], Эйдемиллер. Также изучал Юскис [64]"
            if 'системный семейный подход' in text_lower and 'который изучали яркие представители' not in text_lower:
                # Ищем паттерн после замены
                pattern = r'(системный семейный подход)\s+(Андреева\s+\[5\],\s+Эйдемиллер\.\s+Также\s+изучал\s+Юскис\s+\[64\])'
                replacement = r'\1, который изучали яркие представители \2'
                if re.search(pattern, text, re.IGNORECASE):
                    text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
                    print(f"✓ Перефразировано: системный семейный подход")
            
            # "Психологические теории развития и сценариев Адлер [2]"
            # -> "Психологические теории развития и сценариев, которые изучал Адлер [2]"
            if 'психологические теории развития и сценариев' in text_lower and 'который изучал' not in text_lower:
                pattern = r'(психологические теории развития и сценариев)\s+(Адлер\s+\[2\])'
                replacement = r'\1, которые изучал \2'
                if re.search(pattern, text, re.IGNORECASE):
                    text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
                    print(f"✓ Перефразировано: психологические теории")
            
            # "Теории профессионального самоопределения Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
            # -> "Теории профессионального самоопределения, которые изучали Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
            if 'теории профессионального самоопределения' in text_lower and 'который изучал' not in text_lower and 'который изучали' not in text_lower:
                pattern = r'(теории профессионального самоопределения)\s+(Гинзбург\s+\[18\],\s+Климов\s+\[30\],\s+Прихожан,\s+Толстых\s+\[49\])'
                replacement = r'\1, которые изучали \2'
                if re.search(pattern, text, re.IGNORECASE):
                    text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
                    print(f"✓ Перефразировано: теории профессионального самоопределения")
            
            page_4_count += 1
        
        # Убираем фамилии из скобок в остальных местах
        # Паттерн: [Фамилия, номер] или [Фамилия1, номер1; Фамилия2, номер2]
        pattern_names = r'\[([А-ЯЁ][А-ЯЁа-яё\s,\-;]+?)\]'
        
        def remove_names(match):
            content = match.group(1)
            # Проверяем, есть ли фамилии
            if re.search(r'[А-ЯЁ][а-яё]+', content):
                # Парсим и извлекаем номера
                parts = [p.strip() for p in content.split(';')]
                numbers = []
                for part in parts:
                    number_match = re.search(r'(\d+)\s*$', part)
                    if number_match:
                        numbers.append(int(number_match.group(1)))
                if numbers:
                    unique_numbers = sorted(set(numbers))
                    # Фильтруем неверные номера
                    valid_numbers = [n for n in unique_numbers if min_source <= n <= max_source]
                    if valid_numbers:
                        return f"[{', '.join(map(str, valid_numbers))}]"
                    else:
                        return ""  # Удаляем, если все номера неверные
            return match.group(0)
        
        text = re.sub(pattern_names, remove_names, text)
        
        # Исправляем неверные номера в обычных ссылках
        pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
        
        def fix_numbers(match):
            numbers_str = match.group(1)
            valid_numbers = []
            for num_str in re.split(r'[,\s;]+', numbers_str):
                try:
                    num = int(num_str.strip())
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                except ValueError:
                    pass
            if valid_numbers:
                return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
            return ""
        
        text = re.sub(pattern_numbers, fix_numbers, text)
        
        # Применяем изменения
        if text != original_text:
            para.text = text
            fixed_count += 1
            
            if fixed_count <= 20:
                print(f"\nПараграф {para_idx + 1} (страница 4: {is_page_4}):")
                print(f"  Было: {original_text[:80]}...")
                print(f"  Стало: {text[:80]}...")
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                # Убираем фамилии
                pattern_names = r'\[([А-ЯЁ][А-ЯЁа-яё\s,\-;]+?)\]'
                def remove_names(match):
                    content = match.group(1)
                    if re.search(r'[А-ЯЁ][а-яё]+', content):
                        parts = [p.strip() for p in content.split(';')]
                        numbers = []
                        for part in parts:
                            number_match = re.search(r'(\d+)\s*$', part)
                            if number_match:
                                num = int(number_match.group(1))
                                if min_source <= num <= max_source:
                                    numbers.append(num)
                        if numbers:
                            return f"[{', '.join(map(str, sorted(set(numbers))))}]"
                    return match.group(0)
                
                text = re.sub(pattern_names, remove_names, text)
                
                # Исправляем неверные номера
                pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
                def fix_numbers(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v8_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v8_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    print(f"✓ Обработано параграфов страницы 4: {page_4_count}")

if __name__ == '__main__':
    main()




