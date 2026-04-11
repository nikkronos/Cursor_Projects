# -*- coding: utf-8 -*-
"""
Скрипт для исправления проблем в документе v7 и создания v8
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def load_literature_list(literature_file):
    """Загружает список литературы из файла"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Парсим список литературы (формат: номер. Автор...)
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        # Извлекаем фамилию автора
        author_match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
        if author_match:
            author = author_match.group(1)
        else:
            author_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)', entry)
            if author_match:
                author = author_match.group(1)
            else:
                author = None
        literature[num] = {
            'number': num,
            'entry': entry,
            'author': author
        }
    
    return literature

def find_name_in_literature(name, literature):
    """Находит номер источника по фамилии автора"""
    name_lower = name.lower().strip()
    for num, info in literature.items():
        if info['author'] and info['author'].lower() == name_lower:
            return num
    return None

def fix_page_4_text(text, literature):
    """Исправляет текст на странице 4 - убирает фамилии из скобок и перефразирует"""
    
    fixed_text = text
    
    # Сначала убираем фамилии из скобок: (Андреева 5) -> Андреева [5]
    # Но исключаем случаи типа "(вопрос 12)" - это не ссылки на источники
    simple_pattern = r'\(([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?(?:\s*,\s*[А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)?)\s+(\d+)\)'
    def replace_simple(match):
        names_str = match.group(1).strip()
        num = match.group(2)
        
        # Проверяем, что это действительно фамилия
        exclude_words = ['вопрос', 'его', 'нее', 'него', 'ней']
        if any(word in names_str.lower() for word in exclude_words):
            return match.group(0)  # Оставляем как есть
        
        # Проверяем, что это похоже на фамилию
        if not re.match(r'^[А-ЯЁ][а-яё]+', names_str):
            return match.group(0)
        
        names = [n.strip() for n in names_str.split(',')]
        if len(names) == 1:
            return f"{names[0]} [{num}]"
        else:
            # Несколько авторов: (Эдельмин, Юскис 64) -> Эдельмин, Юскис [64]
            return f"{', '.join(names)} [{num}]"
    
    fixed_text = re.sub(simple_pattern, replace_simple, fixed_text)
    
    # Теперь перефразируем текст согласно требованиям:
    # "системный семейный подход (Андреева 5)" -> 
    # "системный семейный подход, который изучали яркие представители Андреева [5]"
    
    # Ищем "системный семейный подход" с авторами после него
    system_pattern = r'(системный семейный подход)(?:\s*,\s*который\s+изучал[и]?\s+)?(?:яркие\s+представители\s+)?([А-ЯЁа-яё]+(?:\s*,\s*[А-ЯЁа-яё]+)*)\s+\[(\d+)\]'
    def rephrase_system(match):
        approach = match.group(1)
        authors_str = match.group(2)
        num = match.group(3)
        authors = [a.strip() for a in authors_str.split(',')]
        
        if len(authors) == 1:
            return f"{approach}, который изучали яркие представители {authors[0]} [{num}]"
        else:
            # Несколько авторов - формируем список
            authors_list = ', '.join(authors[:-1]) + f". Также изучал {authors[-1]}"
            return f"{approach}, который изучали яркие представители {authors_list} [{num}]"
    
    # Применяем только если еще не перефразировано
    if 'который изучали яркие представители' not in fixed_text:
        fixed_text = re.sub(system_pattern, rephrase_system, fixed_text, flags=re.IGNORECASE)
    
    # Обрабатываем другие подходы и теории
    # "подход Андреева [5]" -> "подход, который изучал Андреева [5]"
    approach_pattern = r'([а-яё]+(?:\s+[а-яё]+){0,3})\s+([А-ЯЁа-яё]+)\s+\[(\d+)\]'
    def rephrase_approach(match):
        approach_words = match.group(1)
        author = match.group(2)
        num = match.group(3)
        # Проверяем, не перефразировано ли уже
        if 'который изучал' in fixed_text or 'который изучали' in fixed_text:
            return match.group(0)
        return f"{approach_words}, который изучал {author} [{num}]"
    
    # Применяем только если это не было уже обработано
    if 'который изучал' not in fixed_text and 'который изучали' not in fixed_text:
        fixed_text = re.sub(approach_pattern, rephrase_approach, fixed_text, flags=re.IGNORECASE)
    
    return fixed_text

def fix_citations_in_text(text, literature):
    """Исправляет все проблемные ссылки в тексте"""
    fixed_text = text
    
    # 1. Исправляем фамилии в скобках: (Андреева 5) -> Андреева [5]
    # Паттерн: (Фамилия номер) или (Фамилия1, Фамилия2 номер)
    # Исключаем: (вопрос 12), (его), (нее) и т.д.
    pattern1 = r'\(([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?(?:\s*,\s*[А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)?)\s+(\d+)\)'
    def fix_name_brackets(match):
        names_str = match.group(1).strip()
        source_num = match.group(2)
        
        # Проверяем, что это действительно фамилия, а не другие слова
        exclude_words = ['вопрос', 'его', 'нее', 'него', 'ней', 'него', 'ней', 'него', 'ней']
        if any(word in names_str.lower() for word in exclude_words):
            return match.group(0)  # Оставляем как есть
        
        # Проверяем, что это похоже на фамилию (начинается с заглавной, содержит только буквы)
        if not re.match(r'^[А-ЯЁ][а-яё]+', names_str):
            return match.group(0)
        
        names = [n.strip() for n in names_str.split(',')]
        if len(names) == 1:
            return f"{names[0]} [{source_num}]"
        else:
            return f"{', '.join(names)} [{source_num}]"
    
    fixed_text = re.sub(pattern1, fix_name_brackets, fixed_text)
    
    # 2. Проверяем и исправляем номера источников, которые выходят за пределы списка
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    
    # Находим все ссылки с номерами в квадратных скобках
    pattern2 = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
    def fix_source_numbers(match):
        numbers_str = match.group(1)
        valid_numbers = []
        for num_str in re.split(r'[,\s;]+', numbers_str):
            num_str = num_str.strip()
            if '-' in num_str:
                # Диапазон: 1-5
                try:
                    start, end = map(int, num_str.split('-'))
                    for num in range(start, end + 1):
                        if min_source <= num <= max_source:
                            valid_numbers.append(num)
                except ValueError:
                    pass
            else:
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                except ValueError:
                    pass
        
        if valid_numbers:
            # Убираем дубликаты и сортируем
            unique_numbers = sorted(set(valid_numbers))
            return f"[{', '.join(map(str, unique_numbers))}]"
        else:
            # Если все номера неверные, удаляем ссылку полностью
            return ""
    
    fixed_text = re.sub(pattern2, fix_source_numbers, fixed_text)
    
    # 3. Удаляем оставшиеся пустые скобки и лишние запятые/точки с запятой
    fixed_text = re.sub(r'\s*,\s*,', ',', fixed_text)  # Двойные запятые
    fixed_text = re.sub(r'\s*;\s*;', ';', fixed_text)  # Двойные точки с запятой
    fixed_text = re.sub(r'\s+', ' ', fixed_text)  # Множественные пробелы
    
    return fixed_text

def process_document(v7_path, v8_path, literature_path):
    """Обрабатывает документ и создает исправленную версию v8"""
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v7 -> v8")
    print("=" * 70)
    
    # Загружаем документ
    doc = Document(str(v7_path))
    print(f"✓ Документ загружен: {v7_path.name}")
    
    # Загружаем список литературы
    literature = load_literature_list(literature_path)
    print(f"✓ Список литературы загружен: {len(literature)} источников")
    
    # Обрабатываем параграфы
    fixed_count = 0
    page_4_count = 0
    
    keywords_page_4 = ['системный семейный подход', 'андреева', 'эдельмин', 'юскис', 
                       'адлер', 'гинзбург', 'климов', 'прихожан', 'толстых']
    
    for para in doc.paragraphs:
        original_text = para.text
        if not original_text.strip():
            continue
        
        # Проверяем, относится ли это к странице 4
        text_lower = original_text.lower()
        is_page_4 = any(keyword in text_lower for keyword in keywords_page_4)
        
        # Исправляем текст
        fixed_text = fix_citations_in_text(original_text, literature)
        
        # Если это страница 4, применяем специальную обработку
        if is_page_4:
            fixed_text = fix_page_4_text(fixed_text, literature)
            page_4_count += 1
        
        # Если текст изменился, обновляем параграф
        if fixed_text != original_text:
            # Очищаем параграф и добавляем новый текст
            para.clear()
            para.add_run(fixed_text)
            fixed_count += 1
            # Показываем что исправлено (первые 5 для отладки)
            if fixed_count <= 5:
                print(f"  Исправлен параграф {fixed_count}: {original_text[:50]}... -> {fixed_text[:50]}...")
    
    if page_4_count > 0:
        print(f"✓ Обработано параграфов страницы 4: {page_4_count}")
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                fixed_text = fix_citations_in_text(original_text, literature)
                
                if fixed_text != original_text:
                    cell.text = fixed_text
                    fixed_count += 1
    
    # Сохраняем документ v8
    doc.save(str(v8_path))
    print(f"\n✓ Документ сохранен: {v8_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    
    return fixed_count

def main():
    version5 = Path(__file__).parent
    version4 = version5.parent / "Version4"
    
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    v8_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v8.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    # Копируем файлы если нужно
    if not v7_path.exists():
        v7_source = version4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
        if v7_source.exists():
            import shutil
            shutil.copy2(v7_source, v7_path)
            print(f"✓ Скопирован документ v7")
        else:
            print(f"ERROR: {v7_source} не найден")
            return
    
    if not literature_path.exists():
        lit_source = version4 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
        if lit_source.exists():
            import shutil
            shutil.copy2(lit_source, literature_path)
            print(f"✓ Скопирован список литературы")
        else:
            print(f"ERROR: {lit_source} не найден")
            return
    
    # Обрабатываем документ
    fixed_count = process_document(v7_path, v8_path, literature_path)
    
    print(f"\n{'='*70}")
    print("ИСПРАВЛЕНИЕ ЗАВЕРШЕНО")
    print(f"{'='*70}")
    print(f"Создан документ: {v8_path.name}")
    print(f"Исправлено элементов: {fixed_count}")

if __name__ == '__main__':
    main()

