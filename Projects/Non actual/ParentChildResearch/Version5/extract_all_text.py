# -*- coding: utf-8 -*-
"""
Скрипт для извлечения всего текста из документа в текстовый файл
"""
from pathlib import Path
from docx import Document

def extract_all_text():
    version5 = Path(__file__).parent
    v7_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_path.exists():
        print(f"ERROR: {v7_path} не найден")
        return
    
    print("Извлекаю текст из документа...")
    doc = Document(str(v7_path))
    
    all_text = []
    
    # Извлекаем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append(f"=== ПАРАГРАФ {para_idx + 1} ===\n{text}\n")
    
    # Извлекаем таблицы
    for table_idx, table in enumerate(doc.tables):
        all_text.append(f"\n=== ТАБЛИЦА {table_idx + 1} ===\n")
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                all_text.append(" | ".join(row_text) + "\n")
    
    # Сохраняем в файл
    output_file = version5 / "ВСЕ_ТЕКСТ_ИЗ_ДОКУМЕНТА.txt"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(all_text))
    
    print(f"✓ Текст сохранен в: {output_file.name}")
    print(f"✓ Всего параграфов: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"✓ Всего таблиц: {len(doc.tables)}")
    
    # Также создаем файл только с проблемными местами
    print("\nИщу проблемные места...")
    problems = []
    
    import re
    pattern = r'\[([А-ЯЁ][а-яё]+[^\]]*)\]'
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        matches = re.finditer(pattern, text)
        for match in matches:
            content = match.group(1)
            # Проверяем, есть ли фамилии (начинаются с заглавной буквы)
            if re.search(r'[А-ЯЁ][а-яё]+', content):
                problems.append({
                    'para': para_idx + 1,
                    'text': text[:200] + '...' if len(text) > 200 else text,
                    'match': match.group(0)
                })
    
    if problems:
        problems_file = version5 / "ПРОБЛЕМНЫЕ_МЕСТА.txt"
        with open(problems_file, 'w', encoding='utf-8') as f:
            f.write("ПРОБЛЕМНЫЕ МЕСТА С ФАМИЛИЯМИ В СКОБКАХ\n")
            f.write("=" * 70 + "\n\n")
            for i, problem in enumerate(problems, 1):
                f.write(f"{i}. Параграф {problem['para']}\n")
                f.write(f"   Найдено: {problem['match']}\n")
                f.write(f"   Текст: {problem['text']}\n\n")
        
        print(f"✓ Проблемные места сохранены в: {problems_file.name}")
        print(f"✓ Найдено проблемных мест: {len(problems)}")
    else:
        print("✓ Проблемных мест не найдено")

if __name__ == '__main__':
    extract_all_text()




