"""
Создаёт Word-файл с таблицей координат (№, X, Y) и готовой диаграммой Рисунка 3
для копирования в работу. Запуск: python create_figure3_word.py
"""
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    defence_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(defence_dir)
    txt_path = os.path.join(defence_dir, "Данные_для_таблицы_Рисунок3.txt")
    img_path = os.path.join(project_root, "Final3", "graphs", "figure3_ovcharova_holland_scatter.png")

    # Читаем данные из существующего txt (если есть)
    rows_data = []  # список (№, X, Y)
    if os.path.exists(txt_path):
        with open(txt_path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("ДАННЫЕ") or line.startswith("=") or line.startswith("Корреляция") or line.startswith("ρ") or line.startswith("Количество") or line.startswith("Таблица") or line.startswith("-"):
                    continue
                parts = line.split("\t")
                if len(parts) >= 3:
                    try:
                        num = parts[0].strip()
                        x_val = parts[1].strip()
                        y_val = parts[2].strip()
                        if num.isdigit() and x_val.replace(".", "").replace(",", "").replace("-", "").isdigit() and y_val.replace(".", "").replace(",", "").replace("-", "").isdigit():
                            rows_data.append((num, x_val, y_val))
                    except Exception:
                        continue
                elif len(parts) == 2 and line[0].isdigit():
                    try:
                        num = str(len(rows_data) + 1)
                        x_val = parts[0].strip()
                        y_val = parts[1].strip()
                        rows_data.append((num, x_val, y_val))
                    except Exception:
                        continue
    else:
        # Данных нет — нужен export_figure3_table_data.py
        print("Сначала запустите: python export_figure3_table_data.py")
        return

    doc = Document()
    doc.add_heading("Рисунок 3. Таблица и диаграмма", level=0)

    # Трактовка таблицы: как её читать
    doc.add_heading("Как читать таблицу координат", level=1)
    interpretation = """
Таблица содержит координаты точек для точечной диаграммы (scatter plot) корреляции между мотивом выбора профессии и типом профессиональных интересов.

Структура таблицы:
• Столбец «№» — порядковый номер наблюдения (от 1 до 50, соответствует одному респонденту-подростку).
• Столбец «X» — балл по мотиву выбора профессии (опросник Овчаровой, вопрос 19: «Позволяет использовать профессиональные умения вне работы»). Значения от 1 до 5 (шкала оценки важности мотива).
• Столбец «Y» — балл по типу профессиональных интересов (опросник Холланда, тип «Предприимчивый (Enterprising)»). Значения от 2 до 13 (сумма баллов по вопросам этого типа).

Как использовать таблицу:
1. Каждая строка — это одна точка на диаграмме.
2. Значение X (столбец 2) определяет горизонтальное положение точки.
3. Значение Y (столбец 3) определяет вертикальное положение точки.
4. Например, строка с №=1 (X=3.0, Y=11.0) означает: респондент №1 оценил важность мотива как 3 балла, а его балл по типу «Предприимчивый» равен 11.

Интерпретация корреляции:
• Положительная корреляция (ρ = 0.283, p = 0.047): чем выше важность мотива «использовать умения вне работы», тем выше балл по типу «Предприимчивый».
• На диаграмме это видно как восходящая линия тренда (красная пунктирная линия).
"""
    doc.add_paragraph(interpretation)
    doc.add_paragraph()

    # Таблица с короткими подписями: только №, X, Y (без длинного текста)
    doc.add_heading("Таблица координат", level=1)
    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "№"
    h[1].text = "X"
    h[2].text = "Y"
    for i, (num, x_val, y_val) in enumerate(rows_data):
        row = table.rows[i + 1].cells
        row[0].text = num
        row[1].text = x_val
        row[2].text = y_val
    # Выравнивание по центру для заголовков и чисел
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_heading("Диаграмма (точечный график)", level=1)
    diagram_desc = """
Диаграмма показывает связь между баллами по мотиву (ось X) и баллами по типу Холланда (ось Y).

Как читать диаграмму:
• По горизонтальной оси (X) отложены значения из столбца «X» таблицы (баллы по мотиву, от 1 до 5).
• По вертикальной оси (Y) отложены значения из столбца «Y» таблицы (баллы по типу Холланда, от 2 до 13).
• Каждая точка соответствует одной строке таблицы (одному респонденту).
• Красная пунктирная линия — линия тренда, показывает общее направление связи.
• В левом верхнем углу указаны: ρ (коэффициент корреляции Спирмена) и p (уровень значимости).

Числовые метки на шкалах позволяют точно определить координаты каждой точки.
"""
    doc.add_paragraph(diagram_desc)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
    else:
        doc.add_paragraph(f"[Файл диаграммы не найден: {img_path}]")

    out_docx = os.path.join(defence_dir, "Рисунок3_таблица_и_диаграмма.docx")
    doc.save(out_docx)
    print(f"Создан файл: {out_docx}")
    print("Откройте его, скопируйте таблицу или диаграмму и вставьте в работу.")

if __name__ == "__main__":
    main()
