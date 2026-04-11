"""
Экспорт сырых данных для значимых корреляций из анализа в Final3/enhanced_analysis.py.

Скрипт:
- повторно запускает enhanced_analysis.py через runpy.run_path
- получает списки значимых корреляций:
  - significant_ovcharova_markovskaya
  - significant_holland
  - significant_ovcharova_holland
- создаёт отдельный документ Word с сырыми значениями, из которых
  считались M, SD, Me и корреляции в «Подробный_анализ_корреляций.docx».
"""

import os
import runpy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np


def calculate_stats(values):
    """
    Локальный расчёт статистики (аналог calculate_statistics из enhanced_analysis.py).
    Возвращает словарь с M, SD, Me, Min, Max и N.
    """
    if not values:
        return {"M": None, "SD": None, "Me": None, "Min": None, "Max": None, "N": 0}

    arr = np.array(values, dtype=float)
    n = len(arr)
    mean = float(np.mean(arr))
    sd = float(np.std(arr, ddof=1)) if n > 1 else 0.0
    med = float(np.median(arr))
    vmin = float(np.min(arr))
    vmax = float(np.max(arr))
    return {"M": mean, "SD": sd, "Me": med, "Min": vmin, "Max": vmax, "N": n}


def add_raw_table(
    doc,
    title,
    left_label,
    right_label,
    left_values,
    right_values,
    corr_value=None,
    p_value=None,
):
    """
    Добавляет в документ таблицу сырых значений для одной корреляции.
    left_values и right_values — это списки чисел одинаковой длины.
    """
    n = min(len(left_values), len(right_values))
    if n == 0:
        return

    doc.add_heading(title, level=3)

    # Основная таблица с парами наблюдений
    table = doc.add_table(rows=n + 1, cols=3)
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "№ наблюдения"
    hdr[1].text = left_label
    hdr[2].text = right_label
    for cell in hdr:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True

    for i in range(n):
        row = table.rows[i + 1].cells
        row[0].text = str(i + 1)
        row[1].text = f"{left_values[i]}"
        row[2].text = f"{right_values[i]}"
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Итоговая таблица со статистикой под сырой таблицей
    stats_left = calculate_stats(left_values)
    stats_right = calculate_stats(right_values)

    stats_table = doc.add_table(rows=1 + 5, cols=3)
    stats_table.style = "Light Grid Accent 1"

    sh = stats_table.rows[0].cells
    sh[0].text = "Показатель"
    sh[1].text = left_label
    sh[2].text = right_label
    for cell in sh:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.bold = True

    def _fmt(val, nd=2):
        return "N/A" if val is None else f"{val:.{nd}f}"

    rows_data = [
        ("N (кол-во наблюдений)", stats_left["N"], stats_right["N"]),
        ("M (среднее)", stats_left["M"], stats_right["M"]),
        ("SD (стандартное отклонение)", stats_left["SD"], stats_right["SD"]),
        ("Me (медиана)", stats_left["Me"], stats_right["Me"]),
        ("Min / Max", None, None),
    ]

    for idx, (name, lv, rv) in enumerate(rows_data, 1):
        row = stats_table.rows[idx].cells
        row[0].text = name
        if name == "Min / Max":
            row[1].text = (
                f"{_fmt(stats_left['Min'])} / {_fmt(stats_left['Max'])}"
                if stats_left["N"] > 0
                else "N/A"
            )
            row[2].text = (
                f"{_fmt(stats_right['Min'])} / {_fmt(stats_right['Max'])}"
                if stats_right["N"] > 0
                else "N/A"
            )
        else:
            row[1].text = str(lv) if name.startswith("N") else _fmt(lv)
            row[2].text = str(rv) if name.startswith("N") else _fmt(rv)
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Добавляем строки про коэффициент корреляции и уровень значимости
    if corr_value is not None:
        row = stats_table.add_row().cells
        row[0].text = "Коэффициент корреляции (Спирмен, ρ)"
        row[1].text = f"{corr_value:.4f}"
        row[2].text = "—"
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if p_value is not None:
        row = stats_table.add_row().cells
        row[0].text = "Уровень значимости (p)"
        # Краткая интерпретация p
        if p_value < 0.001:
            interp = "< 0.001 (очень высокая значимость)"
        elif p_value < 0.01:
            interp = "< 0.01 (высокая значимость)"
        elif p_value < 0.05:
            interp = "< 0.05 (значимая связь)"
        else:
            interp = f"= {p_value:.4f} (не значимо)"
        row[1].text = f"{p_value:.4f} ({interp})"
        row[2].text = "—"
        for cell in row:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()


def main():
    # Базовые пути
    defence_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(defence_dir)
    final3_dir = os.path.join(project_root, "Final3")
    enhanced_path = os.path.join(final3_dir, "enhanced_analysis.py")

    if not os.path.exists(enhanced_path):
        print(f"ОШИБКА: не найден enhanced_analysis.py по пути: {enhanced_path}")
        return

    print("1. Запуск enhanced_analysis.py через runpy...")
    # Запускаем скрипт в отдельном namespace и получаем все его глобальные переменные
    globals_dict = runpy.run_path(enhanced_path)

    # Извлекаем списки значимых корреляций
    significant_ovcharova_markovskaya = globals_dict.get(
        "significant_ovcharova_markovskaya", []
    )
    significant_holland = globals_dict.get("significant_holland", [])
    significant_ovcharova_holland = globals_dict.get(
        "significant_ovcharova_holland", []
    )

    print(
        f"   Найдено значимых корреляций: "
        f"Овчарова-Марковская = {len(significant_ovcharova_markovskaya)}, "
        f"Марковская-Холланд = {len(significant_holland)}, "
        f"Овчарова-Холланд = {len(significant_ovcharova_holland)}"
    )

    # Создаём новый документ с сырыми данными
    doc = Document()
    title = doc.add_heading(
        "СЫРЫЕ ДАННЫЕ ДЛЯ ЗНАЧИМЫХ КОРРЕЛЯЦИЙ", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = (
        "В этом документе перечислены исходные (сырые) значения по шкалам, "
        "из которых рассчитывались среднее (M), стандартное отклонение (SD), "
        "медиана (Me) и коэффициенты корреляции, представленные в документе "
        "«Подробный_анализ_корреляций.docx».\n\n"
        "Для каждой значимой связи приведена таблица с парами наблюдений: "
        "номер наблюдения и значения двух переменных."
    )
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(12)

    # 1. Корреляции Овчарова-Марковская (все значимые, как в разделе 3.1)
    if significant_ovcharova_markovskaya:
        doc.add_heading(
            "1. Сырые данные: корреляции между мотивами Овчаровой и установками Марковской",
            level=1,
        )
        sorted_corrs = sorted(
            significant_ovcharova_markovskaya,
            key=lambda x: abs(x.get("spearman_corr", 0)),
            reverse=True,
        )
        for idx, corr in enumerate(sorted_corrs, 1):
            parent_q = corr.get("parent_question", "")[:120]
            ov_q = corr.get("ovcharova_question", "")
            rho = corr.get("spearman_corr", None)
            p_val = corr.get("spearman_p", None)
            n = corr.get("n", len(corr.get("parent_values", [])))

            header = (
                f"1.{idx}. {parent_q} — {ov_q} "
                f"(ρ = {rho:.4f}, p = {p_val:.4f}, N = {n})"
                if rho is not None and p_val is not None
                else f"1.{idx}. {parent_q} — {ov_q}"
            )

            add_raw_table(
                doc,
                header,
                "Баллы по установке (Марковская)",
                "Баллы по мотиву (Овчарова)",
                corr.get("parent_values", []),
                corr.get("ovcharova_values", []),
                corr_value=rho,
                p_value=p_val,
            )

    # 2. Корреляции Марковская-Холланд (топ-10, как в разделе 3.2)
    if significant_holland:
        doc.add_heading(
            "2. Сырые данные: корреляции между установками Марковской и типами Холланда",
            level=1,
        )
        sorted_corrs = sorted(
            significant_holland,
            key=lambda x: abs(x.get("spearman_corr", 0)),
            reverse=True,
        )[:10]

        for idx, corr in enumerate(sorted_corrs, 1):
            parent_q = corr.get("parent_question", "")[:120]
            h_type_name = corr.get("holland_type_name", "")
            rho = corr.get("spearman_corr", None)
            p_val = corr.get("spearman_p", None)
            n = corr.get("n", len(corr.get("parent_values", [])))

            header = (
                f"2.{idx}. {parent_q} — {h_type_name} "
                f"(ρ = {rho:.4f}, p = {p_val:.4f}, N = {n})"
                if rho is not None and p_val is not None
                else f"2.{idx}. {parent_q} — {h_type_name}"
            )

            add_raw_table(
                doc,
                header,
                "Баллы по установке (Марковская)",
                "Баллы по типу Холланда",
                corr.get("parent_values", []),
                corr.get("holland_values", []),
                corr_value=rho,
                p_value=p_val,
            )

    # 3. Корреляции Овчарова-Холланд (топ-10, как в разделе 3.3)
    if significant_ovcharova_holland:
        doc.add_heading(
            "3. Сырые данные: корреляции между мотивами Овчаровой и типами Холланда",
            level=1,
        )
        sorted_corrs = sorted(
            significant_ovcharova_holland,
            key=lambda x: abs(x.get("spearman_corr", 0)),
            reverse=True,
        )[:10]

        for idx, corr in enumerate(sorted_corrs, 1):
            ov_q = corr.get("ovcharova_question", "")
            h_type_name = corr.get("holland_type_name", "")
            rho = corr.get("spearman_corr", None)
            p_val = corr.get("spearman_p", None)
            n = corr.get("n", len(corr.get("ovcharova_values", [])))

            header = (
                f"3.{idx}. {ov_q} — {h_type_name} "
                f"(ρ = {rho:.4f}, p = {p_val:.4f}, N = {n})"
                if rho is not None and p_val is not None
                else f"3.{idx}. {ov_q} — {h_type_name}"
            )

            add_raw_table(
                doc,
                header,
                "Баллы по мотиву (Овчарова)",
                "Баллы по типу Холланда",
                corr.get("ovcharova_values", []),
                corr.get("holland_values", []),
                corr_value=rho,
                p_value=p_val,
            )

    # Сохраняем документ с сырыми данными в папке Defence1
    output_path = os.path.join(
        defence_dir, "Сырые_данные_значимых_корреляций.docx"
    )
    doc.save(output_path)
    print(f"\nДокумент с сырыми данными сохранён: {output_path}")

    # Дополнительно копируем основной документ с подробным анализом в Defence1
    detailed_doc_path = globals_dict.get(
        "output_file",
        os.path.join(final3_dir, "Подробный_анализ_корреляций.docx"),
    )
    if os.path.exists(detailed_doc_path):
        import shutil

        dst = os.path.join(defence_dir, "Подробный_анализ_корреляций.docx")
        shutil.copy2(detailed_doc_path, dst)
        print(f"Копия подробного анализа сохранена: {dst}")


if __name__ == "__main__":
    main()

