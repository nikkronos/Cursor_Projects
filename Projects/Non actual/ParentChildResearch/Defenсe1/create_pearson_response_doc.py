"""
Формирует документ «Ответ комиссии» по использованию коэффициента Пирсона в исследовании:
- ответы на типичные вопросы комиссии;
- полная таблица всех мест, где рассчитывался Пирсон, с значениями r и p.
Запуск: из папки Defenсe1 выполнить «python create_pearson_response_doc.py».
"""

import os
import runpy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


def main():
    defence_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(defence_dir)
    final3_dir = os.path.join(project_root, "Final3")
    enhanced_path = os.path.join(final3_dir, "enhanced_analysis.py")

    if not os.path.exists(enhanced_path):
        print(f"ОШИБКА: не найден enhanced_analysis.py по пути: {enhanced_path}")
        return

    print("Запуск enhanced_analysis.py для получения verified_correlations...")
    globals_dict = runpy.run_path(enhanced_path)
    verified_correlations = globals_dict.get("verified_correlations", [])
    significant_ovcharova_markovskaya = globals_dict.get("significant_ovcharova_markovskaya", [])
    significant_holland = globals_dict.get("significant_holland", [])
    significant_ovcharova_holland = globals_dict.get("significant_ovcharova_holland", [])

    doc = Document()
    title = doc.add_heading("ОТВЕТ КОМИССИИ: ИСПОЛЬЗОВАНИЕ КОЭФФИЦИЕНТА ПИРСОНА В ИССЛЕДОВАНИИ", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 1. Можно ли обойтись без Пирсона?
    doc.add_heading("1. Можно ли обойтись без коэффициента Пирсона?", 1)
    doc.add_paragraph(
        "Да. Для выводов исследования и оценки статистической значимости достаточен коэффициент корреляции Спирмена (ρ)."
    )
    doc.add_paragraph()

    # 2. Зачем тогда использовали Пирсона?
    doc.add_heading("2. Зачем тогда использовался коэффициент Пирсона?", 1)
    doc.add_paragraph(
        "Он применялся как дополнительный (вспомогательный) метод: "
        "для проверки согласованности результатов и оценки характера связи. "
        "При близких значениях r (Пирсон) и ρ (Спирмен) связь близка к линейной; "
        "при заметном расхождении — связь монотонная, но не обязательно линейная."
    )
    doc.add_paragraph(
        "Основным критерием для всех выводов и интерпретаций в работе был коэффициент Спирмена, "
        "так как данные имеют порядковый характер (шкалы оценок), а не интервальный."
    )
    doc.add_paragraph()

    # 3. Где именно рассчитывался Пирсон?
    doc.add_heading("3. Где в исследовании рассчитывался коэффициент Пирсона?", 1)
    doc.add_paragraph(
        "Коэффициент Пирсона рассчитывался только в одном блоке анализа — "
        "корреляции между установками родителей (ВРР Марковской) и мотивами выбора профессии (опросник Овчаровой)."
    )
    doc.add_paragraph(
        "Конкретно для двух мотивов Овчаровой: "
        "вопрос 12 («Дает возможности для роста профессионального мастерства») и "
        "вопрос 19 («Позволяет использовать профессиональные умения вне работы»). "
        "Для каждой пары «вопрос родителя × мотив 12 или 19» при n ≥ 3 рассчитаны и Спирмен, и Пирсон."
    )
    doc.add_paragraph(
        "Пояснение: n — это количество пар наблюдений (пар родитель–ребёнок), по которым для данной конкретной пары вопросов есть ответы и родителя, и подростка. "
        "Условие «при n ≥ 3» означает: корреляцию считаем только если таких пар не менее трёх (при меньшем числе оценка связи ненадёжна)."
    )
    doc.add_paragraph(
        "Для корреляций с типами Холланда (Марковская–Холланд, Овчарова–Холланд) коэффициент Пирсона не рассчитывался."
    )
    doc.add_paragraph()

    # 3.1. Краткий список всех значимых корреляций
    total_sig = len(significant_ovcharova_markovskaya) + len(significant_holland) + len(significant_ovcharova_holland)
    doc.add_heading(f"3.1. Краткий список всех {total_sig} значимых корреляций в исследовании", 2)
    doc.add_paragraph(
        f"Ниже перечислены все {total_sig} статистически значимых корреляций (p < 0,05), обнаруженных в исследовании."
    )
    items = []
    for idx, c in enumerate(sorted(significant_ovcharova_markovskaya, key=lambda x: abs(x["spearman_corr"]), reverse=True), 1):
        p_short = (c["parent_question"][:60] + "…") if len(str(c["parent_question"])) > 60 else c["parent_question"]
        items.append(f"{idx}. Овчарова–Марковская: «{p_short}» ↔ «{c['ovcharova_question']}» (ρ = {c['spearman_corr']:.3f}, p = {c['spearman_p']:.4f}).")
    next_num = len(items) + 1
    for c in sorted(significant_holland, key=lambda x: abs(x["spearman_corr"]), reverse=True):
        p_short = (c["parent_question"][:60] + "…") if len(str(c["parent_question"])) > 60 else c["parent_question"]
        items.append(f"{next_num}. Марковская–Холланд: «{p_short}» ↔ {c.get('holland_type_name', c.get('holland_type', '?'))} (ρ = {c['spearman_corr']:.3f}, p = {c['spearman_p']:.4f}).")
        next_num += 1
    for c in significant_ovcharova_holland:
        ov = c.get("ovcharova_question", "?")
        items.append(f"{next_num}. Овчарова–Холланд: «{ov}» ↔ {c.get('holland_type_name', '?')} (ρ = {c['spearman_corr']:.3f}, p = {c['spearman_p']:.4f}).")
        next_num += 1
    for line in items:
        doc.add_paragraph(line)
    doc.add_paragraph()

    # Сохраняем компактный список в .txt для вставки в работу
    compact_lines = [
        "Значимые корреляции (p < 0,05):",
        "",
    ]
    for c in sorted(significant_ovcharova_markovskaya, key=lambda x: abs(x["spearman_corr"]), reverse=True):
        p_short = (c["parent_question"][:50] + "…") if len(str(c["parent_question"])) > 50 else c["parent_question"]
        compact_lines.append(f"• ВРР — Овчарова: «{p_short}» ↔ «{c['ovcharova_question']}» (ρ = {c['spearman_corr']:.3f}).")
    for c in sorted(significant_holland, key=lambda x: abs(x["spearman_corr"]), reverse=True):
        p_short = (c["parent_question"][:50] + "…") if len(str(c["parent_question"])) > 50 else c["parent_question"]
        compact_lines.append(f"• ВРР — Холланд: «{p_short}» ↔ {c.get('holland_type_name', c.get('holland_type', '?'))} (ρ = {c['spearman_corr']:.3f}).")
    for c in significant_ovcharova_holland:
        compact_lines.append(f"• Овчарова — Холланд: «{c.get('ovcharova_question', '?')}» ↔ {c.get('holland_type_name', '?')} (ρ = {c['spearman_corr']:.3f}).")
    txt_path = os.path.join(defence_dir, "Список_значимых_корреляций.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("\n".join(compact_lines))
    print(f"Список для вставки в работу сохранён: {txt_path}")

    # 4. Таблица всех значений Пирсона
    doc.add_heading("4. Все пары, для которых рассчитан Пирсон (значения ρ и r)", 1)
    doc.add_paragraph(
        "Ниже приведены все пары, для которых в исследовании рассчитан коэффициент Пирсона, "
        "с значениями коэффициента Спирмена (ρ) и Пирсона (r) и их уровнями значимости (p)."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Вопрос родителя (Марковская)"
    hdr[1].text = "Мотив (Овчарова)"
    hdr[2].text = "N"
    hdr[3].text = "ρ (Спирмен)"
    hdr[4].text = "p (Спирмен)"
    hdr[5].text = "r (Пирсон)"
    hdr[6].text = "p (Пирсон)"

    for c in verified_correlations:
        row_cells = table.add_row().cells
        parent_q = str(c.get("parent_question", ""))
        row_cells[0].text = (parent_q[:80] + "...") if len(parent_q) > 80 else parent_q
        row_cells[1].text = str(c.get("ovcharova_question", ""))[:50]
        row_cells[2].text = str(c.get("n", ""))
        row_cells[3].text = f"{c.get('spearman_corr', 0):.4f}"
        row_cells[4].text = f"{c.get('spearman_p', 0):.4f}"
        row_cells[5].text = f"{c.get('pearson_corr', 0):.4f}"
        row_cells[6].text = f"{c.get('pearson_p', 0):.4f}"

    out_path = os.path.join(defence_dir, "Ответ_комиссии_Пирсон.docx")
    doc.save(out_path)
    print(f"Сохранено: {out_path}")
    print(f"Всего пар с рассчитанным Пирсоном: {len(verified_correlations)}")


if __name__ == "__main__":
    main()
