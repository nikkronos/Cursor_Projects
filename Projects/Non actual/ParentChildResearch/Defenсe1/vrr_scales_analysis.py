"""
Анализ данных по 10 шкалам методики ВРР Марковской
Обработка результатов опросника "Взаимодействие родитель-ребёнок"
"""
import os
import sys
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import warnings
warnings.filterwarnings('ignore')

# ============================================================================
# КОНФИГУРАЦИЯ ШКАЛ ВРР (из Таблицы 11 опросника)
# ============================================================================

# Определение шкал и вопросов (вопросы с * - обратные)
SCALES_CONFIG = {
    1: {
        'name': 'Нетребовательность – требовательность',
        'questions': [1, 13, 25, 37, 49],  # 13 - обратный
        'inverse': [13],
        'divide_by_2': False
    },
    2: {
        'name': 'Мягкость – строгость',
        'questions': [2, 14, 26, 38, 50],  # 14, 38, 50 - обратные
        'inverse': [14, 38, 50],
        'divide_by_2': False
    },
    3: {
        'name': 'Автономность – контроль',
        'questions': [3, 4, 15, 16, 27, 28, 39, 40, 51, 52],  # 3, 4, 39, 52 - обратные
        'inverse': [3, 4, 39, 52],
        'divide_by_2': True  # 10 вопросов, делится на 2
    },
    4: {
        'name': 'Эмоциональная дистанция – близость',
        'questions': [5, 17, 29, 41, 53],  # 41 - обратный
        'inverse': [41],
        'divide_by_2': False
    },
    5: {
        'name': 'Отвержение – принятие',
        'questions': [6, 7, 18, 19, 30, 31, 42, 43, 54, 55],  # 6, 7, 18, 19, 31, 43, 54 - обратные
        'inverse': [6, 7, 18, 19, 31, 43, 54],
        'divide_by_2': True  # 10 вопросов, делится на 2
    },
    6: {
        'name': 'Отсутствие сотрудничества – сотрудничество',
        'questions': [8, 20, 32, 44, 56],
        'inverse': [],
        'divide_by_2': False
    },
    7: {
        'name': 'Несогласие – согласие',
        'questions': [9, 21, 33, 45, 57],  # 9, 57 - обратные
        'inverse': [9, 57],
        'divide_by_2': False
    },
    8: {
        'name': 'Непоследовательность – последовательность',
        'questions': [10, 22, 34, 46, 58],  # Все обратные
        'inverse': [10, 22, 34, 46, 58],
        'divide_by_2': False
    },
    9: {
        'name': 'Авторитетность родителя',
        'questions': [11, 23, 35, 47, 59],
        'inverse': [],
        'divide_by_2': False
    },
    10: {
        'name': 'Удовлетворенность отношениями с ребенком (с родителем)',
        'questions': [12, 24, 36, 48, 60],  # 48 - обратный
        'inverse': [48],
        'divide_by_2': False
    }
}

# ============================================================================
# ФУНКЦИИ ОБРАБОТКИ
# ============================================================================

def inverse_score(value):
    """
    Инвертирует значение для обратных вопросов.
    1→5, 2→4, 3→3, 4→2, 5→1
    """
    if pd.isna(value):
        return value
    try:
        val = int(float(value))
        if val < 1 or val > 5:
            return value
        return 6 - val
    except:
        return value

def calculate_scale_score(row, scale_config, question_cols):
    """
    Рассчитывает балл по шкале для одного респондента.
    
    Parameters:
    -----------
    row : pandas Series
        Строка данных респондента
    scale_config : dict
        Конфигурация шкалы
    question_cols : dict
        Словарь {номер_вопроса: имя_колонки}
    
    Returns:
    --------
    dict : словарь с промежуточными расчетами и итоговым баллом
    """
    scores = []
    raw_scores = []
    processed_scores = []
    question_numbers = []
    
    for q_num in scale_config['questions']:
        if q_num not in question_cols:
            continue
        
        col_name = question_cols[q_num]
        raw_value = row[col_name]
        
        # Преобразуем в число
        try:
            raw_val = int(float(raw_value)) if not pd.isna(raw_value) else None
        except:
            raw_val = None
        
        if raw_val is None:
            continue
        
        # Инвертируем, если вопрос обратный
        if q_num in scale_config['inverse']:
            processed_val = inverse_score(raw_val)
        else:
            processed_val = raw_val
        
        raw_scores.append(raw_val)
        processed_scores.append(processed_val)
        scores.append(processed_val)
        question_numbers.append(q_num)
    
    # Суммируем баллы
    total_sum = sum(scores) if scores else 0
    
    # Делим на 2, если нужно
    if scale_config['divide_by_2']:
        final_score = total_sum / 2.0
    else:
        final_score = total_sum
    
    return {
        'question_numbers': question_numbers,
        'raw_scores': raw_scores,
        'processed_scores': processed_scores,
        'total_sum': total_sum,
        'final_score': final_score,
        'divide_by_2': scale_config['divide_by_2']
    }

def process_vrr_data(df):
    """
    Обрабатывает данные опросника ВРР и рассчитывает баллы по всем шкалам.
    
    Returns:
    --------
    dict : словарь с результатами обработки
    """
    # Создаем словарь соответствия номеров вопросов и имен колонок
    question_cols = {}
    cols_list = df.columns.tolist()
    
    # Ищем колонки, которые являются числами от 1 до 60
    for i, col in enumerate(cols_list):
        col_str = str(col).strip()
        try:
            q_num = int(col_str)
            if 1 <= q_num <= 60:
                question_cols[q_num] = col
        except:
            pass
    
    # Если не нашли, ищем по позиции после служебных колонок
    if len(question_cols) < 60:
        # В файле из newtest вопросы начинаются с 6-й колонки (индекс 6)
        # Служебные колонки: Отметка времени, Number, Фамилия родителя, Фамилия ребенка, Возраст, Согласие
        service_keywords = ['Number', 'Номер', 'Column 66', 'Возраст', 'Подтверждение', 'согласие', 'timestamp']
        idx_start = None
        
        # Ищем последнюю служебную колонку
        for i, col in enumerate(cols_list):
            col_str = str(col).lower()
            if any(keyword.lower() in col_str for keyword in service_keywords):
                idx_start = i + 1
        
        # Если не нашли, пробуем стандартную позицию (6-я колонка, как в enhanced_analysis.py)
        if idx_start is None and len(cols_list) >= 66:
            idx_start = 6
        
        if idx_start is None:
            # Ищем колонку "1" напрямую
            try:
                idx_start = cols_list.index('1')
            except:
                idx_start = None
        
        if idx_start is not None:
            for i in range(60):
                if idx_start + i < len(cols_list):
                    q_num = i + 1
                    if q_num not in question_cols:
                        question_cols[q_num] = cols_list[idx_start + i]
    
    print(f"Найдено колонок с вопросами: {len(question_cols)}")
    
    # Обрабатываем каждого респондента
    results = []
    scale_scores_all = {i: [] for i in range(1, 11)}
    
    print(f"Всего строк в DataFrame: {len(df)}")
    print(f"Начинаем обработку всех строк...")
    
    processed_count = 0
    skipped_count = 0
    
    for idx, row in df.iterrows():
        
        # Проверяем, есть ли данные в вопросах (хотя бы один ответ)
        # Но не пропускаем слишком строго - обрабатываем все строки, где есть хотя бы что-то
        has_data = False
        answers_count = 0
        for q_num in range(1, 61):
            if q_num in question_cols:
                col_name = question_cols[q_num]
                val = row.get(col_name, None)
                if not pd.isna(val):
                    try:
                        int_val = int(float(val))
                        if 1 <= int_val <= 5:  # Валидное значение
                            has_data = True
                            answers_count += 1
                    except:
                        pass
        
        # Пропускаем только строки, где совсем нет валидных ответов
        # Но логируем для отладки
        if not has_data:
            skipped_count += 1
            respondent_num = row.get('Number', row.get('Номер', row.get('Column 66', idx)))
            print(f"Пропущена строка {idx + 1} (индекс {idx}, номер респондента: {respondent_num}): нет валидных ответов")
            # Проверяем, есть ли хотя бы какие-то данные в строке
            non_empty_cols = [col for col in df.columns if not pd.isna(row.get(col, None)) and str(row.get(col, '')).strip()]
            if non_empty_cols:
                print(f"  Но есть данные в {len(non_empty_cols)} колонках (не вопросы)")
            # Проверяем первые несколько вопросов вручную
            sample_questions = []
            for q_num in [1, 2, 3, 4, 5]:
                if q_num in question_cols:
                    col_name = question_cols[q_num]
                    val = row.get(col_name, None)
                    sample_questions.append(f"Q{q_num}={val}")
            print(f"  Примеры значений вопросов: {', '.join(sample_questions)}")
            continue
        
        # Получаем номер респондента
        respondent_id = None
        for col in ['Number', 'Номер', 'Column 66']:
            if col in df.columns:
                val = row.get(col, None)
                if not pd.isna(val):
                    try:
                        respondent_id = int(float(val))
                        break
                    except:
                        pass
        
        if respondent_id is None:
            respondent_id = idx + 1  # Используем индекс строки + 1 для нумерации
        
        # Получаем имена
        parent_name = ''
        child_name = ''
        for col in df.columns:
            col_str = str(col).lower()
            if 'родителя' in col_str or 'родитель' in col_str or 'parent' in col_str:
                val = row.get(col, '')
                if not pd.isna(val):
                    parent_name = str(val).strip()
            if ('ребенка' in col_str or 'ребенок' in col_str or 'child' in col_str) and 'возраст' not in col_str:
                val = row.get(col, '')
                if not pd.isna(val):
                    child_name = str(val).strip()
        
        respondent_data = {
            'id': respondent_id,
            'parent_name': parent_name,
            'child_name': child_name,
            'scales': {}
        }
        
        # Рассчитываем баллы по каждой шкале
        for scale_num in range(1, 11):
            scale_config = SCALES_CONFIG[scale_num]
            scale_result = calculate_scale_score(row, scale_config, question_cols)
            
            respondent_data['scales'][scale_num] = scale_result
            scale_scores_all[scale_num].append(scale_result['final_score'])
        
        results.append(respondent_data)
        processed_count += 1
    
    print(f"Обработано респондентов: {processed_count}")
    if skipped_count > 0:
        print(f"Пропущено строк: {skipped_count}")
    
    return {
        'respondents': results,
        'scale_scores_all': scale_scores_all,
        'question_cols': question_cols
    }

# ============================================================================
# ФУНКЦИИ СОЗДАНИЯ ТАБЛИЦ И ОТЧЕТОВ
# ============================================================================

def create_raw_data_table(processed_data, output_path):
    """
    Создает таблицу с сырыми данными и промежуточными расчетами.
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Таблица: Сырые данные и промежуточные расчеты по 10 шкалам ВРР', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Описание
    doc.add_paragraph(
        'В данной таблице представлены сырые данные ответов респондентов, '
        'обработанные значения (с учетом инверсии обратных вопросов), '
        'суммы баллов по каждой шкале и итоговые баллы.'
    )
    doc.add_paragraph('')
    
    # Для каждого респондента создаем таблицу
    for resp_data in processed_data['respondents']:
        # Заголовок респондента
        doc.add_heading(
            f"Респондент №{resp_data['id']}: {resp_data['parent_name']} / {resp_data['child_name']}", 
            level=1
        )
        
        # Таблица для всех шкал
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Шкала'
        header_cells[1].text = 'Вопросы'
        header_cells[2].text = 'Сырые баллы'
        header_cells[3].text = 'Обработанные баллы'
        header_cells[4].text = 'Сумма'
        header_cells[5].text = 'Делится на 2'
        header_cells[6].text = 'Итоговый балл'
        
        # Форматирование заголовков
        for cell in header_cells:
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Данные по каждой шкале
        for scale_num in range(1, 11):
            scale_config = SCALES_CONFIG[scale_num]
            scale_result = resp_data['scales'][scale_num]
            
            row_cells = table.add_row().cells
            row_cells[0].text = f"{scale_num}. {scale_config['name']}"
            row_cells[1].text = ', '.join(map(str, scale_result['question_numbers']))
            row_cells[2].text = ', '.join(map(str, scale_result['raw_scores']))
            row_cells[3].text = ', '.join(map(str, scale_result['processed_scores']))
            row_cells[4].text = str(scale_result['total_sum'])
            row_cells[5].text = 'Да' if scale_result['divide_by_2'] else 'Нет'
            row_cells[6].text = f"{scale_result['final_score']:.2f}" if scale_result['divide_by_2'] else str(scale_result['final_score'])
            
            # Выравнивание
            for cell in row_cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph('')
    
    doc.save(output_path)
    print(f"Таблица с сырыми данными сохранена: {output_path}")

def create_summary_table(processed_data, output_path):
    """
    Создает сводную таблицу с описательной статистикой по всем шкалам.
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Сводная таблица: Описательная статистика по 10 шкалам ВРР', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Таблица
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки
    header_cells = table.rows[0].cells
    headers = ['Шкала', 'Название', 'M (среднее)', 'SD (стандартное отклонение)', 
               'Me (медиана)', 'Min', 'Max']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Данные по каждой шкале
    for scale_num in range(1, 11):
        scale_config = SCALES_CONFIG[scale_num]
        scores = processed_data['scale_scores_all'][scale_num]
        
        if scores:
            scores_array = np.array([s for s in scores if not pd.isna(s)])
            if len(scores_array) > 0:
                mean_val = np.mean(scores_array)
                std_val = np.std(scores_array, ddof=1) if len(scores_array) > 1 else 0
                median_val = np.median(scores_array)
                min_val = np.min(scores_array)
                max_val = np.max(scores_array)
            else:
                mean_val = std_val = median_val = min_val = max_val = 0
        else:
            mean_val = std_val = median_val = min_val = max_val = 0
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(scale_num)
        row_cells[1].text = scale_config['name']
        row_cells[2].text = f"{mean_val:.2f}"
        row_cells[3].text = f"{std_val:.2f}"
        row_cells[4].text = f"{median_val:.2f}"
        row_cells[5].text = f"{min_val:.2f}"
        row_cells[6].text = f"{max_val:.2f}"
        
        # Выравнивание
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.save(output_path)
    print(f"Сводная таблица сохранена: {output_path}")

def create_description_and_interpretation(processed_data, output_path):
    """
    Создает документ с описанием и интерпретацией проведенного исследования.
    """
    doc = Document()
    
    # Заголовок
    title = doc.add_heading('Описание и интерпретация исследования по 10 шкалам ВРР Марковской', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 1: Описание методики
    doc.add_heading('1. Описание методики', level=1)
    doc.add_paragraph(
        'В исследовании использовалась методика "Взаимодействие родитель-ребёнок" '
        '(ВРР) И.М. Марковской для изучения установок родителей в детско-родительских отношениях. '
        'Опросник состоит из 60 вопросов, которые распределены по 10 шкалам.'
    )
    
    # Список шкал
    doc.add_heading('1.1. Шкалы опросника ВРР', level=2)
    for scale_num in range(1, 11):
        scale_config = SCALES_CONFIG[scale_num]
        p = doc.add_paragraph(f"{scale_num}. {scale_config['name']}", style='List Number')
    
    # Раздел 2: Методика обработки данных
    doc.add_heading('2. Методика обработки данных', level=1)
    
    doc.add_heading('2.1. Распределение вопросов по шкалам', level=2)
    doc.add_paragraph(
        'Каждая шкала включает определенные вопросы из опросника. '
        'Некоторые вопросы являются обратными (помечены в ключе опросника символом *), '
        'что означает необходимость инверсии их значений перед суммированием.'
    )
    
    # Таблица распределения вопросов
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Шкала'
    header_cells[1].text = 'Вопросы'
    header_cells[2].text = 'Обратные вопросы'
    for cell in header_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for scale_num in range(1, 11):
        scale_config = SCALES_CONFIG[scale_num]
        row_cells = table.add_row().cells
        row_cells[0].text = f"{scale_num}. {scale_config['name']}"
        row_cells[1].text = ', '.join(map(str, scale_config['questions']))
        row_cells[2].text = ', '.join(map(str, scale_config['inverse'])) if scale_config['inverse'] else '—'
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.2. Обработка обратных вопросов', level=2)
    doc.add_paragraph(
        'Для обратных вопросов применяется инверсия значений по следующему правилу:'
    )
    doc.add_paragraph('1 → 5', style='List Bullet')
    doc.add_paragraph('2 → 4', style='List Bullet')
    doc.add_paragraph('3 → 3 (остается без изменений)', style='List Bullet')
    doc.add_paragraph('4 → 2', style='List Bullet')
    doc.add_paragraph('5 → 1', style='List Bullet')
    doc.add_paragraph(
        'Это необходимо для того, чтобы все вопросы в рамках одной шкалы '
        'измеряли одно и то же направление характеристики.'
    )
    
    doc.add_heading('2.3. Расчет итоговых баллов', level=2)
    doc.add_paragraph(
        'Итоговый балл по каждой шкале рассчитывается как сумма обработанных '
        '(с учетом инверсии) значений всех вопросов, входящих в данную шкалу.'
    )
    doc.add_paragraph(
        'Особенность расчета для шкал 3 и 5: эти шкалы содержат по 10 вопросов '
        '(в отличие от остальных, которые содержат по 5 вопросов). '
        'Поэтому сумма баллов по шкалам 3 и 5 делится на 2 для получения итогового балла.'
    )
    doc.add_paragraph(
        'Формула расчета:'
    )
    doc.add_paragraph(
        '• Для шкал 1, 2, 4, 6, 7, 8, 9, 10: Итоговый балл = Σ (обработанные значения вопросов)',
        style='List Bullet'
    )
    doc.add_paragraph(
        '• Для шкал 3 и 5: Итоговый балл = Σ (обработанные значения вопросов) / 2',
        style='List Bullet'
    )
    
    # Раздел 3: Описательная статистика
    doc.add_heading('3. Описательная статистика', level=1)
    doc.add_paragraph(
        'Для каждой шкалы были рассчитаны следующие показатели описательной статистики:'
    )
    doc.add_paragraph('• M (среднее арифметическое) - среднее значение баллов по шкале', style='List Bullet')
    doc.add_paragraph('• SD (стандартное отклонение) - мера разброса данных вокруг среднего', style='List Bullet')
    doc.add_paragraph('• Me (медиана) - значение, которое делит упорядоченный ряд пополам', style='List Bullet')
    doc.add_paragraph('• Min и Max - минимальное и максимальное значения', style='List Bullet')
    
    # Добавляем сводную таблицу
    doc.add_heading('3.1. Сводная таблица описательной статистики', level=2)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    
    header_cells = table.rows[0].cells
    headers = ['Шкала', 'Название', 'M', 'SD', 'Me', 'Min', 'Max']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for scale_num in range(1, 11):
        scale_config = SCALES_CONFIG[scale_num]
        scores = processed_data['scale_scores_all'][scale_num]
        
        if scores:
            scores_array = np.array([s for s in scores if not pd.isna(s)])
            if len(scores_array) > 0:
                mean_val = np.mean(scores_array)
                std_val = np.std(scores_array, ddof=1) if len(scores_array) > 1 else 0
                median_val = np.median(scores_array)
                min_val = np.min(scores_array)
                max_val = np.max(scores_array)
            else:
                mean_val = std_val = median_val = min_val = max_val = 0
        else:
            mean_val = std_val = median_val = min_val = max_val = 0
        
        row_cells = table.add_row().cells
        row_cells[0].text = str(scale_num)
        row_cells[1].text = scale_config['name']
        row_cells[2].text = f"{mean_val:.2f}"
        row_cells[3].text = f"{std_val:.2f}"
        row_cells[4].text = f"{median_val:.2f}"
        row_cells[5].text = f"{min_val:.2f}"
        row_cells[6].text = f"{max_val:.2f}"
        
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Раздел 4: Интерпретация
    doc.add_heading('4. Интерпретация результатов', level=1)
    doc.add_paragraph(
        'Результаты обработки данных по 10 шкалам ВРР позволяют оценить различные '
        'аспекты детско-родительских отношений. Каждая шкала отражает определенную '
        'характеристику взаимодействия между родителем и ребенком.'
    )
    doc.add_paragraph(
        'Высокие баллы по шкалам, измеряющим позитивные аспекты отношений '
        '(например, "Эмоциональная близость", "Сотрудничество", "Принятие"), '
        'свидетельствуют о благоприятном характере детско-родительских отношений. '
        'Высокие баллы по шкалам, измеряющим негативные аспекты '
        '(например, "Требовательность", "Строгость", "Контроль"), '
        'могут указывать на более жесткий стиль воспитания.'
    )
    doc.add_paragraph(
        'Интерпретация конкретных результатов должна проводиться с учетом '
        'нормативных данных и контекста конкретного исследования.'
    )
    
    doc.save(output_path)
    print(f"Описание и интерпретация сохранены: {output_path}")

# ============================================================================
# ОСНОВНАЯ ФУНКЦИЯ
# ============================================================================

def main():
    # Пути к файлам
    script_dir = os.path.dirname(os.path.abspath(__file__))
    csv_file = os.path.join(script_dir, 'Опрос_для_родителей.csv')
    
    # Если файл не найден, пытаемся скопировать из других папок
    if not os.path.exists(csv_file):
        # Сначала пробуем из newtest (там полные данные для 50 пар)
        source_file1 = os.path.join(script_dir, '..', 'newtest', 'Опрос для родителей  (Ответы).csv')
        # Затем из Final3
        source_file2 = os.path.join(script_dir, '..', 'Final3', 'Опрос для родителей  (Ответы) - Ответы на форму (1) (1).csv')
        
        source_file = None
        if os.path.exists(source_file1):
            source_file = source_file1
        elif os.path.exists(source_file2):
            source_file = source_file2
        
        if source_file:
            print(f"Копирование файла из {source_file}...")
            import shutil
            shutil.copy2(source_file, csv_file)
            print(f"Файл скопирован в {csv_file}")
        else:
            print(f"ОШИБКА: Файл не найден: {csv_file}")
            print(f"Проверены следующие пути:")
            print(f"  1. {source_file1}")
            print(f"  2. {source_file2}")
            sys.exit(1)
    
    print(f"Загрузка данных из: {csv_file}")
    
    # Загружаем данные
    try:
        df = pd.read_csv(csv_file, encoding='utf-8-sig')
        print(f"Загружено строк: {len(df)}")
        print(f"Колонок: {len(df.columns)}")
        print(f"Первые колонки: {list(df.columns[:10])}")
    except Exception as e:
        print(f"ОШИБКА при загрузке данных: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    
    # Обрабатываем данные
    print("Обработка данных по 10 шкалам ВРР...")
    processed_data = process_vrr_data(df)
    # Сообщение о количестве уже выводится внутри функции
    
    # Создаем выходные файлы
    output_dir = script_dir
    
    # 1. Таблица с сырыми данными
    raw_data_path = os.path.join(output_dir, 'Таблица_сырые_данные_ВРР.docx')
    create_raw_data_table(processed_data, raw_data_path)
    
    # 2. Сводная таблица
    summary_path = os.path.join(output_dir, 'Сводная_таблица_ВРР.docx')
    create_summary_table(processed_data, summary_path)
    
    # 3. Описание и интерпретация
    description_path = os.path.join(output_dir, 'Описание_и_интерпретация_ВРР.docx')
    create_description_and_interpretation(processed_data, description_path)
    
    print("\nОбработка завершена успешно!")
    print(f"Созданные файлы:")
    print(f"  1. {raw_data_path}")
    print(f"  2. {summary_path}")
    print(f"  3. {description_path}")

if __name__ == '__main__':
    main()
