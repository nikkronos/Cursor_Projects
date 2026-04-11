"""Экспорт Главы II в Word документ с графиками и списком литературы"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Определяем пути
script_dir = os.path.dirname(os.path.abspath(__file__))
chapter_file = os.path.join(script_dir, "ГЛАВА_II.md")
charts_dir = os.path.join(script_dir, "charts")
literature_file = os.path.join(script_dir, "Список_литературы_очищенный.txt")
output_file = os.path.join(script_dir, "ГЛАВА_II_ПОЛНАЯ.docx")

print("Создание Word документа...")

# Создаем документ
doc = Document()

# Настройка стилей
def set_font(run, font_name='Times New Roman', size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Настройка параграфа
def set_paragraph_style(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing = spacing
    paragraph_format.first_line_indent = Inches(0.5)  # Абзацный отступ 5 знаков (примерно 1.25 см)

# Читаем текст Главы II
print("Чтение текста Главы II...")
with open(chapter_file, 'r', encoding='utf-8') as f:
    chapter_lines = f.readlines()

# Обрабатываем текст и вставляем графики
current_section = None
in_code_block = False

for line in chapter_lines:
    line = line.rstrip()
    
    # Пропускаем пустые строки в начале
    if not line and not current_section:
        continue
    
    # Заголовки
    if line.startswith('# '):
        # Главный заголовок
        heading = doc.add_heading(line[2:], 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            set_font(run, size=16, bold=True)
        current_section = line[2:]
    
    elif line.startswith('## '):
        # Заголовок 2-го уровня
        heading = doc.add_heading(line[3:], 1)
        for run in heading.runs:
            set_font(run, size=14, bold=True)
        current_section = line[3:]
    
    elif line.startswith('### '):
        # Заголовок 3-го уровня
        heading = doc.add_heading(line[4:], 2)
        for run in heading.runs:
            set_font(run, size=14, bold=True)
        current_section = line[4:]
    
    # Обычный текст
    elif line and not line.startswith('#'):
        # Убираем markdown форматирование
        text = line
        text = text.replace('**', '')  # Убираем жирный
        text = text.replace('*', '')   # Убираем курсив
        
        # Создаем параграф
        para = doc.add_paragraph()
        set_paragraph_style(para)
        
        # Обрабатываем списки
        if text.startswith('- ') or text.startswith('* '):
            text = text[2:]
            para.style = 'List Bullet'
        elif re.match(r'^\d+\.\s+', text):
            para.style = 'List Number'
        
        run = para.add_run(text)
        set_font(run, size=14)
    
    # Пустая строка
    elif not line:
        doc.add_paragraph()

# Добавляем графики после текста главы
print("Вставка графиков...")
doc.add_page_break()
doc.add_heading('ПРИЛОЖЕНИЯ', 1)
heading = doc.paragraphs[-1]
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in heading.runs:
    set_font(run, size=14, bold=True)

doc.add_paragraph()
doc.add_heading('График 1. Распределение коэффициентов корреляции', 2)
chart_path = os.path.join(charts_dir, "1_histogram_correlations.png")
if os.path.exists(chart_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 1. Распределение коэффициентов корреляции между установками родителей и профессиональным выбором подростков")
    set_font(run, size=12)

# График 2: Топ-10 корреляций
doc.add_paragraph()
doc.add_heading('График 2. Топ-10 самых сильных корреляций', 2)
chart_path = os.path.join(charts_dir, "2_top10_correlations.png")
if os.path.exists(chart_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 2. Топ-10 самых сильных корреляций (зеленый = значимо при p<0.05)")
    set_font(run, size=12)

# График 3: Значимые корреляции
doc.add_paragraph()
doc.add_heading('График 3. Статистически значимые корреляции', 2)
chart_path = os.path.join(charts_dir, "3_significant_correlations.png")
if os.path.exists(chart_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 3. Статистически значимые корреляции (p < 0.05)")
    set_font(run, size=12)

# График 4: Сравнение по вопросам
doc.add_paragraph()
doc.add_heading('График 4. Сравнение корреляций по вопросам о профессии', 2)
chart_path = os.path.join(charts_dir, "4_comparison_prof_questions.png")
if os.path.exists(chart_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 4. Распределение корреляций по двум вопросам о профессиональных предпочтениях")
    set_font(run, size=12)

# График 5: Тепловая карта (если есть)
chart_path = os.path.join(charts_dir, "5_heatmap_significant.png")
if os.path.exists(chart_path):
    doc.add_paragraph()
    doc.add_heading('График 5. Тепловая карта значимых корреляций', 2)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 5. Тепловая карта значимых корреляций")
    set_font(run, size=12)

# График 6: Уровни значимости
doc.add_paragraph()
doc.add_heading('График 6. Распределение по уровням значимости', 2)
chart_path = os.path.join(charts_dir, "6_significance_levels.png")
if os.path.exists(chart_path):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(chart_path, width=Inches(6))
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Рис. 6. Распределение корреляций по уровням статистической значимости")
    set_font(run, size=12)

# Добавляем список литературы
print("Добавление списка литературы...")
doc.add_page_break()
doc.add_heading('СПИСОК ЛИТЕРАТУРЫ', 0)
heading = doc.paragraphs[-1]
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in heading.runs:
    set_font(run, size=14, bold=True)

if os.path.exists(literature_file):
    with open(literature_file, 'r', encoding='utf-8') as f:
        literature_lines = f.readlines()
    
    for line in literature_lines:
        line = line.strip()
        if line:
            para = doc.add_paragraph(line)
            set_paragraph_style(para, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            for run in para.runs:
                set_font(run, size=14)

# Сохраняем документ
print(f"Сохранение документа: {output_file}")
doc.save(output_file)

print(f"\n✓ Документ успешно создан: {output_file}")
print(f"  Размер файла: {os.path.getsize(output_file) / 1024:.1f} КБ")
print("\nДокумент содержит:")
print("  - Полный текст Главы II")
print("  - Все 6 графиков")
print("  - Очищенный список литературы")

