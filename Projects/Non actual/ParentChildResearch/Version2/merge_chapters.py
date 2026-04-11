"""Объединение Главы I и Главы II в один документ Word"""
import os
import sys
import re

script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
except ImportError:
    print("Ошибка: библиотека python-docx не установлена")
    print("Установите: pip install python-docx")
    sys.exit(1)

print("=" * 60)
print("ОБЪЕДИНЕНИЕ ГЛАВЫ I И ГЛАВЫ II В ОДИН ДОКУМЕНТ")
print("=" * 60)

# 1. Открываем исходный документ как базу (сохраняет верстку титульного листа и содержания)
print("\n1. Загрузка исходного документа...")
chapter1_file = os.path.join(script_dir, 'A1', 'ВКР Бакалавриат_Ворошилова.docx')

if not os.path.exists(chapter1_file):
    print("   ⚠ Файл Главы I не найден, создаем новый документ")
    doc = Document()
    # Настройка страницы по методичке
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
else:
    # Открываем исходный документ - это сохраняет всю верстку
    doc = Document(chapter1_file)
    print("   ✓ Исходный документ загружен (верстка сохранена)")

# Функции для форматирования
def set_font(run, font_name='Times New Roman', size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def set_paragraph_style(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                       spacing=1.5, first_line_indent=Cm(1.25)):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = spacing
    paragraph_format.first_line_indent = first_line_indent
    paragraph_format.space_after = Pt(0)

def add_heading(doc, text, level=1):
    # Создаем заголовок вручную, не используя стандартные стили Word
    # чтобы избежать ошибок, если стили отсутствуют в документе
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    set_font(run, size=14, bold=True)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Добавляем отступ для подзаголовков
    if level > 1:
        heading_format = heading.paragraph_format
        heading_format.space_before = Pt(12)
        heading_format.space_after = Pt(6)
    return heading

# 2. Добавляем разрыв страницы перед Главой II
doc.add_page_break()

# 3. Читаем Главу II
print("\n2. Добавление Главы II...")
chapter2_file = os.path.join(script_dir, 'ГЛАВА_II_v3.md')

if os.path.exists(chapter2_file):
    with open(chapter2_file, 'r', encoding='utf-8') as f:
        chapter2_lines = f.readlines()
    
    for line in chapter2_lines:
        line = line.strip()
        
        if not line:
            continue
        
        # Заголовки
        if line.startswith('# '):
            add_heading(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_heading(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_heading(doc, line[4:], level=3)
        # Списки
        elif line.startswith('- '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:])
            set_font(run)
            set_paragraph_style(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
            # Добавляем отступ для списка
            p.paragraph_format.left_indent = Cm(1.0)
        # Жирный текст
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line[2:-2])
            set_font(run, bold=True)
            set_paragraph_style(p)
        # Обычный текст
        else:
            # Убираем markdown разметку
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            if text:
                p = doc.add_paragraph()
                run = p.add_run(text)
                set_font(run)
                set_paragraph_style(p)
    
    print("   ✓ Глава II добавлена")
else:
    print("   ⚠ Файл Главы II не найден")

# 4. Добавляем выводы
print("\n3. Добавление выводов...")
conclusions_file = os.path.join(script_dir, 'ВЫВОДЫ.md')
if os.path.exists(conclusions_file):
    doc.add_page_break()
    add_heading(doc, "ВЫВОДЫ", level=1)
    
    with open(conclusions_file, 'r', encoding='utf-8') as f:
        conclusions_lines = f.readlines()
    
    for line in conclusions_lines:
        line = line.strip()
        if line and not line.startswith('#'):
            line = re.sub(r'^\d+\.\s+', '', line)
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run)
            set_paragraph_style(p)
    
    print("   ✓ Выводы добавлены")

# 5. Добавляем заключение
print("\n4. Добавление заключения...")
conclusion_file = os.path.join(script_dir, 'ЗАКЛЮЧЕНИЕ.md')
if os.path.exists(conclusion_file):
    doc.add_page_break()
    add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)
    
    with open(conclusion_file, 'r', encoding='utf-8') as f:
        conclusion_lines = f.readlines()
    
    for line in conclusion_lines:
        line = line.strip()
        if line and not line.startswith('#'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run)
            set_paragraph_style(p)
    
    print("   ✓ Заключение добавлено")

# 6. Добавляем список литературы
print("\n5. Добавление списка литературы...")
lit_file = os.path.join(script_dir, 'СПИСОК_ЛИТЕРАТУРЫ_50_ИСТОЧНИКОВ_ГОСТ.md')
if os.path.exists(lit_file):
    doc.add_page_break()
    add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ", level=1)
    
    with open(lit_file, 'r', encoding='utf-8') as f:
        lit_lines = f.readlines()
    
    for line in lit_lines:
        line = line.strip()
        if line and not line.startswith('#'):
            # СОХРАНЯЕМ нумерацию в начале (не убираем!)
            # Формат: "1. Автор. Название..."
            p = doc.add_paragraph()
            run = p.add_run(line)
            set_font(run, size=14)
            set_paragraph_style(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    
    print("   ✓ Список литературы добавлен")

# 7. Добавляем приложения с графиками
print("\n6. Добавление приложений с графиками...")
charts_dir = os.path.join(script_dir, 'charts')
if os.path.exists(charts_dir):
    doc.add_page_break()
    add_heading(doc, "ПРИЛОЖЕНИЯ", level=1)
    
    chart_files = [
        ('1_histogram_correlations.png', 'Рис. 1. Гистограмма распределения коэффициентов корреляции Спирмена'),
        ('2_top10_correlations.png', 'Рис. 2. Топ-10 самых сильных корреляций'),
        ('3_significant_correlations.png', 'Рис. 3. Статистически значимые корреляции (p < 0.05)'),
        ('4_comparison_prof_questions.png', 'Рис. 4. Сравнение корреляций по двум вопросам о профессии'),
        ('5_heatmap_significant.png', 'Рис. 5. Тепловая карта статистически значимых корреляций'),
        ('6_significance_levels.png', 'Рис. 6. Распределение корреляций по уровням статистической значимости'),
    ]
    
    for chart_file, caption in chart_files:
        chart_path = os.path.join(charts_dir, chart_file)
        if os.path.exists(chart_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_path, width=Inches(6))
            
            p_caption = doc.add_paragraph()
            p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_caption = p_caption.add_run(caption)
            set_font(run_caption, size=12, bold=False)
            
            doc.add_paragraph()
            print(f"   ✓ Добавлен: {chart_file}")
    
    print("   ✓ Приложения добавлены")

# Сохраняем документ
output_file = os.path.join(script_dir, 'ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v4.docx')
doc.save(output_file)

print(f"\n" + "=" * 60)
print("ДОКУМЕНТ СОЗДАН")
print("=" * 60)
print(f"\nФайл: {output_file}")
print(f"Размер: {os.path.getsize(output_file) / 1024:.1f} КБ")
print("\nСодержание:")
print("  - Глава I")
print("  - Глава II")
print("  - Выводы")
print("  - Заключение")
print("  - Список литературы (50 источников)")
print("  - Приложения (6 графиков)")


