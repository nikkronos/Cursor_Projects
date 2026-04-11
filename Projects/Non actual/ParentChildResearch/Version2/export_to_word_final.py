"""Финальный экспорт Главы II в Word с правильным оформлением по методичке"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Определяем пути
script_dir = os.path.dirname(os.path.abspath(__file__))
chapter_file = os.path.join(script_dir, "ГЛАВА_II.md")
conclusions_file = os.path.join(script_dir, "ВЫВОДЫ.md")
conclusion_file = os.path.join(script_dir, "ЗАКЛЮЧЕНИЕ.md")
literature_file = os.path.join(script_dir, "СПИСОК_ЛИТЕРАТУРЫ_ГОСТ.md")
charts_dir = os.path.join(script_dir, "charts")
output_file = os.path.join(script_dir, "ГЛАВА_II_ФИНАЛЬНАЯ.docx")

print("Создание финального Word документа...")

# Создаем документ
doc = Document()

# Настройка страницы по методичке
section = doc.sections[0]
section.page_height = Cm(29.7)  # A4 высота
section.page_width = Cm(21)     # A4 ширина
section.top_margin = Cm(2.0)    # Верх 20 мм
section.bottom_margin = Cm(2.0) # Низ 20 мм
section.left_margin = Cm(3.0)    # Левое 30 мм
section.right_margin = Cm(1.5)  # Правое 15 мм

# Настройка стилей
def set_font(run, font_name='Times New Roman', size=14, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Настройка параграфа по методичке
def set_paragraph_style(paragraph, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, 
                       spacing=1.5, first_line_indent=Cm(1.25)):
    paragraph.alignment = alignment
    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = spacing
    paragraph_format.first_line_indent = first_line_indent  # 5 знаков = 1.25 см
    paragraph_format.space_after = Pt(0)

# Функция для добавления заголовка
def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level)
    for run in heading.runs:
        set_font(run, size=14, bold=True)
    if level == 1:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return heading

# Читаем текст Главы II
print("Чтение текста Главы II...")
with open(chapter_file, 'r', encoding='utf-8') as f:
    chapter_text = f.read()

# Обрабатываем Markdown и конвертируем в Word
lines = chapter_text.split('\n')
current_heading_level = 0
in_list = False

for i, line in enumerate(lines):
    line = line.strip()
    
    # Пропускаем пустые строки
    if not line:
        if not in_list:
            continue
    
    # Заголовки
    if line.startswith('# '):
        add_heading(doc, line[2:].strip(), level=1)
        current_heading_level = 1
    elif line.startswith('## '):
        add_heading(doc, line[3:].strip(), level=2)
        current_heading_level = 2
    elif line.startswith('### '):
        add_heading(doc, line[4:].strip(), level=3)
        current_heading_level = 3
    # Жирный текст
    elif line.startswith('**') and line.endswith('**'):
        p = doc.add_paragraph()
        run = p.add_run(line[2:-2])
        set_font(run, bold=True)
        set_paragraph_style(p)
    # Списки
    elif line.startswith('- '):
        p = doc.add_paragraph(line[2:], style='List Bullet')
        for run in p.runs:
            set_font(run)
        set_paragraph_style(p)
        in_list = True
    # Обычный текст
    else:
        # Убираем markdown разметку
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
        text = re.sub(r'\[(\d+[;\s]*\d*)\]', r'[\1]', text)  # Сохраняем ссылки
        
        if text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font(run)
            set_paragraph_style(p)
            in_list = False

# Добавляем раздел с выводами
print("Добавление выводов...")
doc.add_page_break()
add_heading(doc, "ВЫВОДЫ", level=1)

with open(conclusions_file, 'r', encoding='utf-8') as f:
    conclusions_text = f.read()

for line in conclusions_text.split('\n'):
    line = line.strip()
    if line and not line.startswith('#'):
        # Убираем нумерацию в начале
        line = re.sub(r'^\d+\.\s+', '', line)
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run)
        set_paragraph_style(p)

# Добавляем заключение
print("Добавление заключения...")
doc.add_page_break()
add_heading(doc, "ЗАКЛЮЧЕНИЕ", level=1)

with open(conclusion_file, 'r', encoding='utf-8') as f:
    conclusion_text = f.read()

for line in conclusion_text.split('\n'):
    line = line.strip()
    if line and not line.startswith('#'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run)
        set_paragraph_style(p)

# Добавляем список литературы
print("Добавление списка литературы...")
doc.add_page_break()
add_heading(doc, "СПИСОК ЛИТЕРАТУРЫ", level=1)

with open(literature_file, 'r', encoding='utf-8') as f:
    literature_text = f.read()

for line in literature_text.split('\n'):
    line = line.strip()
    if line and not line.startswith('#'):
        # Убираем нумерацию в начале
        line = re.sub(r'^\d+\.\s+', '', line)
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, size=14)
        set_paragraph_style(p, alignment=WD_ALIGN_PARAGRAPH.LEFT)

# Добавляем приложения с графиками
print("Добавление приложений с графиками...")
doc.add_page_break()
add_heading(doc, "ПРИЛОЖЕНИЯ", level=1)

chart_files = [
    ('1_histogram_correlations.png', 'Рис. 1. Гистограмма распределения коэффициентов корреляции Спирмена между установками родителей и профессиональным выбором подростков'),
    ('2_top10_correlations.png', 'Рис. 2. Топ-10 самых сильных корреляций между установками родителей и профессиональным выбором подростков'),
    ('3_significant_correlations.png', 'Рис. 3. Статистически значимые корреляции (p < 0.05)'),
    ('4_comparison_prof_questions.png', 'Рис. 4. Сравнение корреляций по двум вопросам о профессии (первые 20)'),
    ('5_heatmap_significant.png', 'Рис. 5. Тепловая карта статистически значимых корреляций'),
    ('6_significance_levels.png', 'Рис. 6. Распределение корреляций по уровням статистической значимости'),
]

for chart_file, caption in chart_files:
    chart_path = os.path.join(charts_dir, chart_file)
    if os.path.exists(chart_path):
        # Добавляем график
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(chart_path, width=Inches(6))  # Ширина 6 дюймов
        
        # Добавляем подпись (по методичке: снизу, по центру)
        p_caption = doc.add_paragraph()
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_caption = p_caption.add_run(caption)
        set_font(run_caption, size=12, bold=False)
        
        # Пустая строка после графика
        doc.add_paragraph()
        print(f"  ✓ Добавлен: {chart_file}")
    else:
        print(f"  ⚠ Не найден: {chart_file}")

# Сохраняем документ
doc.save(output_file)
print(f"\n✓ Финальный документ сохранен: {output_file}")
print(f"  Размер: {os.path.getsize(output_file) / 1024:.1f} КБ")
