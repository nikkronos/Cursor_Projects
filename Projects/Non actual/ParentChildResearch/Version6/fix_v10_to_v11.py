# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v10 -> v11:
1. Убрать все фамилии из квадратных скобок, оставить только номера источников
2. Добавить выводы о незначимых связях в раздел ВЫВОДЫ
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    v11_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v11.docx"
    
    if not v10_path.exists():
        print(f"ERROR: {v10_path} не найден")
        print("Попробую найти v9...")
        v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
        if v9_path.exists():
            print(f"Найден v9, используем его как основу")
            v10_path = v9_path
        else:
            return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v10 -> v11")
    print("=" * 70)
    
    doc = Document(str(v10_path))
    print(f"✓ Документ загружен\n")
    
    fixed_count = 0
    max_source = 53
    min_source = 1
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        changes_made = []
        
        # 1. Убираем фамилии из квадратных скобок
        # Ищем все квадратные скобки, которые содержат буквы (фамилии)
        pattern_with_names = r'\[[^\]]*[А-ЯЁа-яё][^\]]*\]'
        
        def extract_numbers_only(match):
            full_match = match.group(0)
            # Извлекаем все номера из скобок
            numbers = re.findall(r'\b(\d+)\b', full_match)
            
            if numbers:
                # Фильтруем номера по диапазону
                valid_numbers = []
                for num_str in numbers:
                    try:
                        num = int(num_str)
                        if min_source <= num <= max_source:
                            valid_numbers.append(num)
                    except ValueError:
                        pass
                
                if valid_numbers:
                    unique_numbers = sorted(set(valid_numbers))
                    return f"[{', '.join(map(str, unique_numbers))}]"
            
            return ""
        
        # Сначала находим все совпадения для отчета
        matches = list(re.finditer(pattern_with_names, text))
        if matches:
            for match in matches:
                old_text = match.group(0)
                # Извлекаем номера
                numbers = re.findall(r'\b(\d+)\b', old_text)
                valid_numbers = []
                for num_str in numbers:
                    try:
                        num = int(num_str)
                        if min_source <= num <= max_source:
                            valid_numbers.append(num)
                    except ValueError:
                        pass
                
                if valid_numbers:
                    new_value = f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    if old_text != new_value:
                        changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_value}")
        
        # Заменяем все совпадения
        new_text = re.sub(pattern_with_names, extract_numbers_only, text)
        if new_text != text:
            text = new_text
            print(f"✓ Параграф {para_idx + 1}: убраны фамилии из квадратных скобок")
            if changes_made:
                for change in changes_made[:3]:  # Показываем первые 3
                    print(f"  {change}")
        
        # 2. Исправляем неверные номера источников (если остались)
        number_fixes = {
            '[60]': '[51]',  # Франкл
            '[64]': '[53]',  # Эйдемиллер, Юстицкис
        }
        
        for wrong, correct in number_fixes.items():
            if wrong in text:
                text = text.replace(wrong, correct)
                changes_made.append(f"БЫЛО: {wrong}\n     СТАЛО: {correct}")
                print(f"✓ Параграф {para_idx + 1}: исправлен номер {wrong} -> {correct}")
        
        # 3. Проверяем и исправляем номера в обычных квадратных скобках [N, N, N]
        pattern_numbers = r'\[(\d+(?:[,\s;]+\s*\d+)*)\]'
        
        def fix_numbers_in_brackets(match):
            numbers_str = match.group(1)
            valid_numbers = []
            
            for num_str in re.split(r'[,\s;]+', numbers_str):
                num_str = num_str.strip()
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                except ValueError:
                    pass
            
            if valid_numbers:
                unique_numbers = sorted(set(valid_numbers))
                return f"[{', '.join(map(str, unique_numbers))}]"
            return ""
        
        new_text = re.sub(pattern_numbers, fix_numbers_in_brackets, text)
        if new_text != text and new_text != original_text:
            text = new_text
        
        # Применяем изменения
        if text != original_text:
            para.text = text
            fixed_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Убираем фамилии из скобок
                pattern_with_names = r'\[[^\]]*[А-ЯЁа-яё][^\]]*\]'
                
                def extract_numbers_cell(match):
                    full_match = match.group(0)
                    numbers = re.findall(r'\b(\d+)\b', full_match)
                    valid_numbers = []
                    for num_str in numbers:
                        try:
                            num = int(num_str)
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_with_names, extract_numbers_cell, text)
                
                # Исправляем номера
                pattern_numbers = r'\[(\d+(?:[,\s;]+\s*\d+)*)\]'
                def fix_numbers_cell(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers_cell, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # 4. Добавляем выводы о незначимых связях в раздел ВЫВОДЫ
    print("\n" + "=" * 70)
    print("ДОБАВЛЕНИЕ ВЫВОДОВ О НЕЗНАЧИМЫХ СВЯЗЯХ")
    print("=" * 70)
    
    conclusions_added = False
    
    # Сначала проверяем, есть ли уже такой вывод
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if "незначим" in text.lower() and "117" in text and "122" in text:
            print("✓ Выводы о незначимых связях уже есть в документе")
            conclusions_added = True
            break
    
    if not conclusions_added:
        # Ищем параграф с упоминанием "5 из 122" в разделе ВЫВОДЫ
        target_para_idx = None
        max_point_num = 0
        
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            
            # Ищем упоминание "5 из 122" или "статистически значимых корреляций"
            if ("5 из 122" in text or "статистически значимых корреляций выявлено относительно немного" in text) and \
               ("ВЫВОДЫ" in doc.paragraphs[max(0, para_idx-20):para_idx+1][-1].text or para_idx > 250):
                target_para_idx = para_idx
                
                # Ищем максимальный номер пункта в выводах
                for check_idx in range(max(0, para_idx - 30), para_idx + 1):
                    check_text = doc.paragraphs[check_idx].text
                    point_match = re.search(r'^(\d+)\.\s', check_text.strip())
                    if point_match:
                        max_point_num = max(max_point_num, int(point_match.group(1)))
                break
        
        if target_para_idx is not None:
            # Вставляем новый пункт после найденного параграфа
            new_point_num = max_point_num + 1
            new_para = doc.add_paragraph()
            new_para.text = f"{new_point_num}. Анализ статистической значимости выявленных корреляций показал, что большинство проверенных взаимосвязей между параметрами детско-родительских отношений и характеристиками профессионального выбора подростков не достигли уровня статистической значимости. Из 122 рассчитанных корреляций статистически значимыми оказались только 5 (4,1%), что означает, что 117 взаимосвязей (95,9%) не подтвердились на статистическом уровне. Это указывает на то, что влияние детско-родительских отношений на профессиональный выбор подростков носит избирательный и специфический характер, проявляясь лишь в определенных аспектах взаимодействия. Отсутствие значимых связей между большинством параметров может быть связано с опосредованным характером влияния семейных факторов, необходимостью учета дополнительных переменных (таких как индивидуальные особенности подростка, социальный контекст, образовательная среда) или недостаточной чувствительностью применяемых методик к выявлению слабых, но существующих связей. Данный результат подчеркивает сложность и многогранность процесса профессионального самоопределения, который не может быть сведен к простым линейным зависимостям между отдельными параметрами детско-родительских отношений и профессиональными предпочтениями."
            
            # Вставляем после целевого параграфа
            doc.paragraphs[target_para_idx]._element.addnext(new_para._element)
            print(f"✓ Добавлен вывод о незначимых связях (пункт {new_point_num}) после параграфа {target_para_idx + 1}")
            conclusions_added = True
        else:
            # Если не нашли, ищем раздел ВЫВОДЫ и добавляем в конец
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text == "ВЫВОДЫ":
                    # Ищем последний пункт выводов перед ЗАКЛЮЧЕНИЕМ
                    last_point_num = 0
                    insert_idx = None
                    
                    for next_idx in range(para_idx + 1, len(doc.paragraphs)):
                        next_text = doc.paragraphs[next_idx].text.strip()
                        if next_text == "ЗАКЛЮЧЕНИЕ" or "ЗАКЛЮЧЕНИЕ" in next_text:
                            insert_idx = next_idx
                            break
                        # Ищем номера пунктов
                        point_match = re.search(r'^(\d+)\.\s', next_text)
                        if point_match:
                            last_point_num = max(last_point_num, int(point_match.group(1)))
                    
                    if insert_idx is not None:
                        new_point_num = last_point_num + 1
                        new_para = doc.add_paragraph()
                        new_para.text = f"{new_point_num}. Анализ статистической значимости выявленных корреляций показал, что большинство проверенных взаимосвязей между параметрами детско-родительских отношений и характеристиками профессионального выбора подростков не достигли уровня статистической значимости. Из 122 рассчитанных корреляций статистически значимыми оказались только 5 (4,1%), что означает, что 117 взаимосвязей (95,9%) не подтвердились на статистическом уровне. Это указывает на то, что влияние детско-родительских отношений на профессиональный выбор подростков носит избирательный и специфический характер, проявляясь лишь в определенных аспектах взаимодействия. Отсутствие значимых связей между большинством параметров может быть связано с опосредованным характером влияния семейных факторов, необходимостью учета дополнительных переменных (таких как индивидуальные особенности подростка, социальный контекст, образовательная среда) или недостаточной чувствительностью применяемых методик к выявлению слабых, но существующих связей. Данный результат подчеркивает сложность и многогранность процесса профессионального самоопределения, который не может быть сведен к простым линейным зависимостям между отдельными параметрами детско-родительских отношений и профессиональными предпочтениями."
                        doc.paragraphs[insert_idx]._element.addprevious(new_para._element)
                        print(f"✓ Добавлен вывод о незначимых связях (пункт {new_point_num}) перед ЗАКЛЮЧЕНИЕМ")
                        conclusions_added = True
                    break
    
    # Сохраняем документ
    doc.save(str(v11_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v11_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    if conclusions_added:
        print(f"✓ Добавлены выводы о незначимых связях")

if __name__ == '__main__':
    fix_document()

