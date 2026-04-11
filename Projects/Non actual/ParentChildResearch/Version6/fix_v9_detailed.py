# -*- coding: utf-8 -*-
"""
Детальный скрипт для исправления документа v9 -> v10
Показывает, что именно было изменено
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v9 -> v10")
    print("=" * 70)
    
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен\n")
    
    fixed_count = 0
    max_source = 53
    min_source = 1
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        changes_made = []
        
        # 1. Параграф 77 (страница 7): [Ананьев, 3; Кон, 31; Дружинин, 22]
        # Пробуем разные варианты формата
        patterns_77 = [
            r'\[Ананьев\s*,\s*3\s*;\s*Кон\s*,\s*31\s*;\s*Дружинин\s*,\s*22\]',
            r'Ананьев\s*\[3\]\s*,\s*Кон\s*\[31\]\s*,\s*Дружинин\s*\[22\]',
        ]
        for pattern in patterns_77:
            match = re.search(pattern, text)
            if match:
                old_text = match.group(0)
                new_text = 'как отмечают Ананьев [3], Кон [31] и Дружинин [22]'
                text = re.sub(pattern, new_text, text)
                changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_text}")
                print(f"✓ Параграф {para_idx + 1} (стр.7): Ананьев, Кон, Дружинин")
                break
        
        # 2. Параграф 92 (страница 10): [30; 18; 48] или [18, 30, 48]
        patterns_92 = [
            r'\[30\s*;\s*18\s*;\s*48\]',
            r'\[18\s*,\s*30\s*,\s*48\]',
            r'\[30\s*,\s*18\s*,\s*48\]',
        ]
        for pattern in patterns_92:
            match = re.search(pattern, text)
            if match:
                old_text = match.group(0)
                new_text = 'как подчеркивают исследователи [18, 30, 48]'
                text = re.sub(pattern, new_text, text)
                changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_text}")
                print(f"✓ Параграф {para_idx + 1} (стр.10): [18, 30, 48]")
                break
        
        # 3. Параграф 105 (страница 13): [Ананьев, 3; Божович, 12; Кон, 31]
        patterns_105 = [
            r'\[Ананьев\s*,\s*3\s*;\s*Божович\s*,\s*12\s*;\s*Кон\s*,\s*31\]',
            r'Ананьев\s*\[3\]\s*,\s*Божович\s*\[12\]\s*,\s*Кон\s*\[31\]',
        ]
        for pattern in patterns_105:
            match = re.search(pattern, text)
            if match:
                old_text = match.group(0)
                new_text = 'как отмечают Ананьев [3], Божович [12] и Кон [31]'
                text = re.sub(pattern, new_text, text)
                changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_text}")
                print(f"✓ Параграф {para_idx + 1} (стр.13): Ананьев, Божович, Кон")
                break
        
        # 4. Параграф 127 (страница 16): [Выготский, 16; Кле, 29; ...]
        patterns_127 = [
            r'\[Выготский\s*,\s*16\s*;\s*Кле\s*,\s*29\s*;\s*Кривцова\s*,\s*32\s*;\s*Кулагина\s*,\s*Колюцкий\s*,\s*33\s*;\s*Мухина\s*,\s*44\s*;\s*Ремшмидт\s*,\s*51\]',
            r'Выготский\s*\[16\]\s*,\s*Кле\s*\[29\]\s*,\s*Кривцова\s*\[32\]\s*,\s*Кулагина\s*,\s*Колюцкий\s*\[33\]\s*,\s*Мухина\s*\[44\]\s*,\s*Ремшмидт\s*\[51\]',
        ]
        for pattern in patterns_127:
            match = re.search(pattern, text)
            if match:
                old_text = match.group(0)
                new_text = 'как подчеркивают Выготский [16], Кле [29], Кривцова [32], Кулагина и Колюцкий [33], Мухина [44], Ремшмидт [51]'
                text = re.sub(pattern, new_text, text)
                changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_text}")
                print(f"✓ Параграф {para_idx + 1} (стр.16): Выготский, Кле, Кривцова и др.")
                break
        
        # 5. Другие перечисления авторов (3+ авторов)
        pattern_enum = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)(?:\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)){2,}\]'
        
        def replace_enum(match):
            full_match = match.group(0)
            pairs = re.findall(r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)', full_match)
            
            if len(pairs) >= 3:
                authors_list = []
                for author, num in pairs:
                    authors_list.append(f"{author} [{num}]")
                
                if len(authors_list) > 1:
                    result = ", ".join(authors_list[:-1]) + " и " + authors_list[-1]
                    return f"как отмечают {result}"
            
            return full_match
        
        new_text = re.sub(pattern_enum, replace_enum, text)
        if new_text != text:
            # Найдено перечисление, но не одно из известных - показываем что изменилось
            old_match = re.search(pattern_enum, text)
            if old_match:
                old_text = old_match.group(0)
                new_match = re.search(r'как отмечают .+', new_text)
                if new_match:
                    changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_match.group(0)}")
            text = new_text
            print(f"✓ Параграф {para_idx + 1}: перефразировано перечисление авторов")
        
        # 6. Исправляем неверные номера источников
        # [Шнейдер, 63; Эйдемиллер, Юстицкис, 64] -> [Эйдемиллер, Юстицкис, 53]
        pattern_wrong = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*63\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*64\]'
        match = re.search(pattern_wrong, text)
        if match:
            old_text = match.group(0)
            new_text = f"[{match.group(2)}, 53]"
            text = re.sub(pattern_wrong, new_text, text)
            changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: {new_text}")
            print(f"✓ Параграф {para_idx + 1}: исправлено [Шнейдер, 63; Эйдемиллер, Юстицкис, 64] -> [Эйдемиллер, Юстицкис, 53]")
        
        # [Автор, 63] - удаляем
        pattern_63 = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*63\]'
        match = re.search(pattern_63, text)
        if match:
            old_text = match.group(0)
            text = re.sub(pattern_63, '', text)
            changes_made.append(f"БЫЛО: {old_text}\n     СТАЛО: (удалено)")
            print(f"✓ Параграф {para_idx + 1}: удалено [Автор, 63]")
        
        # [64] -> [53]
        if '[64]' in text:
            text = text.replace('[64]', '[53]')
            changes_made.append("БЫЛО: [64]\n     СТАЛО: [53]")
            print(f"✓ Параграф {para_idx + 1}: исправлен номер [64] -> [53]")
        
        # [60] -> [51]
        if '[60]' in text:
            text = text.replace('[60]', '[51]')
            changes_made.append("БЫЛО: [60]\n     СТАЛО: [51]")
            print(f"✓ Параграф {para_idx + 1}: исправлен номер [60] -> [51]")
        
        # Показываем детали изменений
        if changes_made:
            print(f"\n  Детали изменений в параграфе {para_idx + 1}:")
            for change in changes_made:
                print(f"  {change}")
            print()
        
        # Применяем изменения
        if text != original_text:
            para.text = text
            fixed_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Применяем те же исправления
                if '[64]' in text:
                    text = text.replace('[64]', '[53]')
                if '[60]' in text:
                    text = text.replace('[60]', '[51]')
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v10_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v10_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")

if __name__ == '__main__':
    fix_document()



