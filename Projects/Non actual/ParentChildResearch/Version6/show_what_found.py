# -*- coding: utf-8 -*-
"""Показывает, что именно найдено в документе v9"""
from pathlib import Path
from docx import Document
import re

def show_found():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    doc = Document(str(v9_path))
    
    print("=" * 70)
    print("ПОИСК ПРОБЛЕМНЫХ МЕСТ В ДОКУМЕНТЕ v9")
    print("=" * 70 + "\n")
    
    found_count = 0
    
    # Ищем все перечисления авторов в скобках
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # Паттерн 1: [Автор, N; Автор, N; Автор, N]
        pattern1 = r'\[[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+(?:\s*;\s*[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+){2,}\]'
        matches1 = list(re.finditer(pattern1, text))
        
        # Паттерн 2: [N; N; N] или [N, N, N]
        pattern2 = r'\[\d+\s*[;,\s]+\d+\s*[;,\s]+\d+\]'
        matches2 = list(re.finditer(pattern2, text))
        
        if matches1 or matches2:
            found_count += 1
            print(f"ПАРАГРАФ {para_idx + 1}:")
            
            if matches1:
                for match in matches1:
                    print(f"  Найдено перечисление авторов: {match.group(0)}")
            
            if matches2:
                for match in matches2:
                    print(f"  Найдено перечисление номеров: {match.group(0)}")
            
            # Показываем контекст
            print(f"  Контекст (первые 300 символов):")
            print(f"  {text[:300]}...")
            print()
    
    print(f"\nВсего найдено проблемных мест: {found_count}")
    
    # Также ищем конкретные авторы
    print("\n" + "=" * 70)
    print("ПОИСК КОНКРЕТНЫХ АВТОРОВ:")
    print("=" * 70 + "\n")
    
    authors_to_find = [
        ('Ананьев', 'Кон', 'Дружинин'),
        ('Ананьев', 'Божович', 'Кон'),
        ('Выготский', 'Кле', 'Кривцова'),
    ]
    
    for author_group in authors_to_find:
        print(f"Ищем: {', '.join(author_group)}")
        found = False
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            if all(author in text for author in author_group):
                # Проверяем, есть ли номера рядом
                if re.search(r'\[\d+\]', text):
                    print(f"  Параграф {para_idx + 1}: найдено")
                    print(f"    {text[:200]}...")
                    found = True
                    break
        if not found:
            print(f"  Не найдено")
        print()

if __name__ == '__main__':
    show_found()



