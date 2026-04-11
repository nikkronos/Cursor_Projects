# -*- coding: utf-8 -*-
"""
Скрипт для поиска конкретных проблемных мест в документе v9
"""
import re
from pathlib import Path
from docx import Document

def find_problematic_places():
    """Ищет конкретные проблемные места, указанные пользователем"""
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("=" * 70)
    print("ПОИСК ПРОБЛЕМНЫХ МЕСТ В ДОКУМЕНТЕ v9")
    print("=" * 70)
    
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен\n")
    
    # Ищем конкретные проблемные места
    problematic_patterns = [
        {
            'name': 'Страница 7 - Ананьев, Кон, Дружинин',
            'patterns': [
                r'\[Ананьев.*3.*Кон.*31.*Дружинин.*22\]',
                r'Ананьев.*\[3\].*Кон.*\[31\].*Дружинин.*\[22\]',
                r'\[Ананьев,\s*3[;\s]+Кон,\s*31[;\s]+Дружинин,\s*22\]',
            ]
        },
        {
            'name': 'Страница 10 - [18, 30, 48] или [30; 18; 48]',
            'patterns': [
                r'\[18[,\s;]\s*30[,\s;]\s*48\]',
                r'\[30[,\s;]\s*18[,\s;]\s*48\]',
            ]
        },
        {
            'name': 'Страница 13 - Ананьев, Божович, Кон',
            'patterns': [
                r'\[Ананьев.*3.*Божович.*12.*Кон.*31\]',
                r'Ананьев.*\[3\].*Божович.*\[12\].*Кон.*\[31\]',
                r'\[Ананьев,\s*3[;\s]+Божович,\s*12[;\s]+Кон,\s*31\]',
            ]
        },
        {
            'name': 'Страница 16 - Выготский, Кле, Кривцова и др.',
            'patterns': [
                r'\[Выготский.*16.*Кле.*29.*Кривцова.*32.*Кулагина.*Колюцкий.*33.*Мухина.*44.*Ремшмидт.*51\]',
                r'Выготский.*\[16\].*Кле.*\[29\].*Кривцова.*\[32\].*Кулагина.*Колюцкий.*\[33\].*Мухина.*\[44\].*Ремшмидт.*\[51\]',
            ]
        }
    ]
    
    found_places = []
    
    # Анализируем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        for problem in problematic_patterns:
            for pattern in problem['patterns']:
                if re.search(pattern, text, re.IGNORECASE | re.DOTALL):
                    found_places.append({
                        'para': para_idx + 1,
                        'name': problem['name'],
                        'text': text[:300] + '...' if len(text) > 300 else text,
                        'full_text': text
                    })
                    break
    
    # Выводим результаты
    if found_places:
        print(f"Найдено проблемных мест: {len(found_places)}\n")
        for i, place in enumerate(found_places, 1):
            print(f"{i}. {place['name']}")
            print(f"   Параграф: {place['para']}")
            print(f"   Текст: {place['text']}")
            print()
    else:
        print("⚠ Проблемные места не найдены!")
        print("\nПроверяю альтернативные форматы...\n")
        
        # Ищем любые перечисления авторов
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():
                continue
            
            # Ищем паттерны типа [Автор, N; Автор, N; Автор, N]
            pattern = r'\[[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+(?:\s*;\s*[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+){2,}\]'
            matches = re.finditer(pattern, text)
            for match in matches:
                print(f"Параграф {para_idx + 1}: найдено перечисление авторов")
                print(f"  Текст: {text[:200]}...")
                print(f"  Совпадение: {match.group(0)}")
                print()
    
    # Сохраняем результаты
    if found_places:
        report_file = version6 / "НАЙДЕННЫЕ_ПРОБЛЕМНЫЕ_МЕСТА.txt"
        with open(report_file, 'w', encoding='utf-8') as f:
            f.write("НАЙДЕННЫЕ ПРОБЛЕМНЫЕ МЕСТА В ДОКУМЕНТЕ v9\n")
            f.write("=" * 70 + "\n\n")
            for i, place in enumerate(found_places, 1):
                f.write(f"{i}. {place['name']}\n")
                f.write(f"   Параграф: {place['para']}\n")
                f.write(f"   Полный текст:\n{place['full_text']}\n\n")
        print(f"✓ Результаты сохранены в: {report_file.name}")

if __name__ == '__main__':
    find_problematic_places()



