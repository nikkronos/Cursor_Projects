# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v14 -> v15:
1. Страница 4: добавить источник к Эйдемиллер
2. Страница 4: добавить источник к Прихожан
3. Страница 4: к [30, 40] добавить фамилии и перефразировать
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    # Пробуем найти v14
    v14_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v14.docx"
    if not v14_path.exists():
        # Пробуем найти v13
        v13_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v13.docx"
        if v13_path.exists():
            v14_path = v13_path
            print(f"ВНИМАНИЕ: v14 не найден, используем v13")
        else:
            print(f"ERROR: Не найден документ v14 или v13")
            return
    
    v15_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v15.docx"
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v14 -> v15")
    print("=" * 70)
    
    doc = Document(str(v14_path))
    print(f"✓ Документ загружен: {v14_path.name}\n")
    
    fixed_count = 0
    
    # Номера источников:
    # Эйдемиллер, Юстицкис - 53
    # Прихожан, Толстых - 41 (но в тексте может быть [49], нужно проверить)
    # 30 - Леонтьев А.Н.
    # 40 - Поваренков Ю.П.
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        
        # 1. Добавить источник к Эйдемиллер
        # Ищем: "Андреева [5], Эйдемиллер" (без источника после Эйдемиллер)
        # Проверяем, есть ли уже источник у Эйдемиллер
        has_eidemiller_source = ('[53]' in text and 'Эйдемиллер' in text and 'Юстицкис' in text) or \
                                ('Эйдемиллер, Юстицкис' in text)
        
        if not has_eidemiller_source and 'Эйдемиллер' in text and 'Андреева' in text:
            # Простая замена: ", Эйдемиллер" -> ", Эйдемиллер, Юстицкис [53]"
            # Но только если после Эйдемиллер нет источника
            if ', Эйдемиллер' in text:
                # Проверяем, что после "Эйдемиллер" нет источника
                eidemiller_pos = text.find(', Эйдемиллер')
                if eidemiller_pos != -1:
                    after_eidemiller = text[eidemiller_pos + len(', Эйдемиллер'):eidemiller_pos + len(', Эйдемиллер') + 20]
                    # Если после "Эйдемиллер" нет [53] или Юстицкис, добавляем
                    if '[53]' not in after_eidemiller and 'Юстицкис' not in after_eidemiller:
                        text = text.replace(', Эйдемиллер', ', Эйдемиллер, Юстицкис [53]', 1)
                        print(f"✓ Параграф {para_idx + 1} (стр.4): добавлен источник к Эйдемиллер -> [53]")
                        fixed_count += 1
        
        # 2. Исправить/добавить источник к Прихожан
        # Ищем: "Прихожан, Толстых [49]" - исправляем на [41]
        # По списку литературы Прихожан, Толстых - номер 41
        pattern_prikhozhan = r'Прихожан\s*,\s*Толстых\s*\[49\]'
        if re.search(pattern_prikhozhan, text):
            text = re.sub(pattern_prikhozhan, 'Прихожан, Толстых [41]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.4): исправлен номер источника Прихожан, Толстых [49] -> [41]")
            fixed_count += 1
        
        # Также ищем: "Гинзбург [18], Климов [30], Прихожан, Толстых [49]"
        pattern_prikhozhan2 = r'(Гинзбург\s*\[18\]\s*,\s*Климов\s*\[30\]\s*,\s*)Прихожан\s*,\s*Толстых\s*\[49\]'
        if re.search(pattern_prikhozhan2, text):
            text = re.sub(pattern_prikhozhan2, r'\1Прихожан, Толстых [41]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.4): исправлен номер источника Прихожан, Толстых [49] -> [41]")
            fixed_count += 1
        
        # 3. К [30, 40] добавить фамилии и перефразировать
        # Ищем: "Социально-когнитивный и карьерный подход [30, 40]" или "[30; 40]"
        pattern_30_40 = r'Социально-когнитивный\s+и\s+карьерный\s+подход\s*\[30\s*[;,\s]+\s*40\]'
        if re.search(pattern_30_40, text):
            text = re.sub(pattern_30_40, 
                         'Социально-когнитивный и карьерный подход, который изучали Леонтьев [30] и Поваренков [40], объясняющий влияние родительских моделей, самоэффективности и ожиданий результатов на профессиональные решения',
                         text)
            print(f"✓ Параграф {para_idx + 1} (стр.4): добавлены фамилии к [30, 40] и перефразировано")
            fixed_count += 1
        
        # Также ищем: "подход [30; 40], объясняющий"
        pattern_30_40_2 = r'(подход\s*)\[30\s*[;,\s]+\s*40\]\s*,\s*объясняющий'
        if re.search(pattern_30_40_2, text) and 'Леонтьев' not in text:
            text = re.sub(pattern_30_40_2, 
                         r'\1который изучали Леонтьев [30] и Поваренков [40], объясняющий',
                         text)
            print(f"✓ Параграф {para_idx + 1} (стр.4): добавлены фамилии к [30, 40] и перефразировано")
            fixed_count += 1
        
        # Применяем изменения
        if text != original_text:
            para.text = text
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Применяем те же исправления
                if 'Эйдемиллер' in text and '[53]' not in text:
                    text = re.sub(r'Эйдемиллер(?!\s*[,\[\d])', 'Эйдемиллер, Юстицкис [53]', text)
                
                if 'Прихожан, Толстых [49]' in text:
                    text = text.replace('Прихожан, Толстых [49]', 'Прихожан, Толстых [41]')
                
                if '[30, 40]' in text and 'Леонтьев' not in text:
                    text = text.replace('[30, 40]', 'Леонтьев [30] и Поваренков [40]')
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v15_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v15_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")

if __name__ == '__main__':
    fix_document()
