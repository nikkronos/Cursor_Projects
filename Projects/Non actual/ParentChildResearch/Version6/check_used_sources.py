# -*- coding: utf-8 -*-
"""
Скрипт для проверки, какие источники действительно используются в документе
"""
from pathlib import Path
from docx import Document
import re

def find_all_citations_in_text(text: str) -> set:
    """
    Находит все номера источников в тексте.
    Возвращает множество номеров.
    """
    numbers = set()
    
    # Паттерн для поиска ссылок: [числа, разделённые запятыми, точками с запятой или дефисами]
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    
    for match in re.finditer(pattern, text):
        citation_text = match.group(1)
        # Извлекаем все числа
        nums = re.findall(r'\d+', citation_text)
        for num_str in nums:
            numbers.add(int(num_str))
    
    return numbers

def check_used_sources():
    version6 = Path(__file__).parent
    
    v15_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v15.docx"
    if not v15_path.exists():
        print(f"ERROR: Документ v15 не найден: {v15_path}")
        return
    
    print("=" * 70)
    print("ПРОВЕРКА ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ В ДОКУМЕНТЕ")
    print("=" * 70)
    
    doc = Document(str(v15_path))
    print(f"✓ Документ загружен: {v15_path.name}\n")
    
    used_numbers = set()
    
    # Собираем все использованные номера из параграфов
    for para in doc.paragraphs:
        numbers = find_all_citations_in_text(para.text)
        used_numbers.update(numbers)
    
    # Собираем из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                numbers = find_all_citations_in_text(cell.text)
                used_numbers.update(numbers)
    
    print(f"📊 Найдено использованных номеров источников: {len(used_numbers)}")
    print(f"   Номера: {sorted(used_numbers)}\n")
    
    # Читаем исходный список литературы
    lit_path = version6 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    if lit_path.exists():
        with open(lit_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Парсим все номера из списка
        pattern = r'^(\d+)\.\s+'
        matches = re.finditer(pattern, content, re.MULTILINE)
        all_numbers = set()
        for match in matches:
            all_numbers.add(int(match.group(1)))
        
        print(f"📚 Всего источников в списке литературы: {len(all_numbers)}")
        print(f"   Номера: {sorted(all_numbers)}\n")
        
        # Находим неиспользуемые
        unused = all_numbers - used_numbers
        if unused:
            print(f"⚠ НЕИСПОЛЬЗУЕМЫЕ ИСТОЧНИКИ ({len(unused)}):")
            for num in sorted(unused):
                # Находим название источника
                pattern_entry = rf'^{num}\.\s+(.+?)(?=^\d+\.\s+|$)'
                match = re.search(pattern_entry, content, re.MULTILINE | re.DOTALL)
                if match:
                    entry = match.group(1).strip()[:80]
                    print(f"   {num}. {entry}...")
            print()
        
        # Находим используемые, но отсутствующие в списке
        missing = used_numbers - all_numbers
        if missing:
            print(f"❌ ИСПОЛЬЗУЮТСЯ В ТЕКСТЕ, НО ОТСУТСТВУЮТ В СПИСКЕ ({len(missing)}):")
            for num in sorted(missing):
                print(f"   Номер {num}")
            print()
    
    return used_numbers

if __name__ == '__main__':
    check_used_sources()
