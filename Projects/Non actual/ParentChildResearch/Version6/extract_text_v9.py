# -*- coding: utf-8 -*-
"""Извлечение всего текста из документа v9 для анализа"""
from pathlib import Path
from docx import Document

def extract_text():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("Извлекаю текст из документа v9...")
    doc = Document(str(v9_path))
    
    all_text = []
    
    # Извлекаем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append(f"=== ПАРАГРАФ {para_idx + 1} ===\n{text}\n")
    
    # Извлекаем таблицы
    for table_idx, table in enumerate(doc.tables):
        all_text.append(f"\n=== ТАБЛИЦА {table_idx + 1} ===\n")
        for row_idx, row in enumerate(table.rows):
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                all_text.append(" | ".join(row_text) + "\n")
    
    # Сохраняем в файл
    output_file = version6 / "ВСЕ_ТЕКСТ_ИЗ_v9.txt"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write('\n'.join(all_text))
    
    print(f"✓ Текст сохранен в: {output_file.name}")
    print(f"✓ Всего параграфов: {len([p for p in doc.paragraphs if p.text.strip()])}")
    print(f"✓ Всего таблиц: {len(doc.tables)}")

if __name__ == '__main__':
    extract_text()



