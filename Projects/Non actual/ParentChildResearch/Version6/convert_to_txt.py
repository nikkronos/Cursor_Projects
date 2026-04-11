# -*- coding: utf-8 -*-
"""Конвертация документа v9 в текстовый файл для удобной работы"""
from pathlib import Path
from docx import Document

def convert_to_txt():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("Конвертирую документ v9 в текстовый файл...")
    doc = Document(str(v9_path))
    
    output_file = version6 / "ВКР_v9_ТЕКСТ.txt"
    
    with open(output_file, 'w', encoding='utf-8') as f:
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                f.write(f"=== ПАРАГРАФ {para_idx + 1} ===\n")
                f.write(f"{text}\n\n")
        
        # Также извлекаем таблицы
        for table_idx, table in enumerate(doc.tables):
            f.write(f"\n=== ТАБЛИЦА {table_idx + 1} ===\n")
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    f.write(" | ".join(row_text) + "\n")
            f.write("\n")
    
    print(f"✓ Текст сохранен в: {output_file.name}")
    print(f"✓ Всего параграфов: {len([p for p in doc.paragraphs if p.text.strip()])}")

if __name__ == '__main__':
    convert_to_txt()



