# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v16:
1. Удаление неиспользуемых источников из списка литературы
2. Пересоздание отсортированного списка только с используемыми источниками
3. Обновление номеров в документе v16
"""
from pathlib import Path
from docx import Document
import re
from typing import Dict, List, Tuple

def extract_author_surname(entry: str) -> str:
    """
    Извлекает фамилию первого автора из записи библиографии.
    Возвращает фамилию для сортировки (в верхнем регистре).
    """
    # Убираем номер в начале (например, "1. " или "10. ")
    entry = re.sub(r'^\d+\.\s*', '', entry.strip())
    
    # Ищем первую фамилию
    match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
    if match:
        surname = match.group(1)
        return surname.upper()
    
    # Для английских фамилий
    match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z]\.?)*)', entry)
    if match:
        surname = match.group(1).split()[0]
        return surname.upper()
    
    # Если не нашлось, пробуем найти любую заглавную букву
    match = re.search(r'([А-ЯЁA-Z])', entry)
    if match:
        return match.group(1)
    
    return 'Я'  # Если не удалось определить, ставим в конец

def parse_literature_list(lit_path: Path) -> Dict[int, Dict]:
    """
    Парсит список литературы и возвращает словарь {номер: {entry, old_number}}
    """
    literature = {}
    
    if not lit_path.exists():
        print(f"ERROR: Файл списка литературы не найден: {lit_path}")
        return literature
    
    with open(lit_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_number = None
    current_entry = []
    
    for line in lines:
        line = line.strip()
        
        # Пропускаем заголовки и комментарии
        if line.startswith('#') or line.startswith('*') or not line:
            continue
        
        # Ищем строку с номером: "N. Автор..."
        match = re.match(r'^(\d+)\.\s*(.+)', line)
        if match:
            # Сохраняем предыдущую запись, если есть
            if current_number is not None:
                entry_text = ' '.join(current_entry)
                literature[current_number] = {
                    'entry': entry_text,
                    'old_number': current_number
                }
            
            # Начинаем новую запись
            current_number = int(match.group(1))
            current_entry = [match.group(2)]
        else:
            # Продолжение текущей записи
            if current_number is not None and line:
                current_entry.append(line)
    
    # Сохраняем последнюю запись
    if current_number is not None:
        entry_text = ' '.join(current_entry)
        literature[current_number] = {
            'entry': entry_text,
            'old_number': current_number
        }
    
    return literature

def find_all_citations_in_text(text: str) -> set:
    """
    Находит все номера источников в тексте.
    Возвращает множество номеров.
    """
    numbers = set()
    
    # Паттерн для поиска ссылок: [числа, разделённые запятыми, точками с запятой или дефисами]
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    
    for match in re.finditer(pattern, text):
        citation_text = match.group(1)
        # Извлекаем все числа
        nums = re.findall(r'\d+', citation_text)
        for num_str in nums:
            numbers.add(int(num_str))
    
    return numbers

def get_used_sources_from_document(doc: Document) -> set:
    """
    Извлекает все использованные номера источников из документа.
    """
    used_numbers = set()
    
    # Собираем все использованные номера из параграфов
    for para in doc.paragraphs:
        numbers = find_all_citations_in_text(para.text)
        used_numbers.update(numbers)
    
    # Собираем из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                numbers = find_all_citations_in_text(cell.text)
                used_numbers.update(numbers)
    
    return used_numbers

def sort_literature_by_alphabet(literature: Dict[int, Dict]) -> List[Tuple[str, int, Dict]]:
    """
    Сортирует источники по алфавиту (по фамилии первого автора).
    Возвращает список кортежей (фамилия_для_сортировки, старый_номер, данные)
    """
    # Правильный порядок русского алфавита для сортировки
    russian_alphabet_order = {
        'А': 0, 'Б': 1, 'В': 2, 'Г': 3, 'Д': 4, 'Е': 5, 'Ё': 5.5,
        'Ж': 6, 'З': 7, 'И': 8, 'Й': 9, 'К': 10, 'Л': 11, 'М': 12,
        'Н': 13, 'О': 14, 'П': 15, 'Р': 16, 'С': 17, 'Т': 18, 'У': 19,
        'Ф': 20, 'Х': 21, 'Ц': 22, 'Ч': 23, 'Ш': 24, 'Щ': 25,
        'Ъ': 26, 'Ы': 27, 'Ь': 28, 'Э': 29, 'Ю': 30, 'Я': 31
    }
    
    def get_sort_key(surname):
        """
        Возвращает кортеж для правильной сортировки по русскому алфавиту.
        """
        if not surname:
            return (3, 999, surname)  # Пустые в самом конце
        
        first_char = surname[0].upper()
        
        # Русские буквы
        if first_char in russian_alphabet_order:
            pos = russian_alphabet_order[first_char]
            # Сортируем сначала по позиции в алфавите, потом по полной фамилии
            return (0, pos, surname.upper())
        
        # Английские буквы
        elif first_char.isalpha() and ord(first_char) < 128:
            return (1, ord(first_char), surname.upper())
        
        # Остальное (кириллица не из основного алфавита или другие символы)
        else:
            return (2, ord(first_char) if first_char else 999, surname.upper())
    
    sorted_items = []
    
    for old_num, data in literature.items():
        entry = data['entry']
        surname = extract_author_surname(entry)
        sorted_items.append((surname, old_num, data))
    
    # Сортируем с правильным ключом
    sorted_items.sort(key=lambda x: get_sort_key(x[0]))
    
    return sorted_items

def create_number_mapping(sorted_items: List[Tuple]) -> Dict[int, int]:
    """
    Создаёт маппинг старый_номер -> новый_номер
    """
    mapping = {}
    
    for new_num, (first_letter, old_num, data) in enumerate(sorted_items, start=1):
        mapping[old_num] = new_num
        data['new_number'] = new_num
    
    return mapping

def find_all_citations_in_text_for_replacement(text: str) -> List[Tuple[int, int, str]]:
    """
    Находит все ссылки на источники в тексте.
    Возвращает список (start_pos, end_pos, citation_text)
    """
    citations = []
    
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    
    for match in re.finditer(pattern, text):
        start = match.start()
        end = match.end()
        citation_text = match.group(0)
        citations.append((start, end, citation_text))
    
    return citations

def replace_citation_numbers(citation_text: str, mapping: Dict[int, int]) -> str:
    """
    Заменяет номера в ссылке согласно маппингу.
    """
    numbers = re.findall(r'\d+', citation_text)
    
    if not numbers:
        return citation_text
    
    new_numbers = []
    for num_str in numbers:
        old_num = int(num_str)
        if old_num in mapping:
            new_numbers.append(str(mapping[old_num]))
        else:
            new_numbers.append(num_str)
    
    # Восстанавливаем формат с разделителями
    content = citation_text[1:-1]
    
    if ',' in content:
        sep = ', '
        result = '[' + sep.join(new_numbers) + ']'
    elif ';' in content:
        sep = '; '
        result = '[' + sep.join(new_numbers) + ']'
    elif '-' in content:
        sep = '-'
        result = '[' + sep.join(new_numbers) + ']'
    else:
        result = '[' + new_numbers[0] + ']'
    
    return result

def update_document_citations(doc: Document, mapping: Dict[int, int]) -> int:
    """
    Обновляет все ссылки на источники в документе.
    """
    replaced_count = 0
    
    # Обрабатываем параграфы
    for para in doc.paragraphs:
        text = para.text
        citations = find_all_citations_in_text_for_replacement(text)
        
        if citations:
            for start, end, citation_text in reversed(citations):
                new_citation = replace_citation_numbers(citation_text, mapping)
                if new_citation != citation_text:
                    text = text[:start] + new_citation + text[end:]
                    replaced_count += 1
            
            if text != para.text:
                para.text = text
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                citations = find_all_citations_in_text_for_replacement(text)
                
                if citations:
                    for start, end, citation_text in reversed(citations):
                        new_citation = replace_citation_numbers(citation_text, mapping)
                        if new_citation != citation_text:
                            text = text[:start] + new_citation + text[end:]
                            replaced_count += 1
                    
                    if text != cell.text:
                        cell.text = text
    
    return replaced_count

def update_literature_list_in_document(doc: Document, sorted_items: List[Tuple]):
    """
    Обновляет список литературы в документе Word.
    Ищет раздел со списком литературы и заменяет его отсортированным списком.
    """
    # Ищем начало списка литературы (обычно это заголовок или текст "Список литературы")
    literature_start_idx = -1
    literature_end_idx = -1
    
    for idx, para in enumerate(doc.paragraphs):
        text_lower = para.text.lower()
        # Ищем начало списка литературы
        if literature_start_idx == -1:
            if any(keyword in text_lower for keyword in ['список литературы', 'литература', 'библиографический список']):
                literature_start_idx = idx
                continue
        else:
            # Если нашли начало, ищем конец списка (пустая строка или новый раздел)
            if para.text.strip() == '':
                # Пропускаем одну пустую строку
                if idx < len(doc.paragraphs) - 1:
                    continue
            # Если встретили новый заглавный заголовок, это конец списка
            if para.style.name.startswith('Heading') or para.style.name.startswith('Заголовок'):
                if idx > literature_start_idx + 10:  # Убеждаемся, что прошли достаточно далеко
                    literature_end_idx = idx
                    break
    
    # Если не нашли конец явно, ищем по паттерну номеров
    if literature_start_idx >= 0 and literature_end_idx == -1:
        # Ищем последний параграф с номером источника
        for idx in range(literature_start_idx, len(doc.paragraphs)):
            para = doc.paragraphs[idx]
            # Если это не номер источника и не пустая строка после списка
            if not re.match(r'^\d+\.\s+', para.text) and para.text.strip() != '':
                # Проверяем, не начался ли новый раздел
                if idx > literature_start_idx + 10:
                    literature_end_idx = idx
                    break
    
    # Если нашли список литературы в документе
    if literature_start_idx >= 0:
        print(f"  Найден список литературы в документе: параграфы {literature_start_idx}-{literature_end_idx if literature_end_idx > 0 else len(doc.paragraphs)-1}")
        
        # Удаляем старые параграфы со списком литературы (с конца, чтобы не сбить индексы)
        if literature_end_idx > literature_start_idx:
            # Удаляем параграфы от конца до начала
            for idx in range(literature_end_idx - 1, literature_start_idx - 1, -1):
                if idx < len(doc.paragraphs):
                    p = doc.paragraphs[idx]
                    p.clear()
                    # Удаляем параграф
                    p_element = p._element
                    p_element.getparent().remove(p_element)
        else:
            # Если не нашли конец явно, удаляем все параграфы от начала до конца документа
            # Но это опасно, так что ограничимся 200 параграфами
            max_remove = min(literature_start_idx + 200, len(doc.paragraphs))
            for idx in range(max_remove - 1, literature_start_idx - 1, -1):
                para = doc.paragraphs[idx]
                if re.match(r'^\d+\.\s+', para.text):
                    para.clear()
                    para_element = para._element
                    para_element.getparent().remove(para_element)
        
        # Вставляем новый список литературы после заголовка
        insert_idx = literature_start_idx + 1
        for new_num, (surname, old_num, data) in enumerate(sorted_items, start=1):
            entry_text = f"{new_num}. {data['entry']}"
            # Вставляем новый параграф
            new_para = doc.paragraphs[insert_idx].insert_paragraph_before(entry_text)
            insert_idx += 1
            # Добавляем пустую строку после каждого источника
            empty_para = doc.paragraphs[insert_idx].insert_paragraph_before("")
            insert_idx += 1
    
    else:
        print("  ⚠ Список литературы в документе не найден автоматически")
        print("  Список будет обновлён только в .md файле")


def write_sorted_literature_list(lit_path: Path, sorted_items: List[Tuple], max_source: int, update_original=False):
    """
    Записывает отсортированный список литературы в файл.
    """
    if update_original:
        output_path = lit_path
    else:
        output_path = lit_path.parent / f"{lit_path.stem}_sorted.md"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ\n\n")
        f.write("*Сформирован на основе реального использования в тексте*\n\n")
        f.write(f"*Всего источников: {max_source}*\n")
        f.write("*Отсортировано по алфавиту*\n\n")
        
        for new_num, (first_letter, old_num, data) in enumerate(sorted_items, start=1):
            f.write(f"{new_num}. {data['entry']}\n\n")
    
    return output_path

def fix_document():
    version6 = Path(__file__).parent
    
    # Пути к файлам
    v15_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v15.docx"
    if not v15_path.exists():
        print(f"ERROR: Документ v15 не найден: {v15_path}")
        return
    
    lit_path = version6 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    if not lit_path.exists():
        print(f"ERROR: Список литературы не найден: {lit_path}")
        return
    
    v16_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v16.docx"
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v16: УДАЛЕНИЕ НЕИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ")
    print("=" * 70)
    
    # Загружаем документ v15 (оригинал)
    doc = Document(str(v15_path))
    print(f"✓ Документ загружен: {v15_path.name}\n")
    
    # Извлекаем используемые источники
    print("🔍 Извлечение используемых источников из документа...")
    used_numbers = get_used_sources_from_document(doc)
    print(f"✓ Найдено использованных источников: {len(used_numbers)}")
    print(f"   Номера: {sorted(used_numbers)}\n")
    
    # Парсим список литературы
    print("📚 Парсинг списка литературы...")
    all_literature = parse_literature_list(lit_path)
    print(f"✓ Найдено источников в списке: {len(all_literature)}\n")
    
    # Фильтруем: оставляем только используемые
    print("✂ Фильтрация: оставляем только используемые источники...")
    used_literature = {num: data for num, data in all_literature.items() if num in used_numbers}
    unused_literature = {num: data for num, data in all_literature.items() if num not in used_numbers}
    
    print(f"✓ Используемых источников: {len(used_literature)}")
    print(f"✓ Неиспользуемых источников (будут удалены): {len(unused_literature)}")
    
    if unused_literature:
        print("\n  Удаляемые источники:")
        for num in sorted(unused_literature.keys()):
            entry = unused_literature[num]['entry'][:60]
            print(f"    {num}. {entry}...")
    print()
    
    # Сортируем по алфавиту
    print("🔤 Сортировка списка литературы по алфавиту...")
    sorted_items = sort_literature_by_alphabet(used_literature)
    print(f"✓ Отсортировано {len(sorted_items)} источников\n")
    
    # Создаём маппинг
    print("🔄 Создание маппинга старых номеров на новые...")
    mapping = create_number_mapping(sorted_items)
    
    # Выводим изменения (первые 10)
    print("  Примеры изменений номеров:")
    count = 0
    for old_num, new_num in sorted(mapping.items()):
        if old_num != new_num:
            if count < 10:
                entry = used_literature[old_num]['entry'][:50] + "..."
                print(f"    {old_num} -> {new_num}: {entry}")
            count += 1
    if count > 10:
        print(f"    ... и ещё {count - 10} изменений")
    print()
    
    # Обновляем ссылки в документе
    print("📝 Обновление ссылок в документе...")
    replaced_count = update_document_citations(doc, mapping)
    print(f"✓ Заменено ссылок: {replaced_count}\n")
    
    # Обновляем список литературы в самом документе Word
    print("📄 Обновление списка литературы в документе Word...")
    update_literature_list_in_document(doc, sorted_items)
    print()
    
    # Сохраняем отсортированный список литературы
    print("💾 Сохранение отсортированного списка литературы...")
    max_source = len(sorted_items)
    # Сохраняем в sorted.md
    sorted_lit_path = write_sorted_literature_list(lit_path, sorted_items, max_source, update_original=False)
    print(f"✓ Сохранён отсортированный список: {sorted_lit_path.name}")
    # Также обновляем исходный файл
    original_lit_path = write_sorted_literature_list(lit_path, sorted_items, max_source, update_original=True)
    print(f"✓ Обновлён исходный файл: {original_lit_path.name}\n")
    
    # Сохраняем документ
    doc.save(str(v16_path))
    
    print("=" * 70)
    print("РЕЗУЛЬТАТЫ")
    print("=" * 70)
    print(f"✓ Документ сохранён: {v16_path.name}")
    print(f"✓ Отсортированный список литературы: {sorted_lit_path.name}")
    print(f"✓ Заменено ссылок в документе: {replaced_count}")
    print(f"✓ Всего источников (после удаления неиспользуемых): {max_source}")
    print(f"✓ Удалено неиспользуемых источников: {len(unused_literature)}")

if __name__ == '__main__':
    fix_document()
