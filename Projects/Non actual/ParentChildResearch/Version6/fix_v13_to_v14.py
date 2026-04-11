# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v13 -> v14:
1) Стр.5: добавить фамилии к источникам [16], [36, 40], [42] (фамилии — В ТЕКСТЕ, не в скобках)
2) Стр.20: для [22, 35] добавить контекст (с фамилиями)
3) Стр.20: если встречается "Фромм" без ссылки — добавить ссылку; если источника нет в списке литературы документа — добавить его

Важно: В квадратных скобках должны быть ТОЛЬКО числа.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Dict, Optional, Tuple

from docx import Document


def _extract_surnames(entry: str) -> str:
    """
    Очень простой экстрактор фамилий из библиографической записи.
    Берём всё до первой точки, затем берём фамилии (первое слово) из кусочков, разделённых запятыми.
    """
    head = entry.strip()
    # обрезаем после первой точки (часто после инициалов/названия)
    if "." in head:
        head = head.split(".", 1)[0].strip()
    # иногда до первой точки нет — тогда берём всю строку
    # делим по запятым, берём первое слово (фамилию)
    parts = [p.strip() for p in head.split(",") if p.strip()]
    surnames = []
    for p in parts:
        tok = p.split()[0].strip()
        # отфильтруем мусор (типа "Holland" тоже ок)
        if tok:
            surnames.append(tok)
    # если фамилий несколько — аккуратно объединим
    if not surnames:
        return ""
    if len(surnames) == 1:
        return surnames[0]
    if len(surnames) == 2:
        return f"{surnames[0]} и {surnames[1]}"
    return ", ".join(surnames[:-1]) + " и " + surnames[-1]


def _parse_doc_bibliography(doc: Document) -> Dict[int, str]:
    """
    Парсит список литературы из документа.
    Ищем заголовок "СПИСОК ЛИТЕРАТУРЫ" (или "Список литературы"), далее собираем строки вида "N. ..."
    """
    bib: Dict[int, str] = {}
    in_bib = False

    for para in doc.paragraphs:
        t = (para.text or "").strip()
        if not t:
            continue

        if re.fullmatch(r"СПИСОК\s+ЛИТЕРАТУРЫ", t.upper()):
            in_bib = True
            continue

        if in_bib:
            # выход из библиографии по началу приложений/конца
            if re.fullmatch(r"ПРИЛОЖЕНИЯ", t.upper()):
                break
            m = re.match(r"^(\d+)\.\s+(.*)$", t)
            if m:
                num = int(m.group(1))
                entry = m.group(2).strip()
                bib[num] = entry
            else:
                # иногда запись переносится на несколько строк — не поддерживаем склейку,
                # но и не выходим: просто игнорируем
                pass

    return bib


def _find_bib_number_by_keyword(bib: Dict[int, str], keyword: str) -> Optional[int]:
    kw = keyword.lower()
    for num, entry in bib.items():
        if kw in entry.lower():
            return num
    return None


def _replace_first(pattern: str, repl: str, text: str, flags: int = 0) -> Tuple[str, bool]:
    m = re.search(pattern, text, flags)
    if not m:
        return text, False
    return re.sub(pattern, repl, text, count=1, flags=flags), True


def fix_document() -> None:
    version6 = Path(__file__).parent

    # Берём исходник: предпочтительно v13, иначе v12, иначе v11
    candidates = [
        version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v13.docx",
        version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v12.docx",
        version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v11.docx",
    ]
    src_path = next((p for p in candidates if p.exists()), None)
    if not src_path:
        print("ERROR: Не найден исходный документ v13/v12/v11 в Version6")
        return

    dst_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v14.docx"

    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА -> v14")
    print("=" * 70)
    print(f"✓ Исходник: {src_path.name}")

    doc = Document(str(src_path))

    bib = _parse_doc_bibliography(doc)
    if not bib:
        print("⚠ Не удалось распарсить список литературы в документе (буду использовать фамилии-шаблоны).")

    def surname_for(num: int, fallback: str = "") -> str:
        if num in bib:
            s = _extract_surnames(bib[num])
            return s or fallback
        return fallback

    s16 = surname_for(16, "авторы")
    s36 = surname_for(36, "авторы")
    s40 = surname_for(40, "авторы")
    s42 = surname_for(42, "авторы")
    s22 = surname_for(22, "авторы")
    s35 = surname_for(35, "авторы")

    # --- 3) Фромм: источник и ссылка ---
    fromm_num = _find_bib_number_by_keyword(bib, "Фромм") if bib else None
    if fromm_num is None:
        # Попробуем добавить в список литературы (если нашли секцию)
        # Берём запись из нашего md-списка (каноническая)
        fromm_entry = (
            "Фромм А. Азбука для родителей / А. Фромм; пер. с англ. И.Г. Константиновой. — Л.: Лениздат, 1991. — 320 с."
        )
        if bib:
            new_num = max(bib.keys()) + 1
        else:
            new_num = 54

        inserted = False
        for i, para in enumerate(doc.paragraphs):
            if re.fullmatch(r"ПРИЛОЖЕНИЯ", (para.text or "").strip().upper()):
                # вставим перед приложениями (в конец списка литературы обычно)
                new_para = doc.add_paragraph()
                new_para.text = f"{new_num}. {fromm_entry}"
                para._element.addprevious(new_para._element)
                inserted = True
                break

        if inserted:
            fromm_num = new_num
            print(f"✓ Добавлен источник Фромм в список литературы под номером {fromm_num}")
        else:
            # не нашли куда вставить — просто будем использовать номер 52 (как в md), но без гарантии
            fromm_num = 52
            print("⚠ Не удалось вставить источник Фромм в список литературы автоматически; использую ссылку [52].")
    else:
        print(f"✓ Источник Фромм найден в списке литературы документа: №{fromm_num}")

    fixed = 0

    for idx, para in enumerate(doc.paragraphs):
        t0 = para.text or ""
        t = t0
        changed = False

        # --- 1) Стр.5: добавление фамилий в контекст ---
        # Возрастной и культурно-исторический подход [16] -> ... подход, представленный в работах {s16} [16]
        if "Возрастной" in t and "культурно-истор" in t and re.search(r"\[16\]", t):
            t, ok = _replace_first(
                r"подход\s*\[16\]",
                f"подход, представленный в работах {s16} [16]",
                t,
                flags=re.IGNORECASE,
            )
            if ok:
                print(f"✓ Параграф {idx+1}: добавлен автор к [16] ({s16})")
                changed = True

        # Ценностно-смысловой подход [36, 40] -> ... подход, представленный в работах {s36} [36] и {s40} [40]
        if "Ценностно-смысловой" in t and re.search(r"\[\s*36\s*[,;]\s*40\s*\]", t):
            t, ok = _replace_first(
                r"подход\s*\[\s*36\s*[,;]\s*40\s*\]",
                f"подход, представленный в работах {s36} [36] и {s40} [40]",
                t,
                flags=re.IGNORECASE,
            )
            if ok:
                print(f"✓ Параграф {idx+1}: добавлены авторы к [36, 40] ({s36}; {s40})")
                changed = True

        # Интегративный подход [42] -> ... подход, описанный {s42} [42]
        if "Интегративный" in t and re.search(r"\[42\]", t):
            t, ok = _replace_first(
                r"подход\s*\[42\]",
                f"подход, описанный {s42} [42]",
                t,
                flags=re.IGNORECASE,
            )
            if ok:
                print(f"✓ Параграф {idx+1}: добавлен автор к [42] ({s42})")
                changed = True

        # --- 2) Стр.20: [22, 35] добавить контекст ---
        # Ищем конкретно место, где рядом с [22, 35] говорится про результаты/опросник/семью
        if re.search(r"\[\s*22\s*,\s*35\s*\]", t) and ("опросник" in t.lower() or "выборке" in t.lower() or "сем" in t.lower()):
            # заменим саму ссылку на фамилии с номерами
            t2, ok = _replace_first(
                r"\[\s*22\s*,\s*35\s*\]",
                f"{s22} [22] и {s35} [35]",
                t,
            )
            if ok:
                # добавим короткий вводный контекст, если ссылка стояла просто после утверждения
                # например: "... в выборке [22, 35]." -> "... в выборке (что согласуется с данными ...)."
                if f"{s22} [22] и {s35} [35]" in t2 and "(" not in t2:
                    t2 = t2.replace(f"{s22} [22] и {s35} [35]", f"(см. {s22} [22] и {s35} [35])", 1)
                t = t2
                print(f"✓ Параграф {idx+1}: добавлен контекст к [22, 35] ({s22}; {s35})")
                changed = True

        # --- 3) Фромм: добавить ссылку рядом с фамилией, если нет ---
        if "Фромм" in t:
            # если после "Фромм" в пределах 15 символов нет [число], добавим
            if not re.search(r"Фромм.{0,15}\[\d+", t):
                t = re.sub(r"(Фромм)(?!\s*\[\d+\])", rf"\1 [{fromm_num}]", t, count=1)
                print(f"✓ Параграф {idx+1}: добавлена ссылка к 'Фромм' -> [{fromm_num}]")
                changed = True

        # Безопасность: в квадратных скобках должны быть только числа/разделители.
        # Если в результате где-то остались буквы внутри [] — удалим буквы, оставим числа.
        def _numbers_only_brackets(m: re.Match) -> str:
            inside = m.group(1)
            nums = re.findall(r"\b\d+\b", inside)
            if not nums:
                return ""  # убрать пустые
            # сохраняем порядок примерно: уникальные по первому вхождению
            seen = set()
            out = []
            for n in nums:
                if n not in seen:
                    out.append(n)
                    seen.add(n)
            return "[" + ", ".join(out) + "]"

        t_clean = re.sub(r"\[([^\]]*[A-Za-zА-Яа-яЁё][^\]]*)\]", _numbers_only_brackets, t)
        if t_clean != t:
            t = t_clean
            changed = True

        if changed and t != t0:
            para.text = t
            fixed += 1

    doc.save(str(dst_path))

    print("\n" + "=" * 70)
    print("РЕЗУЛЬТАТЫ")
    print("=" * 70)
    print(f"✓ Документ сохранен: {dst_path.name}")
    print(f"✓ Изменено параграфов: {fixed}")


if __name__ == "__main__":
    fix_document()



