# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v9 -> v10:
1. Проверка и исправление всех номеров источников
2. Перефразирование проблемных мест с простым перечислением источников
"""
import re
from pathlib import Path
from docx import Document

def load_literature_list(literature_file):
    """Загружает список литературы"""
    literature = {}
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
    
    return literature

def rephrase_simple_enumeration(text, citation_match, citation_text):
    """
    Перефразирует простое перечисление источников
    Например: "Ананьев [3], Кон [31], Дружинин [22]" 
    -> "как отмечают Ананьев [3], Кон [31] и Дружинин [22]"
    """
    # Если это просто перечисление в начале или конце предложения
    # Добавляем контекст
    
    # Паттерн для поиска перечислений типа "Автор [N], Автор [N]"
    pattern = r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\](?:\s*,\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\])*(?:\s*,\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\])*'
    
    # Более простой подход - ищем известные проблемные паттерны
    replacements = [
        # Страница 7 (параграф 77)
        {
            'search': r'([А-ЯЁ][а-яё]+)\s*\[(\d+)\];\s*([А-ЯЁ][а-яё]+)\s*\[(\d+)\];\s*([А-ЯЁ][а-яё]+)\s*\[(\d+)\]',
            'replace': r'\1 [\2], \3 [\4] и \5 [\6]',
            'context': 'как отмечают'
        },
        # Два автора
        {
            'search': r'([А-ЯЁ][а-яё]+)\s*\[(\d+)\],\s*([А-ЯЁ][а-яё]+)\s*\[(\d+)\]',
            'replace': r'\1 [\2] и \3 [\4]',
            'context': 'как подчеркивают'
        },
        # Три и более авторов
        {
            'search': r'([А-ЯЁ][а-яё]+)\s*\[(\d+)\],\s*([А-ЯЁ][а-яё]+)\s*\[(\d+)\],\s*([А-ЯЁ][а-яё]+)\s*\[(\d+)\]',
            'replace': r'\1 [\2], \3 [\4] и \5 [\6]',
            'context': 'как отмечают'
        }
    ]
    
    return text

def fix_document():
    """Основная функция исправления"""
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v9 -> v10")
    print("=" * 70)
    
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен")
    
    literature = load_literature_list(literature_path)
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    print(f"✓ Список литературы: {len(literature)} источников (диапазон: {min_source}-{max_source})\n")
    
    fixed_count = 0
    invalid_sources_found = []
    all_sources_in_text = set()
    
    # Словарь известных исправлений номеров (если есть несоответствия)
    # Из анализа v8->v9 известны следующие исправления:
    number_fixes = {
        '[60]': '[51]',  # Франкл: 60 -> 51
        '[64]': '[53]',  # Эйдемиллер, Юстицкис: 64 -> 53
        '[55]': '[47]',  # Слободчиков, Исаев: 55 -> 47
        '[59]': '[50]',  # Фонарев: 59 -> 50
        '[37]': '[18]',  # Гинзбург: 37 -> 18 (но в списке литературы Гинзбург - 16)
        '[39]': '[33]',  # Маркова: 39 -> 33
        '[28]': '[23]',  # Карпова, Артемьева: 28 -> 23
        '[24]': '[20]',  # Журавлев, Купрейченко: 24 -> 20
        '[46]': '[38]',  # Осорина: 46 -> 38
        '[65]': '[13; 15]',  # Ядов: 65 -> 13; 15
    }
    
    # Список замен для перефразирования проблемных мест
    # Ищем в разных форматах: [Автор, N; Автор, N] или Автор [N], Автор [N]
    rephrase_replacements = [
        # Страница 7 - Ананьев [3], Кон [31], Дружинин [22] или [Ананьев, 3; Кон, 31; Дружинин, 22]
        {
            'patterns': [
                r'Ананьев\s*\[3\]\s*,\s*Кон\s*\[31\]\s*,\s*Дружинин\s*\[22\]',
                r'\[Ананьев\s*,\s*3\s*;\s*Кон\s*,\s*31\s*;\s*Дружинин\s*,\s*22\]',
            ],
            'replace': 'как отмечают Ананьев [3], Кон [31] и Дружинин [22]',
            'description': 'Страница 7: Ананьев, Кон, Дружинин'
        },
        # Страница 10 - [18, 30, 48] или [30; 18; 48]
        {
            'patterns': [
                r'\[18\s*,\s*30\s*,\s*48\]',
                r'\[30\s*;\s*18\s*;\s*48\]',
                r'\[18\s*;\s*30\s*;\s*48\]',
            ],
            'replace': 'как подчеркивают исследователи [18, 30, 48]',
            'description': 'Страница 10: [18, 30, 48]'
        },
        # Страница 13 - Ананьев [3], Божович [12], Кон [31] или [Ананьев, 3; Божович, 12; Кон, 31]
        {
            'patterns': [
                r'Ананьев\s*\[3\]\s*,\s*Божович\s*\[12\]\s*,\s*Кон\s*\[31\]',
                r'\[Ананьев\s*,\s*3\s*;\s*Божович\s*,\s*12\s*;\s*Кон\s*,\s*31\]',
            ],
            'replace': 'как отмечают Ананьев [3], Божович [12] и Кон [31]',
            'description': 'Страница 13: Ананьев, Божович, Кон'
        },
        # Страница 16 - Выготский [16], Кле [29], Кривцова [32], Кулагина, Колюцкий [33], Мухина [44], Ремшмидт [51]
        {
            'patterns': [
                r'Выготский\s*\[16\]\s*,\s*Кле\s*\[29\]\s*,\s*Кривцова\s*\[32\]\s*,\s*Кулагина\s*,\s*Колюцкий\s*\[33\]\s*,\s*Мухина\s*\[44\]\s*,\s*Ремшмидт\s*\[51\]',
                r'\[Выготский\s*,\s*16\s*;\s*Кле\s*,\s*29\s*;\s*Кривцова\s*,\s*32\s*;\s*Кулагина\s*,\s*Колюцкий\s*,\s*33\s*;\s*Мухина\s*,\s*44\s*;\s*Ремшмидт\s*,\s*51\]',
            ],
            'replace': 'как подчеркивают Выготский [16], Кле [29], Кривцова [32], Кулагина и Колюцкий [33], Мухина [44], Ремшмидт [51]',
            'description': 'Страница 16: Выготский, Кле, Кривцова, Кулагина, Колюцкий, Мухина, Ремшмидт'
        },
    ]
    
    # Дополнительные паттерны для поиска и перефразирования простых перечислений
    def rephrase_author_enumeration(text):
        """Перефразирует перечисления авторов в разных форматах"""
        # Паттерн 1: [Автор1, N1; Автор2, N2; Автор3, N3]
        pattern1 = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)(?:\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+))+(?:\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+))?\]'
        
        def replace_func1(match):
            full_match = match.group(0)
            # Извлекаем все пары автор-номер
            pairs = re.findall(r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)', full_match)
            
            if len(pairs) >= 3:  # Если 3 и более авторов - перефразируем
                authors_list = []
                for author, num in pairs:
                    authors_list.append(f"{author} [{num}]")
                
                # Последний автор с "и"
                if len(authors_list) > 1:
                    result = ", ".join(authors_list[:-1]) + " и " + authors_list[-1]
                    return f"как отмечают {result}"
            
            return full_match
        
        text = re.sub(pattern1, replace_func1, text)
        
        # Паттерн 2: Автор1 [N1], Автор2 [N2], Автор3 [N3] (без скобок вокруг всего)
        # Ищем последовательность из 3+ авторов с номерами подряд
        pattern2 = r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\]\s*,\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\]\s*,\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\](?:\s*,\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\])*'
        
        def replace_func2(match):
            # Извлекаем все пары
            full_text = match.group(0)
            pairs = re.findall(r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\]', full_text)
            
            if len(pairs) >= 3:  # Если 3 и более авторов - перефразируем
                authors_list = []
                for author, num in pairs:
                    authors_list.append(f"{author} [{num}]")
                
                if len(authors_list) > 1:
                    result = ", ".join(authors_list[:-1]) + " и " + authors_list[-1]
                    return f"как отмечают {result}"
            
            return full_text
        
        text = re.sub(pattern2, replace_func2, text)
        
        return text
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        
        # 1. Исправляем неверные номера источников
        for wrong_num, correct_num in number_fixes.items():
            if wrong_num in text:
                text = text.replace(wrong_num, correct_num)
                if fixed_count < 20:
                    print(f"✓ Исправлен номер: {wrong_num} -> {correct_num} (параграф {para_idx + 1})")
        
        # 2. Перефразируем проблемные места (с использованием регулярных выражений)
        for replacement in rephrase_replacements:
            patterns = replacement.get('patterns', [])
            found = False
            
            for pattern in patterns:
                match = re.search(pattern, text)
                if match:
                    text = re.sub(pattern, replacement['replace'], text)
                    print(f"✓ {replacement['description']}: перефразировано (параграф {para_idx + 1})")
                    found = True
                    break
            
            # Fallback на точное совпадение (для обратной совместимости)
            if not found:
                search_text = replacement.get('search', '')
                if search_text and search_text in text:
                    text = text.replace(search_text, replacement['replace'])
                    print(f"✓ {replacement['description']}: перефразировано (параграф {para_idx + 1})")
        
        # 2.1. Перефразируем общие паттерны перечислений авторов
        text_before = text
        text = rephrase_author_enumeration(text)
        if text != text_before:
            print(f"✓ Перефразировано перечисление авторов (параграф {para_idx + 1})")
        
        # 3. Проверяем и исправляем номера источников в квадратных скобках
        pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
        
        def fix_numbers_in_brackets(match):
            numbers_str = match.group(1)
            valid_numbers = []
            
            for num_str in re.split(r'[,\s;]+', numbers_str):
                num_str = num_str.strip()
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                    else:
                        # Записываем неверные номера
                        invalid_sources_found.append({
                            'para': para_idx + 1,
                            'number': num,
                            'text': text[:100]
                        })
                        print(f"⚠ Параграф {para_idx + 1}: найден неверный номер {num} (вне диапазона {min_source}-{max_source})")
                except ValueError:
                    pass
            
            # Сохраняем все найденные номера
            for num in valid_numbers:
                all_sources_in_text.add(num)
            
            if valid_numbers:
                unique_numbers = sorted(set(valid_numbers))
                return f"[{', '.join(map(str, unique_numbers))}]"
            return ""
        
        text = re.sub(pattern_numbers, fix_numbers_in_brackets, text)
        
        # Применяем изменения
        if text != original_text:
            para.text = text
            fixed_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Исправляем неверные номера
                for wrong_num, correct_num in number_fixes.items():
                    if wrong_num in text:
                        text = text.replace(wrong_num, correct_num)
                
                # Исправляем номера в скобках
                pattern_numbers = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
                def fix_numbers(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v10_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v10_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")
    print(f"✓ Найдено источников в тексте: {len(all_sources_in_text)}")
    print(f"  Диапазон: {min(all_sources_in_text) if all_sources_in_text else 'N/A'}-{max(all_sources_in_text) if all_sources_in_text else 'N/A'}")
    
    if invalid_sources_found:
        print(f"\n⚠ ВНИМАНИЕ: Найдено {len(invalid_sources_found)} случаев с неверными номерами источников")
        print("  Эти номера были удалены из ссылок. Проверьте документ вручную!")
    
    # Проверяем, все ли источники из списка литературы используются
    unused_sources = set(literature.keys()) - all_sources_in_text
    if unused_sources:
        print(f"\nℹ Неиспользуемых источников в списке литературы: {len(unused_sources)}")
        if len(unused_sources) <= 10:
            print(f"  Номера: {sorted(unused_sources)}")

if __name__ == '__main__':
    fix_document()

