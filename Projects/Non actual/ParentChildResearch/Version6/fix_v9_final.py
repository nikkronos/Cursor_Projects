# -*- coding: utf-8 -*-
"""
Финальный скрипт для исправления документа v9 -> v10
Сначала находит проблемные места, затем исправляет их
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v9 -> v10")
    print("=" * 70)
    
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен\n")
    
    fixed_count = 0
    
    # Загружаем список литературы для проверки номеров
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    max_source = 53
    min_source = 1
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        changed = False
        
        # 1. Исправляем конкретные проблемные места (гибкий поиск)
        
        # Параграф 77: [Ананьев, 3; Кон, 31; Дружинин, 22]
        pattern1 = r'\[Ананьев\s*,\s*3\s*;\s*Кон\s*,\s*31\s*;\s*Дружинин\s*,\s*22\]'
        if re.search(pattern1, text):
            text = re.sub(pattern1, 'как отмечают Ананьев [3], Кон [31] и Дружинин [22]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.7): Ананьев, Кон, Дружинин - исправлено")
            changed = True
        
        # Параграф 92: [30; 18; 48] или [18, 30, 48]
        pattern2 = r'\[(?:30\s*;\s*18\s*;\s*48|18\s*,\s*30\s*,\s*48|30\s*,\s*18\s*,\s*48)\]'
        if re.search(pattern2, text):
            text = re.sub(pattern2, 'как подчеркивают исследователи [18, 30, 48]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.10): [18, 30, 48] - исправлено")
            changed = True
        
        # Параграф 105: [Ананьев, 3; Божович, 12; Кон, 31]
        pattern3 = r'\[Ананьев\s*,\s*3\s*;\s*Божович\s*,\s*12\s*;\s*Кон\s*,\s*31\]'
        if re.search(pattern3, text):
            text = re.sub(pattern3, 'как отмечают Ананьев [3], Божович [12] и Кон [31]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.13): Ананьев, Божович, Кон - исправлено")
            changed = True
        
        # Параграф 127: [Выготский, 16; Кле, 29; Кривцова, 32; Кулагина, Колюцкий, 33; Мухина, 44; Ремшмидт, 51]
        pattern4 = r'\[Выготский\s*,\s*16\s*;\s*Кле\s*,\s*29\s*;\s*Кривцова\s*,\s*32\s*;\s*Кулагина\s*,\s*Колюцкий\s*,\s*33\s*;\s*Мухина\s*,\s*44\s*;\s*Ремшмидт\s*,\s*51\]'
        if re.search(pattern4, text):
            text = re.sub(pattern4, 'как подчеркивают Выготский [16], Кле [29], Кривцова [32], Кулагина и Колюцкий [33], Мухина [44], Ремшмидт [51]', text)
            print(f"✓ Параграф {para_idx + 1} (стр.16): Выготский, Кле, Кривцова и др. - исправлено")
            changed = True
        
        # 2. Исправляем другие проблемные места с перечислениями (3+ авторов)
        # Ищем паттерн [Автор, N; Автор, N; Автор, N; ...]
        pattern_enum = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)(?:\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)){2,}\]'
        
        def replace_enum(match):
            full_match = match.group(0)
            # Извлекаем все пары автор-номер
            pairs = re.findall(r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)', full_match)
            
            if len(pairs) >= 3:  # Если 3 и более авторов
                authors_list = []
                for author, num in pairs:
                    authors_list.append(f"{author} [{num}]")
                
                if len(authors_list) > 1:
                    result = ", ".join(authors_list[:-1]) + " и " + authors_list[-1]
                    return f"как отмечают {result}"
            
            return full_match
        
        new_text = re.sub(pattern_enum, replace_enum, text)
        if new_text != text:
            text = new_text
            print(f"✓ Параграф {para_idx + 1}: перефразировано перечисление авторов")
            changed = True
        
        # 3. Исправляем неверные номера источников в перечислениях авторов
        # [Шнейдер, 63; Эйдемиллер, Юстицкис, 64] -> [Эйдемиллер, Юстицкис, 53]
        pattern_wrong = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*63\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*64\]'
        if re.search(pattern_wrong, text):
            text = re.sub(pattern_wrong, r'[\2, 53]', text)
            print(f"✓ Параграф {para_idx + 1}: исправлено [Шнейдер, 63; Эйдемиллер, Юстицкис, 64] -> [Эйдемиллер, Юстицкис, 53]")
            changed = True
        
        # [Автор, 63] - удаляем полностью
        pattern_63 = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*63\]'
        if re.search(pattern_63, text):
            text = re.sub(pattern_63, '', text)
            print(f"✓ Параграф {para_idx + 1}: удалено [Автор, 63]")
            changed = True
        
        # Замены известных неверных номеров в обычных ссылках
        number_fixes = {
            '[60]': '[51]',  # Франкл
            '[64]': '[53]',  # Эйдемиллер, Юстицкис
        }
        
        for wrong, correct in number_fixes.items():
            if wrong in text:
                text = text.replace(wrong, correct)
                print(f"✓ Параграф {para_idx + 1}: исправлен номер {wrong} -> {correct}")
                changed = True
        
        # 4. Проверяем и исправляем номера в квадратных скобках [N, N, N]
        pattern_numbers = r'\[(\d+(?:[,\s;]+\s*\d+)*)\]'
        
        def fix_numbers(match):
            numbers_str = match.group(1)
            valid_numbers = []
            
            for num_str in re.split(r'[,\s;]+', numbers_str):
                num_str = num_str.strip()
                try:
                    num = int(num_str)
                    if min_source <= num <= max_source:
                        valid_numbers.append(num)
                except ValueError:
                    pass
            
            if valid_numbers:
                unique_numbers = sorted(set(valid_numbers))
                return f"[{', '.join(map(str, unique_numbers))}]"
            return ""
        
        new_text = re.sub(pattern_numbers, fix_numbers, text)
        if new_text != text:
            text = new_text
            changed = True
        
        # Применяем изменения
        if changed:
            para.text = text
            fixed_count += 1
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                # Применяем те же исправления
                for wrong, correct in number_fixes.items():
                    if wrong in text and correct:
                        text = text.replace(wrong, correct)
                
                # Исправляем номера
                pattern_numbers = r'\[(\d+(?:[,\s;]+\s*\d+)*)\]'
                def fix_numbers_cell(match):
                    numbers_str = match.group(1)
                    valid_numbers = []
                    for num_str in re.split(r'[,\s;]+', numbers_str):
                        try:
                            num = int(num_str.strip())
                            if min_source <= num <= max_source:
                                valid_numbers.append(num)
                        except ValueError:
                            pass
                    if valid_numbers:
                        return f"[{', '.join(map(str, sorted(set(valid_numbers))))}]"
                    return ""
                
                text = re.sub(pattern_numbers, fix_numbers_cell, text)
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v10_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v10_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")

if __name__ == '__main__':
    fix_document()

