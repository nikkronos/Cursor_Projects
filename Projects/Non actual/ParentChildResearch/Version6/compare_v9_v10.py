# -*- coding: utf-8 -*-
"""Сравнение документов v9 и v10 для проверки изменений"""
from pathlib import Path
from docx import Document

def compare_documents():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    if not v10_path.exists():
        print(f"ERROR: {v10_path} не найден. Сначала запустите fix_v9_to_v10.py")
        return
    
    print("=" * 70)
    print("СРАВНЕНИЕ ДОКУМЕНТОВ v9 и v10")
    print("=" * 70 + "\n")
    
    doc_v9 = Document(str(v9_path))
    doc_v10 = Document(str(v10_path))
    
    changes_count = 0
    
    # Сравниваем параграфы
    max_paras = max(len(doc_v9.paragraphs), len(doc_v10.paragraphs))
    
    for para_idx in range(max_paras):
        if para_idx < len(doc_v9.paragraphs):
            text_v9 = doc_v9.paragraphs[para_idx].text
        else:
            text_v9 = ""
        
        if para_idx < len(doc_v10.paragraphs):
            text_v10 = doc_v10.paragraphs[para_idx].text
        else:
            text_v10 = ""
        
        if text_v9 != text_v10:
            changes_count += 1
            print(f"ИЗМЕНЕНИЕ в параграфе {para_idx + 1}:")
            print(f"  v9: {text_v9[:150]}...")
            print(f"  v10: {text_v10[:150]}...")
            print()
            
            if changes_count >= 10:
                print("... (показано первых 10 изменений)")
                break
    
    print(f"\nВсего найдено изменений: {changes_count}")
    
    if changes_count == 0:
        print("\n⚠ ИЗМЕНЕНИЙ НЕ НАЙДЕНО!")
        print("Возможные причины:")
        print("1. Проблемные места уже исправлены в v9")
        print("2. Формат в документе отличается от ожидаемого")
        print("3. Номера параграфов не соответствуют страницам")

if __name__ == '__main__':
    compare_documents()



