# -*- coding: utf-8 -*-
"""
Простой скрипт для исправления документа v9 -> v10
Работает напрямую с текстом, делает точные замены
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    v10_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v10.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v9 -> v10")
    print("=" * 70)
    
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен\n")
    
    fixed_count = 0
    
    # Точные замены для проблемных мест
    replacements = [
        # Параграф 77 (страница 7)
        {
            'old': '[Ананьев, 3; Кон, 31; Дружинин, 22]',
            'new': 'как отмечают Ананьев [3], Кон [31] и Дружинин [22]',
            'desc': 'Параграф 77: Ананьев, Кон, Дружинин'
        },
        # Параграф 92 (страница 10)
        {
            'old': '[30; 18; 48]',
            'new': 'как подчеркивают исследователи [18, 30, 48]',
            'desc': 'Параграф 92: [30; 18; 48]'
        },
        # Параграф 105 (страница 13)
        {
            'old': '[Ананьев, 3; Божович, 12; Кон, 31]',
            'new': 'как отмечают Ананьев [3], Божович [12] и Кон [31]',
            'desc': 'Параграф 105: Ананьев, Божович, Кон'
        },
        # Параграф 127 (страница 16)
        {
            'old': '[Выготский, 16; Кле, 29; Кривцова, 32; Кулагина, Колюцкий, 33; Мухина, 44; Ремшмидт, 51]',
            'new': 'как подчеркивают Выготский [16], Кле [29], Кривцова [32], Кулагина и Колюцкий [33], Мухина [44], Ремшмидт [51]',
            'desc': 'Параграф 127: Выготский, Кле, Кривцова и др.'
        },
    ]
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        
        # Делаем точные замены
        for rep in replacements:
            if rep['old'] in text:
                text = text.replace(rep['old'], rep['new'])
                print(f"✓ {rep['desc']}: исправлено (параграф {para_idx + 1})")
                fixed_count += 1
        
        # Применяем изменения
        if text != original_text:
            para.text = text
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                for rep in replacements:
                    if rep['old'] in text:
                        text = text.replace(rep['old'], rep['new'])
                        fixed_count += 1
                
                if text != original_text:
                    cell.text = text
    
    # Сохраняем документ
    doc.save(str(v10_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v10_path.name}")
    print(f"✓ Исправлено мест: {fixed_count}")

if __name__ == '__main__':
    fix_document()



