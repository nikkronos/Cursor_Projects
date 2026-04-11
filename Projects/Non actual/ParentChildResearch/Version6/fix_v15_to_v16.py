# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v15 -> v16:
1. Проверка синхронизации номеров источников между текстом и списком литературы
2. Сортировка списка литературы по алфавиту (от А до Я)
3. Обновление всех номеров источников в тексте документа
"""
from pathlib import Path
from docx import Document
import re
from typing import Dict, List, Tuple

def extract_author_surname(entry: str) -> str:
    """
    Извлекает фамилию первого автора из записи библиографии.
    Возвращает фамилию для сортировки (в верхнем регистре).
    """
    # Убираем номер в начале (например, "1. " или "10. ")
    entry = re.sub(r'^\d+\.\s*', '', entry.strip())
    
    # Ищем первую фамилию
    # Формат обычно: "Фамилия И.О. Название..." или "Фамилия, Имя Название..."
    # Или для нескольких авторов: "Фамилия1 И.О., Фамилия2 И.О. Название..."
    
    # Паттерн для русских фамилий: начинается с заглавной буквы, содержит буквы
    match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
    if match:
        surname = match.group(1)
        return surname.upper()
    
    # Для английских фамилий
    match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z]\.?)*)', entry)
    if match:
        surname = match.group(1).split()[0]
        return surname.upper()
    
    # Если не нашлось, пробуем найти любую заглавную букву
    match = re.search(r'([А-ЯЁA-Z])', entry)
    if match:
        return match.group(1)
    
    return 'Я'  # Если не удалось определить, ставим в конец

def parse_literature_list(lit_path: Path) -> Dict[int, Dict]:
    """
    Парсит список литературы и возвращает словарь {номер: {entry, old_number}}
    """
    literature = {}
    
    if not lit_path.exists():
        print(f"ERROR: Файл списка литературы не найден: {lit_path}")
        return literature
    
    with open(lit_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_number = None
    current_entry = []
    
    for line in lines:
        line = line.strip()
        
        # Пропускаем заголовки и комментарии
        if line.startswith('#') or line.startswith('*') or not line:
            continue
        
        # Ищем строку с номером: "N. Автор..."
        match = re.match(r'^(\d+)\.\s*(.+)', line)
        if match:
            # Сохраняем предыдущую запись, если есть
            if current_number is not None:
                entry_text = ' '.join(current_entry)
                literature[current_number] = {
                    'entry': entry_text,
                    'old_number': current_number
                }
            
            # Начинаем новую запись
            current_number = int(match.group(1))
            current_entry = [match.group(2)]
        else:
            # Продолжение текущей записи
            if current_number is not None and line:
                current_entry.append(line)
    
    # Сохраняем последнюю запись
    if current_number is not None:
        entry_text = ' '.join(current_entry)
        literature[current_number] = {
            'entry': entry_text,
            'old_number': current_number
        }
    
    return literature

def sort_literature_by_alphabet(literature: Dict[int, Dict]) -> List[Tuple[str, int, Dict]]:
    """
    Сортирует источники по алфавиту (по фамилии первого автора).
    Возвращает список кортежей (фамилия_для_сортировки, старый_номер, данные)
    """
    sorted_items = []
    
    for old_num, data in literature.items():
        entry = data['entry']
        surname = extract_author_surname(entry)
        sorted_items.append((surname, old_num, data))
    
    # Сортируем по фамилии (учитывая русский алфавит)
    # Русские буквы идут перед английскими
    def sort_key(x):
        surname = x[0]
        # Если начинается с русской буквы
        if surname and surname[0] in 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ':
            return (0, surname)  # Русские сначала
        elif surname and surname[0].isalpha():
            return (1, surname)  # Английские потом
        else:
            return (2, surname)  # Остальное в конце
    
    sorted_items.sort(key=sort_key)
    
    return sorted_items

def create_number_mapping(sorted_items: List[Tuple]) -> Dict[int, int]:
    """
    Создаёт маппинг старый_номер -> новый_номер
    """
    mapping = {}
    
    for new_num, (first_letter, old_num, data) in enumerate(sorted_items, start=1):
        mapping[old_num] = new_num
        data['new_number'] = new_num
    
    return mapping

def find_all_citations_in_text(text: str) -> List[Tuple[int, int, str]]:
    """
    Находит все ссылки на источники в тексте.
    Возвращает список (start_pos, end_pos, citation_text)
    Поддерживает форматы: [N], [N, M], [N; M], [N-M], [N, M, K]
    """
    citations = []
    
    # Паттерн для поиска ссылок: [числа, разделённые запятыми, точками с запятой или дефисами]
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    
    for match in re.finditer(pattern, text):
        start = match.start()
        end = match.end()
        citation_text = match.group(0)
        citations.append((start, end, citation_text))
    
    return citations

def replace_citation_numbers(citation_text: str, mapping: Dict[int, int]) -> str:
    """
    Заменяет номера в ссылке согласно маппингу.
    Пример: "[1, 2, 3]" -> "[5, 7, 12]" (если 1->5, 2->7, 3->12)
    Сохраняет оригинальные разделители и пробелы.
    """
    # Извлекаем все числа из ссылки
    numbers = re.findall(r'\d+', citation_text)
    
    if not numbers:
        return citation_text
    
    # Заменяем каждое число
    new_numbers = []
    for num_str in numbers:
        old_num = int(num_str)
        if old_num in mapping:
            new_numbers.append(str(mapping[old_num]))
        else:
            # Если номер не найден в маппинге, оставляем как есть
            new_numbers.append(num_str)
    
    # Восстанавливаем формат с разделителями
    # Определяем разделители из оригинальной ссылки
    content = citation_text[1:-1]  # Убираем [ и ]
    
    # Пробуем сохранить оригинальный формат
    # Если есть запятые, используем запятые
    if ',' in content:
        sep = ', '
        result = '[' + sep.join(new_numbers) + ']'
    # Если есть точки с запятой, используем их
    elif ';' in content:
        sep = '; '
        result = '[' + sep.join(new_numbers) + ']'
    # Если есть дефисы, используем дефисы
    elif '-' in content:
        sep = '-'
        result = '[' + sep.join(new_numbers) + ']'
    else:
        # Если только одно число
        result = '[' + new_numbers[0] + ']'
    
    return result

def update_document_citations(doc: Document, mapping: Dict[int, int]) -> int:
    """
    Обновляет все ссылки на источники в документе.
    Возвращает количество заменённых ссылок.
    """
    replaced_count = 0
    
    # Обрабатываем параграфы
    for para in doc.paragraphs:
        text = para.text
        citations = find_all_citations_in_text(text)
        
        if citations:
            # Заменяем с конца, чтобы не сбить позиции
            for start, end, citation_text in reversed(citations):
                new_citation = replace_citation_numbers(citation_text, mapping)
                if new_citation != citation_text:
                    text = text[:start] + new_citation + text[end:]
                    replaced_count += 1
            
            if text != para.text:
                para.text = text
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                citations = find_all_citations_in_text(text)
                
                if citations:
                    for start, end, citation_text in reversed(citations):
                        new_citation = replace_citation_numbers(citation_text, mapping)
                        if new_citation != citation_text:
                            text = text[:start] + new_citation + text[end:]
                            replaced_count += 1
                    
                    if text != cell.text:
                        cell.text = text
    
    return replaced_count

def write_sorted_literature_list(lit_path: Path, sorted_items: List[Tuple], max_source: int):
    """
    Записывает отсортированный список литературы в файл.
    """
    output_path = lit_path.parent / f"{lit_path.stem}_sorted.md"
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ\n\n")
        f.write("*Сформирован на основе реального использования в тексте*\n\n")
        f.write(f"*Всего источников: {max_source}*\n")
        f.write("*Отсортировано по алфавиту*\n\n")
        
        for new_num, (first_letter, old_num, data) in enumerate(sorted_items, start=1):
            f.write(f"{new_num}. {data['entry']}\n\n")
    
    return output_path

def check_citation_sync(doc: Document, literature: Dict[int, Dict]) -> List[str]:
    """
    Проверяет синхронизацию номеров источников между текстом и списком литературы.
    Возвращает список проблем.
    """
    issues = []
    used_numbers = set()
    
    # Собираем все использованные номера из документа
    for para in doc.paragraphs:
        citations = find_all_citations_in_text(para.text)
        for _, _, citation_text in citations:
            numbers = re.findall(r'\d+', citation_text)
            for num_str in numbers:
                used_numbers.add(int(num_str))
    
    # Проверяем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                citations = find_all_citations_in_text(cell.text)
                for _, _, citation_text in citations:
                    numbers = re.findall(r'\d+', citation_text)
                    for num_str in numbers:
                        used_numbers.add(int(num_str))
    
    # Проверяем, все ли использованные номера есть в списке литературы
    for num in sorted(used_numbers):
        if num not in literature:
            issues.append(f"⚠ Номер {num} используется в тексте, но отсутствует в списке литературы")
    
    # Проверяем, все ли номера из списка используются
    literature_numbers = set(literature.keys())
    unused = literature_numbers - used_numbers
    if unused:
        issues.append(f"ℹ Номера из списка литературы не используются в тексте: {sorted(unused)}")
    
    return issues

def fix_document():
    version6 = Path(__file__).parent
    
    # Пути к файлам
    v15_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v15.docx"
    if not v15_path.exists():
        print(f"ERROR: Документ v15 не найден: {v15_path}")
        return
    
    lit_path = version6 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    if not lit_path.exists():
        print(f"ERROR: Список литературы не найден: {lit_path}")
        return
    
    v16_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v16.docx"
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v15 -> v16")
    print("=" * 70)
    print("1. Проверка синхронизации источников")
    print("2. Сортировка списка литературы по алфавиту")
    print("3. Обновление номеров в тексте документа")
    print("=" * 70)
    
    # Загружаем документ
    doc = Document(str(v15_path))
    print(f"✓ Документ загружен: {v15_path.name}\n")
    
    # Парсим список литературы
    print("📚 Парсинг списка литературы...")
    literature = parse_literature_list(lit_path)
    print(f"✓ Найдено источников: {len(literature)}\n")
    
    # Проверяем синхронизацию
    print("🔍 Проверка синхронизации номеров...")
    issues = check_citation_sync(doc, literature)
    if issues:
        for issue in issues:
            print(f"  {issue}")
    else:
        print("  ✓ Все номера синхронизированы")
    print()
    
    # Сортируем по алфавиту
    print("🔤 Сортировка списка литературы по алфавиту...")
    sorted_items = sort_literature_by_alphabet(literature)
    print(f"✓ Отсортировано {len(sorted_items)} источников\n")
    
    # Создаём маппинг
    print("🔄 Создание маппинга старых номеров на новые...")
    mapping = create_number_mapping(sorted_items)
    
    # Выводим изменения (первые 10)
    print("  Примеры изменений номеров:")
    count = 0
    for old_num, new_num in sorted(mapping.items()):
        if old_num != new_num:
            if count < 10:
                entry = literature[old_num]['entry'][:50] + "..."
                print(f"    {old_num} -> {new_num}: {entry}")
            count += 1
    if count > 10:
        print(f"    ... и ещё {count - 10} изменений")
    print()
    
    # Обновляем ссылки в документе
    print("📝 Обновление ссылок в документе...")
    replaced_count = update_document_citations(doc, mapping)
    print(f"✓ Заменено ссылок: {replaced_count}\n")
    
    # Сохраняем отсортированный список литературы
    print("💾 Сохранение отсортированного списка литературы...")
    max_source = len(sorted_items)
    sorted_lit_path = write_sorted_literature_list(lit_path, sorted_items, max_source)
    print(f"✓ Сохранён: {sorted_lit_path.name}\n")
    
    # Сохраняем документ
    doc.save(str(v16_path))
    
    print("=" * 70)
    print("РЕЗУЛЬТАТЫ")
    print("=" * 70)
    print(f"✓ Документ сохранён: {v16_path.name}")
    print(f"✓ Отсортированный список литературы: {sorted_lit_path.name}")
    print(f"✓ Заменено ссылок в документе: {replaced_count}")
    print(f"✓ Всего источников: {max_source}")

if __name__ == '__main__':
    fix_document()
