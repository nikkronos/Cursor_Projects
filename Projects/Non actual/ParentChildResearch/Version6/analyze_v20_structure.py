# -*- coding: utf-8 -*-
"""
Скрипт для анализа структуры документа v20
Извлекает список литературы из конца документа
"""
from pathlib import Path
from docx import Document
import re

def analyze_document():
    version6 = Path(__file__).parent
    v20_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
    
    if not v20_path.exists():
        print(f"ERROR: Документ v20 не найден: {v20_path}")
        return
    
    doc = Document(str(v20_path))
    print(f"=" * 70)
    print(f"АНАЛИЗ ДОКУМЕНТА v20")
    print(f"=" * 70)
    print(f"Всего параграфов: {len(doc.paragraphs)}\n")
    
    # Ищем начало списка литературы
    lit_start_idx = None
    for i, para in enumerate(doc.paragraphs):
        text_lower = para.text.strip().lower()
        if any(keyword in text_lower for keyword in ['список литературы', 'литература', 'библиографический список', 'литературы']):
            lit_start_idx = i
            print(f"✓ Найден заголовок списка литературы на позиции {i}:")
            print(f"  '{para.text[:80]}'\n")
            break
    
    if lit_start_idx is None:
        print("⚠ Заголовок списка литературы не найден, проверяю последние 50 параграфов...\n")
        lit_start_idx = max(0, len(doc.paragraphs) - 50)
    
    # Показываем последние параграфы
    print("Последние параграфы документа (для анализа структуры списка литературы):")
    print("-" * 70)
    
    start = max(0, len(doc.paragraphs) - 40) if lit_start_idx is None else lit_start_idx
    
    for i in range(start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.text.strip():
            preview = para.text[:100].replace('\n', ' ')
            print(f"{i:4d}: {preview}")
    
    print("\n" + "=" * 70)
    print("Поиск записей в формате 'N. Автор...':")
    print("-" * 70)
    
    # Ищем записи вида "N. Автор..."
    bibliography_items = []
    for i in range(start, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        
        # Паттерн для номера с точкой: "1. ", "10. ", "100. "
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            entry_start = match.group(2)[:60]
            bibliography_items.append((i, num, entry_start))
            print(f"Позиция {i:4d}, Номер {num:3d}: {entry_start}")
    
    print(f"\n✓ Найдено записей в формате 'N. Автор...': {len(bibliography_items)}")
    
    if bibliography_items:
        print(f"\nДиапазон номеров: {min(b[1] for b in bibliography_items)} - {max(b[1] for b in bibliography_items)}")
        print(f"\nПервые 5 записей:")
        for idx, num, entry in bibliography_items[:5]:
            print(f"  {num}. {entry}")
        print(f"\nПоследние 5 записей:")
        for idx, num, entry in bibliography_items[-5:]:
            print(f"  {num}. {entry}")

if __name__ == '__main__':
    analyze_document()
