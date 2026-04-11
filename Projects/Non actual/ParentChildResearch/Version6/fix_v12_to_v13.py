# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v12 -> v13:
1. На странице 10 перед "[18, 30, 48]" добавить фамилии исследователей
"""
from pathlib import Path
from docx import Document
import re

def fix_document():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    # Пробуем найти v12
    v12_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v12.docx"
    if not v12_path.exists():
        # Пробуем найти v11
        v11_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v11.docx"
        if v11_path.exists():
            v12_path = v11_path
            print(f"ВНИМАНИЕ: v12 не найден, используем v11")
        else:
            print(f"ERROR: Не найден документ v12 или v11")
            return
    
    v13_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v13.docx"
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v12 -> v13")
    print("=" * 70)
    
    doc = Document(str(v12_path))
    print(f"✓ Документ загружен: {v12_path.name}\n")
    
    fixed_count = 0
    
    # Авторы для номеров 18, 30, 48:
    # 18 - Дружинин В.Н.
    # 30 - Леонтьев А.Н.
    # 48 - Фельдштейн Д.И.
    
    # Обрабатываем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        original_text = para.text
        if not original_text.strip():
            continue
        
        text = original_text
        
        # Ищем "[18, 30, 48]" или "[30, 18, 48]" или "как подчеркивают исследователи [18, 30, 48]"
        # Нужно заменить на "как подчеркивают Дружинин [18], Леонтьев [30] и Фельдштейн [48]"
        
        patterns = [
            r'как подчеркивают исследователи\s*\[18\s*,\s*30\s*,\s*48\]',
            r'как подчеркивают исследователи\s*\[30\s*,\s*18\s*,\s*48\]',
            r'\[18\s*,\s*30\s*,\s*48\]',
            r'\[30\s*,\s*18\s*,\s*48\]',
        ]
        
        for pattern in patterns:
            match = re.search(pattern, text)
            if match:
                old_text = match.group(0)
                new_text = 'как подчеркивают Дружинин [18], Леонтьев [30] и Фельдштейн [48]'
                text = re.sub(pattern, new_text, text)
                
                print(f"✓ Параграф {para_idx + 1} (стр.10): добавлены фамилии исследователей")
                print(f"  БЫЛО: {old_text}")
                print(f"  СТАЛО: {new_text}")
                fixed_count += 1
                break
        
        # Применяем изменения
        if text != original_text:
            para.text = text
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                original_text = cell.text
                if not original_text.strip():
                    continue
                
                text = original_text
                
                for pattern in patterns:
                    match = re.search(pattern, text)
                    if match:
                        new_text = 'как подчеркивают Дружинин [18], Леонтьев [30] и Фельдштейн [48]'
                        text = re.sub(pattern, new_text, text)
                        break
                
                if text != original_text:
                    cell.text = text
                    fixed_count += 1
    
    # Сохраняем документ
    doc.save(str(v13_path))
    
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ")
    print(f"{'='*70}")
    print(f"✓ Документ сохранен: {v13_path.name}")
    print(f"✓ Исправлено параграфов/ячеек: {fixed_count}")

if __name__ == '__main__':
    fix_document()



