# -*- coding: utf-8 -*-
"""
Скрипт для полного анализа документа v9:
1. Извлечение всех источников из текста
2. Проверка соответствия номеров источников списку литературы
3. Поиск проблемных мест с простым перечислением источников
"""
import re
from pathlib import Path
from docx import Document
from collections import defaultdict

def load_literature_list(literature_file):
    """Загружает список литературы и создает словарь номер->автор"""
    literature = {}
    author_to_number = {}
    
    with open(literature_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Паттерн для извлечения номеров и записей
    pattern = r'^(\d+)\.\s+([^\n]+)'
    matches = re.finditer(pattern, content, re.MULTILINE)
    
    for match in matches:
        num = int(match.group(1))
        entry = match.group(2).strip()
        literature[num] = {'number': num, 'entry': entry}
        
        # Извлекаем автора (первая часть до точки или запятой)
        author_match = re.match(r'^([^./]+)', entry)
        if author_match:
            author = author_match.group(1).strip()
            # Убираем инициалы и оставляем только фамилию
            author_parts = author.split()
            if author_parts:
                # Берем первую часть (фамилию)
                surname = author_parts[0]
                author_to_number[surname] = num
                # Также добавляем варианты с запятыми
                if ',' in surname:
                    for part in surname.split(','):
                        author_to_number[part.strip()] = num
    
    return literature, author_to_number

def extract_all_citations(text):
    """Извлекает все ссылки на источники из текста"""
    citations = []
    
    # 1. Квадратные скобки с номерами: [5], [1, 2], [1; 2]
    pattern1 = r'\[(\d+(?:[,\s;]\s*\d+)*)\]'
    for match in re.finditer(pattern1, text):
        numbers_str = match.group(1)
        numbers = []
        for num_str in re.split(r'[,\s;]+', numbers_str):
            num_str = num_str.strip()
            if num_str:
                try:
                    numbers.append(int(num_str))
                except ValueError:
                    pass
        if numbers:
            citations.append({
                'type': 'square_brackets_numbers',
                'full_match': match.group(0),
                'numbers': numbers,
                'position': match.span()
            })
    
    # 2. Фамилии с номерами в скобках: [Ананьев, 3], [Кон, 31]
    pattern2 = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)\]'
    for match in re.finditer(pattern2, text):
        authors = [a.strip() for a in match.group(1).split(',')]
        number = int(match.group(2))
        citations.append({
            'type': 'square_brackets_author_number',
            'full_match': match.group(0),
            'authors': authors,
            'number': number,
            'position': match.span()
        })
    
    # 3. Фамилии с номерами через точку с запятой: [Ананьев, 3; Кон, 31]
    pattern3 = r'\[([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)(?:\s*;\s*([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+))*\]'
    for match in re.finditer(pattern3, text):
        full_text = match.group(0)
        # Разбираем все пары автор-номер
        pairs = re.findall(r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*,\s*(\d+)', full_text)
        citations.append({
            'type': 'square_brackets_multiple_authors',
            'full_match': full_text,
            'pairs': [(a.strip(), int(n)) for a, n in pairs],
            'position': match.span()
        })
    
    # 4. Фамилии без скобок с номерами: Ананьев [3], Кон [31]
    pattern4 = r'([А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*)\s*\[(\d+)\]'
    for match in re.finditer(pattern4, text):
        authors = [a.strip() for a in match.group(1).split(',')]
        number = int(match.group(2))
        citations.append({
            'type': 'author_square_brackets',
            'full_match': match.group(0),
            'authors': authors,
            'number': number,
            'position': match.span()
        })
    
    return citations

def analyze_document():
    """Основная функция анализа"""
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    literature_path = version5 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    if not literature_path.exists():
        print(f"ERROR: {literature_path} не найден")
        return
    
    print("=" * 70)
    print("АНАЛИЗ ДОКУМЕНТА v9")
    print("=" * 70)
    
    # Загружаем документ
    doc = Document(str(v9_path))
    print(f"✓ Документ загружен: {v9_path.name}")
    
    # Загружаем список литературы
    literature, author_to_number = load_literature_list(literature_path)
    max_source = max(literature.keys())
    min_source = min(literature.keys())
    print(f"✓ Список литературы: {len(literature)} источников (диапазон: {min_source}-{max_source})\n")
    
    # Собираем все цитаты
    all_citations = []
    problematic_places = []
    invalid_numbers = []
    used_sources = set()
    
    # Анализируем параграфы
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        citations = extract_all_citations(text)
        for citation in citations:
            citation['para_idx'] = para_idx + 1
            citation['text'] = text[:200] + '...' if len(text) > 200 else text
            all_citations.append(citation)
            
            # Проверяем номера источников
            if 'numbers' in citation:
                for num in citation['numbers']:
                    used_sources.add(num)
                    if num < min_source or num > max_source:
                        invalid_numbers.append({
                            'para': para_idx + 1,
                            'number': num,
                            'citation': citation['full_match'],
                            'text': text[:150]
                        })
            
            if 'number' in citation:
                num = citation['number']
                used_sources.add(num)
                if num < min_source or num > max_source:
                    invalid_numbers.append({
                        'para': para_idx + 1,
                        'number': num,
                        'citation': citation['full_match'],
                        'text': text[:150]
                    })
            
            # Ищем проблемные места - простые перечисления
            if citation['type'] in ['square_brackets_multiple_authors', 'author_square_brackets']:
                # Проверяем, является ли это простым перечислением
                context_before = text[max(0, citation['position'][0] - 50):citation['position'][0]]
                context_after = text[citation['position'][1]:min(len(text), citation['position'][1] + 50)]
                
                # Если перед цитатой нет содержательного текста или это просто перечисление
                if len(citation.get('pairs', [])) >= 3 or (citation.get('authors') and len(citation.get('authors', [])) >= 2):
                    problematic_places.append({
                        'para': para_idx + 1,
                        'type': 'simple_enumeration',
                        'citation': citation['full_match'],
                        'text': text[:300],
                        'context_before': context_before,
                        'context_after': context_after
                    })
    
    # Анализируем таблицы
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                text = cell.text
                if not text.strip():
                    continue
                
                citations = extract_all_citations(text)
                for citation in citations:
                    citation['table'] = table_idx + 1
                    citation['row'] = row_idx + 1
                    citation['cell'] = cell_idx + 1
                    all_citations.append(citation)
    
    # Выводим результаты
    print(f"\n{'='*70}")
    print("РЕЗУЛЬТАТЫ АНАЛИЗА")
    print(f"{'='*70}\n")
    
    print(f"Всего найдено цитат: {len(all_citations)}")
    print(f"Используемых источников: {len(used_sources)} (диапазон: {min(used_sources) if used_sources else 0}-{max(used_sources) if used_sources else 0})")
    print(f"Неиспользуемых источников: {len(literature) - len(used_sources)}")
    
    if invalid_numbers:
        print(f"\n⚠ НАЙДЕНО НЕВЕРНЫХ НОМЕРОВ ИСТОЧНИКОВ: {len(invalid_numbers)}")
        for item in invalid_numbers[:10]:  # Показываем первые 10
            print(f"  Параграф {item['para']}: номер {item['number']} (вне диапазона {min_source}-{max_source})")
            print(f"    Цитата: {item['citation']}")
            print(f"    Текст: {item['text']}...\n")
        if len(invalid_numbers) > 10:
            print(f"  ... и еще {len(invalid_numbers) - 10} случаев\n")
    
    if problematic_places:
        print(f"\n⚠ НАЙДЕНО ПРОБЛЕМНЫХ МЕСТ (простое перечисление): {len(problematic_places)}")
        for item in problematic_places[:15]:  # Показываем первые 15
            print(f"\n  Параграф {item['para']}:")
            print(f"    Тип: {item['type']}")
            print(f"    Цитата: {item['citation']}")
            print(f"    Контекст: ...{item['context_before']}[ЦИТАТА]{item['context_after']}...")
            print(f"    Полный текст: {item['text']}...")
        if len(problematic_places) > 15:
            print(f"\n  ... и еще {len(problematic_places) - 15} случаев")
    
    # Сохраняем детальный отчет
    report_file = version6 / "АНАЛИЗ_v9_ОТЧЕТ.txt"
    with open(report_file, 'w', encoding='utf-8') as f:
        f.write("=" * 70 + "\n")
        f.write("ПОЛНЫЙ ОТЧЕТ ПО АНАЛИЗУ ДОКУМЕНТА v9\n")
        f.write("=" * 70 + "\n\n")
        
        f.write(f"Всего найдено цитат: {len(all_citations)}\n")
        f.write(f"Используемых источников: {len(used_sources)}\n")
        f.write(f"Неиспользуемых источников: {len(literature) - len(used_sources)}\n\n")
        
        if invalid_numbers:
            f.write(f"\nНЕВЕРНЫЕ НОМЕРА ИСТОЧНИКОВ ({len(invalid_numbers)}):\n")
            f.write("-" * 70 + "\n")
            for item in invalid_numbers:
                f.write(f"Параграф {item['para']}: номер {item['number']}\n")
                f.write(f"  Цитата: {item['citation']}\n")
                f.write(f"  Текст: {item['text']}\n\n")
        
        if problematic_places:
            f.write(f"\nПРОБЛЕМНЫЕ МЕСТА - ПРОСТОЕ ПЕРЕЧИСЛЕНИЕ ({len(problematic_places)}):\n")
            f.write("-" * 70 + "\n")
            for item in problematic_places:
                f.write(f"Параграф {item['para']}:\n")
                f.write(f"  Тип: {item['type']}\n")
                f.write(f"  Цитата: {item['citation']}\n")
                f.write(f"  Полный текст: {item['text']}\n\n")
    
    print(f"\n✓ Детальный отчет сохранен: {report_file.name}")
    
    return {
        'citations': all_citations,
        'invalid_numbers': invalid_numbers,
        'problematic_places': problematic_places,
        'used_sources': used_sources,
        'literature': literature
    }

if __name__ == '__main__':
    analyze_document()



