# -*- coding: utf-8 -*-
"""Проверка текста документа v9 для поиска проблемных мест"""
from pathlib import Path
from docx import Document
import re

def check_text():
    version6 = Path(__file__).parent
    version5 = version6.parent / "Version5"
    
    v9_path = version5 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v9.docx"
    
    if not v9_path.exists():
        print(f"ERROR: {v9_path} не найден")
        return
    
    doc = Document(str(v9_path))
    
    # Ищем конкретные проблемные места
    search_terms = [
        ('Ананьев', 'Кон', 'Дружинин'),
        ('18', '30', '48'),
        ('Ананьев', 'Божович', 'Кон'),
        ('Выготский', 'Кле', 'Кривцова', 'Кулагина', 'Колюцкий', 'Мухина', 'Ремшмидт'),
    ]
    
    print("Поиск проблемных мест в документе v9...\n")
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # Ищем первое проблемное место
        if 'Ананьев' in text and 'Кон' in text and 'Дружинин' in text:
            if '[3]' in text or '[31]' in text or '[22]' in text:
                print(f"ПАРАГРАФ {para_idx + 1} (возможно страница 7):")
                print(f"{text[:400]}...")
                print()
        
        # Ищем второе проблемное место
        if '[18' in text and '30' in text and '48]' in text:
            print(f"ПАРАГРАФ {para_idx + 1} (возможно страница 10):")
            print(f"{text[:400]}...")
            print()
        
        # Ищем третье проблемное место
        if 'Ананьев' in text and 'Божович' in text and 'Кон' in text:
            if '[3]' in text or '[12]' in text or '[31]' in text:
                print(f"ПАРАГРАФ {para_idx + 1} (возможно страница 13):")
                print(f"{text[:400]}...")
                print()
        
        # Ищем четвертое проблемное место
        if 'Выготский' in text and 'Кле' in text and 'Кривцова' in text:
            if 'Кулагина' in text or 'Колюцкий' in text:
                print(f"ПАРАГРАФ {para_idx + 1} (возможно страница 16):")
                print(f"{text[:400]}...")
                print()
    
    # Также ищем все перечисления авторов в скобках
    print("\n" + "="*70)
    print("Все перечисления авторов в скобках:")
    print("="*70 + "\n")
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        if not text.strip():
            continue
        
        # Ищем паттерн [Автор, N; Автор, N; ...]
        pattern = r'\[[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+(?:\s*;\s*[А-ЯЁ][а-яё]+(?:\s*,\s*[А-ЯЁ][а-яё]+)*\s*,\s*\d+)+?\]'
        matches = list(re.finditer(pattern, text))
        if matches:
            print(f"Параграф {para_idx + 1}:")
            for match in matches:
                print(f"  Найдено: {match.group(0)}")
            print(f"  Текст: {text[:200]}...")
            print()

if __name__ == '__main__':
    check_text()



