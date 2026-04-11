"""
Анализ дополнительных файлов и сводка результатов по двум связкам:
1. Опрос для родителей (Ответы) - Ответы на форму (1) (1).csv
2. Опрос ученика (Ответы) - Ответы на форму (1) (1).csv
"""
import os
import sys
import pandas as pd
import numpy as np
from scipy.stats import spearmanr, pearsonr
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import warnings
warnings.filterwarnings('ignore')

# Функции для обработки данных (копируем из основного скрипта)
def classify_profession_to_holland_types(profession_text):
    """Классифицирует профессию по типам Холланда (R, I, A, S, E, C)."""
    if pd.isna(profession_text) or not str(profession_text).strip():
        return []
    
    prof_lower = str(profession_text).lower()
    types = []
    
    r_keywords = ['инженер', 'техник', 'механик', 'автомеханик', 'ремонт', 'конструктор', 
                  'оператор', 'фермер', 'водитель', 'электроник', 'чертёжник', 'технолог',
                  'слесарь', 'сварщик', 'токарь', 'фрезеровщик', 'монтажник', 'строитель']
    if any(kw in prof_lower for kw in r_keywords):
        types.append('R')
    
    i_keywords = ['учёный', 'исследователь', 'аналитик', 'химик', 'физик', 'биолог', 
                  'микробиолог', 'археолог', 'криминалист', 'синоптик', 'эколог',
                  'врач', 'психиатр', 'ветеринар', 'провизор', 'философ', 'политолог']
    if any(kw in prof_lower for kw in i_keywords):
        types.append('I')
    
    a_keywords = ['художник', 'дизайнер', 'актёр', 'режиссёр', 'композитор', 'писатель',
                  'поэт', 'фотограф', 'декоратор', 'визажист', 'стилист', 'искусствовед',
                  'реставратор', 'литературный критик', 'видеооператор', 'видеограф']
    if any(kw in prof_lower for kw in a_keywords):
        types.append('A')
    
    s_keywords = ['психолог', 'учитель', 'преподаватель', 'тренер', 'врач', 'терапевт',
                  'логопед', 'социальный', 'молодёжью', 'консультант', 'карьерный',
                  'профилактики', 'правонарушений', 'социолог', 'журналист', 'корреспондент']
    if any(kw in prof_lower for kw in s_keywords):
        types.append('S')
    
    e_keywords = ['менеджер', 'предприниматель', 'директор', 'руководитель', 'трейдер',
                  'адвокат', 'риэлтор', 'продюсер', 'лидер', 'движения', 'общественного',
                  'маркетинг', 'сетевому', 'финансовый', 'генеральный', 'отдела']
    if any(kw in prof_lower for kw in e_keywords):
        types.append('E')
    
    c_keywords = ['бухгалтер', 'оператор ввода', 'ввода данных', 'кассир', 'продавец',
                  'редактор', 'корректор', 'статистический', 'экономист', 'контролёр',
                  'качества', 'логистике', 'доставке', 'офис-менеджер']
    if any(kw in prof_lower for kw in c_keywords):
        types.append('C')
    
    return types

def calculate_holland_scores(students_df, holland_cols, pairs):
    """Рассчитывает баллы по типам Холланда (R, I, A, S, E, C) для каждого респондента."""
    holland_scores = {
        'R': [],
        'I': [],
        'A': [],
        'S': [],
        'E': [],
        'C': []
    }
    
    for pair in pairs:
        scores = {'R': 0, 'I': 0, 'A': 0, 'S': 0, 'E': 0, 'C': 0}
        
        for h_col in holland_cols:
            try:
                profession = students_df.iloc[pair['student_idx']][h_col]
                if pd.notna(profession):
                    types = classify_profession_to_holland_types(profession)
                    for h_type in types:
                        if h_type in scores:
                            scores[h_type] += 1
            except:
                continue
        
        for h_type in holland_scores:
            holland_scores[h_type].append(scores[h_type])
    
    return holland_scores

def calculate_statistics(values):
    """Рассчитывает статистические показатели для массива значений."""
    if len(values) == 0:
        return {'M': None, 'SD': None, 'Me': None, 'percentages': {}}
    
    values_array = np.array(values)
    
    M = np.mean(values_array)
    SD = np.std(values_array, ddof=1) if len(values_array) > 1 else 0
    Me = np.median(values_array)
    
    unique, counts = np.unique(values_array, return_counts=True)
    percentages = {int(val): (count / len(values_array)) * 100 
                   for val, count in zip(unique, counts)}
    
    return {'M': M, 'SD': SD, 'Me': Me, 'percentages': percentages}

def explain_significance_level(p_value):
    """Объясняет уровень значимости простым языком."""
    if p_value < 0.001:
        return ("***", "Очень высокая значимость (p < 0.001). Вероятность ошибки менее 0.1%.")
    elif p_value < 0.01:
        return ("**", "Высокая значимость (p < 0.01). Вероятность ошибки менее 1%.")
    elif p_value < 0.05:
        return ("*", "Значимость (p < 0.05). Вероятность ошибки менее 5%.")
    else:
        return ("", f"Незначимо (p = {p_value:.4f}). Вероятность ошибки {p_value*100:.1f}%.")

script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)

# Создаём папку Extra, если её нет
os.makedirs(script_dir, exist_ok=True)

print("=" * 80)
print("АНАЛИЗ ДОПОЛНИТЕЛЬНЫХ ФАЙЛОВ")
print("=" * 80)

# Пути к файлам
parents_file = os.path.join(parent_dir, "Опрос для родителей  (Ответы) - Ответы на форму (1) (1).csv")
students_file = os.path.join(parent_dir, "Опрос ученика (Ответы) - Ответы на форму (1) (1).csv")

print("\n1. Загрузка данных...")
try:
    parents_df = pd.read_csv(parents_file, encoding='utf-8-sig')
    students_df = pd.read_csv(students_file, encoding='utf-8-sig')
    print(f"   Родителей: {len(parents_df)}")
    print(f"   Учеников: {len(students_df)}")
except Exception as e:
    print(f"   ОШИБКА при загрузке файлов: {e}")
    sys.exit(1)

# Определяем структуру данных родителей
# Ищем столбцы с вопросами ВРР Марковской (обычно начинаются с "1.", "2." и т.д.)
parent_cols = parents_df.columns.tolist()
parent_questions = []

# Ищем вопросы ВРР Марковской (60 вопросов)
for col in parent_cols:
    col_str = str(col).strip()
    # Проверяем, является ли столбец вопросом (начинается с цифры и точки)
    if any(col_str.startswith(f"{i}. ") for i in range(1, 61)):
        parent_questions.append(col)

print(f"   Найдено вопросов родителей (ВРР Марковской): {len(parent_questions)}")

# Определяем структуру данных учеников
student_cols = students_df.columns.tolist()

# Находим вопросы Овчаровой
q12_col = None
q19_col = None
ovcharova_all = {}

for col in student_cols:
    col_str = str(col)
    for i in range(1, 21):
        if col_str.startswith(f"{i}. "):
            if any(x in col_str for x in ['Требует', 'Нравится', 'Предполагает', 'Соответствует', 
                                          'Позволяет', 'Дает', 'Является', 'Близка', 'Избрана', 'Единственно', 'Способствует']):
                ovcharova_all[i] = col
                if i == 12 and 'Дает возможности для роста' in col_str:
                    q12_col = col
                if i == 19 and 'Позволяет использовать профессиональные умения' in col_str:
                    q19_col = col
                break

if not q12_col or not q19_col:
    print("   ⚠ Вопросы 12 и 19 Овчаровой не найдены, ищем альтернативные варианты...")
    # Попробуем найти по другим признакам
    for col in student_cols:
        col_str = str(col)
        if '12' in col_str and ('Дает' in col_str or 'роста' in col_str or 'мастерства' in col_str):
            q12_col = col
        if '19' in col_str and ('Позволяет' in col_str or 'умения' in col_str or 'вне работы' in col_str):
            q19_col = col

print(f"   Вопрос 12 Овчаровой: {q12_col if q12_col else 'НЕ НАЙДЕН'}")
print(f"   Вопрос 19 Овчаровой: {q19_col if q19_col else 'НЕ НАЙДЕН'}")

# Находим столбцы с фамилиями
parent_name_col = None
parent_child_name_col = None
student_name_col = None

for col in parent_cols:
    if 'фамилия' in str(col).lower() and 'родителя' in str(col).lower():
        parent_name_col = col
    if 'фамилия' in str(col).lower() and 'ребенка' in str(col).lower():
        parent_child_name_col = col

for col in student_cols:
    if 'фамилия' in str(col).lower() and 'имя' in str(col).lower():
        student_name_col = col
        break

print(f"   Столбец фамилии родителя: {parent_name_col}")
print(f"   Столбец фамилии ребенка (в файле родителей): {parent_child_name_col}")
print(f"   Столбец фамилии ученика: {student_name_col}")

# Находим столбец с номером для сопоставления пар
parent_number_col = None
student_number_col = None

for col in parent_cols:
    if 'номер' in str(col).lower() or 'number' in str(col).lower() or col == 'Column 66':
        parent_number_col = col
        break

for col in student_cols:
    if 'номер' in str(col).lower() or 'number' in str(col).lower() or col == 'Column 127' or 'Порядковый номер' in str(col):
        student_number_col = col
        break

print(f"   Столбец номера родителей: {parent_number_col}")
print(f"   Столбец номера учеников: {student_number_col}")

# Функция для извлечения фамилии
def extract_surname(full_name):
    """Извлекает фамилию из полного имени."""
    if pd.isna(full_name) or not str(full_name).strip():
        return "Неизвестно"
    name_parts = str(full_name).strip().split()
    return name_parts[0] if len(name_parts) > 0 else str(full_name)

# Сопоставляем пары с фамилиями
print("\n2. Сопоставление пар родитель-ребенок...")
pairs = []

if parent_number_col and student_number_col:
    parents_sorted = parents_df.sort_values(parent_number_col).reset_index(drop=True)
    students_sorted = students_df.sort_values(student_number_col).reset_index(drop=True)
    
    for i in range(min(len(parents_sorted), len(students_sorted))):
        try:
            p_num = parents_sorted.iloc[i][parent_number_col]
            s_num = students_sorted.iloc[i][student_number_col]
            
            # Преобразуем в числа, если возможно
            try:
                p_num = float(p_num) if pd.notna(p_num) else None
                s_num = float(s_num) if pd.notna(s_num) else None
            except:
                p_num = str(p_num).strip() if pd.notna(p_num) else None
                s_num = str(s_num).strip() if pd.notna(s_num) else None
            
            if p_num is not None and s_num is not None and p_num == s_num:
                # Извлекаем фамилии
                parent_name = parents_sorted.iloc[i][parent_name_col] if parent_name_col else None
                child_name_from_parent = parents_sorted.iloc[i][parent_child_name_col] if parent_child_name_col else None
                student_name = students_sorted.iloc[i][student_name_col] if student_name_col else None
                
                parent_surname = extract_surname(parent_name)
                child_surname = extract_surname(child_name_from_parent) if child_name_from_parent else extract_surname(student_name)
                
                pairs.append({
                    'pair_id': len(pairs) + 1, 
                    'parent_idx': i, 
                    'student_idx': i, 
                    'number': p_num,
                    'parent_surname': parent_surname,
                    'child_surname': child_surname,
                    'parent_full_name': str(parent_name) if pd.notna(parent_name) else "Неизвестно",
                    'child_full_name': str(child_name_from_parent) if pd.notna(child_name_from_parent) else (str(student_name) if pd.notna(student_name) else "Неизвестно")
                })
        except Exception as e:
            print(f"   Ошибка при обработке пары {i}: {e}")
            pass
else:
    # Если нет столбцов с номерами, используем порядковый номер и сопоставляем по фамилиям
    parents_sorted = parents_df.reset_index(drop=True)
    students_sorted = students_df.reset_index(drop=True)
    
    for i in range(min(len(parents_df), len(students_df))):
        try:
            parent_name = parents_sorted.iloc[i][parent_name_col] if parent_name_col else None
            child_name_from_parent = parents_sorted.iloc[i][parent_child_name_col] if parent_child_name_col else None
            student_name = students_sorted.iloc[i][student_name_col] if student_name_col else None
            
            parent_surname = extract_surname(parent_name)
            child_surname_from_parent = extract_surname(child_name_from_parent) if child_name_from_parent else None
            student_surname = extract_surname(student_name)
            
            # Сопоставляем по фамилиям или используем порядковый номер
            if child_surname_from_parent and student_surname and child_surname_from_parent == student_surname:
                pairs.append({
                    'pair_id': len(pairs) + 1, 
                    'parent_idx': i, 
                    'student_idx': i, 
                    'number': i + 1,
                    'parent_surname': parent_surname,
                    'child_surname': child_surname_from_parent,
                    'parent_full_name': str(parent_name) if pd.notna(parent_name) else "Неизвестно",
                    'child_full_name': str(child_name_from_parent) if pd.notna(child_name_from_parent) else (str(student_name) if pd.notna(student_name) else "Неизвестно")
                })
            else:
                # Если фамилии не совпадают, всё равно добавляем по порядку
                pairs.append({
                    'pair_id': len(pairs) + 1, 
                    'parent_idx': i, 
                    'student_idx': i, 
                    'number': i + 1,
                    'parent_surname': parent_surname,
                    'child_surname': student_surname if student_surname else "Неизвестно",
                    'parent_full_name': str(parent_name) if pd.notna(parent_name) else "Неизвестно",
                    'child_full_name': str(student_name) if pd.notna(student_name) else "Неизвестно"
                })
        except Exception as e:
            print(f"   Ошибка при обработке пары {i}: {e}")
            pass

print(f"   Найдено пар: {len(pairs)}")
for pair in pairs:
    print(f"   Пара {pair['pair_id']}: {pair['parent_surname']} - {pair['child_surname']}")

# Рассчитываем корреляции Овчарова-Марковская для каждой пары отдельно
print("\n3. Расчёт корреляций Овчарова-Марковская...")
correlations_by_pair = {}  # Словарь: pair_id -> список корреляций

if q12_col and q19_col and len(parent_questions) > 0:
    for pair in pairs:
        pair_id = pair['pair_id']
        correlations_by_pair[pair_id] = {
            'pair_info': pair,
            'correlations': []
        }
        
        for parent_q in parent_questions:
            try:
                p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                s12_val = students_sorted.iloc[pair['student_idx']][q12_col] if q12_col else None
                s19_val = students_sorted.iloc[pair['student_idx']][q19_col] if q19_col else None
                
                p_num = float(p_val) if pd.notna(p_val) else None
                s12_num = float(s12_val) if pd.notna(s12_val) and s12_val != '' else None
                s19_num = float(s19_val) if pd.notna(s19_val) and s19_val != '' else None
                
                # Для одной пары корреляцию рассчитать нельзя, но сохраняем данные
                if p_num is not None and s12_num is not None:
                    correlations_by_pair[pair_id]['correlations'].append({
                        'parent_question': parent_q,
                        'ovcharova_question': '12. Дает возможности для роста профессионального мастерства',
                        'parent_value': p_num,
                        'ovcharova_value': s12_num,
                        'pair_id': pair_id
                    })
                
                if p_num is not None and s19_num is not None:
                    correlations_by_pair[pair_id]['correlations'].append({
                        'parent_question': parent_q,
                        'ovcharova_question': '19. Позволяет использовать профессиональные умения вне работы',
                        'parent_value': p_num,
                        'ovcharova_value': s19_num,
                        'pair_id': pair_id
                    })
            except:
                pass
    
    # Теперь рассчитываем корреляции по всем парам вместе (для статистики)
    correlations_ovcharova_markovskaya = []
    
    for parent_q in parent_questions:
        p_vals_12, q12_vals = [], []
        p_vals_19, q19_vals = [], []
        
        for pair in pairs:
            try:
                p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                s12_val = students_sorted.iloc[pair['student_idx']][q12_col] if q12_col else None
                s19_val = students_sorted.iloc[pair['student_idx']][q19_col] if q19_col else None
                
                p_num = float(p_val) if pd.notna(p_val) else None
                s12_num = float(s12_val) if pd.notna(s12_val) and s12_val != '' else None
                s19_num = float(s19_val) if pd.notna(s19_val) and s19_val != '' else None
                
                if p_num is not None and s12_num is not None:
                    p_vals_12.append(p_num)
                    q12_vals.append(s12_num)
                
                if p_num is not None and s19_num is not None:
                    p_vals_19.append(p_num)
                    q19_vals.append(s19_num)
            except:
                pass
        
        if len(p_vals_12) >= 3:
            try:
                rho12, p12 = spearmanr(p_vals_12, q12_vals)
                r12, p12_p = pearsonr(p_vals_12, q12_vals)
                correlations_ovcharova_markovskaya.append({
                    'parent_question': parent_q,
                    'ovcharova_question': '12. Дает возможности для роста профессионального мастерства',
                    'parent_values': p_vals_12,
                    'ovcharova_values': q12_vals,
                    'spearman_corr': rho12,
                    'spearman_p': p12,
                    'pearson_corr': r12,
                    'pearson_p': p12_p,
                    'n': len(p_vals_12)
                })
            except:
                pass
        
        if len(p_vals_19) >= 3:
            try:
                rho19, p19 = spearmanr(p_vals_19, q19_vals)
                r19, p19_p = pearsonr(p_vals_19, q19_vals)
                correlations_ovcharova_markovskaya.append({
                    'parent_question': parent_q,
                    'ovcharova_question': '19. Позволяет использовать профессиональные умения вне работы',
                    'parent_values': p_vals_19,
                    'ovcharova_values': q19_vals,
                    'spearman_corr': rho19,
                    'spearman_p': p19,
                    'pearson_corr': r19,
                    'pearson_p': p19_p,
                    'n': len(p_vals_19)
                })
            except:
                pass

print(f"   Рассчитано корреляций (по всем парам): {len(correlations_ovcharova_markovskaya)}")
print(f"   Данных по парам: {len(correlations_by_pair)}")

# Находим вопросы Холланда
print("\n4. Поиск вопросов Холланда...")
holland_cols = []
for col in student_cols:
    col_str = str(col).strip()
    if 'Что тебе ближе' in col_str or col_str.startswith('1. Что тебе ближе') or \
       (col_str.startswith('1.') and 'ближе' in col_str.lower()):
        holland_cols.append(col)

print(f"   Найдено вопросов Холланда: {len(holland_cols)}")

# Рассчитываем баллы по типам Холланда
print("\n5. Расчёт баллов по типам Холланда...")
holland_scores = None
ovcharova_holland_correlations = []

if len(holland_cols) > 0:
    holland_scores = calculate_holland_scores(students_sorted, holland_cols, pairs)
    for h_type in ['R', 'I', 'A', 'S', 'E', 'C']:
        scores = holland_scores[h_type]
        print(f"   Тип {h_type}: средний балл = {np.mean(scores):.2f}, мин = {min(scores)}, макс = {max(scores)}")
    
    # Рассчитываем корреляции Овчарова-Холланд
    print("\n6. Расчёт корреляций Овчарова-Холланд...")
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный (Realistic)',
        'I': 'Исследовательский (Investigative)',
        'A': 'Артистический (Artistic)',
        'S': 'Социальный (Social)',
        'E': 'Предприимчивый (Enterprising)',
        'C': 'Конвенциональный (Conventional)'
    }
    
    if q12_col and q19_col:
        for ovcharova_q_col in [q12_col, q19_col]:
            for h_type in holland_types:
                o_vals = []
                h_vals = []
                
                for i, pair in enumerate(pairs):
                    try:
                        o_val = students_sorted.iloc[pair['student_idx']][ovcharova_q_col]
                        h_score = holland_scores[h_type][i]
                        
                        o_num = float(o_val) if pd.notna(o_val) and o_val != '' else None
                        h_num = float(h_score) if pd.notna(h_score) else None
                        
                        if o_num is not None and h_num is not None:
                            o_vals.append(o_num)
                            h_vals.append(h_num)
                    except:
                        pass
                
                if len(o_vals) >= 3:
                    try:
                        rho, p = spearmanr(o_vals, h_vals)
                        ovcharova_holland_correlations.append({
                            'ovcharova_question': ovcharova_q_col,
                            'holland_type': h_type,
                            'holland_type_name': holland_type_names[h_type],
                            'spearman_corr': rho,
                            'spearman_p': p,
                            'n': len(o_vals),
                            'ovcharova_values': o_vals,
                            'holland_values': h_vals
                        })
                    except:
                        pass

print(f"   Корреляций Овчарова-Холланд: {len(ovcharova_holland_correlations)}")

# Фильтруем значимые корреляции
significant_ovcharova_markovskaya = [c for c in correlations_ovcharova_markovskaya if c['spearman_p'] < 0.05]
significant_ovcharova_holland = [c for c in ovcharova_holland_correlations if c['spearman_p'] < 0.05]

print(f"\n7. Статистика значимых корреляций:")
print(f"   Овчарова-Марковская: {len(significant_ovcharova_markovskaya)}")
print(f"   Овчарова-Холланд: {len(significant_ovcharova_holland)}")

# Создаём документ со сводкой
print("\n8. Создание документа со сводкой результатов...")
doc = Document()

def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Заголовок
title = doc.add_heading('СВОДКА РЕЗУЛЬТАТОВ АНАЛИЗА ДОПОЛНИТЕЛЬНЫХ ФАЙЛОВ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Информация о данных
doc.add_heading('1. ИНФОРМАЦИЯ О ДАННЫХ', 1)
data_info = f"""
В анализе использованы следующие файлы:
• Опрос для родителей (Ответы) - Ответы на форму (1) (1).csv
• Опрос ученика (Ответы) - Ответы на форму (1) (1).csv

Количество данных:
• Родителей: {len(parents_df)}
• Учеников: {len(students_df)}
• Пар для анализа: {len(pairs)}
• Вопросов родителей (ВРР Марковской): {len(parent_questions)}
• Вопросов Холланда: {len(holland_cols)}
"""
doc.add_paragraph(data_info)

# Список пар
doc.add_paragraph()
doc.add_heading('2. СПИСОК ПАР РОДИТЕЛЬ-РЕБЕНОК', 1)
pairs_list = doc.add_table(rows=1, cols=4)
pairs_list.style = 'Light Grid Accent 1'

hdr_cells = pairs_list.rows[0].cells
hdr_cells[0].text = '№ пары'
hdr_cells[1].text = 'Фамилия родителя'
hdr_cells[2].text = 'Фамилия ребенка'
hdr_cells[3].text = 'Полные имена'

for pair in pairs:
    row = pairs_list.add_row().cells
    row[0].text = str(pair['pair_id'])
    row[1].text = pair['parent_surname']
    row[2].text = pair['child_surname']
    row[3].text = f"{pair['parent_full_name']} / {pair['child_full_name']}"

# Разделяем результаты по парам
doc.add_heading('3. РЕЗУЛЬТАТЫ ПО ПАРАМ', 1)

# Для каждой пары создаём отдельный раздел
for pair in pairs:
    pair_id = pair['pair_id']
    parent_surname = pair['parent_surname']
    child_surname = pair['child_surname']
    parent_full = pair['parent_full_name']
    child_full = pair['child_full_name']
    
    doc.add_heading(f'3.{pair_id}. Пара {pair_id}: {parent_surname} - {child_surname}', 1)
    doc.add_paragraph(f"Родитель: {parent_full}")
    doc.add_paragraph(f"Ребёнок: {child_full}")
    doc.add_paragraph()
    
    # Собираем данные для этой пары
    pair_parent_vals = {}
    pair_ovcharova_vals = {}
    pair_holland_scores = {}
    
    try:
        # Данные родителя (все вопросы)
        for parent_q in parent_questions:
            try:
                p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                p_num = float(p_val) if pd.notna(p_val) else None
                if p_num is not None:
                    pair_parent_vals[parent_q] = p_num
            except:
                pass
        
        # Данные ребёнка по Овчаровой
        if q12_col:
            try:
                s12_val = students_sorted.iloc[pair['student_idx']][q12_col]
                s12_num = float(s12_val) if pd.notna(s12_val) and s12_val != '' else None
                if s12_num is not None:
                    pair_ovcharova_vals['12. Дает возможности для роста профессионального мастерства'] = s12_num
            except:
                pass
        
        if q19_col:
            try:
                s19_val = students_sorted.iloc[pair['student_idx']][q19_col]
                s19_num = float(s19_val) if pd.notna(s19_val) and s19_val != '' else None
                if s19_num is not None:
                    pair_ovcharova_vals['19. Позволяет использовать профессиональные умения вне работы'] = s19_num
            except:
                pass
        
        # Данные по Холланду
        if holland_scores:
            pair_idx = pair_id - 1  # pair_id начинается с 1, индексы с 0
            if 0 <= pair_idx < len(pairs):
                for h_type in ['R', 'I', 'A', 'S', 'E', 'C']:
                    try:
                        if pair_idx < len(holland_scores[h_type]):
                            pair_holland_scores[h_type] = holland_scores[h_type][pair_idx]
                    except:
                        pass
    except Exception as e:
        print(f"   Ошибка при обработке данных пары {pair_id}: {e}")
    
    # Таблица с данными родителя
    if len(pair_parent_vals) > 0:
        doc.add_heading(f'3.{pair_id}.1. Данные родителя (ВРР Марковской)', 2)
        parent_table = doc.add_table(rows=1, cols=3)
        parent_table.style = 'Light Grid Accent 1'
        
        hdr_cells = parent_table.rows[0].cells
        hdr_cells[0].text = '№ вопроса'
        hdr_cells[1].text = 'Вопрос'
        hdr_cells[2].text = 'Ответ'
        
        for idx, (q, val) in enumerate(list(pair_parent_vals.items())[:20], 1):  # Показываем первые 20
            row = parent_table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q[:60] + "..." if len(q) > 60 else q
            row[2].text = str(int(val)) if val == int(val) else f"{val:.1f}"
        
        if len(pair_parent_vals) > 20:
            doc.add_paragraph(f"... и ещё {len(pair_parent_vals) - 20} вопросов")
    
    # Таблица с данными ребёнка по Овчаровой
    if len(pair_ovcharova_vals) > 0:
        doc.add_heading(f'3.{pair_id}.2. Данные ребёнка по мотивам выбора профессии (Овчарова)', 2)
        ovcharova_table = doc.add_table(rows=1, cols=2)
        ovcharova_table.style = 'Light Grid Accent 1'
        
        hdr_cells = ovcharova_table.rows[0].cells
        hdr_cells[0].text = 'Мотив'
        hdr_cells[1].text = 'Оценка'
        
        for motive, val in pair_ovcharova_vals.items():
            row = ovcharova_table.add_row().cells
            row[0].text = motive
            row[1].text = str(int(val)) if val == int(val) else f"{val:.1f}"
    
    # Таблица с данными по Холланду
    if len(pair_holland_scores) > 0:
        doc.add_heading(f'3.{pair_id}.3. Данные ребёнка по типам профессиональных интересов (Холланд)', 2)
        holland_table = doc.add_table(rows=1, cols=2)
        holland_table.style = 'Light Grid Accent 1'
        
        hdr_cells = holland_table.rows[0].cells
        hdr_cells[0].text = 'Тип Холланда'
        hdr_cells[1].text = 'Балл'
        
        holland_type_names_short = {
            'R': 'Реалистичный',
            'I': 'Исследовательский',
            'A': 'Артистический',
            'S': 'Социальный',
            'E': 'Предприимчивый',
            'C': 'Конвенциональный'
        }
        
        for h_type in ['R', 'I', 'A', 'S', 'E', 'C']:
            if h_type in pair_holland_scores:
                row = holland_table.add_row().cells
                row[0].text = f"{h_type} - {holland_type_names_short[h_type]}"
                row[1].text = str(int(pair_holland_scores[h_type]))
    
    doc.add_paragraph()

# Общая сводка корреляций
doc.add_heading('4. ОБЩАЯ СВОДКА КОРРЕЛЯЦИЙ', 1)

# Сводка корреляций Овчарова-Марковская
doc.add_heading('4.1. КОРРЕЛЯЦИИ ОВЧАРОВА-МАРКОВСКАЯ', 2)

doc.add_paragraph(f"Всего рассчитано корреляций: {len(correlations_ovcharova_markovskaya)}")
doc.add_paragraph(f"Статистически значимых корреляций (p < 0.05): {len(significant_ovcharova_markovskaya)}")

if len(significant_ovcharova_markovskaya) > 0:
    doc.add_heading('4.1.1. Значимые корреляции', 3)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Вопрос родителя (Марковская)'
    hdr_cells[2].text = 'Мотив подростка (Овчарова)'
    hdr_cells[3].text = 'Корреляция Спирмена'
    hdr_cells[4].text = 'Уровень значимости (p)'
    
    for idx, corr in enumerate(sorted(significant_ovcharova_markovskaya, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True), 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = corr['parent_question'][:80] + "..." if len(corr['parent_question']) > 80 else corr['parent_question']
        row[2].text = corr['ovcharova_question'][:60] + "..." if len(corr['ovcharova_question']) > 60 else corr['ovcharova_question']
        row[3].text = f"{corr['spearman_corr']:.4f}"
        
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
    
    # Подробные таблицы для каждой значимой связи
    doc.add_heading('4.1.2. Подробные расчёты для каждой значимой связи', 3)
    
    for idx, corr in enumerate(sorted(significant_ovcharova_markovskaya, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True), 1):
        doc.add_heading(f'Связь {idx}', 3)
        
        parent_stats = calculate_statistics(corr['parent_values'])
        ovcharova_stats = calculate_statistics(corr['ovcharova_values'])
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        
        # Таблица с подробными расчётами
        detail_table = doc.add_table(rows=1, cols=6)
        detail_table.style = 'Light Grid Accent 1'
        
        hdr_cells = detail_table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Родительские установки (Марковская)'
        hdr_cells[2].text = 'Мотивы выбора профессии (Овчарова)'
        hdr_cells[3].text = 'Корреляция'
        hdr_cells[4].text = 'Уровень значимости'
        hdr_cells[5].text = 'N'
        
        row = detail_table.add_row().cells
        row[0].text = 'M (среднее)'
        row[1].text = f"{parent_stats['M']:.2f}" if parent_stats['M'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['M']:.2f}" if ovcharova_stats['M'] is not None else "N/A"
        row[3].text = f"{corr['spearman_corr']:.4f}"
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
        row[5].text = str(corr['n'])
        
        row = detail_table.add_row().cells
        row[0].text = 'SD (стандартное отклонение)'
        row[1].text = f"{parent_stats['SD']:.2f}" if parent_stats['SD'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['SD']:.2f}" if ovcharova_stats['SD'] is not None else "N/A"
        row[3].text = ""
        row[4].text = sig_explanation
        row[5].text = ""
        
        row = detail_table.add_row().cells
        row[0].text = 'Me (медиана)'
        row[1].text = f"{parent_stats['Me']:.2f}" if parent_stats['Me'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['Me']:.2f}" if ovcharova_stats['Me'] is not None else "N/A"
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        
        # Проценты
        if parent_stats['percentages']:
            row = detail_table.add_row().cells
            row[0].text = '% (проценты)'
            parent_pct = ', '.join([f"{val}: {pct:.1f}%" for val, pct in sorted(parent_stats['percentages'].items())])
            row[1].text = parent_pct[:100] + "..." if len(parent_pct) > 100 else parent_pct
            if ovcharova_stats['percentages']:
                ovcharova_pct = ', '.join([f"{val}: {pct:.1f}%" for val, pct in sorted(ovcharova_stats['percentages'].items())])
                row[2].text = ovcharova_pct[:100] + "..." if len(ovcharova_pct) > 100 else ovcharova_pct
            else:
                row[2].text = "N/A"
            row[3].text = ""
            row[4].text = ""
            row[5].text = ""
        
        doc.add_paragraph(f"\nВопрос родителя: {corr['parent_question']}")
        doc.add_paragraph(f"Мотив подростка: {corr['ovcharova_question']}")
        doc.add_paragraph()

# Сводка корреляций Овчарова-Холланд
doc.add_heading('4.2. КОРРЕЛЯЦИИ ОВЧАРОВА-ХОЛЛАНД', 2)

doc.add_paragraph(f"Всего рассчитано корреляций: {len(ovcharova_holland_correlations)}")
doc.add_paragraph(f"Статистически значимых корреляций (p < 0.05): {len(significant_ovcharova_holland)}")

if len(significant_ovcharova_holland) > 0:
    doc.add_heading('4.2.1. Значимые корреляции', 3)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '№'
    hdr_cells[1].text = 'Мотив подростка (Овчарова)'
    hdr_cells[2].text = 'Тип Холланда'
    hdr_cells[3].text = 'Корреляция Спирмена'
    hdr_cells[4].text = 'Уровень значимости (p)'
    
    for idx, corr in enumerate(sorted(significant_ovcharova_holland, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True), 1):
        row = table.add_row().cells
        row[0].text = str(idx)
        ovcharova_q = str(corr.get('ovcharova_question', ''))
        if '12' in ovcharova_q:
            row[1].text = 'Вопрос 12'
        elif '19' in ovcharova_q:
            row[1].text = 'Вопрос 19'
        else:
            row[1].text = ovcharova_q[:60] + "..." if len(ovcharova_q) > 60 else ovcharova_q
        row[2].text = corr['holland_type_name']
        row[3].text = f"{corr['spearman_corr']:.4f}"
        
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
    
    # Подробные таблицы для каждой значимой связи
    doc.add_heading('4.2.2. Подробные расчёты для каждой значимой связи', 3)
    
    for idx, corr in enumerate(sorted(significant_ovcharova_holland, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True), 1):
        doc.add_heading(f'Связь {idx}', 3)
        
        ovcharova_stats = calculate_statistics(corr.get('ovcharova_values', []))
        holland_stats = calculate_statistics(corr.get('holland_values', []))
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        
        # Таблица с подробными расчётами
        detail_table = doc.add_table(rows=1, cols=6)
        detail_table.style = 'Light Grid Accent 1'
        
        hdr_cells = detail_table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Мотивы выбора профессии (Овчарова)'
        hdr_cells[2].text = 'Тип Холланда'
        hdr_cells[3].text = 'Корреляция'
        hdr_cells[4].text = 'Уровень значимости'
        hdr_cells[5].text = 'N'
        
        row = detail_table.add_row().cells
        row[0].text = 'M (среднее)'
        row[1].text = f"{ovcharova_stats['M']:.2f}" if ovcharova_stats['M'] is not None else "N/A"
        row[2].text = f"{holland_stats['M']:.2f}" if holland_stats['M'] is not None else "N/A"
        row[3].text = f"{corr['spearman_corr']:.4f}"
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
        row[5].text = str(corr['n'])
        
        row = detail_table.add_row().cells
        row[0].text = 'SD (стандартное отклонение)'
        row[1].text = f"{ovcharova_stats['SD']:.2f}" if ovcharova_stats['SD'] is not None else "N/A"
        row[2].text = f"{holland_stats['SD']:.2f}" if holland_stats['SD'] is not None else "N/A"
        row[3].text = ""
        row[4].text = sig_explanation
        row[5].text = ""
        
        row = detail_table.add_row().cells
        row[0].text = 'Me (медиана)'
        row[1].text = f"{ovcharova_stats['Me']:.2f}" if ovcharova_stats['Me'] is not None else "N/A"
        row[2].text = f"{holland_stats['Me']:.2f}" if holland_stats['Me'] is not None else "N/A"
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        
        ovcharova_q = str(corr.get('ovcharova_question', ''))
        if '12' in ovcharova_q:
            q_name = 'Вопрос 12'
        elif '19' in ovcharova_q:
            q_name = 'Вопрос 19'
        else:
            q_name = ovcharova_q
        
        doc.add_paragraph(f"\nМотив подростка: {q_name}")
        doc.add_paragraph(f"Тип Холланда: {corr['holland_type_name']}")
        doc.add_paragraph()

# Сохраняем документ
output_file = os.path.join(script_dir, "Сводка_результатов_дополнительных_файлов.docx")
doc.save(output_file)
print(f"\n9. Документ сохранён: {output_file}")

print("\n" + "=" * 80)
print("АНАЛИЗ ДОПОЛНИТЕЛЬНЫХ ФАЙЛОВ ЗАВЕРШЁН!")
print("=" * 80)
