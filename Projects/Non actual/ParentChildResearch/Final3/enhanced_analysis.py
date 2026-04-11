"""
Улучшенный анализ корреляций с подробными расчётами и визуализациями:
1. Подробные расчёты с объяснением "Уровень значимости"
2. Указание M (mean), SD (standard deviation), Me (median), % (проценты)
3. Таблицы для каждой значимой связи
4. Визуализации:
   - Столбчатые диаграммы по 6 типам Холланда с error bars (SD)
   - Тепловая карта значимых корреляций
   - Столбчатая диаграмма по мотивам
   - Мини-матрица корреляций между Овчаровой и Марковской
"""
import os
import sys
import pandas as pd
import numpy as np
from scipy.stats import spearmanr, pearsonr, rankdata
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import warnings
warnings.filterwarnings('ignore')

# Импорты для визуализации
import matplotlib
matplotlib.use('Agg')  # Для работы без GUI
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib import rcParams
from PIL import Image

# Настройка шрифтов для русского языка
rcParams['font.family'] = 'DejaVu Sans'
rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'Tahoma']
plt.rcParams['axes.unicode_minus'] = False  # Для корректного отображения минусов

# ============================================================================
# ФУНКЦИИ ДЛЯ ОБРАБОТКИ ДАННЫХ ХОЛЛАНДА
# ============================================================================

def classify_profession_to_holland_types(profession_text):
    """
    Классифицирует профессию по типам Холланда (R, I, A, S, E, C).
    Возвращает список типов, к которым относится профессия.
    """
    if pd.isna(profession_text) or not str(profession_text).strip():
        return []
    
    prof_lower = str(profession_text).lower()
    types = []
    
    # R (Realistic) - реалистичный
    r_keywords = ['инженер', 'техник', 'механик', 'автомеханик', 'ремонт', 'конструктор', 
                  'оператор', 'фермер', 'водитель', 'электроник', 'чертёжник', 'технолог',
                  'слесарь', 'сварщик', 'токарь', 'фрезеровщик', 'монтажник', 'строитель']
    if any(kw in prof_lower for kw in r_keywords):
        types.append('R')
    
    # I (Investigative) - исследовательский
    i_keywords = ['учёный', 'исследователь', 'аналитик', 'химик', 'физик', 'биолог', 
                  'микробиолог', 'археолог', 'криминалист', 'синоптик', 'эколог',
                  'врач', 'психиатр', 'ветеринар', 'провизор', 'философ', 'политолог']
    if any(kw in prof_lower for kw in i_keywords):
        types.append('I')
    
    # A (Artistic) - артистический
    a_keywords = ['художник', 'дизайнер', 'актёр', 'режиссёр', 'композитор', 'писатель',
                  'поэт', 'фотограф', 'декоратор', 'визажист', 'стилист', 'искусствовед',
                  'реставратор', 'литературный критик', 'видеооператор', 'видеограф']
    if any(kw in prof_lower for kw in a_keywords):
        types.append('A')
    
    # S (Social) - социальный
    s_keywords = ['психолог', 'учитель', 'преподаватель', 'тренер', 'врач', 'терапевт',
                  'логопед', 'социальный', 'молодёжью', 'консультант', 'карьерный',
                  'профилактики', 'правонарушений', 'социолог', 'журналист', 'корреспондент']
    if any(kw in prof_lower for kw in s_keywords):
        types.append('S')
    
    # E (Enterprising) - предприимчивый
    e_keywords = ['менеджер', 'предприниматель', 'директор', 'руководитель', 'трейдер',
                  'адвокат', 'риэлтор', 'продюсер', 'лидер', 'движения', 'общественного',
                  'маркетинг', 'сетевому', 'финансовый', 'генеральный', 'отдела']
    if any(kw in prof_lower for kw in e_keywords):
        types.append('E')
    
    # C (Conventional) - конвенциональный
    c_keywords = ['бухгалтер', 'оператор ввода', 'ввода данных', 'кассир', 'продавец',
                  'редактор', 'корректор', 'статистический', 'экономист', 'контролёр',
                  'качества', 'логистике', 'доставке', 'офис-менеджер']
    if any(kw in prof_lower for kw in c_keywords):
        types.append('C')
    
    return types

def calculate_holland_scores(students_df, holland_cols, pairs):
    """
    Рассчитывает баллы по типам Холланда (R, I, A, S, E, C) для каждого респондента.
    """
    holland_scores = {
        'R': [],
        'I': [],
        'A': [],
        'S': [],
        'E': [],
        'C': []
    }
    
    for pair in pairs:
        scores = {'R': 0, 'I': 0, 'A': 0, 'S': 0, 'E': 0, 'C': 0}
        
        for h_col in holland_cols:
            try:
                profession = students_df.iloc[pair['student_idx']][h_col]
                if pd.notna(profession):
                    types = classify_profession_to_holland_types(profession)
                    for h_type in types:
                        if h_type in scores:
                            scores[h_type] += 1
            except:
                continue
        
        for h_type in holland_scores:
            holland_scores[h_type].append(scores[h_type])
    
    return holland_scores

# ============================================================================
# ФУНКЦИИ ДЛЯ РАСЧЁТА СТАТИСТИКИ
# ============================================================================

def calculate_statistics(values):
    """
    Рассчитывает статистические показатели для массива значений.
    Возвращает словарь с M (mean), SD (standard deviation), Me (median), % (проценты).
    """
    if len(values) == 0:
        return {'M': None, 'SD': None, 'Me': None, 'percentages': {}}
    
    values_array = np.array(values)
    
    # M - среднее арифметическое
    M = np.mean(values_array)
    
    # SD - стандартное отклонение
    SD = np.std(values_array, ddof=1) if len(values_array) > 1 else 0
    
    # Me - медиана
    Me = np.median(values_array)
    
    # % - проценты (доля каждого значения в процентах)
    unique, counts = np.unique(values_array, return_counts=True)
    percentages = {int(val): (count / len(values_array)) * 100 
                   for val, count in zip(unique, counts)}
    
    return {'M': M, 'SD': SD, 'Me': Me, 'percentages': percentages}

def explain_significance_level(p_value):
    """
    Объясняет уровень значимости простым языком.
    
    Уровень значимости (p-value) - это вероятность того, что мы ошибочно 
    считаем связь между переменными реальной, когда на самом деле её нет.
    
    Чем меньше p-value, тем увереннее мы можем сказать, что связь реальна.
    """
    if p_value < 0.001:
        return ("***", "Очень высокая значимость (p < 0.001). Вероятность ошибки менее 0.1%. "
                "Мы можем быть очень уверены, что связь между переменными реальна.")
    elif p_value < 0.01:
        return ("**", "Высокая значимость (p < 0.01). Вероятность ошибки менее 1%. "
                "Мы можем быть уверены, что связь между переменными реальна.")
    elif p_value < 0.05:
        return ("*", "Значимость (p < 0.05). Вероятность ошибки менее 5%. "
                "Мы можем считать, что связь между переменными реальна.")
    else:
        return ("", f"Незначимо (p = {p_value:.4f}). Вероятность ошибки {p_value*100:.1f}%. "
                "Мы не можем уверенно сказать, что связь между переменными реальна.")

# ============================================================================
# ФУНКЦИИ ДЛЯ ВИЗУАЛИЗАЦИИ
# ============================================================================

def create_holland_bar_chart_with_error_bars(holland_scores, save_path):
    """
    Создаёт столбчатую диаграмму по 6 типам Холланда с error bars (SD).
    """
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный',
        'I': 'Исследовательский',
        'A': 'Артистический',
        'S': 'Социальный',
        'E': 'Предприимчивый',
        'C': 'Конвенциональный'
    }
    
    means = [np.mean(holland_scores[t]) for t in holland_types]
    stds = [np.std(holland_scores[t], ddof=1) for t in holland_types]
    labels = [holland_type_names[t] for t in holland_types]
    
    fig, ax = plt.subplots(figsize=(12, 6))
    x_pos = np.arange(len(labels))
    
    bars = ax.bar(x_pos, means, yerr=stds, capsize=5, alpha=0.7, 
                  color=['#2E86AB', '#A23B72', '#F18F01', '#C73E1D', '#6A994E', '#BC4749'],
                  edgecolor='black', linewidth=1.5)
    
    ax.set_xlabel('Типы профессиональных интересов по Холланду', fontsize=12, fontweight='bold')
    ax.set_ylabel('Средний балл (M ± SD)', fontsize=12, fontweight='bold')
    ax.set_title('Распределение баллов по типам профессиональных интересов Холланда', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(labels, rotation=15, ha='right')
    ax.grid(True, alpha=0.3, axis='y', linestyle='--')
    
    # Добавляем значения на столбцы
    for i, (mean, std) in enumerate(zip(means, stds)):
        ax.text(i, mean + std + 0.1, f'{mean:.2f}±{std:.2f}', 
               ha='center', va='bottom', fontsize=10, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_correlation_heatmap(correlations, title, save_path, method_name='Овчарова-Марковская'):
    """
    Создаёт тепловую карту значимых корреляций.
    """
    if len(correlations) == 0:
        print(f"   Нет значимых корреляций для {title}")
        return
    
    # Создаём матрицу корреляций
    # Для Овчаровой-Марковской: строки - вопросы Овчаровой, столбцы - вопросы Марковской
    # Для Холланда: строки - типы Холланда, столбцы - вопросы/типы
    
    # Определяем уникальные значения для строк и столбцов
    if method_name == 'Овчарова-Марковская':
        row_labels = sorted(set([c.get('ovcharova_question', '') for c in correlations]))
        col_labels = sorted(set([c.get('parent_question', '')[:50] for c in correlations]))
    elif 'Холланд' in method_name:
        row_labels = sorted(set([c.get('holland_type_name', c.get('holland_type', '')) for c in correlations]))
        if 'Овчарова' in method_name:
            col_labels = sorted(set([c.get('ovcharova_question', '')[:50] for c in correlations]))
        else:
            col_labels = sorted(set([c.get('parent_question', '')[:50] for c in correlations]))
    else:
        return
    
    # Создаём матрицу
    matrix = np.zeros((len(row_labels), len(col_labels)))
    
    for corr in correlations:
        if method_name == 'Овчарова-Марковская':
            row_idx = row_labels.index(corr.get('ovcharova_question', ''))
            col_idx = col_labels.index(corr.get('parent_question', '')[:50])
        elif 'Холланд' in method_name:
            row_idx = row_labels.index(corr.get('holland_type_name', corr.get('holland_type', '')))
            if 'Овчарова' in method_name:
                col_idx = col_labels.index(corr.get('ovcharova_question', '')[:50])
            else:
                col_idx = col_labels.index(corr.get('parent_question', '')[:50])
        else:
            continue
        
        matrix[row_idx, col_idx] = corr.get('spearman_corr', 0)
    
    # Создаём тепловую карту
    fig, ax = plt.subplots(figsize=(max(12, len(col_labels)*0.5), max(8, len(row_labels)*0.5)))
    
    # Обрезаем длинные названия для столбцов
    col_labels_short = [label[:30] + '...' if len(label) > 30 else label for label in col_labels]
    
    sns.heatmap(matrix, annot=True, fmt='.3f', cmap='RdYlBu_r', center=0,
                xticklabels=col_labels_short, yticklabels=row_labels,
                cbar_kws={'label': 'Коэффициент корреляции Спирмена'},
                linewidths=0.5, linecolor='gray', ax=ax)
    
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel('', fontsize=10)
    ax.set_ylabel('', fontsize=10)
    
    plt.xticks(rotation=45, ha='right')
    plt.yticks(rotation=0)
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_ovcharova_holland_bar_chart(correlations, save_path):
    """
    Создаёт столбчатую диаграмму всех корреляций Овчарова-Холланд (значимых и незначимых).
    Показывает все рассчитанные корреляции.
    """
    if len(correlations) == 0:
        print(f"   Нет корреляций для диаграммы Овчарова-Холланд")
        return
    
    # Группируем по типам Холланда и вопросам Овчаровой
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный',
        'I': 'Исследовательский',
        'A': 'Артистический',
        'S': 'Социальный',
        'E': 'Предприимчивый',
        'C': 'Конвенциональный'
    }
    
    # Создаём структуру данных для диаграммы
    # Группируем по типам Холланда
    data_by_type = {h_type: [] for h_type in holland_types}
    
    for corr in correlations:
        h_type = corr.get('holland_type', '')
        if h_type in data_by_type:
            ovcharova_q = str(corr.get('ovcharova_question', ''))
            # Упрощаем название вопроса - проверяем по содержимому
            if '12' in ovcharova_q or 'Дает возможности' in ovcharova_q or 'роста профессионального' in ovcharova_q:
                q_name = 'Вопрос 12'
            elif '19' in ovcharova_q or 'Позволяет использовать' in ovcharova_q or 'профессиональные умения вне работы' in ovcharova_q:
                q_name = 'Вопрос 19'
            else:
                # Если не удалось определить, используем первые символы
                q_name = ovcharova_q[:30] if len(ovcharova_q) > 0 else 'Неизвестно'
            
            data_by_type[h_type].append({
                'question': q_name,
                'correlation': corr.get('spearman_corr', 0),
                'p_value': corr.get('spearman_p', 1),
                'significant': corr.get('spearman_p', 1) < 0.05
            })
    
    # Создаём диаграмму
    fig, ax = plt.subplots(figsize=(14, 8))
    
    # Подготовка данных для группированной столбчатой диаграммы
    x_pos_base = np.arange(len(holland_types))
    width = 0.35  # Ширина столбцов
    
    # Определяем позиции для каждого вопроса
    questions = ['Вопрос 12', 'Вопрос 19']
    colors = ['#2E86AB', '#A23B72']
    
    for q_idx, question in enumerate(questions):
        correlations_by_type = []
        significance_by_type = []
        
        for h_type in holland_types:
            # Находим корреляцию для этого вопроса и типа
            corr_value = 0
            is_significant = False
            for item in data_by_type[h_type]:
                if item['question'] == question:
                    corr_value = item['correlation']
                    is_significant = item['significant']
                    break
            correlations_by_type.append(corr_value)
            significance_by_type.append(is_significant)
        
        # Позиции для этого вопроса
        x_pos = x_pos_base + q_idx * width
        
        # Цвет столбцов: основной цвет для значимых, серый для незначимых
        bar_colors = [colors[q_idx] if is_sig else '#CCCCCC' 
                     for is_sig in significance_by_type]
        
        bars = ax.bar(x_pos, correlations_by_type, width, 
                     label=question, color=bar_colors, alpha=0.7,
                     edgecolor='black', linewidth=1)
        
        # Добавляем значения на столбцы
        for i, (val, is_sig) in enumerate(zip(correlations_by_type, significance_by_type)):
            if abs(val) > 0.001:  # Показываем только ненулевые значения
                marker = '*' if is_sig else ''
                ax.text(x_pos[i], val + (0.02 if val >= 0 else -0.02), 
                       f'{val:.3f}{marker}', 
                       ha='center', va='bottom' if val >= 0 else 'top', 
                       fontsize=8, fontweight='bold' if is_sig else 'normal')
    
    ax.set_xlabel('Типы профессиональных интересов по Холланду', 
                 fontsize=12, fontweight='bold')
    ax.set_ylabel('Коэффициент корреляции Спирмена', 
                 fontsize=12, fontweight='bold')
    ax.set_title('Корреляции между мотивами Овчаровой и типами Холланда\n(все рассчитанные корреляции)', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x_pos_base + width / 2)
    ax.set_xticklabels([holland_type_names[t] for t in holland_types], 
                      rotation=15, ha='right')
    ax.legend(loc='upper right')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
    ax.grid(True, alpha=0.3, axis='y', linestyle='--')
    
    # Добавляем пояснение
    ax.text(0.02, 0.98, '* - значимые корреляции (p < 0.05)', 
           transform=ax.transAxes, fontsize=9, 
           verticalalignment='top', bbox=dict(boxstyle='round', 
           facecolor='wheat', alpha=0.5))
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_motives_bar_chart(ovcharova_data, save_path, use_median=True):
    """
    Создаёт столбчатую диаграмму по мотивам выбора профессии.
    Использует медиану (Me) или среднее (M) в зависимости от параметра.
    """
    motives = []
    values = []
    labels = []
    
    for motive_name, values_list in ovcharova_data.items():
        if len(values_list) > 0:
            motives.append(motive_name)
            if use_median:
                values.append(np.median(values_list))
                labels.append('Me')
            else:
                values.append(np.mean(values_list))
                labels.append('M')
    
    fig, ax = plt.subplots(figsize=(14, 6))
    x_pos = np.arange(len(motives))
    
    bars = ax.bar(x_pos, values, alpha=0.7, color='#2E86AB', 
                  edgecolor='black', linewidth=1.5)
    
    ax.set_xlabel('Мотивы выбора профессии (опросник Овчаровой)', 
                 fontsize=12, fontweight='bold')
    ax.set_ylabel(f'Значение ({labels[0] if labels else "M"})', 
                 fontsize=12, fontweight='bold')
    ax.set_title('Выраженность мотивов выбора профессии у подростков', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(motives, rotation=45, ha='right', fontsize=9)
    ax.grid(True, alpha=0.3, axis='y', linestyle='--')
    
    # Добавляем значения на столбцы
    for i, val in enumerate(values):
        ax.text(i, val + 0.05, f'{val:.2f}', 
               ha='center', va='bottom', fontsize=9, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_ovcharova_markovskaya_pie_chart(correlations, save_path):
    """
    Создаёт круговую диаграмму значимых корреляций между вопросами Овчаровой и Марковской.
    Показывает распределение корреляций по силе связи.
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    if len(significant) == 0:
        print("   Нет значимых корреляций для круговой диаграммы")
        return
    
    # Группируем корреляции по силе связи
    strong_positive = [c for c in significant if c.get('spearman_corr', 0) > 0.5]
    moderate_positive = [c for c in significant if 0.3 < c.get('spearman_corr', 0) <= 0.5]
    weak_positive = [c for c in significant if 0 < c.get('spearman_corr', 0) <= 0.3]
    weak_negative = [c for c in significant if -0.3 <= c.get('spearman_corr', 0) < 0]
    moderate_negative = [c for c in significant if -0.5 <= c.get('spearman_corr', 0) < -0.3]
    strong_negative = [c for c in significant if c.get('spearman_corr', 0) < -0.5]
    
    # Подготовка данных для круговой диаграммы
    labels = []
    sizes = []
    colors_list = []
    
    if len(strong_positive) > 0:
        labels.append(f'Сильная положительная\n(ρ > 0.5)\n{len(strong_positive)} связей')
        sizes.append(len(strong_positive))
        colors_list.append('#2E86AB')
    
    if len(moderate_positive) > 0:
        labels.append(f'Умеренная положительная\n(0.3 < ρ ≤ 0.5)\n{len(moderate_positive)} связей')
        sizes.append(len(moderate_positive))
        colors_list.append('#6A994E')
    
    if len(weak_positive) > 0:
        labels.append(f'Слабая положительная\n(0 < ρ ≤ 0.3)\n{len(weak_positive)} связей')
        sizes.append(len(weak_positive))
        colors_list.append('#A7C957')
    
    if len(weak_negative) > 0:
        labels.append(f'Слабая отрицательная\n(-0.3 ≤ ρ < 0)\n{len(weak_negative)} связей')
        sizes.append(len(weak_negative))
        colors_list.append('#F18F01')
    
    if len(moderate_negative) > 0:
        labels.append(f'Умеренная отрицательная\n(-0.5 ≤ ρ < -0.3)\n{len(moderate_negative)} связей')
        sizes.append(len(moderate_negative))
        colors_list.append('#C73E1D')
    
    if len(strong_negative) > 0:
        labels.append(f'Сильная отрицательная\n(ρ < -0.5)\n{len(strong_negative)} связей')
        sizes.append(len(strong_negative))
        colors_list.append('#BC4749')
    
    if len(sizes) == 0:
        print("   Нет данных для круговой диаграммы")
        return
    
    # Создаём круговую диаграмму
    fig, ax = plt.subplots(figsize=(12, 10))
    
    wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors_list, autopct='%1.1f%%',
                                      startangle=90, textprops={'fontsize': 10, 'fontweight': 'bold'})
    
    ax.set_title('Распределение значимых корреляций между мотивами Овчаровой и установками Марковской\nпо силе связи', 
                fontsize=14, fontweight='bold', pad=20)
    
    # Добавляем общую информацию
    total_text = f'Всего значимых корреляций: {len(significant)}'
    ax.text(0, -1.3, total_text, ha='center', fontsize=11, fontweight='bold',
           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_markovskaya_holland_bar_chart(correlations, save_path):
    """
    Создаёт столбчатую диаграмму значимых корреляций между установками Марковской и типами Холланда.
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    if len(significant) == 0:
        print("   Нет значимых корреляций для столбчатой диаграммы Марковская-Холланд")
        return
    
    # Группируем по типам Холланда
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный',
        'I': 'Исследовательский',
        'A': 'Артистический',
        'S': 'Социальный',
        'E': 'Предприимчивый',
        'C': 'Конвенциональный'
    }
    
    # Собираем данные по типам
    data_by_type = {h_type: [] for h_type in holland_types}
    
    for corr in significant:
        h_type = corr.get('holland_type', '')
        if h_type in data_by_type:
            data_by_type[h_type].append({
                'correlation': corr.get('spearman_corr', 0),
                'parent_question': corr.get('parent_question', '')[:50],
                'p_value': corr.get('spearman_p', 1)
            })
    
    # Сортируем по силе корреляции и берём топ-5 для каждого типа
    top_correlations = []
    for h_type in holland_types:
        type_corrs = sorted(data_by_type[h_type], key=lambda x: abs(x['correlation']), reverse=True)[:5]
        for corr in type_corrs:
            top_correlations.append({
                'holland_type': h_type,
                'holland_name': holland_type_names[h_type],
                'correlation': corr['correlation'],
                'parent_question': corr['parent_question'],
                'p_value': corr['p_value']
            })
    
    if len(top_correlations) == 0:
        print("   Нет данных для столбчатой диаграммы")
        return
    
    # Сортируем все корреляции по силе
    top_correlations = sorted(top_correlations, key=lambda x: abs(x['correlation']), reverse=True)[:20]
    
    # Создаём диаграмму
    fig, ax = plt.subplots(figsize=(16, 8))
    
    labels = []
    correlations_vals = []
    colors_list = []
    
    for item in top_correlations:
        # Создаём короткую метку
        h_type = item['holland_type']
        q_short = item['parent_question'][:30] + "..." if len(item['parent_question']) > 30 else item['parent_question']
        labels.append(f"{h_type}: {q_short}")
        correlations_vals.append(item['correlation'])
        
        # Цвет в зависимости от знака и типа
        if item['correlation'] > 0:
            colors_list.append('#2E86AB')
        else:
            colors_list.append('#C73E1D')
    
    x_pos = np.arange(len(labels))
    bars = ax.barh(x_pos, correlations_vals, color=colors_list, alpha=0.7, edgecolor='black', linewidth=1)
    
    ax.set_yticks(x_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel('Коэффициент корреляции Спирмена', fontsize=12, fontweight='bold')
    ax.set_title('Топ значимых корреляций между установками Марковской и типами Холланда', 
                fontsize=14, fontweight='bold', pad=20)
    ax.axvline(x=0, color='black', linestyle='-', linewidth=0.5)
    ax.grid(True, alpha=0.3, axis='x', linestyle='--')
    
    # Добавляем значения
    for i, (val, item) in enumerate(zip(correlations_vals, top_correlations)):
        marker = '*' if item['p_value'] < 0.05 else ''
        ax.text(val + (0.02 if val >= 0 else -0.02), i, 
               f'{val:.3f}{marker}', 
               ha='left' if val >= 0 else 'right', va='center', 
               fontsize=8, fontweight='bold')
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_comparative_holland_table(correlations, save_path):
    """
    Создаёт сравнительную столбчатую таблицу (визуализацию) по топ-связям установок и профессионального типа по Холланду.
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    if len(significant) == 0:
        print("   Нет значимых корреляций для сравнительной таблицы")
        return
    
    # Группируем по типам Холланда
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный',
        'I': 'Исследовательский',
        'A': 'Артистический',
        'S': 'Социальный',
        'E': 'Предприимчивый',
        'C': 'Конвенциональный'
    }
    
    # Собираем топ-3 корреляции для каждого типа
    top_by_type = {}
    for h_type in holland_types:
        type_corrs = [c for c in significant if c.get('holland_type', '') == h_type]
        type_corrs_sorted = sorted(type_corrs, key=lambda x: abs(x.get('spearman_corr', 0)), reverse=True)[:3]
        top_by_type[h_type] = type_corrs_sorted
    
    # Создаём визуализацию
    fig, ax = plt.subplots(figsize=(16, 10))
    
    # Подготовка данных
    x_pos_base = np.arange(len(holland_types))
    width = 0.25
    
    # Для каждого места (1-е, 2-е, 3-е) создаём столбцы
    positions = [x_pos_base - width, x_pos_base, x_pos_base + width]
    colors = ['#2E86AB', '#6A994E', '#A7C957']
    labels_pos = ['1-е место', '2-е место', '3-е место']
    
    for pos_idx, (x_pos, color, label) in enumerate(zip(positions, colors, labels_pos)):
        values = []
        for h_type in holland_types:
            if len(top_by_type[h_type]) > pos_idx:
                values.append(top_by_type[h_type][pos_idx].get('spearman_corr', 0))
            else:
                values.append(0)
        
        bars = ax.bar(x_pos, values, width, label=label, color=color, alpha=0.7, 
                     edgecolor='black', linewidth=1)
        
        # Добавляем значения
        for i, val in enumerate(values):
            if abs(val) > 0.001:
                ax.text(x_pos[i], val + (0.02 if val >= 0 else -0.02), 
                       f'{val:.3f}', ha='center', va='bottom' if val >= 0 else 'top',
                       fontsize=8, fontweight='bold')
    
    ax.set_xlabel('Типы профессиональных интересов по Холланду', fontsize=12, fontweight='bold')
    ax.set_ylabel('Коэффициент корреляции Спирмена', fontsize=12, fontweight='bold')
    ax.set_title('Сравнительная таблица топ-3 значимых корреляций установок Марковской с типами Холланда', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xticks(x_pos_base)
    ax.set_xticklabels([holland_type_names[t] for t in holland_types], rotation=15, ha='right')
    ax.legend(loc='upper right')
    ax.axhline(y=0, color='black', linestyle='-', linewidth=0.5)
    ax.grid(True, alpha=0.3, axis='y', linestyle='--')
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_figure1_ovcharova_markovskaya_heatmap(correlations, save_path):
    """
    Рисунок 1: Тепловая карта корреляций между ВРР Марковской и мотивами Овчаровой.
    Строки: пункты ВРР (только те, что в значимых связях)
    Столбцы: мотив 12 и мотив 19
    Ячейки: ρ со знаком + пометка * (p<0.05)
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    # Группируем по вопросам Овчаровой (12 и 19)
    q12_corrs = []
    q19_corrs = []
    
    for corr in significant:
        ovcharova_q = str(corr.get('ovcharova_question', ''))
        if '12' in ovcharova_q or 'Дает возможности' in ovcharova_q or 'роста профессионального' in ovcharova_q:
            q12_corrs.append(corr)
        elif '19' in ovcharova_q or 'Позволяет использовать' in ovcharova_q or 'профессиональные умения вне работы' in ovcharova_q:
            q19_corrs.append(corr)
    
    # Получаем уникальные вопросы ВРР только из корреляций для вопросов 12 и 19
    all_parent_questions = set()
    for corr in q12_corrs + q19_corrs:
        parent_q = corr.get('parent_question', '')
        if parent_q:
            all_parent_questions.add(parent_q)
    
    parent_questions_sorted = sorted(list(all_parent_questions))
    
    # Создаём матрицу
    matrix = np.full((len(parent_questions_sorted), 2), np.nan)
    annotations = [['' for _ in range(2)] for _ in range(len(parent_questions_sorted))]
    
    # Заполняем матрицу для вопроса 12
    for corr in q12_corrs:
        parent_q = corr.get('parent_question', '')
        if parent_q in parent_questions_sorted:
            row_idx = parent_questions_sorted.index(parent_q)
            rho = corr.get('spearman_corr', 0)
            p_val = corr.get('spearman_p', 1)
            matrix[row_idx, 0] = rho
            marker = '*' if p_val < 0.05 else ''
            annotations[row_idx][0] = f'{rho:.3f}{marker}'
    
    # Заполняем матрицу для вопроса 19
    for corr in q19_corrs:
        parent_q = corr.get('parent_question', '')
        if parent_q in parent_questions_sorted:
            row_idx = parent_questions_sorted.index(parent_q)
            rho = corr.get('spearman_corr', 0)
            p_val = corr.get('spearman_p', 1)
            matrix[row_idx, 1] = rho
            marker = '*' if p_val < 0.05 else ''
            annotations[row_idx][1] = f'{rho:.3f}{marker}'
    
    if len(parent_questions_sorted) == 0:
        print("   Нет данных для Рисунка 1")
        return
    
    # Создаём тепловую карту
    fig, ax = plt.subplots(figsize=(10, max(8, len(parent_questions_sorted) * 0.4)))
    
    # Обрезаем длинные названия для отображения
    parent_questions_short = [q[:60] + "..." if len(q) > 60 else q for q in parent_questions_sorted]
    
    # Заменяем NaN на 0 для визуализации
    matrix_plot = np.nan_to_num(matrix, nan=0)
    
    sns.heatmap(matrix_plot, annot=np.array(annotations), fmt='', cmap='RdYlBu_r', center=0,
                xticklabels=['Мотив 12', 'Мотив 19'],
                yticklabels=parent_questions_short,
                cbar_kws={'label': 'Коэффициент корреляции Спирмена (ρ)'},
                linewidths=0.5, linecolor='gray', ax=ax, vmin=-1, vmax=1)
    
    ax.set_title('Рисунок 1. Корреляции между ВРР Марковской и мотивами Овчаровой\n(значимые связи, p < 0.05)', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel('Мотивы выбора профессии (Овчарова)', fontsize=12, fontweight='bold')
    ax.set_ylabel('Пункты ВРР Марковской', fontsize=12, fontweight='bold')
    
    plt.xticks(rotation=0, fontsize=11)
    plt.yticks(rotation=0, fontsize=9)
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_figure2_markovskaya_holland_heatmap(correlations, save_path):
    """
    Рисунок 2: Тепловая карта корреляций между ВРР Марковской и профтипами Холланда (все значимые).
    Строки: пункты ВРР (только те, что имеют значимые связи)
    Столбцы: 6 типов Холланда
    Ячейки: ρ + * (p<0.05), пустые = "—"
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    if len(significant) == 0:
        print("   Нет данных для Рисунка 2")
        return
    
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный',
        'I': 'Исследовательский',
        'A': 'Артистический',
        'S': 'Социальный',
        'E': 'Предприимчивый',
        'C': 'Конвенциональный'
    }
    
    # Получаем уникальные вопросы ВРР
    all_parent_questions = set()
    for corr in significant:
        parent_q = corr.get('parent_question', '')
        if parent_q:
            all_parent_questions.add(parent_q)
    
    parent_questions_sorted = sorted(list(all_parent_questions))
    
    # Создаём матрицу
    matrix = np.full((len(parent_questions_sorted), len(holland_types)), np.nan)
    annotations = [['—' for _ in range(len(holland_types))] for _ in range(len(parent_questions_sorted))]
    
    # Заполняем матрицу
    for corr in significant:
        parent_q = corr.get('parent_question', '')
        h_type = corr.get('holland_type', '')
        
        if parent_q in parent_questions_sorted and h_type in holland_types:
            row_idx = parent_questions_sorted.index(parent_q)
            col_idx = holland_types.index(h_type)
            rho = corr.get('spearman_corr', 0)
            p_val = corr.get('spearman_p', 1)
            matrix[row_idx, col_idx] = rho
            marker = '*' if p_val < 0.05 else ''
            annotations[row_idx][col_idx] = f'{rho:.3f}{marker}'
    
    # Создаём тепловую карту
    fig, ax = plt.subplots(figsize=(12, max(8, len(parent_questions_sorted) * 0.4)))
    
    # Обрезаем длинные названия
    parent_questions_short = [q[:50] + "..." if len(q) > 50 else q for q in parent_questions_sorted]
    holland_labels = [holland_type_names[h] for h in holland_types]
    
    # Заменяем NaN на 0 для визуализации, но сохраняем информацию о пустых ячейках
    matrix_plot = np.nan_to_num(matrix, nan=0)
    
    sns.heatmap(matrix_plot, annot=np.array(annotations), fmt='', cmap='RdYlBu_r', center=0,
                xticklabels=holland_labels,
                yticklabels=parent_questions_short,
                cbar_kws={'label': 'Коэффициент корреляции Спирмена (ρ)'},
                linewidths=0.5, linecolor='gray', ax=ax, vmin=-1, vmax=1,
                mask=np.isnan(matrix))  # Маскируем пустые ячейки
    
    ax.set_title('Рисунок 2. Корреляции между ВРР Марковской и профтипами Холланда\n(значимые связи, p < 0.05)', 
                fontsize=14, fontweight='bold', pad=20)
    ax.set_xlabel('Типы профессиональных интересов (Холланд)', fontsize=12, fontweight='bold')
    ax.set_ylabel('Пункты ВРР Марковской', fontsize=12, fontweight='bold')
    
    plt.xticks(rotation=15, ha='right', fontsize=10)
    plt.yticks(rotation=0, fontsize=9)
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

def create_figure3_ovcharova_holland_scatter(correlations, save_path):
    """
    Рисунок 3: Точечная диаграмма корреляции между мотивами выбора профессии и профтипом Холланда (1 значимая).
    X: балл по мотиву
    Y: балл по типу Холланда
    Подпись: ρ = …; p = …
    """
    # Фильтруем только значимые корреляции
    significant = [c for c in correlations if c.get('spearman_p', 1) < 0.05]
    
    if len(significant) == 0:
        print("   Нет данных для Рисунка 3")
        return
    
    # Берём первую значимую корреляцию (или все, если их несколько)
    for corr in significant:
        ovcharova_values = corr.get('ovcharova_values', [])
        holland_values = corr.get('holland_values', [])
        rho = corr.get('spearman_corr', 0)
        p_val = corr.get('spearman_p', 1)
        ovcharova_q = corr.get('ovcharova_question', '')
        holland_type = corr.get('holland_type', '')
        holland_name = corr.get('holland_type_name', holland_type)
        
        if len(ovcharova_values) == 0 or len(holland_values) == 0:
            continue
        
        # Создаём scatter plot
        fig, ax = plt.subplots(figsize=(10, 8))
        
        ax.scatter(ovcharova_values, holland_values, alpha=0.6, s=100, color='#2E86AB', edgecolors='black', linewidth=1)
        
        # Добавляем линию тренда
        z = np.polyfit(ovcharova_values, holland_values, 1)
        p = np.poly1d(z)
        x_trend = np.linspace(min(ovcharova_values), max(ovcharova_values), 100)
        ax.plot(x_trend, p(x_trend), "r--", alpha=0.8, linewidth=2, label='Линия тренда')
        
        # Без текстовых подписей осей и заголовка, но с числовыми метками на шкалах
        ax.set_xlabel('')
        ax.set_ylabel('')
        ax.set_title('')
        # Числовые метки на шкалах оставляем (labelbottom=True, labelleft=True по умолчанию)
        
        # Добавляем подпись с ρ и p
        sig_text = f'ρ = {rho:.3f}; p = {p_val:.4f}'
        if p_val < 0.001:
            sig_text += ' ***'
        elif p_val < 0.01:
            sig_text += ' **'
        elif p_val < 0.05:
            sig_text += ' *'
        
        ax.text(0.05, 0.95, sig_text, transform=ax.transAxes, fontsize=12,
               verticalalignment='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
        
        ax.grid(True, alpha=0.3, linestyle='--')
        ax.legend(loc='lower right')
        plt.tight_layout()
        plt.savefig(save_path, dpi=300, bbox_inches='tight')
        plt.close()
        print(f"   Сохранено: {save_path}")
        return  # Сохраняем только первую значимую корреляцию
    
    print("   Нет данных для Рисунка 3")

def create_figure4_rating_all_significant(all_significant_correlations, save_path):
    """
    Рисунок 4: Рейтинг значимых связей.
    Горизонтальная столбчатая диаграмма по |ρ| (24 строки).
    Y: все значимые пары
    X: ρ (нулевая линия по центру)
    Сортировка по модулю ρ от сильных к слабым
    """
    if len(all_significant_correlations) == 0:
        print("   Нет данных для Рисунка 4")
        return
    
    # Подготовка данных с подписями
    data_list = []
    for corr in all_significant_correlations:
        rho = corr.get('spearman_corr', 0)
        p_val = corr.get('spearman_p', 1)
        
        # Формируем подпись в зависимости от типа корреляции
        if 'ovcharova_question' in corr and 'parent_question' in corr:
            # Овчарова-Марковская
            ovcharova_q = str(corr.get('ovcharova_question', ''))[:30]
            parent_q = str(corr.get('parent_question', ''))[:40]
            label = f"Овч-Марк: {ovcharova_q} ↔ {parent_q}"
        elif 'parent_question' in corr and 'holland_type' in corr:
            # Марковская-Холланд
            parent_q = str(corr.get('parent_question', ''))[:40]
            holland_type = corr.get('holland_type_name', corr.get('holland_type', ''))
            label = f"Марк-Холл: {parent_q} ↔ {holland_type}"
        elif 'ovcharova_question' in corr and 'holland_type' in corr:
            # Овчарова-Холланд
            ovcharova_q = str(corr.get('ovcharova_question', ''))[:30]
            holland_type = corr.get('holland_type_name', corr.get('holland_type', ''))
            label = f"Овч-Холл: {ovcharova_q} ↔ {holland_type}"
        else:
            label = "Неизвестно"
        
        data_list.append({
            'label': label,
            'rho': rho,
            'abs_rho': abs(rho),
            'p_val': p_val
        })
    
    # Сортируем по модулю ρ от сильных к слабым
    data_list.sort(key=lambda x: x['abs_rho'], reverse=True)
    
    # Ограничиваем до 24 самых сильных
    data_list = data_list[:24]
    
    # Извлекаем данные для графика
    labels = [d['label'] for d in data_list]
    rhos = [d['rho'] for d in data_list]
    
    # Создаём горизонтальную столбчатую диаграмму
    fig, ax = plt.subplots(figsize=(14, max(10, len(labels) * 0.5)))
    
    y_pos = np.arange(len(labels))
    colors = ['#2E86AB' if r > 0 else '#C73E1D' for r in rhos]
    
    bars = ax.barh(y_pos, rhos, color=colors, alpha=0.7, edgecolor='black', linewidth=1)
    
    # Добавляем значения на столбцы
    for i, (val, p_val) in enumerate(zip(rhos, [d['p_val'] for d in data_list])):
        marker = '*' if p_val < 0.05 else ''
        ax.text(val + (0.02 if val >= 0 else -0.02), i, 
               f'{val:.3f}{marker}', 
               ha='left' if val >= 0 else 'right', va='center', 
               fontsize=8, fontweight='bold')
    
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=9)
    ax.set_xlabel('Коэффициент корреляции Спирмена (ρ)', fontsize=12, fontweight='bold')
    ax.set_title('Рисунок 4. Рейтинг значимых связей\n(отсортировано по модулю ρ от сильных к слабым, p < 0.05)', 
                fontsize=14, fontweight='bold', pad=20)
    ax.axvline(x=0, color='black', linestyle='-', linewidth=1)
    ax.grid(True, alpha=0.3, axis='x', linestyle='--')
    
    # Добавляем легенду
    from matplotlib.patches import Patch
    legend_elements = [
        Patch(facecolor='#2E86AB', alpha=0.7, label='Положительная корреляция'),
        Patch(facecolor='#C73E1D', alpha=0.7, label='Отрицательная корреляция')
    ]
    ax.legend(handles=legend_elements, loc='lower right')
    
    ax.text(0.02, 0.98, '* - значимые корреляции (p < 0.05)', 
           transform=ax.transAxes, fontsize=9, 
           verticalalignment='top', bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.5))
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    print(f"   Сохранено: {save_path}")

# ============================================================================
# ОСНОВНОЙ КОД
# ============================================================================

script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)
newtest_dir = os.path.join(parent_dir, "newtest")

print("=" * 80)
print("УЛУЧШЕННЫЙ АНАЛИЗ КОРРЕЛЯЦИЙ С ПОДРОБНЫМИ РАСЧЁТАМИ")
print("=" * 80)

# Загружаем данные
parents_file = os.path.join(newtest_dir, "Опрос для родителей  (Ответы).csv")
students_file = os.path.join(newtest_dir, "Опрос ученика (Ответы) Новый.csv")

print("\n1. Загрузка данных...")
try:
    if not os.path.exists(parents_file):
        print(f"ОШИБКА: Файл не найден: {parents_file}")
        sys.exit(1)
    if not os.path.exists(students_file):
        print(f"ОШИБКА: Файл не найден: {students_file}")
        sys.exit(1)
    
    parents_df = pd.read_csv(parents_file, encoding='utf-8-sig')
    students_df = pd.read_csv(students_file, encoding='utf-8-sig')
except Exception as e:
    print(f"ОШИБКА при загрузке данных: {e}")
    sys.exit(1)

print(f"   Родителей: {len(parents_df)}")
print(f"   Учеников: {len(students_df)}")

# Определяем структуру
parent_questions = parents_df.columns[6:66].tolist()

# Находим вопросы Овчаровой
student_cols = students_df.columns.tolist()
q12_col = None
q19_col = None
ovcharova_all = {}

for col in student_cols:
    col_str = str(col)
    for i in range(1, 21):
        if col_str.startswith(f"{i}. "):
            if any(x in col_str for x in ['Требует', 'Нравится', 'Предполагает', 'Соответствует', 
                                          'Позволяет', 'Дает', 'Является', 'Близка', 'Избрана', 'Единственно', 'Способствует']):
                ovcharova_all[i] = col
                if i == 12 and 'Дает возможности для роста' in col_str:
                    q12_col = col
                if i == 19 and 'Позволяет использовать профессиональные умения' in col_str:
                    q19_col = col
                break

if not q12_col or not q19_col:
    print("ОШИБКА: Не найдены вопросы 12 и 19!")
    sys.exit(1)

# Сопоставляем пары
if 'Number' in parents_df.columns and 'Number' in students_df.columns:
    parents_sorted = parents_df.sort_values('Number').reset_index(drop=True)
    students_sorted = students_df.sort_values('Number').reset_index(drop=True)
    pairs = []
    for i in range(min(len(parents_sorted), len(students_sorted))):
        p_num = parents_sorted.iloc[i]['Number'] if pd.notna(parents_sorted.iloc[i]['Number']) else None
        s_num = students_sorted.iloc[i]['Number'] if pd.notna(students_sorted.iloc[i]['Number']) else None
        if p_num is not None and s_num is not None and p_num == s_num:
            pairs.append({'pair_id': len(pairs) + 1, 'parent_idx': i, 'student_idx': i, 'number': p_num})
else:
    pairs = [{'pair_id': i+1, 'parent_idx': i, 'student_idx': i, 'number': i+1} 
             for i in range(min(len(parents_df), len(students_df)))]
    parents_sorted = parents_df.reset_index(drop=True)
    students_sorted = students_df.reset_index(drop=True)

print(f"\n2. Найдено пар: {len(pairs)}")

# Перепроверяем корреляции
print("\n3. Перепроверка корреляций...")
verified_correlations = []

for parent_q in parent_questions:
    p_vals_12, q12_vals = [], []
    p_vals_19, q19_vals = [], []
    
    for pair in pairs:
        try:
            p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
            s12_val = students_sorted.iloc[pair['student_idx']][q12_col]
            s19_val = students_sorted.iloc[pair['student_idx']][q19_col]
            
            p_num = float(p_val) if pd.notna(p_val) else None
            s12_num = float(s12_val) if pd.notna(s12_val) else None
            s19_num = float(s19_val) if pd.notna(s19_val) else None
            
            if p_num is not None and s12_num is not None:
                p_vals_12.append(p_num)
                q12_vals.append(s12_num)
            
            if p_num is not None and s19_num is not None:
                p_vals_19.append(p_num)
                q19_vals.append(s19_num)
        except:
            pass
    
    if len(p_vals_12) >= 3:
        try:
            rho12, p12 = spearmanr(p_vals_12, q12_vals)
            r12, p12_p = pearsonr(p_vals_12, q12_vals)
            verified_correlations.append({
                'parent_question': parent_q,
                'ovcharova_question': '12. Дает возможности для роста профессионального мастерства',
                'parent_values': p_vals_12,
                'ovcharova_values': q12_vals,
                'spearman_corr': rho12,
                'spearman_p': p12,
                'pearson_corr': r12,
                'pearson_p': p12_p,
                'n': len(p_vals_12)
            })
        except:
            pass
    
    if len(p_vals_19) >= 3:
        try:
            rho19, p19 = spearmanr(p_vals_19, q19_vals)
            r19, p19_p = pearsonr(p_vals_19, q19_vals)
            verified_correlations.append({
                'parent_question': parent_q,
                'ovcharova_question': '19. Позволяет использовать профессиональные умения вне работы',
                'parent_values': p_vals_19,
                'ovcharova_values': q19_vals,
                'spearman_corr': rho19,
                'spearman_p': p19,
                'pearson_corr': r19,
                'pearson_p': p19_p,
                'n': len(p_vals_19)
            })
        except:
            pass

print(f"   Перепроверено корреляций: {len(verified_correlations)}")

# Находим вопросы Холланда
print("\n4. Поиск вопросов Холланда...")
holland_cols = []
for col in student_cols:
    col_str = str(col).strip()
    if 'Что тебе ближе' in col_str or col_str.startswith('1. Что тебе ближе') or \
       (col_str.startswith('1.') and 'ближе' in col_str.lower()):
        holland_cols.append(col)

print(f"   Найдено вопросов Холланда: {len(holland_cols)}")

# Рассчитываем баллы по типам Холланда
print("\n5. Расчёт баллов по типам Холланда...")
holland_scores = None

if len(holland_cols) > 0:
    holland_scores = calculate_holland_scores(students_sorted, holland_cols, pairs)
    for h_type in ['R', 'I', 'A', 'S', 'E', 'C']:
        scores = holland_scores[h_type]
        print(f"   Тип {h_type}: средний балл = {np.mean(scores):.2f}, мин = {min(scores)}, макс = {max(scores)}")
else:
    print("   ⚠ Столбцы Холланда не найдены")

# Рассчитываем корреляции с Холландом
print("\n6. Расчёт корреляций с типами Холланда...")
holland_correlations = []
ovcharova_holland_correlations = []

if holland_scores is not None:
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный (Realistic)',
        'I': 'Исследовательский (Investigative)',
        'A': 'Артистический (Artistic)',
        'S': 'Социальный (Social)',
        'E': 'Предприимчивый (Enterprising)',
        'C': 'Конвенциональный (Conventional)'
    }
    
    # Корреляции родители-Холланд
    for parent_q in parent_questions:
        for h_type in holland_types:
            p_vals = []
            h_vals = []
            
            for i, pair in enumerate(pairs):
                try:
                    p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                    h_score = holland_scores[h_type][i]
                    
                    p_num = float(p_val) if pd.notna(p_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if p_num is not None and h_num is not None:
                        p_vals.append(p_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(p_vals) >= 3:
                try:
                    rho, p = spearmanr(p_vals, h_vals)
                    holland_correlations.append({
                        'parent_question': parent_q,
                        'holland_type': h_type,
                        'holland_type_name': holland_type_names[h_type],
                        'spearman_corr': rho,
                        'spearman_p': p,
                        'n': len(p_vals),
                        'parent_values': p_vals,
                        'holland_values': h_vals
                    })
                except:
                    pass
    
    # Корреляции Овчарова-Холланд
    for ovcharova_q_col in [q12_col, q19_col]:
        for h_type in holland_types:
            o_vals = []
            h_vals = []
            
            for i, pair in enumerate(pairs):
                try:
                    o_val = students_sorted.iloc[pair['student_idx']][ovcharova_q_col]
                    h_score = holland_scores[h_type][i]
                    
                    o_num = float(o_val) if pd.notna(o_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if o_num is not None and h_num is not None:
                        o_vals.append(o_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(o_vals) >= 3:
                try:
                    rho, p = spearmanr(o_vals, h_vals)
                    ovcharova_holland_correlations.append({
                        'ovcharova_question': ovcharova_q_col,
                        'holland_type': h_type,
                        'holland_type_name': holland_type_names[h_type],
                        'spearman_corr': rho,
                        'spearman_p': p,
                        'n': len(o_vals),
                        'ovcharova_values': o_vals,
                        'holland_values': h_vals
                    })
                except:
                    pass

print(f"   Корреляций родители-Холланд: {len(holland_correlations)}")
print(f"   Корреляций Овчарова-Холланд: {len(ovcharova_holland_correlations)}")

# Фильтруем значимые корреляции
significant_ovcharova_markovskaya = [c for c in verified_correlations if c['spearman_p'] < 0.05]
significant_holland = [c for c in holland_correlations if c['spearman_p'] < 0.05]
significant_ovcharova_holland = [c for c in ovcharova_holland_correlations if c['spearman_p'] < 0.05]

print(f"\n7. Статистика значимых корреляций:")
print(f"   Овчарова-Марковская: {len(significant_ovcharova_markovskaya)}")
print(f"   Марковская-Холланд: {len(significant_holland)}")
print(f"   Овчарова-Холланд: {len(significant_ovcharova_holland)}")

# Создаём папку для графиков
graphs_dir = os.path.join(script_dir, "graphs")
os.makedirs(graphs_dir, exist_ok=True)

# Создаём визуализации
print("\n8. Создание визуализаций...")

# 1. Столбчатая диаграмма по типам Холланда с error bars
if holland_scores is not None:
    create_holland_bar_chart_with_error_bars(
        holland_scores, 
        os.path.join(graphs_dir, "holland_types_bar_chart.png")
    )

# 2. Столбчатая диаграмма всех корреляций Овчарова-Холланд (значимых и незначимых)
if len(ovcharova_holland_correlations) > 0:
    create_ovcharova_holland_bar_chart(
        ovcharova_holland_correlations,  # Используем ВСЕ корреляции, а не только значимые
        os.path.join(graphs_dir, "ovcharova_holland_bar_chart.png")
    )

# 3. Тепловая карта значимых корреляций Марковская-Холланд
if len(significant_holland) > 0:
    create_correlation_heatmap(
        significant_holland,
        "Значимые корреляции между установками Марковской и типами Холланда",
        os.path.join(graphs_dir, "markovskaya_holland_heatmap.png"),
        "Марковская-Холланд"
    )

# 4. Столбчатая диаграмма по мотивам
ovcharova_data = {}
for q_num, col in ovcharova_all.items():
    values = []
    for pair in pairs:
        try:
            val = students_sorted.iloc[pair['student_idx']][col]
            num_val = float(val) if pd.notna(val) else None
            if num_val is not None:
                values.append(num_val)
        except:
            pass
    if len(values) > 0:
        ovcharova_data[f"Вопрос {q_num}"] = values

if len(ovcharova_data) > 0:
    create_motives_bar_chart(
        ovcharova_data,
        os.path.join(graphs_dir, "motives_bar_chart.png"),
        use_median=True  # Используем медиану, как рекомендовано
    )

# 5. Круговая диаграмма корреляций Овчарова-Марковская
if len(significant_ovcharova_markovskaya) > 0:
    create_ovcharova_markovskaya_pie_chart(
        verified_correlations,  # Функция сама отфильтрует значимые
        os.path.join(graphs_dir, "ovcharova_markovskaya_pie_chart.png")
    )

# 6. Столбчатая диаграмма значимых корреляций Марковская-Холланд
if len(significant_holland) > 0:
    create_markovskaya_holland_bar_chart(
        significant_holland,
        os.path.join(graphs_dir, "markovskaya_holland_bar_chart.png")
    )

# 7. Сравнительная столбчатая таблица по топ-связям Марковская-Холланд
if len(significant_holland) > 0:
    create_comparative_holland_table(
        significant_holland,
        os.path.join(graphs_dir, "comparative_holland_table.png")
    )

# 8. Рисунок 1: Тепловая карта Овчарова-Марковская (значимые)
if len(significant_ovcharova_markovskaya) > 0:
    create_figure1_ovcharova_markovskaya_heatmap(
        verified_correlations,  # Передаём все корреляции, функция сама отфильтрует значимые
        os.path.join(graphs_dir, "figure1_ovcharova_markovskaya_heatmap.png")
    )

# 9. Рисунок 2: Тепловая карта Марковская-Холланд (значимые)
if len(significant_holland) > 0:
    create_figure2_markovskaya_holland_heatmap(
        holland_correlations,  # Передаём все корреляции, функция сама отфильтрует значимые
        os.path.join(graphs_dir, "figure2_markovskaya_holland_heatmap.png")
    )

# 10. Рисунок 3: Scatter plot Овчарова-Холланд (значимые)
if len(significant_ovcharova_holland) > 0:
    create_figure3_ovcharova_holland_scatter(
        ovcharova_holland_correlations,  # Передаём все корреляции, функция сама отфильтрует значимые
        os.path.join(graphs_dir, "figure3_ovcharova_holland_scatter.png")
    )

# 11. Рисунок 4: Рейтинг всех значимых связей
all_significant = (significant_ovcharova_markovskaya + 
                   significant_holland + 
                   significant_ovcharova_holland)
if len(all_significant) > 0:
    create_figure4_rating_all_significant(
        all_significant,
        os.path.join(graphs_dir, "figure4_rating_all_significant.png")
    )

# Создаём документ Word с подробными расчётами
print("\n9. Создание документа с подробными расчётами...")
doc = Document()

def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Заголовок
title = doc.add_heading('ПОДРОБНЫЙ АНАЛИЗ КОРРЕЛЯЦИЙ С ДЕТАЛЬНЫМИ РАСЧЁТАМИ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Общая статистика корреляционного анализа
doc.add_heading('0. ОБЩАЯ СТАТИСТИКА КОРРЕЛЯЦИОННОГО АНАЛИЗА', 1)

# Создаём таблицу с общей статистикой
stats_table = doc.add_table(rows=1, cols=2)
stats_table.style = 'Light Grid Accent 1'
stats_table.autofit = True

# Заголовки
hdr_cells = stats_table.rows[0].cells
hdr_cells[0].text = 'Показатель'
hdr_cells[1].text = 'Значение'

# Данные о выборке
row = stats_table.add_row().cells
row[0].text = 'Размер выборки'
row[1].text = f'{len(pairs)} пар (родитель-ребёнок)'

row = stats_table.add_row().cells
row[0].text = 'Количество родителей'
row[1].text = str(len(parents_df))

row = stats_table.add_row().cells
row[0].text = 'Количество учеников'
row[1].text = str(len(students_df))

row = stats_table.add_row().cells
row[0].text = 'Количество вопросов родителей (Марковская)'
row[1].text = str(len(parent_questions))

row = stats_table.add_row().cells
row[0].text = 'Количество мотивов подростков (Овчарова)'
row[1].text = str(len(ovcharova_all))

row = stats_table.add_row().cells
row[0].text = 'Количество типов Холланда'
row[1].text = '6 (R, I, A, S, E, C)'

doc.add_paragraph()

# Статистика по корреляциям
doc.add_heading('0.1. Статистика рассчитанных корреляций', 2)

corr_stats_table = doc.add_table(rows=1, cols=3)
corr_stats_table.style = 'Light Grid Accent 1'
corr_stats_table.autofit = True

hdr_cells = corr_stats_table.rows[0].cells
hdr_cells[0].text = 'Тип корреляции'
hdr_cells[1].text = 'Всего рассчитано'
hdr_cells[2].text = 'Статистически значимых (p < 0.05)'

row = corr_stats_table.add_row().cells
row[0].text = 'Овчарова-Марковская'
row[1].text = str(len(verified_correlations))
row[2].text = f'{len(significant_ovcharova_markovskaya)} ({len(significant_ovcharova_markovskaya)/len(verified_correlations)*100:.1f}%)' if len(verified_correlations) > 0 else '0'

row = corr_stats_table.add_row().cells
row[0].text = 'Марковская-Холланд'
row[1].text = str(len(holland_correlations))
row[2].text = f'{len(significant_holland)} ({len(significant_holland)/len(holland_correlations)*100:.1f}%)' if len(holland_correlations) > 0 else '0'

row = corr_stats_table.add_row().cells
row[0].text = 'Овчарова-Холланд'
row[1].text = str(len(ovcharova_holland_correlations))
row[2].text = f'{len(significant_ovcharova_holland)} ({len(significant_ovcharova_holland)/len(ovcharova_holland_correlations)*100:.1f}%)' if len(ovcharova_holland_correlations) > 0 else '0'

row = corr_stats_table.add_row().cells
row[0].text = 'ВСЕГО'
total_correlations = len(verified_correlations) + len(holland_correlations) + len(ovcharova_holland_correlations)
total_significant = len(significant_ovcharova_markovskaya) + len(significant_holland) + len(significant_ovcharova_holland)
row[1].text = str(total_correlations)
row[2].text = f'{total_significant} ({total_significant/total_correlations*100:.1f}%)' if total_correlations > 0 else '0'

doc.add_paragraph()

# Дополнительная информация
summary_text = f"""
ИТОГОВАЯ СВОДКА:

В ходе корреляционного анализа было рассчитано {total_correlations} корреляций между различными переменными исследования.

Из них статистически значимыми (p < 0.05) оказались {total_significant} корреляций, что составляет {total_significant/total_correlations*100:.1f}% от общего числа рассчитанных корреляций.

Распределение значимых корреляций:
• Овчарова-Марковская: {len(significant_ovcharova_markovskaya)} значимых корреляций из {len(verified_correlations)} рассчитанных
• Марковская-Холланд: {len(significant_holland)} значимых корреляций из {len(holland_correlations)} рассчитанных
• Овчарова-Холланд: {len(significant_ovcharova_holland)} значимых корреляций из {len(ovcharova_holland_correlations)} рассчитанных

Наибольшее количество значимых связей обнаружено между установками родителей (Марковская) и типами профессиональных интересов подростков (Холланд) - {len(significant_holland)} значимых корреляций.
"""
doc.add_paragraph(summary_text)

doc.add_paragraph()

# Объяснение уровня значимости
doc.add_heading('1. ОБЪЯСНЕНИЕ УРОВНЯ ЗНАЧИМОСТИ', 1)
significance_explanation = """
Уровень значимости (p-value) - это очень важный показатель в статистике, который показывает, насколько мы можем быть уверены в том, что найденная связь между переменными реальна, а не случайна.

Простыми словами:
• Если p-value меньше 0.05 (p < 0.05), это означает, что вероятность того, что мы ошибочно считаем связь реальной, составляет менее 5%. 
  В этом случае мы говорим, что связь "статистически значима" и можем быть уверены, что она реальна.

• Если p-value больше 0.05 (p > 0.05), это означает, что вероятность ошибки слишком высока (более 5%), 
  и мы не можем уверенно сказать, что связь между переменными реальна. Возможно, это просто случайное совпадение.

Чем меньше p-value, тем увереннее мы можем сказать, что связь реальна:
• p < 0.001 (***) - очень высокая значимость, вероятность ошибки менее 0.1%
• p < 0.01 (**) - высокая значимость, вероятность ошибки менее 1%
• p < 0.05 (*) - значимость, вероятность ошибки менее 5%
• p ≥ 0.05 - незначимо, связь может быть случайной

В нашем исследовании мы считаем связь значимой, если p < 0.05.
"""
doc.add_paragraph(significance_explanation)

# Объяснение статистических показателей
doc.add_heading('2. ОБЪЯСНЕНИЕ СТАТИСТИЧЕСКИХ ПОКАЗАТЕЛЕЙ', 1)
statistics_explanation = """
В таблицах используются следующие статистические показатели:

• M (mean) - среднее арифметическое. Это сумма всех значений, делённая на их количество. 
  Показывает "типичное" значение в выборке.

• SD (standard deviation) - стандартное отклонение. Показывает, насколько сильно значения 
  разбросаны вокруг среднего. Чем больше SD, тем больше разброс данных.

• Me (median) - медиана. Это значение, которое находится ровно посередине упорядоченного ряда. 
  Медиана более устойчива к выбросам (экстремальным значениям), чем среднее.

• % (проценты) - доля респондентов или ответов в процентах. Показывает, какой процент 
  респондентов выбрал то или иное значение.
"""
doc.add_paragraph(statistics_explanation)

# 2.1. Использование коэффициента Пирсона в исследовании (ответ комиссии)
doc.add_heading('2.1. Использование коэффициента Пирсона в исследовании', 2)
pearson_explanation = """
Можно ли обойтись без коэффициента Пирсона? Да. Для выводов исследования и оценки статистической значимости достаточен коэффициент корреляции Спирмена (ρ).

Зачем тогда использовался коэффициент Пирсона? Он применялся как дополнительный (вспомогательный) метод:
• для проверки согласованности результатов: при близких значениях r (Пирсон) и ρ (Спирмен) связь близка к линейной;
• при заметном расхождении r и ρ связь является монотонной, но не обязательно линейной.

Основным критерием для всех выводов и интерпретаций в работе был коэффициент Спирмена, так как данные имеют порядковый характер (шкалы оценок), а не интервальный.

Где именно рассчитывался коэффициент Пирсона? Только в одном блоке анализа — корреляции между установками родителей (ВРР Марковской) и мотивами выбора профессии (опросник Овчаровой). Конкретно для двух мотивов: вопрос 12 («Дает возможности для роста профессионального мастерства») и вопрос 19 («Позволяет использовать профессиональные умения вне работы»). Для каждой пары «вопрос родителя × мотив 12 или 19» при n ≥ 3 рассчитаны и Спирмен, и Пирсон. Для корреляций с типами Холланда (Марковская–Холланд, Овчарова–Холланд) коэффициент Пирсона не рассчитывался.
"""
doc.add_paragraph(pearson_explanation)

# Таблица всех мест, где рассчитан Пирсон (все verified_correlations)
doc.add_paragraph("Ниже приведены все пары, для которых рассчитан коэффициент Пирсона, с значениями ρ (Спирмен) и r (Пирсон):")
pearson_table = doc.add_table(rows=1, cols=7)
pearson_table.style = 'Light Grid Accent 1'
pearson_hdr = pearson_table.rows[0].cells
pearson_hdr[0].text = 'Вопрос родителя (Марковская)'
pearson_hdr[1].text = 'Мотив (Овчарова)'
pearson_hdr[2].text = 'N'
pearson_hdr[3].text = 'ρ (Спирмен)'
pearson_hdr[4].text = 'p (Спирмен)'
pearson_hdr[5].text = 'r (Пирсон)'
pearson_hdr[6].text = 'p (Пирсон)'
for c in verified_correlations:
    row_cells = pearson_table.add_row().cells
    row_cells[0].text = (c['parent_question'][:80] + '...') if len(str(c['parent_question'])) > 80 else str(c['parent_question'])
    row_cells[1].text = str(c['ovcharova_question'])[:50]
    row_cells[2].text = str(c['n'])
    row_cells[3].text = f"{c['spearman_corr']:.4f}"
    row_cells[4].text = f"{c['spearman_p']:.4f}"
    row_cells[5].text = f"{c['pearson_corr']:.4f}"
    row_cells[6].text = f"{c['pearson_p']:.4f}"
doc.add_paragraph()

# Таблицы для значимых связей
doc.add_heading('3. ТАБЛИЦЫ ДЛЯ ЗНАЧИМЫХ СВЯЗЕЙ', 1)

# 3.1. Корреляции Овчарова-Марковская
if len(significant_ovcharova_markovskaya) > 0:
    doc.add_heading('3.1. Корреляции между мотивами Овчаровой и установками Марковской', 2)
    
    for idx, corr in enumerate(sorted(significant_ovcharova_markovskaya, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True), 1):
        doc.add_heading(f'Связь {idx}', 3)
        
        # Статистика для родительских установок
        parent_stats = calculate_statistics(corr['parent_values'])
        # Статистика для мотивов
        ovcharova_stats = calculate_statistics(corr['ovcharova_values'])
        
        # Уровень значимости
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        
        # Создаём таблицу
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Родительские установки (Марковская)'
        hdr_cells[2].text = 'Мотивы выбора профессии (Овчарова)'
        hdr_cells[3].text = 'Корреляция'
        hdr_cells[4].text = 'Уровень значимости'
        hdr_cells[5].text = 'N'
        
        # Данные
        row = table.add_row().cells
        row[0].text = 'M (среднее)'
        row[1].text = f"{parent_stats['M']:.2f}" if parent_stats['M'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['M']:.2f}" if ovcharova_stats['M'] is not None else "N/A"
        row[3].text = f"{corr['spearman_corr']:.4f}"
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
        row[5].text = str(corr['n'])
        
        row = table.add_row().cells
        row[0].text = 'SD (стандартное отклонение)'
        row[1].text = f"{parent_stats['SD']:.2f}" if parent_stats['SD'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['SD']:.2f}" if ovcharova_stats['SD'] is not None else "N/A"
        row[3].text = ""
        row[4].text = sig_explanation
        row[5].text = ""
        
        row = table.add_row().cells
        row[0].text = 'Me (медиана)'
        row[1].text = f"{parent_stats['Me']:.2f}" if parent_stats['Me'] is not None else "N/A"
        row[2].text = f"{ovcharova_stats['Me']:.2f}" if ovcharova_stats['Me'] is not None else "N/A"
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""

        # Коэффициент Пирсона (рассчитывался только для пар Овчарова–Марковская)
        if 'pearson_corr' in corr and 'pearson_p' in corr:
            row = table.add_row().cells
            row[0].text = 'r (Пирсон)'
            row[1].text = ""
            row[2].text = ""
            row[3].text = f"{corr['pearson_corr']:.4f}"
            row[4].text = f"p={corr['pearson_p']:.4f}"
            row[5].text = ""

        # Проценты
        if parent_stats['percentages']:
            row = table.add_row().cells
            row[0].text = '% (проценты)'
            parent_pct = ', '.join([f"{val}: {pct:.1f}%" for val, pct in sorted(parent_stats['percentages'].items())])
            row[1].text = parent_pct[:100] + "..." if len(parent_pct) > 100 else parent_pct
            if ovcharova_stats['percentages']:
                ovcharova_pct = ', '.join([f"{val}: {pct:.1f}%" for val, pct in sorted(ovcharova_stats['percentages'].items())])
                row[2].text = ovcharova_pct[:100] + "..." if len(ovcharova_pct) > 100 else ovcharova_pct
            else:
                row[2].text = "N/A"
            row[3].text = ""
            row[4].text = ""
            row[5].text = ""
        
        # Описание связи
        doc.add_paragraph(f"\nВопрос родителя: {corr['parent_question'][:100]}...")
        doc.add_paragraph(f"Мотив подростка: {corr['ovcharova_question']}")
        doc.add_paragraph()

# 3.2. Корреляции Марковская-Холланд
if len(significant_holland) > 0:
    doc.add_heading('3.2. Корреляции между установками Марковской и типами Холланда', 2)
    
    for idx, corr in enumerate(sorted(significant_holland, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True)[:10], 1):
        doc.add_heading(f'Связь {idx}', 3)
        
        parent_stats = calculate_statistics(corr.get('parent_values', []))
        holland_stats = calculate_statistics(corr.get('holland_values', []))
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Родительские установки (Марковская)'
        hdr_cells[2].text = 'Тип Холланда'
        hdr_cells[3].text = 'Корреляция'
        hdr_cells[4].text = 'Уровень значимости'
        hdr_cells[5].text = 'N'
        
        row = table.add_row().cells
        row[0].text = 'M (среднее)'
        row[1].text = f"{parent_stats['M']:.2f}" if parent_stats['M'] is not None else "N/A"
        row[2].text = f"{holland_stats['M']:.2f}" if holland_stats['M'] is not None else "N/A"
        row[3].text = f"{corr['spearman_corr']:.4f}"
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
        row[5].text = str(corr['n'])
        
        row = table.add_row().cells
        row[0].text = 'SD (стандартное отклонение)'
        row[1].text = f"{parent_stats['SD']:.2f}" if parent_stats['SD'] is not None else "N/A"
        row[2].text = f"{holland_stats['SD']:.2f}" if holland_stats['SD'] is not None else "N/A"
        row[3].text = ""
        row[4].text = sig_explanation
        row[5].text = ""
        
        row = table.add_row().cells
        row[0].text = 'Me (медиана)'
        row[1].text = f"{parent_stats['Me']:.2f}" if parent_stats['Me'] is not None else "N/A"
        row[2].text = f"{holland_stats['Me']:.2f}" if holland_stats['Me'] is not None else "N/A"
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        
        doc.add_paragraph(f"\nВопрос родителя: {corr['parent_question'][:100]}...")
        doc.add_paragraph(f"Тип Холланда: {corr['holland_type_name']}")
        doc.add_paragraph()

# 3.3. Корреляции Овчарова-Холланд
if len(significant_ovcharova_holland) > 0:
    doc.add_heading('3.3. Корреляции между мотивами Овчаровой и типами Холланда', 2)
    
    for idx, corr in enumerate(sorted(significant_ovcharova_holland, 
                                     key=lambda x: abs(x['spearman_corr']), reverse=True)[:10], 1):
        doc.add_heading(f'Связь {idx}', 3)
        
        ovcharova_stats = calculate_statistics(corr.get('ovcharova_values', []))
        holland_stats = calculate_statistics(corr.get('holland_values', []))
        sig_symbol, sig_explanation = explain_significance_level(corr['spearman_p'])
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Мотивы выбора профессии (Овчарова)'
        hdr_cells[2].text = 'Тип Холланда'
        hdr_cells[3].text = 'Корреляция'
        hdr_cells[4].text = 'Уровень значимости'
        hdr_cells[5].text = 'N'
        
        row = table.add_row().cells
        row[0].text = 'M (среднее)'
        row[1].text = f"{ovcharova_stats['M']:.2f}" if ovcharova_stats['M'] is not None else "N/A"
        row[2].text = f"{holland_stats['M']:.2f}" if holland_stats['M'] is not None else "N/A"
        row[3].text = f"{corr['spearman_corr']:.4f}"
        row[4].text = f"{sig_symbol} p={corr['spearman_p']:.4f}"
        row[5].text = str(corr['n'])
        
        row = table.add_row().cells
        row[0].text = 'SD (стандартное отклонение)'
        row[1].text = f"{ovcharova_stats['SD']:.2f}" if ovcharova_stats['SD'] is not None else "N/A"
        row[2].text = f"{holland_stats['SD']:.2f}" if holland_stats['SD'] is not None else "N/A"
        row[3].text = ""
        row[4].text = sig_explanation
        row[5].text = ""
        
        row = table.add_row().cells
        row[0].text = 'Me (медиана)'
        row[1].text = f"{ovcharova_stats['Me']:.2f}" if ovcharova_stats['Me'] is not None else "N/A"
        row[2].text = f"{holland_stats['Me']:.2f}" if holland_stats['Me'] is not None else "N/A"
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        
        doc.add_paragraph(f"\nМотив подростка: {corr.get('ovcharova_question', 'N/A')}")
        doc.add_paragraph(f"Тип Холланда: {corr['holland_type_name']}")
        doc.add_paragraph()

# Раздел о статистическом анализе
doc.add_heading('4. СТАТИСТИЧЕСКИЙ АНАЛИЗ ДАННЫХ', 1)

statistical_analysis_text = """
В данном исследовании для анализа полученных данных использовались методы описательной и инференциальной статистики.

4.1. МЕТОДЫ ОПИСАТЕЛЬНОЙ СТАТИСТИКИ

Описательная статистика позволяет описать и обобщить основные характеристики данных без вывода о генеральной совокупности.

4.1.1. Среднее арифметическое (M, mean)

Среднее арифметическое - это сумма всех значений, делённая на их количество. Формула: M = (Σx) / n, где Σx - сумма всех значений, n - количество значений.

В нашем исследовании среднее арифметическое использовалось для:
• Описания типичных значений ответов родителей по вопросам ВРР Марковской
• Описания типичных значений мотивов выбора профессии у подростков (опросник Овчаровой)
• Описания типичных баллов по типам профессиональных интересов Холланда

Например, если среднее значение ответа родителя на вопрос равен 3.5, это означает, что "типичный" родитель в нашей выборке отвечает на этот вопрос примерно на 3.5 балла по 5-балльной шкале.

4.1.2. Проценты (%)

Проценты показывают долю респондентов, выбравших то или иное значение, в процентах от общего числа респондентов. Формула: % = (количество выбравших значение / общее количество) × 100%.

В нашем исследовании проценты использовались для:
• Описания распределения ответов - какой процент родителей выбрал каждое значение (1, 2, 3, 4, 5)
• Описания распределения мотивов - какой процент подростков оценил каждый мотив определённым баллом
• Понимания структуры данных - какие значения наиболее распространены

Например, если 40% родителей ответили "3" на вопрос, это означает, что 40% всех родителей в выборке выбрали именно это значение.

4.1.3. Стандартное отклонение (SD, standard deviation)

Стандартное отклонение показывает, насколько сильно значения разбросаны вокруг среднего. Формула: SD = √(Σ(x - M)² / (n - 1)), где x - каждое значение, M - среднее, n - количество значений.

Чем больше SD, тем больше разброс данных. Если SD маленькое, значит большинство значений близки к среднему. Если SD большое, значит значения сильно различаются.

В нашем исследовании SD использовалось для:
• Оценки однородности ответов - если SD маленькое, значит родители отвечают похоже
• Понимания вариативности данных - насколько различаются ответы респондентов
• Построения error bars на графиках - показывают разброс данных вокруг среднего

4.1.4. Медиана (Me, median)

Медиана - это значение, которое находится ровно посередине упорядоченного ряда значений. Если значений нечётное количество, медиана - это среднее значение. Если чётное - среднее между двумя центральными значениями.

Медиана более устойчива к выбросам (экстремальным значениям), чем среднее арифметическое. Если в данных есть несколько очень больших или очень маленьких значений, медиана даст более "типичное" представление о данных.

В нашем исследовании медиана использовалась для:
• Описания мотивов выбора профессии - показывает "типичный" мотив, не искажённый экстремальными значениями
• Сравнения с средним арифметическим - если медиана сильно отличается от среднего, значит есть выбросы

4.2. МЕТОДЫ ИНФЕРЕНЦИАЛЬНОЙ СТАТИСТИКИ

Инференциальная статистика позволяет делать выводы о генеральной совокупности на основе выборки.

4.2.1. Корреляционный анализ

Корреляционный анализ позволяет определить, существует ли связь между двумя переменными и насколько она сильна.

В нашем исследовании использовался коэффициент корреляции Спирмена (ρ, rho), который:
• Подходит для порядковых данных (как наши шкалы оценок)
• Показывает силу и направление связи (от -1 до +1)
• Не требует нормального распределения данных

Интерпретация коэффициента корреляции Спирмена:
• ρ от 0.7 до 1.0 (или от -0.7 до -1.0) - сильная связь
• ρ от 0.3 до 0.7 (или от -0.3 до -0.7) - умеренная связь
• ρ от 0 до 0.3 (или от 0 до -0.3) - слабая связь
• ρ = 0 - отсутствие связи

Положительная корреляция означает, что при увеличении одной переменной увеличивается и другая. Отрицательная - при увеличении одной переменной уменьшается другая.

4.2.2. Уровень значимости (p-value)

Уровень значимости показывает вероятность того, что мы ошибочно считаем связь реальной, когда на самом деле её нет (ошибка первого рода).

В нашем исследовании использовался стандартный уровень значимости p < 0.05, что означает:
• Вероятность ошибки менее 5%
• Мы можем быть уверены, что связь реальна, а не случайна
• Результат статистически значим

4.3. ПОЧЕМУ ИМЕННО ЭТИ МЕТОДЫ?

Мы выбрали именно эти методы потому что:

1. Наши данные порядковые (шкалы оценок), а не интервальные - поэтому используем медиану и корреляцию Спирмена
2. Нужно описать типичные значения - используем среднее арифметическое и медиану
3. Нужно понять разброс данных - используем стандартное отклонение
4. Нужно понять распределение ответов - используем проценты
5. Нужно проверить наличие связи - используем корреляционный анализ
6. Нужно убедиться, что связь реальна - используем уровень значимости

Эти методы позволяют нам:
• Описать характеристики нашей выборки
• Найти связи между переменными
• Убедиться, что связи статистически значимы
• Сделать выводы о влиянии установок родителей на мотивы выбора профессии подростками
"""
doc.add_paragraph(statistical_analysis_text)

# Раздел результатов корреляционного анализа с основными графиками
doc.add_heading('5. РЕЗУЛЬТАТЫ КОРРЕЛЯЦИОННОГО АНАЛИЗА', 1)

# Рисунок 1: Овчарова-Марковская
if len(significant_ovcharova_markovskaya) > 0:
    doc.add_heading('5.1. Результаты корреляционного анализа между ВРР и мотивами', 2)
    description_1 = f"""
В ходе корреляционного анализа между установками родителей (ВРР Марковской) и мотивами выбора профессии подростков (опросник Овчаровой) было обнаружено {len(significant_ovcharova_markovskaya)} статистически значимых связей (p < 0.05).

Тепловая карта (Рисунок 1) представляет все значимые корреляции между пунктами ВРР Марковской (строки) и двумя мотивами выбора профессии по Овчаровой - мотивом 12 (возможности для профессионального роста) и мотивом 19 (использование профессиональных умений вне работы) (столбцы).

В ячейках тепловой карты отражены:
• Коэффициент корреляции Спирмена (ρ) со знаком (положительный или отрицательный)
• Пометка * для всех значимых корреляций (p < 0.05)

Цветовая гамма тепловой карты позволяет быстро оценить силу и направление связи:
• Синие оттенки - положительные корреляции (чем синее, тем сильнее связь)
• Красные оттенки - отрицательные корреляции (чем краснее, тем сильнее связь)

Тепловая карта позволяет увидеть, какие именно установки родителей наиболее связаны с мотивами профессионального роста и использования умений вне работы у подростков.
"""
    doc.add_paragraph(description_1)
    figure1_path = os.path.join(graphs_dir, "figure1_ovcharova_markovskaya_heatmap.png")
    if os.path.exists(figure1_path):
        doc.add_picture(figure1_path, width=Inches(6))

# Рисунок 2: Марковская-Холланд
if len(significant_holland) > 0:
    doc.add_heading('5.2. Результаты корреляционного анализа между ВРР и профтипами', 2)
    description_2 = f"""
В ходе корреляционного анализа между установками родителей (ВРР Марковской) и типами профессиональных интересов подростков (опросник Холланда) было обнаружено {len(significant_holland)} статистически значимых связей (p < 0.05).

Тепловая карта (Рисунок 2) представляет все значимые корреляции между пунктами ВРР Марковской (строки) и шестью типами профессиональных интересов по Холланду (столбцы):
• R - Реалистичный
• I - Исследовательский
• A - Артистический
• S - Социальный
• E - Предприимчивый
• C - Конвенциональный

В ячейках тепловой карты отражены:
• Коэффициент корреляции Спирмена (ρ) со знаком
• Пометка * для всех значимых корреляций (p < 0.05)
• Символ "—" в пустых ячейках, где значимых связей не обнаружено

Тепловая карта позволяет увидеть структуру связей между установками родителей и профессиональными интересами подростков, выявить, какие типы интересов наиболее связаны с определёнными родительскими установками.
"""
    doc.add_paragraph(description_2)
    figure2_path = os.path.join(graphs_dir, "figure2_markovskaya_holland_heatmap.png")
    if os.path.exists(figure2_path):
        doc.add_picture(figure2_path, width=Inches(6))

# Рисунок 3: Овчарова-Холланд
if len(significant_ovcharova_holland) > 0:
    doc.add_heading('5.3. Связь мотива и профтипа', 2)
    description_3 = f"""
В ходе корреляционного анализа между мотивами выбора профессии подростков (опросник Овчаровой) и типами профессиональных интересов (опросник Холланда) было обнаружено {len(significant_ovcharova_holland)} статистически значимая связь (p < 0.05).

Точечная диаграмма (Рисунок 3) показывает характер связи между баллом по мотиву выбора профессии (ось X) и баллом по типу профессионального интереса (ось Y).

На графике отражены:
• Каждая точка представляет одного респондента (подростка)
• Пунктирная линия - линия тренда, показывающая общее направление связи
• Подпись на графике содержит: ρ (коэффициент корреляции Спирмена), p (уровень значимости) и символы значимости (*, **, ***)

Точечная диаграмма позволяет визуально оценить характер связи между мотивом и типом интереса, увидеть, как изменяется один показатель при изменении другого.
"""
    doc.add_paragraph(description_3)
    figure3_path = os.path.join(graphs_dir, "figure3_ovcharova_holland_scatter.png")
    if os.path.exists(figure3_path):
        doc.add_picture(figure3_path, width=Inches(6))

# Рисунок 4: Рейтинг всех значимых связей
all_significant_count = len(significant_ovcharova_markovskaya) + len(significant_holland) + len(significant_ovcharova_holland)
if all_significant_count > 0:
    doc.add_heading('5.4. Сводная диаграмма по всем значимым корреляциям', 2)
    description_4 = f"""
Для получения целостной картины всех обнаруженных значимых связей была создана сводная диаграмма (Рисунок 4), представляющая рейтинг всех {all_significant_count} значимых корреляций, отсортированных по модулю коэффициента корреляции от сильных к слабым.

Горизонтальная столбчатая диаграмма (Рисунок 4) показывает:
• По оси Y: все значимые пары переменных (все три типа связей: Овчарова-Марковская, Марковская-Холланд, Овчарова-Холланд)
• По оси X: коэффициент корреляции Спирмена (ρ) с нулевой линией по центру
• Цветовое кодирование: синий - положительные корреляции, красный - отрицательные
• Значения корреляций и пометки значимости (*) на каждом столбце

Диаграмма отсортирована по модулю коэффициента корреляции (|ρ|) от сильных связей к слабым, что позволяет:
• Сразу увидеть наиболее сильные связи в исследовании
• Сравнить силу связей между разными типами переменных
• Оценить общую картину значимых связей в исследовании
• Понять направление влияния (положительное или отрицательное) для каждой связи

Эта диаграмма служит как "главная картинка" исследования, показывающая, что сильнее всего связано и в какую сторону.
"""
    doc.add_paragraph(description_4)
    figure4_path = os.path.join(graphs_dir, "figure4_rating_all_significant.png")
    if os.path.exists(figure4_path):
        doc.add_picture(figure4_path, width=Inches(6))

# Добавляем графики в документ
doc.add_heading('6. ДОПОЛНИТЕЛЬНЫЕ ВИЗУАЛИЗАЦИИ', 1)

if holland_scores is not None:
    doc.add_heading('5.1. Столбчатая диаграмма по типам Холланда с error bars (SD)', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
На графике представлены средние баллы (M) по каждому из шести типов профессиональных интересов по опроснику Холланда: Реалистичный (R), Исследовательский (I), Артистический (A), Социальный (S), Предприимчивый (E), Конвенциональный (C).

Как читать график:
• По оси X расположены шесть типов профессиональных интересов Холланда
• По оси Y отложены средние баллы (M) с указанием стандартного отклонения (SD)
• Каждый столбец показывает средний балл по соответствующему типу
• "Усы" (error bars) сверху и снизу от столбца показывают стандартное отклонение (SD)

Какие показатели отражены:
• M (среднее арифметическое) - типичный балл по каждому типу
• SD (стандартное отклонение) - разброс баллов вокруг среднего, показан в виде error bars

Почему отражены именно эти показатели:
• Среднее арифметическое (M) показывает типичный уровень интереса к каждому типу профессий
• Стандартное отклонение (SD) показывает, насколько однородны интересы подростков - маленькое SD означает, что большинство подростков имеют похожие интересы, большое SD - что интересы сильно различаются

Что даёт этот график:
• Позволяет увидеть, какие типы профессиональных интересов наиболее выражены у подростков
• Показывает, насколько различаются интересы подростков (по величине SD)
• Помогает понять структуру профессиональных интересов в выборке
• Позволяет сравнить выраженность разных типов интересов между собой
"""
    doc.add_paragraph(explanation)
    holland_img_path = os.path.join(graphs_dir, "holland_types_bar_chart.png")
    if os.path.exists(holland_img_path):
        doc.add_picture(holland_img_path, width=Inches(6))

if len(ovcharova_holland_correlations) > 0:
    doc.add_heading('5.2. Столбчатая диаграмма корреляций Овчарова-Холланд', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
График показывает все рассчитанные корреляции между мотивами выбора профессии (опросник Овчаровой, вопросы 12 и 19) и типами профессиональных интересов (опросник Холланда, типы R, I, A, S, E, C). Всего 12 корреляций (2 вопроса × 6 типов).

Как читать график:
• По оси X расположены шесть типов Холланда
• По оси Y отложены коэффициенты корреляции Спирмена (от -1 до +1)
• Два столбца для каждого типа - один для вопроса 12, другой для вопроса 19
• Значимые корреляции (p < 0.05) отмечены звёздочкой (*) и выделены цветом
• Незначимые корреляции показаны серым цветом

Какие показатели отражены:
• Коэффициент корреляции Спирмена (ρ) - сила и направление связи
• Уровень значимости - отмечен звёздочкой для значимых корреляций

Почему отражены именно эти показатели:
• Все корреляции (значимые и незначимые) показывают полную картину связей
• Значимые корреляции выделены, чтобы сразу видеть реальные связи
• Группировка по типам Холланда позволяет сравнить связи разных типов

Что даёт этот график:
• Позволяет увидеть все связи между мотивами и интересами
• Показывает, какие типы Холланда наиболее связаны с мотивами
• Помогает понять, какие мотивы важны для разных типов профессиональных интересов
• Позволяет сравнить силу связей между вопросами 12 и 19
"""
    doc.add_paragraph(explanation)
    bar_chart_path = os.path.join(graphs_dir, "ovcharova_holland_bar_chart.png")
    if os.path.exists(bar_chart_path):
        doc.add_picture(bar_chart_path, width=Inches(6))

if len(ovcharova_data) > 0:
    doc.add_heading('5.3. Столбчатая диаграмма по мотивам выбора профессии', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
График показывает выраженность различных мотивов выбора профессии у подростков по опроснику Овчаровой (20 вопросов). Используется медиана (Me) как показатель, устойчивый к выбросам.

Как читать график:
• По оси X расположены все 20 мотивов выбора профессии (вопросы опросника Овчаровой)
• По оси Y отложены медианные значения (Me) оценок важности каждого мотива
• Высота столбца показывает, насколько важен этот мотив для подростков

Какие показатели отражены:
• Me (медиана) - значение, которое находится посередине упорядоченного ряда оценок
• Медиана выбрана вместо среднего, так как она более устойчива к экстремальным значениям

Почему отражены именно эти показатели:
• Медиана (Me) показывает типичную важность мотива, не искажённую выбросами
• Если бы использовали среднее (M), один подросток с очень высокой оценкой мог бы исказить результат
• Медиана даёт более надёжное представление о том, что важно для большинства подростков

Что даёт этот график:
• Позволяет увидеть, какие мотивы наиболее важны для подростков
• Показывает иерархию мотивов - от самых важных до наименее важных
• Помогает понять приоритеты подростков при выборе профессии
• Позволяет сравнить выраженность разных мотивов
"""
    doc.add_paragraph(explanation)
    motives_path = os.path.join(graphs_dir, "motives_bar_chart.png")
    if os.path.exists(motives_path):
        doc.add_picture(motives_path, width=Inches(6))

if len(significant_ovcharova_markovskaya) > 0:
    doc.add_heading('5.4. Круговая диаграмма корреляций Овчарова-Марковская', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
Круговая диаграмма показывает распределение значимых корреляций между мотивами выбора профессии (Овчарова) и установками родителей (Марковская) по силе связи. Корреляции разделены на категории: сильные положительные (ρ > 0.5), умеренные положительные (0.3 < ρ ≤ 0.5), слабые положительные (0 < ρ ≤ 0.3), слабые отрицательные (-0.3 ≤ ρ < 0), умеренные отрицательные (-0.5 ≤ ρ < -0.3), сильные отрицательные (ρ < -0.5).

Как читать график:
• Каждый сегмент круга представляет одну категорию корреляций
• Размер сегмента пропорционален количеству корреляций в этой категории
• Процент показывает долю корреляций каждой категории от общего числа
• Цвета различают категории по силе и направлению связи

Какие показатели отражены:
• Количество корреляций в каждой категории
• Процент от общего числа значимых корреляций
• Сила связи (сильная, умеренная, слабая)
• Направление связи (положительная, отрицательная)

Почему отражены именно эти показатели:
• Группировка по силе связи позволяет понять структуру найденных связей
• Процент показывает, какая доля связей относится к каждой категории
• Разделение на положительные и отрицательные показывает направление влияния

Что даёт этот график:
• Позволяет быстро увидеть общую картину связей - сколько сильных, сколько слабых
• Показывает, преобладают ли положительные или отрицательные связи
• Помогает понять структуру влияния установок родителей на мотивы подростков
• Даёт наглядное представление о распределении силы связей
"""
    doc.add_paragraph(explanation)
    pie_chart_path = os.path.join(graphs_dir, "ovcharova_markovskaya_pie_chart.png")
    if os.path.exists(pie_chart_path):
        doc.add_picture(pie_chart_path, width=Inches(6))

if len(significant_holland) > 0:
    doc.add_heading('5.5. Столбчатая диаграмма значимых корреляций Марковская-Холланд', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
График показывает топ-20 самых сильных значимых корреляций между установками родителей (ВРР Марковской) и типами профессиональных интересов подростков (опросник Холланда). Корреляции отсортированы по силе связи (по абсолютному значению коэффициента корреляции).

Как читать график:
• По оси Y расположены вопросы родителей и типы Холланда (например, "R: Вопрос родителя...")
• По оси X отложены коэффициенты корреляции Спирмена
• Положительные корреляции показаны синим цветом (столбцы справа от нуля)
• Отрицательные корреляции показаны красным цветом (столбцы слева от нуля)
• Значимые корреляции отмечены звёздочкой (*)

Какие показатели отражены:
• Коэффициент корреляции Спирмена (ρ) - сила и направление связи
• Уровень значимости (p < 0.05) - отмечен звёздочкой
• Вопросы родителей (первые 30 символов)
• Типы Холланда (R, I, A, S, E, C)

Почему отражены именно эти показатели:
• Топ-20 корреляций показывают самые сильные и важные связи
• Горизонтальная ориентация позволяет удобно читать длинные названия вопросов
• Цветовое кодирование (синий/красный) сразу показывает направление связи
• Сортировка по силе позволяет увидеть наиболее важные связи первыми

Что даёт этот график:
• Позволяет увидеть самые сильные связи между установками родителей и интересами подростков
• Показывает, какие установки родителей наиболее связаны с какими типами интересов
• Помогает понять направление влияния (положительное или отрицательное)
• Даёт наглядное представление о структуре связей
"""
    doc.add_paragraph(explanation)
    markovskaya_holland_bar_path = os.path.join(graphs_dir, "markovskaya_holland_bar_chart.png")
    if os.path.exists(markovskaya_holland_bar_path):
        doc.add_picture(markovskaya_holland_bar_path, width=Inches(6))

if len(significant_holland) > 0:
    doc.add_heading('5.6. Сравнительная столбчатая таблица топ-связей Марковская-Холланд', 2)
    explanation = """
ПОЛНОЕ ОПИСАНИЕ ГРАФИКА:

Что отражено на графике:
Сравнительная столбчатая таблица показывает топ-3 самые сильные значимые корреляции между установками родителей (Марковская) и каждым из шести типов профессиональных интересов (Холланд). Для каждого типа Холланда показаны три столбца: 1-е место (самая сильная связь), 2-е место, 3-е место.

Как читать график:
• По оси X расположены шесть типов Холланда (R, I, A, S, E, C)
• По оси Y отложены коэффициенты корреляции Спирмена
• Для каждого типа показаны три столбца разного цвета:
  - Синий: 1-е место (самая сильная связь)
  - Зелёный: 2-е место
  - Светло-зелёный: 3-е место
• Высота столбца показывает силу корреляции

Какие показатели отражены:
• Коэффициент корреляции Спирмена (ρ) для каждой связи
• Ранжирование связей по силе (1-е, 2-е, 3-е место)
• Типы Холланда (R, I, A, S, E, C)

Почему отражены именно эти показатели:
• Топ-3 для каждого типа позволяет сравнить силу связей между типами
• Ранжирование показывает, какие установки родителей наиболее важны для каждого типа интересов
• Группировка по типам Холланда позволяет увидеть структуру влияния

Что даёт этот график:
• Позволяет сравнить силу связей между разными типами Холланда
• Показывает, какие установки родителей наиболее важны для каждого типа интересов
• Помогает понять, есть ли различия в силе влияния установок на разные типы интересов
• Даёт наглядное сравнение топ-связей для всех типов одновременно
"""
    doc.add_paragraph(explanation)
    comparative_table_path = os.path.join(graphs_dir, "comparative_holland_table.png")
    if os.path.exists(comparative_table_path):
        doc.add_picture(comparative_table_path, width=Inches(6))

# Сохраняем документ
output_file = os.path.join(script_dir, "Подробный_анализ_корреляций.docx")
doc.save(output_file)
print(f"\n10. Документ сохранён: {output_file}")

print("\n" + "=" * 80)
print("АНАЛИЗ ЗАВЕРШЁН!")
print("=" * 80)
