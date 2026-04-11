"""Скрипт для чтения Word и PDF документов"""
from docx import Document
import os

# Определяем пути
script_dir = os.path.dirname(os.path.abspath(__file__))
project_root = os.path.dirname(os.path.dirname(script_dir))

# Читаем Word документ
print("=" * 80)
print("ЧТЕНИЕ ДОКУМЕНТА: Отчет_Анализ_Корреляций.docx")
print("=" * 80)

doc_path = os.path.join(project_root, "newtest", "Отчет_Анализ_Корреляций.docx")
print(f"Путь к документу: {doc_path}")
doc = Document(doc_path)

print(f"\nВсего параграфов: {len(doc.paragraphs)}")
print("\nСодержимое документа:\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i+1}] {text[:150]}")

# Ищем таблицы
print("\n" + "=" * 80)
print("ТАБЛИЦЫ В ДОКУМЕНТЕ")
print("=" * 80)
print(f"Всего таблиц: {len(doc.tables)}")

for i, table in enumerate(doc.tables):
    print(f"\nТаблица {i+1}:")
    print(f"  Строк: {len(table.rows)}, Столбцов: {len(table.columns) if table.rows else 0}")
    if table.rows:
        # Первая строка (заголовки)
        if len(table.rows) > 0:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  Заголовки: {headers}")
        # Первые 3 строки данных
        for row_idx in range(min(3, len(table.rows))):
            row_data = [cell.text.strip()[:50] for cell in table.rows[row_idx].cells]
            print(f"  Строка {row_idx+1}: {row_data}")
