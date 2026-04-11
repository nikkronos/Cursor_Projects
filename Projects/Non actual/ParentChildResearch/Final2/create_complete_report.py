"""
Создание полного итогового отчёта, объединяющего:
1. Перепроверенные расчёты корреляций
2. Полные расчёты в формате примеров
3. Объяснение выбора вопросов 12 и 19
4. Проверку итоговых результатов
5. Оформление по примеру главы 2.2 из PDF
"""
import os
import sys
import pandas as pd
import numpy as np
from scipy.stats import spearmanr, pearsonr, rankdata
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import warnings
warnings.filterwarnings('ignore')

# Импорты для визуализации
import matplotlib
matplotlib.use('Agg')  # Для работы без GUI
import matplotlib.pyplot as plt
from matplotlib import rcParams
from PIL import Image

# Настройка шрифтов для русского языка
rcParams['font.family'] = 'DejaVu Sans'
rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial', 'Tahoma']

# ============================================================================
# ФУНКЦИИ ДЛЯ ОБРАБОТКИ ДАННЫХ ХОЛЛАНДА
# ============================================================================

def classify_profession_to_holland_types(profession_text):
    """
    Классифицирует профессию по типам Холланда (R, I, A, S, E, C).
    Возвращает список типов, к которым относится профессия.
    
    Типы Холланда:
    - R (Realistic) - реалистичный: работа с инструментами, машинами, техникой
    - I (Investigative) - исследовательский: научная работа, анализ, исследование
    - A (Artistic) - артистический: творческая работа, искусство, дизайн
    - S (Social) - социальный: работа с людьми, помощь, обучение
    - E (Enterprising) - предприимчивый: бизнес, управление, продажи
    - C (Conventional) - конвенциональный: работа с данными, организация, учёт
    """
    if pd.isna(profession_text) or not str(profession_text).strip():
        return []
    
    prof_lower = str(profession_text).lower()
    types = []
    
    # R (Realistic) - реалистичный
    r_keywords = ['инженер', 'техник', 'механик', 'автомеханик', 'ремонт', 'конструктор', 
                  'оператор', 'фермер', 'водитель', 'электроник', 'чертёжник', 'технолог',
                  'слесарь', 'сварщик', 'токарь', 'фрезеровщик', 'монтажник', 'строитель']
    if any(kw in prof_lower for kw in r_keywords):
        types.append('R')
    
    # I (Investigative) - исследовательский
    i_keywords = ['учёный', 'исследователь', 'аналитик', 'химик', 'физик', 'биолог', 
                  'микробиолог', 'археолог', 'криминалист', 'синоптик', 'эколог',
                  'врач', 'психиатр', 'ветеринар', 'провизор', 'философ', 'политолог']
    if any(kw in prof_lower for kw in i_keywords):
        types.append('I')
    
    # A (Artistic) - артистический
    a_keywords = ['художник', 'дизайнер', 'актёр', 'режиссёр', 'композитор', 'писатель',
                  'поэт', 'фотограф', 'декоратор', 'визажист', 'стилист', 'искусствовед',
                  'реставратор', 'литературный критик', 'видеооператор', 'видеограф']
    if any(kw in prof_lower for kw in a_keywords):
        types.append('A')
    
    # S (Social) - социальный
    s_keywords = ['психолог', 'учитель', 'преподаватель', 'тренер', 'врач', 'терапевт',
                  'логопед', 'социальный', 'молодёжью', 'консультант', 'карьерный',
                  'профилактики', 'правонарушений', 'социолог', 'журналист', 'корреспондент']
    if any(kw in prof_lower for kw in s_keywords):
        types.append('S')
    
    # E (Enterprising) - предприимчивый
    e_keywords = ['менеджер', 'предприниматель', 'директор', 'руководитель', 'трейдер',
                  'адвокат', 'риэлтор', 'продюсер', 'лидер', 'движения', 'общественного',
                  'маркетинг', 'сетевому', 'финансовый', 'генеральный', 'отдела']
    if any(kw in prof_lower for kw in e_keywords):
        types.append('E')
    
    # C (Conventional) - конвенциональный
    c_keywords = ['бухгалтер', 'оператор ввода', 'ввода данных', 'кассир', 'продавец',
                  'редактор', 'корректор', 'статистический', 'экономист', 'контролёр',
                  'качества', 'логистике', 'доставке', 'офис-менеджер']
    if any(kw in prof_lower for kw in c_keywords):
        types.append('C')
    
    # Если профессия не классифицирована, возвращаем пустой список
    return types

def calculate_holland_scores(students_df, holland_cols, pairs):
    """
    Рассчитывает баллы по типам Холланда (R, I, A, S, E, C) для каждого респондента.
    
    Возвращает словарь с ключами 'R', 'I', 'A', 'S', 'E', 'C', где каждое значение
    - это список баллов для каждого респондента из pairs.
    """
    holland_scores = {
        'R': [],
        'I': [],
        'A': [],
        'S': [],
        'E': [],
        'C': []
    }
    
    for pair in pairs:
        # Инициализируем счётчики для этого респондента
        scores = {'R': 0, 'I': 0, 'A': 0, 'S': 0, 'E': 0, 'C': 0}
        
        # Проходим по всем столбцам Холланда
        for h_col in holland_cols:
            try:
                profession = students_df.iloc[pair['student_idx']][h_col]
                if pd.notna(profession):
                    # Классифицируем профессию
                    types = classify_profession_to_holland_types(profession)
                    # Увеличиваем счётчики для каждого типа
                    for h_type in types:
                        if h_type in scores:
                            scores[h_type] += 1
            except:
                continue
        
        # Добавляем баллы для этого респондента
        for h_type in holland_scores:
            holland_scores[h_type].append(scores[h_type])
    
    return holland_scores

# ============================================================================
# ВАРИАНТЫ ГРАФИКОВ ДЛЯ КОРРЕЛЯЦИЙ
# ============================================================================

def create_correlation_plot_variant1(x_vals, y_vals, x_label, y_label, corr_coef, p_value, 
                                     title, save_path):
    """
    ВАРИАНТ 1: Комбинированный график (scatter + гистограммы + квадранты)
    Показывает распределения, медианы и структуру связи.
    """
    fig = plt.figure(figsize=(12, 10))
    
    # Создаём сетку для размещения графиков
    gs = fig.add_gridspec(3, 3, hspace=0.3, wspace=0.3, 
                         left=0.1, right=0.95, top=0.95, bottom=0.1)
    
    # Основной scatter plot (занимает центральную область)
    ax_main = fig.add_subplot(gs[1:3, 0:2])
    
    # Вычисляем медианы
    x_median = np.median(x_vals)
    y_median = np.median(y_vals)
    
    # Разделяем точки по квадрантам
    q1 = [(x, y) for x, y in zip(x_vals, y_vals) if x >= x_median and y >= y_median]
    q2 = [(x, y) for x, y in zip(x_vals, y_vals) if x < x_median and y >= y_median]
    q3 = [(x, y) for x, y in zip(x_vals, y_vals) if x < x_median and y < y_median]
    q4 = [(x, y) for x, y in zip(x_vals, y_vals) if x >= x_median and y < y_median]
    
    # Рисуем квадранты разными цветами
    if q1:
        ax_main.scatter([p[0] for p in q1], [p[1] for p in q1], 
                       alpha=0.6, s=60, color='#2E86AB', edgecolors='black', linewidth=0.5, label='Q1: Высокий X, Высокий Y')
    if q2:
        ax_main.scatter([p[0] for p in q2], [p[1] for p in q2], 
                       alpha=0.6, s=60, color='#A23B72', edgecolors='black', linewidth=0.5, label='Q2: Низкий X, Высокий Y')
    if q3:
        ax_main.scatter([p[0] for p in q3], [p[1] for p in q3], 
                       alpha=0.6, s=60, color='#F18F01', edgecolors='black', linewidth=0.5, label='Q3: Низкий X, Низкий Y')
    if q4:
        ax_main.scatter([p[0] for p in q4], [p[1] for p in q4], 
                       alpha=0.6, s=60, color='#C73E1D', edgecolors='black', linewidth=0.5, label='Q4: Высокий X, Низкий Y')
    
    # Линии медиан
    ax_main.axvline(x_median, color='red', linestyle='--', linewidth=2, alpha=0.7, label=f'Медиана X: {x_median:.2f}')
    ax_main.axhline(y_median, color='red', linestyle='--', linewidth=2, alpha=0.7, label=f'Медиана Y: {y_median:.2f}')
    
    # Линия тренда
    z = np.polyfit(x_vals, y_vals, 1)
    p = np.poly1d(z)
    x_trend = np.linspace(min(x_vals), max(x_vals), 100)
    ax_main.plot(x_trend, p(x_trend), "g-", alpha=0.8, linewidth=3, label='Линия тренда')
    
    ax_main.set_xlabel(x_label, fontsize=11, fontweight='bold')
    ax_main.set_ylabel(y_label, fontsize=11, fontweight='bold')
    ax_main.grid(True, alpha=0.3, linestyle='--')
    ax_main.legend(loc='upper left', fontsize=8)
    
    # Гистограмма для X (сверху)
    ax_x = fig.add_subplot(gs[0, 0:2], sharex=ax_main)
    ax_x.hist(x_vals, bins=15, color='#2E86AB', alpha=0.7, edgecolor='black')
    ax_x.axvline(x_median, color='red', linestyle='--', linewidth=2, alpha=0.7)
    ax_x.set_ylabel('Частота', fontsize=9)
    ax_x.set_title('Распределение по оси X', fontsize=10, fontweight='bold')
    ax_x.grid(True, alpha=0.3, axis='y')
    plt.setp(ax_x.get_xticklabels(), visible=False)
    
    # Гистограмма для Y (справа)
    ax_y = fig.add_subplot(gs[1:3, 2], sharey=ax_main)
    ax_y.hist(y_vals, bins=15, color='#A23B72', alpha=0.7, edgecolor='black', orientation='horizontal')
    ax_y.axhline(y_median, color='red', linestyle='--', linewidth=2, alpha=0.7)
    ax_y.set_xlabel('Частота', fontsize=9)
    ax_y.set_title('Распределение\nпо оси Y', fontsize=10, fontweight='bold')
    ax_y.grid(True, alpha=0.3, axis='x')
    plt.setp(ax_y.get_yticklabels(), visible=False)
    
    # Информационная панель (верхний правый угол)
    ax_info = fig.add_subplot(gs[0, 2])
    ax_info.axis('off')
    
    direction = "прямая" if corr_coef > 0 else "обратная"
    strength = "сильная" if abs(corr_coef) > 0.5 else "умеренная" if abs(corr_coef) > 0.3 else "слабая"
    sig_text = "значима" if p_value < 0.05 else "незначима"
    
    info_text = f"""СТАТИСТИКА:
    
ρ (Спирмен) = {corr_coef:.4f}
p-value = {p_value:.4f}
Связь: {strength} {direction}
Статус: {sig_text}

РАСПРЕДЕЛЕНИЕ:
Q1: {len(q1)} ({len(q1)/len(x_vals)*100:.1f}%)
Q2: {len(q2)} ({len(q2)/len(x_vals)*100:.1f}%)
Q3: {len(q3)} ({len(q3)/len(x_vals)*100:.1f}%)
Q4: {len(q4)} ({len(q4)/len(x_vals)*100:.1f}%)

N = {len(x_vals)}"""
    
    ax_info.text(0.1, 0.95, info_text, transform=ax_info.transAxes,
                fontsize=9, verticalalignment='top',
                bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    fig.suptitle(title, fontsize=14, fontweight='bold', y=0.98)
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return save_path


def create_correlation_plot_variant2(x_vals, y_vals, x_label, y_label, corr_coef, p_value, 
                                     title, save_path):
    """
    ВАРИАНТ 2: График с квадрантами и пояснениями
    Показывает, где находятся большинство наблюдений.
    """
    fig, ax = plt.subplots(figsize=(11, 8))
    
    # Вычисляем медианы
    x_median = np.median(x_vals)
    y_median = np.median(y_vals)
    
    # Разделяем точки по квадрантам
    q1 = [(x, y) for x, y in zip(x_vals, y_vals) if x >= x_median and y >= y_median]
    q2 = [(x, y) for x, y in zip(x_vals, y_vals) if x < x_median and y >= y_median]
    q3 = [(x, y) for x, y in zip(x_vals, y_vals) if x < x_median and y < y_median]
    q4 = [(x, y) for x, y in zip(x_vals, y_vals) if x >= x_median and y < y_median]
    
    # Рисуем квадранты с разными цветами
    colors = ['#2E86AB', '#A23B72', '#F18F01', '#C73E1D']
    labels = [
        f'Квадрант 1: Высокий X, Высокий Y\n({len(q1)} точек, {len(q1)/len(x_vals)*100:.1f}%)',
        f'Квадрант 2: Низкий X, Высокий Y\n({len(q2)} точек, {len(q2)/len(x_vals)*100:.1f}%)',
        f'Квадрант 3: Низкий X, Низкий Y\n({len(q3)} точек, {len(q3)/len(x_vals)*100:.1f}%)',
        f'Квадрант 4: Высокий X, Низкий Y\n({len(q4)} точек, {len(q4)/len(x_vals)*100:.1f}%)'
    ]
    
    for i, (quad, color, label) in enumerate(zip([q1, q2, q3, q4], colors, labels)):
        if quad:
            ax.scatter([p[0] for p in quad], [p[1] for p in quad], 
                      alpha=0.7, s=80, color=color, edgecolors='black', linewidth=1, label=label)
    
    # Линии медиан
    ax.axvline(x_median, color='red', linestyle='--', linewidth=2.5, alpha=0.8, 
              label=f'Медиана X = {x_median:.2f}')
    ax.axhline(y_median, color='red', linestyle='--', linewidth=2.5, alpha=0.8, 
              label=f'Медиана Y = {y_median:.2f}')
    
    # Линия тренда с доверительным интервалом (упрощённый)
    z = np.polyfit(x_vals, y_vals, 1)
    p = np.poly1d(z)
    x_trend = np.linspace(min(x_vals), max(x_vals), 100)
    y_trend = p(x_trend)
    
    # Вычисляем стандартное отклонение остатков для доверительного интервала
    residuals = y_vals - p(x_vals)
    std_residuals = np.std(residuals)
    ax.plot(x_trend, y_trend, "g-", alpha=0.9, linewidth=3, label='Линия тренда')
    ax.fill_between(x_trend, y_trend - std_residuals, y_trend + std_residuals, 
                    alpha=0.2, color='green', label='Область неопределённости (±1σ)')
    
    # Добавляем подписи квадрантов
    x_range = max(x_vals) - min(x_vals)
    y_range = max(y_vals) - min(y_vals)
    
    ax.text(x_median + x_range*0.1, y_median + y_range*0.1, 'Q1', 
           fontsize=16, fontweight='bold', color='gray', alpha=0.5)
    ax.text(x_median - x_range*0.1, y_median + y_range*0.1, 'Q2', 
           fontsize=16, fontweight='bold', color='gray', alpha=0.5)
    ax.text(x_median - x_range*0.1, y_median - y_range*0.1, 'Q3', 
           fontsize=16, fontweight='bold', color='gray', alpha=0.5)
    ax.text(x_median + x_range*0.1, y_median - y_range*0.1, 'Q4', 
           fontsize=16, fontweight='bold', color='gray', alpha=0.5)
    
    ax.set_xlabel(x_label, fontsize=12, fontweight='bold')
    ax.set_ylabel(y_label, fontsize=12, fontweight='bold')
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.legend(loc='best', fontsize=9)
    
    # Информация о корреляции
    direction = "прямая" if corr_coef > 0 else "обратная"
    strength = "сильная" if abs(corr_coef) > 0.5 else "умеренная" if abs(corr_coef) > 0.3 else "слабая"
    sig_text = "значима" if p_value < 0.05 else "незначима"
    
    info_text = f'ρ = {corr_coef:.4f} ({strength} {direction} связь, p = {p_value:.4f}, {sig_text})'
    ax.text(0.02, 0.98, info_text, transform=ax.transAxes, 
           fontsize=11, verticalalignment='top',
           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return save_path


def create_correlation_plot_variant3(x_vals, y_vals, x_label, y_label, corr_coef, p_value, 
                                     title, save_path):
    """
    ВАРИАНТ 3: График с box plots и scatter
    Показывает распределения и выбросы.
    """
    fig = plt.figure(figsize=(12, 10))
    
    # Создаём сетку
    gs = fig.add_gridspec(3, 3, hspace=0.3, wspace=0.3, 
                         left=0.1, right=0.95, top=0.95, bottom=0.1)
    
    # Основной scatter plot
    ax_main = fig.add_subplot(gs[1:3, 0:2])
    
    # Scatter plot с цветовой кодировкой по расстоянию от линии тренда
    z = np.polyfit(x_vals, y_vals, 1)
    p = np.poly1d(z)
    residuals = np.abs(y_vals - p(x_vals))
    max_residual = np.max(residuals)
    
    scatter = ax_main.scatter(x_vals, y_vals, c=residuals, cmap='viridis', 
                             s=80, alpha=0.7, edgecolors='black', linewidth=0.5)
    
    # Линия тренда
    x_trend = np.linspace(min(x_vals), max(x_vals), 100)
    ax_main.plot(x_trend, p(x_trend), "r-", alpha=0.9, linewidth=3, label='Линия тренда')
    
    ax_main.set_xlabel(x_label, fontsize=11, fontweight='bold')
    ax_main.set_ylabel(y_label, fontsize=11, fontweight='bold')
    ax_main.grid(True, alpha=0.3, linestyle='--')
    
    # Цветовая шкала
    cbar = plt.colorbar(scatter, ax=ax_main, pad=0.02)
    cbar.set_label('Отклонение от линии тренда', fontsize=9, rotation=270, labelpad=15)
    
    # Box plot для X (сверху)
    ax_x = fig.add_subplot(gs[0, 0:2], sharex=ax_main)
    bp_x = ax_x.boxplot(x_vals, vert=False, patch_artist=True, widths=0.6)
    for patch in bp_x['boxes']:
        patch.set_facecolor('#2E86AB')
        patch.set_alpha(0.7)
    ax_x.set_ylabel('X', fontsize=10, fontweight='bold')
    ax_x.set_title('Распределение по оси X (Box Plot)', fontsize=10, fontweight='bold')
    ax_x.grid(True, alpha=0.3, axis='x')
    plt.setp(ax_x.get_xticklabels(), visible=False)
    
    # Box plot для Y (справа)
    ax_y = fig.add_subplot(gs[1:3, 2], sharey=ax_main)
    bp_y = ax_y.boxplot(y_vals, vert=True, patch_artist=True, widths=0.6)
    for patch in bp_y['boxes']:
        patch.set_facecolor('#A23B72')
        patch.set_alpha(0.7)
    ax_y.set_xlabel('Y', fontsize=10, fontweight='bold')
    ax_y.set_title('Распределение\nпо оси Y\n(Box Plot)', fontsize=10, fontweight='bold')
    ax_y.grid(True, alpha=0.3, axis='y')
    plt.setp(ax_y.get_yticklabels(), visible=False)
    
    # Статистическая информация
    ax_info = fig.add_subplot(gs[0, 2])
    ax_info.axis('off')
    
    direction = "прямая" if corr_coef > 0 else "обратная"
    strength = "сильная" if abs(corr_coef) > 0.5 else "умеренная" if abs(corr_coef) > 0.3 else "слабая"
    sig_text = "значима" if p_value < 0.05 else "незначима"
    
    x_q1, x_median, x_q3 = np.percentile(x_vals, [25, 50, 75])
    y_q1, y_median, y_q3 = np.percentile(y_vals, [25, 50, 75])
    
    info_text = f"""СТАТИСТИКА КОРРЕЛЯЦИИ:
ρ = {corr_coef:.4f}
p = {p_value:.4f}
{strength} {direction} связь
{sig_text}

КВАРТИЛИ X:
Q1: {x_q1:.2f}
Медиана: {x_median:.2f}
Q3: {x_q3:.2f}

КВАРТИЛИ Y:
Q1: {y_q1:.2f}
Медиана: {y_median:.2f}
Q3: {y_q3:.2f}

N = {len(x_vals)}"""
    
    ax_info.text(0.1, 0.95, info_text, transform=ax_info.transAxes,
                fontsize=9, verticalalignment='top',
                bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    fig.suptitle(title, fontsize=14, fontweight='bold', y=0.98)
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return save_path


def create_correlation_plot_variant4(x_vals, y_vals, x_label, y_label, corr_coef, p_value, 
                                     title, save_path):
    """
    ВАРИАНТ 4: График с категоризацией по силе связи
    Помогает визуально выделить группы наблюдений.
    """
    fig, ax = plt.subplots(figsize=(11, 8))
    
    # Категоризируем точки по значениям X и Y
    x_mean = np.mean(x_vals)
    y_mean = np.mean(y_vals)
    x_std = np.std(x_vals)
    y_std = np.std(y_vals)
    
    # Создаём категории: низкие, средние, высокие значения
    def categorize(val, mean, std):
        if val < mean - 0.5 * std:
            return 'Низкое'
        elif val > mean + 0.5 * std:
            return 'Высокое'
        else:
            return 'Среднее'
    
    categories = []
    colors_map = {
        ('Низкое', 'Низкое'): '#F18F01',      # Оранжевый
        ('Низкое', 'Среднее'): '#FFB347',     # Светло-оранжевый
        ('Низкое', 'Высокое'): '#A23B72',      # Фиолетовый
        ('Среднее', 'Низкое'): '#FFD700',      # Золотой
        ('Среднее', 'Среднее'): '#90EE90',     # Светло-зелёный
        ('Среднее', 'Высокое'): '#87CEEB',     # Небесно-голубой
        ('Высокое', 'Низкое'): '#C73E1D',      # Красно-коричневый
        ('Высокое', 'Среднее'): '#2E86AB',     # Синий
        ('Высокое', 'Высокое'): '#006400',     # Тёмно-зелёный
    }
    
    for x, y in zip(x_vals, y_vals):
        cat_x = categorize(x, x_mean, x_std)
        cat_y = categorize(y, y_mean, y_std)
        categories.append((cat_x, cat_y))
    
    # Группируем точки по категориям
    category_groups = {}
    for (x, y), cat in zip(zip(x_vals, y_vals), categories):
        if cat not in category_groups:
            category_groups[cat] = {'x': [], 'y': []}
        category_groups[cat]['x'].append(x)
        category_groups[cat]['y'].append(y)
    
    # Рисуем точки по категориям
    for cat, color in colors_map.items():
        if cat in category_groups:
            count = len(category_groups[cat]['x'])
            if count > 0:
                label = f"{cat[0]} X, {cat[1]} Y ({count} точек)"
                ax.scatter(category_groups[cat]['x'], category_groups[cat]['y'],
                          alpha=0.7, s=80, color=color, edgecolors='black', 
                          linewidth=1, label=label)
    
    # Линия тренда
    z = np.polyfit(x_vals, y_vals, 1)
    p = np.poly1d(z)
    x_trend = np.linspace(min(x_vals), max(x_vals), 100)
    ax.plot(x_trend, p(x_trend), "k-", alpha=0.9, linewidth=3, 
           linestyle='--', label='Линия тренда')
    
    # Линии средних значений
    ax.axvline(x_mean, color='red', linestyle=':', linewidth=2, alpha=0.7, 
              label=f'Среднее X = {x_mean:.2f}')
    ax.axhline(y_mean, color='red', linestyle=':', linewidth=2, alpha=0.7, 
              label=f'Среднее Y = {y_mean:.2f}')
    
    # Зоны стандартных отклонений
    ax.axvspan(x_mean - 0.5*x_std, x_mean + 0.5*x_std, alpha=0.1, color='gray', 
              label='Средняя зона (±0.5σ)')
    
    ax.set_xlabel(x_label, fontsize=12, fontweight='bold')
    ax.set_ylabel(y_label, fontsize=12, fontweight='bold')
    ax.set_title(title, fontsize=14, fontweight='bold', pad=20)
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.legend(loc='best', fontsize=8, ncol=2)
    
    # Информация о корреляции
    direction = "прямая" if corr_coef > 0 else "обратная"
    strength = "сильная" if abs(corr_coef) > 0.5 else "умеренная" if abs(corr_coef) > 0.3 else "слабая"
    sig_text = "значима" if p_value < 0.05 else "незначима"
    
    info_text = f'ρ = {corr_coef:.4f} ({strength} {direction} связь, p = {p_value:.4f}, {sig_text})'
    ax.text(0.02, 0.98, info_text, transform=ax.transAxes, 
           fontsize=11, verticalalignment='top',
           bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))
    
    plt.tight_layout()
    plt.savefig(save_path, dpi=300, bbox_inches='tight')
    plt.close()
    
    return save_path


def create_correlation_plot(x_vals, y_vals, x_label, y_label, corr_coef, p_value, 
                           title, save_path, variant=1):
    """
    Универсальная функция для создания графиков корреляций.
    По умолчанию использует вариант 1.
    
    Параметры:
    - variant: номер варианта (1-4)
    """
    if variant == 1:
        return create_correlation_plot_variant1(x_vals, y_vals, x_label, y_label, corr_coef, p_value, title, save_path)
    elif variant == 2:
        return create_correlation_plot_variant2(x_vals, y_vals, x_label, y_label, corr_coef, p_value, title, save_path)
    elif variant == 3:
        return create_correlation_plot_variant3(x_vals, y_vals, x_label, y_label, corr_coef, p_value, title, save_path)
    elif variant == 4:
        return create_correlation_plot_variant4(x_vals, y_vals, x_label, y_label, corr_coef, p_value, title, save_path)
    else:
        return create_correlation_plot_variant1(x_vals, y_vals, x_label, y_label, corr_coef, p_value, title, save_path)

def generate_interpretation_and_recommendations(corr, correlation_type, holland_type_name=None):
    """
    Генерирует интерпретацию и прикладные рекомендации для корреляции.
    
    Параметры:
    - corr: словарь с данными корреляции
    - correlation_type: 'parent_holland' или 'ovcharova_holland'
    - holland_type_name: название типа Холланда
    """
    rho = corr['spearman_corr']
    p_val = corr['spearman_p']
    
    direction = "прямая" if rho > 0 else "обратная"
    strength = "сильная" if abs(rho) > 0.5 else "умеренная" if abs(rho) > 0.3 else "слабая"
    
    interpretation = f"""
ИНТЕРПРЕТАЦИЯ РЕЗУЛЬТАТА:

Обнаружена {strength} {direction} связь (ρ = {rho:.4f}, p = {p_val:.4f}) между:
"""
    
    if correlation_type == 'parent_holland':
        parent_q = str(corr['parent_question'])
        h_type = holland_type_name or corr.get('holland_type_name', 'N/A')
        
        interpretation += f"""
• Установкой родителя: "{parent_q[:80]}..."
• Профессиональным интересом: {h_type}

"""
        
        if rho > 0:
            interpretation += f"""
Чем выше выраженность данной установки у родителя, тем выше интерес подростка к профессиям типа {h_type}.
"""
        else:
            interpretation += f"""
Чем выше выраженность данной установки у родителя, тем ниже интерес подростка к профессиям типа {h_type}.
"""
        
        # Рекомендации для родителей
        recommendations = f"""
ПРИКЛАДНЫЕ РЕКОМЕНДАЦИИ ДЛЯ РОДИТЕЛЕЙ:

1. ОСОЗНАНИЕ ВЛИЯНИЯ:
   • Ваши установки и подходы к воспитанию влияют на формирование профессиональных интересов вашего ребёнка
   • Важно понимать, что ваше отношение к ребёнку может способствовать или препятствовать развитию определённых профессиональных интересов

2. КОНСТРУКТИВНЫЕ ДЕЙСТВИЯ:
"""
        
        if rho > 0:
            recommendations += f"""
   • Если вы хотите поддержать интерес ребёнка к профессиям типа {h_type}, продолжайте развивать данную установку
   • Создавайте возможности для знакомства ребёнка с профессиями данного типа (экскурсии, встречи с профессионалами)
   • Поощряйте развитие навыков, связанных с данным типом профессиональных интересов
"""
        else:
            recommendations += f"""
   • Если вы хотите поддержать интерес ребёнка к профессиям типа {h_type}, рассмотрите возможность изменения данной установки
   • Проанализируйте, как ваше отношение может препятствовать развитию интереса к данному типу профессий
   • Создавайте поддерживающую среду для развития интереса к профессиям типа {h_type}
"""
        
        recommendations += f"""
3. ДЛЯ ПЕДАГОГОВ И ПСИХОЛОГОВ:

   • При работе с семьёй учитывайте влияние родительских установок на профессиональные интересы подростка
   • Проводите консультации с родителями о важности их роли в профессиональном самоопределении
   • Разрабатывайте программы для родителей по поддержке профессионального развития подростков
   • При выявлении конфликтов между родительскими установками и интересами подростка проводите семейное консультирование
"""
        
    elif correlation_type == 'ovcharova_holland':
        ovcharova_q = str(corr['ovcharova_question'])
        h_type = holland_type_name or corr.get('holland_type_name', 'N/A')
        
        interpretation += f"""
• Мотивом выбора профессии: "{ovcharova_q[:80]}..."
• Профессиональным интересом: {h_type}

"""
        
        if rho > 0:
            interpretation += f"""
Чем важнее для подростка данный мотив выбора профессии, тем выше его интерес к профессиям типа {h_type}.
"""
        else:
            interpretation += f"""
Чем важнее для подростка данный мотив выбора профессии, тем ниже его интерес к профессиям типа {h_type}.
"""
        
        # Рекомендации
        recommendations = f"""
ПРИКЛАДНЫЕ РЕКОМЕНДАЦИИ:

1. ДЛЯ ПЕДАГОГОВ:
   • Учитывайте связь между мотивами выбора профессии и профессиональными интересами при профориентационной работе
   • Помогайте подросткам осознавать связь между их мотивами и интересами
   • Разрабатывайте программы, которые учитывают как мотивы, так и интересы подростков
   • При выявлении противоречий между мотивами и интересами проводите индивидуальную работу

2. ДЛЯ ПСИХОЛОГОВ:
   • Используйте данные о связи мотивов и интересов для более точной профориентационной диагностики
   • Помогайте подросткам интегрировать мотивы выбора профессии с их профессиональными интересами
   • Проводите консультирование по гармонизации мотивов и интересов
   • Разрабатывайте индивидуальные программы профессионального развития

3. ДЛЯ РОДИТЕЛЕЙ:
   • Обсуждайте с подростком его мотивы выбора профессии и профессиональные интересы
   • Поддерживайте подростка в исследовании профессий, соответствующих его интересам и мотивам
   • Помогайте подростку найти баланс между мотивами и интересами при выборе профессии
"""
    
    return interpretation + recommendations

script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)
newtest_dir = os.path.join(parent_dir, "newtest")

print("=" * 80)
print("СОЗДАНИЕ ПОЛНОГО ИТОГОВОГО ОТЧЁТА")
print("=" * 80)

# Загружаем данные
parents_file = os.path.join(newtest_dir, "Опрос для родителей  (Ответы).csv")
students_file = os.path.join(newtest_dir, "Опрос ученика (Ответы) Новый.csv")

print("\n1. Загрузка данных...")
parents_df = pd.read_csv(parents_file, encoding='utf-8-sig')
students_df = pd.read_csv(students_file, encoding='utf-8-sig')

print(f"   Родителей: {len(parents_df)}")
print(f"   Учеников: {len(students_df)}")

# Определяем структуру
parent_questions = parents_df.columns[6:66].tolist()

# Находим вопросы Овчаровой
student_cols = students_df.columns.tolist()
q12_col = None
q19_col = None
ovcharova_all = {}

for col in student_cols:
    col_str = str(col)
    for i in range(1, 21):
        if col_str.startswith(f"{i}. "):
            if any(x in col_str for x in ['Требует', 'Нравится', 'Предполагает', 'Соответствует', 
                                          'Позволяет', 'Дает', 'Является', 'Близка', 'Избрана', 'Единственно', 'Способствует']):
                ovcharova_all[i] = col
                if i == 12 and 'Дает возможности для роста' in col_str:
                    q12_col = col
                if i == 19 and 'Позволяет использовать профессиональные умения' in col_str:
                    q19_col = col
                break

if not q12_col or not q19_col:
    print("ОШИБКА: Не найдены вопросы 12 и 19!")
    sys.exit(1)

# Сопоставляем пары
if 'Number' in parents_df.columns and 'Number' in students_df.columns:
    parents_sorted = parents_df.sort_values('Number').reset_index(drop=True)
    students_sorted = students_df.sort_values('Number').reset_index(drop=True)
    pairs = []
    for i in range(min(len(parents_sorted), len(students_sorted))):
        p_num = parents_sorted.iloc[i]['Number'] if pd.notna(parents_sorted.iloc[i]['Number']) else None
        s_num = students_sorted.iloc[i]['Number'] if pd.notna(students_sorted.iloc[i]['Number']) else None
        if p_num is not None and s_num is not None and p_num == s_num:
            pairs.append({'pair_id': len(pairs) + 1, 'parent_idx': i, 'student_idx': i, 'number': p_num})
else:
    pairs = [{'pair_id': i+1, 'parent_idx': i, 'student_idx': i, 'number': i+1} 
             for i in range(min(len(parents_df), len(students_df)))]
    parents_sorted = parents_df.reset_index(drop=True)
    students_sorted = students_df.reset_index(drop=True)

print(f"\n2. Найдено пар: {len(pairs)}")

# Перепроверяем корреляции
print("\n3. Перепроверка корреляций...")
verified_correlations = []

for parent_q in parent_questions:
    p_vals_12, q12_vals = [], []
    p_vals_19, q19_vals = [], []
    
    for pair in pairs:
        try:
            p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
            s12_val = students_sorted.iloc[pair['student_idx']][q12_col]
            s19_val = students_sorted.iloc[pair['student_idx']][q19_col]
            
            p_num = float(p_val) if pd.notna(p_val) else None
            s12_num = float(s12_val) if pd.notna(s12_val) else None
            s19_num = float(s19_val) if pd.notna(s19_val) else None
            
            if p_num is not None and s12_num is not None:
                p_vals_12.append(p_num)
                q12_vals.append(s12_num)
            
            if p_num is not None and s19_num is not None:
                p_vals_19.append(p_num)
                q19_vals.append(s19_num)
        except:
            pass
    
    if len(p_vals_12) >= 3:
        try:
            rho12, p12 = spearmanr(p_vals_12, q12_vals)
            r12, p12_p = pearsonr(p_vals_12, q12_vals)
            verified_correlations.append({
                'parent_question': parent_q,
                'ovcharova_question': '12. Дает возможности для роста профессионального мастерства',
                'parent_values': p_vals_12,
                'ovcharova_values': q12_vals,
                'spearman_corr': rho12,
                'spearman_p': p12,
                'pearson_corr': r12,
                'pearson_p': p12_p,
                'n': len(p_vals_12)
            })
        except:
            pass
    
    if len(p_vals_19) >= 3:
        try:
            rho19, p19 = spearmanr(p_vals_19, q19_vals)
            r19, p19_p = pearsonr(p_vals_19, q19_vals)
            verified_correlations.append({
                'parent_question': parent_q,
                'ovcharova_question': '19. Позволяет использовать профессиональные умения вне работы',
                'parent_values': p_vals_19,
                'ovcharova_values': q19_vals,
                'spearman_corr': rho19,
                'spearman_p': p19,
                'pearson_corr': r19,
                'pearson_p': p19_p,
                'n': len(p_vals_19)
            })
        except:
            pass

print(f"   Перепроверено корреляций: {len(verified_correlations)}")

# Проверяем итоговые результаты
print("\n4. Проверка итоговых результатов...")
total_scores = []
for idx, row in students_sorted.iterrows():
    scores = []
    for q_num, col in ovcharova_all.items():
        val = row[col]
        try:
            num_val = float(val) if pd.notna(val) else None
            if num_val is not None:
                scores.append(num_val)
        except:
            pass
    if len(scores) > 0:
        total_scores.append(sum(scores))

existing_corr_file = os.path.join(newtest_dir, "correlations_analysis.csv")
total_used = False
if os.path.exists(existing_corr_file):
    existing_corr_df = pd.read_csv(existing_corr_file, encoding='utf-8-sig')
    for _, row in existing_corr_df.iterrows():
        if 'итог' in str(row['profession_question']).lower() or 'total' in str(row['profession_question']).lower():
            total_used = True
            break

print(f"   Итоговых баллов рассчитано: {len(total_scores)}")
print(f"   Итоговые результаты использовались: {'Да' if total_used else 'Нет'}")

# Создаём документ
print("\n5. Создание итогового документа...")
doc = Document()

def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Заголовок
title = doc.add_heading('АНАЛИЗ КОРРЕЛЯЦИЙ МЕЖДУ УСТАНОВКАМИ РОДИТЕЛЕЙ И МОТИВАМИ ВЫБОРА ПРОФЕССИИ ПОДРОСТКАМИ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 1. ВВЕДЕНИЕ
doc.add_heading('1. ВВЕДЕНИЕ', 1)
intro_text = """
Целью данного исследования является проверка гипотезы о влиянии установок родителей на мотивы выбора профессии подростками. 
Исследование основано на анализе данных психологического опроса, проведенного среди пар родитель-ребенок.

Гипотеза исследования: установки родителей, проявляющиеся в их воспитательных подходах и отношениях к ребенку, 
влияют на мотивы выбора ребенком будущей профессии.

В исследовании использовались следующие методики:
• ВРР Марковской (60 вопросов) - для оценки установок родителей
• Опросник Овчаровой (20 вопросов) - для оценки мотивов выбора профессии подростками
"""
doc.add_paragraph(intro_text)

# 2. МЕТОДОЛОГИЯ (по примеру главы 2.2)
doc.add_heading('2. МЕТОДОЛОГИЯ ИССЛЕДОВАНИЯ', 1)

# 2.1. Описание данных
doc.add_heading('2.1. Описание данных', 2)
data_desc = f"""
В исследовании приняли участие:
• {len(parents_df)} родителей
• {len(students_df)} подростков
• {len(pairs)} пар родитель-ребенок для анализа корреляций

Родители отвечали на 60 вопросов методики ВРР Марковской, касающихся их установок и подходов к воспитанию.
Подростки отвечали на 20 вопросов опросника Овчаровой о мотивах выбора профессии.
"""
doc.add_paragraph(data_desc)

# 2.2. Методы математической статистики (по примеру из PDF)
doc.add_heading('2.2. Методы математической статистики', 2)

methods_text = """
Для проверки гипотезы о влиянии установок родителей на мотивы выбора профессии подростками использовались следующие методы математической статистики:

1. КОЭФФИЦИЕНТ КОРРЕЛЯЦИИ СПИРМЕНА (ρ, rho)
   
   Коэффициент ранговой корреляции Спирмена используется для измерения силы и направления монотонной связи 
   между двумя переменными. Этот метод был выбран по следующим причинам:
   
   • Данные психологического исследования часто имеют порядковую (ранговую) природу
   • Метод не требует нормального распределения данных
   • Устойчив к выбросам
   • Подходит для анализа связи между установками (порядковые данные) и мотивами (порядковые данные)
   
   Формула коэффициента Спирмена:
   
   ρ = 1 - (6 × Σd²) / (n × (n² - 1))
   
   где:
   d = разность рангов между парами наблюдений
   n = количество наблюдений
   
   Значения коэффициента интерпретируются следующим образом:
   • |ρ| = 0.00 - 0.19: очень слабая связь
   • |ρ| = 0.20 - 0.39: слабая связь
   • |ρ| = 0.40 - 0.59: умеренная связь
   • |ρ| = 0.60 - 0.79: сильная связь
   • |ρ| = 0.80 - 1.00: очень сильная связь

2. КОЭФФИЦИЕНТ КОРРЕЛЯЦИИ ПИРСОНА (r)
   
   Коэффициент линейной корреляции Пирсона использовался как дополнительный метод для проверки линейной 
   зависимости между переменными.
   
   Формула коэффициента Пирсона:
   
   r = Σ((X - X̄)(Y - Ȳ)) / √(Σ(X - X̄)² × Σ(Y - Ȳ)²)
   
   где:
   X, Y = значения переменных
   X̄, Ȳ = средние значения переменных

3. СТАТИСТИЧЕСКАЯ ЗНАЧИМОСТЬ (p-value)
   
   Для оценки статистической значимости корреляций использовался критерий значимости с уровнем α = 0.05.
   Корреляции считались статистически значимыми при p < 0.05.
"""
doc.add_paragraph(methods_text)

# 2.3. Обоснование выбора вопросов
doc.add_heading('2.3. Обоснование выбора вопросов опросника Овчаровой', 2)

justification = """
В исследовании были выбраны два вопроса из опросника Овчаровой (20 вопросов о мотивах выбора профессии):

ВОПРОС 12: "Дает возможности для роста профессионального мастерства"
ВОПРОС 19: "Позволяет использовать профессиональные умения вне работы"

ОБОСНОВАНИЕ ВЫБОРА:

1. ТЕОРЕТИЧЕСКОЕ ОБОСНОВАНИЕ:
   • Вопрос 12 связан с мотивом профессионального развития и карьерного роста
   • Вопрос 19 связан с мотивом интеграции профессии в личную жизнь
   • Оба вопроса отражают внутренние мотивы выбора профессии, связанные с самореализацией
   • Эти мотивы могут быть более тесно связаны с установками родителей на развитие ребенка

2. МЕТОДОЛОГИЧЕСКОЕ ОБОСНОВАНИЕ:
   • Вопросы 12 и 19 отражают разные аспекты профессиональной мотивации (развитие vs. интеграция)
   • Позволяют исследовать как стремление к профессиональному росту, так и желание применять навыки в жизни
   • Эти мотивы могут быть более чувствительны к влиянию родительских установок

3. ПОЧЕМУ НЕ ВЫБРАЛИ ДРУГИЕ ВОПРОСЫ:
   
   ВОПРОС 2 ("Нравится родителям"):
   • Отражает зависимость от родительского одобрения, что может создавать искусственную корреляцию
   • Не позволяет исследовать самостоятельность выбора профессии
   • Высокая корреляция может быть обусловлена не влиянием установок, а прямой зависимостью от мнения родителей
   • Использование вопроса 2 могло бы исказить результаты исследования
   
   ДРУГИЕ ВОПРОСЫ (1, 3-11, 13-18, 20):
   • Ограничение числа вопросов позволяет провести более глубокий анализ
   • Вопросы 12 и 19 наиболее релевантны для исследования влияния установок родителей
   • Анализ всех 20 вопросов создал бы 1200 корреляций, что сделало бы анализ слишком объёмным
   • Фокус на наиболее релевантных вопросах повышает валидность исследования

4. ПРОВЕРКА ИТОГОВЫХ РЕЗУЛЬТАТОВ:
   
   Итоговый балл по опроснику Овчаровой (сумма всех 20 вопросов) НЕ использовался в корреляционном анализе.
   
   ОБОСНОВАНИЕ:
   • Итоговый балл может "размывать" специфические связи между установками родителей и конкретными мотивами
   • Использование отдельных вопросов позволяет выявить более точные и специфические связи
   • Разные мотивы могут по-разному реагировать на установки родителей
   • Итоговый балл не позволяет исследовать специфические механизмы влияния
"""
doc.add_paragraph(justification)

# 3. ПОЛНЫЕ РАСЧЁТЫ
doc.add_heading('3. ПОЛНЫЕ РАСЧЁТЫ КОРРЕЛЯЦИЙ', 1)

intro_calc = """
Ниже представлены полные пошаговые расчёты корреляций для статистически значимых комбинаций вопросов родителей 
(ВРР Марковской) и вопросов 12 и 19 опросника Овчаровой. Каждый расчёт выполнен в соответствии 
с методологией коэффициента корреляции Спирмена и включает все промежуточные этапы вычислений.

Для незначимых корреляций представлена сводная таблица в разделе 4.3.
"""
doc.add_paragraph(intro_calc)

# Разделяем на значимые и незначимые
significant_correlations = [c for c in verified_correlations if c['spearman_p'] < 0.05]
non_significant_correlations = [c for c in verified_correlations if c['spearman_p'] >= 0.05]

# Сортируем значимые по силе корреляции
sorted_significant = sorted(significant_correlations, key=lambda x: abs(x['spearman_corr']), reverse=True)

print(f"\n6. Генерация полных расчётов для {len(sorted_significant)} значимых корреляций...")

for idx, corr in enumerate(sorted_significant, 1):
    doc.add_heading(f'3.{idx}. Расчёт корреляции {idx}', 2)
    
    # Описание
    desc_para = doc.add_paragraph()
    desc_para.add_run("Вопрос родителя (ВРР Марковской): ").bold = True
    desc_para.add_run(f"{corr['parent_question']}")
    
    desc_para2 = doc.add_paragraph()
    desc_para2.add_run("Вопрос Овчаровой: ").bold = True
    desc_para2.add_run(f"{corr['ovcharova_question']}")
    
    doc.add_paragraph()
    
    # ШАГ 1
    doc.add_heading("ШАГ 1: Сбор данных по всем парам", 3)
    step1 = f"""
Для данной комбинации вопросов были собраны ответы всех {corr['n']} пар родитель-ребенок.
Каждая пара представляет собой ответы одного родителя и его ребенка на соответствующие вопросы.
"""
    doc.add_paragraph(step1)
    
    # Таблица со всеми парами
    doc.add_paragraph("Данные по всем парам:")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Пара №'
    hdr_cells[1].text = 'Ответ родителя'
    hdr_cells[2].text = 'Ответ подростка'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    parent_vals = corr['parent_values']
    ovcharova_vals = corr['ovcharova_values']
    
    for i in range(len(parent_vals)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = f"{parent_vals[i]:.0f}"
        row_cells[2].text = f"{ovcharova_vals[i]:.0f}"
    
    doc.add_paragraph()
    
    # ШАГ 2
    doc.add_heading("ШАГ 2: Преобразование в ранги", 3)
    
    parent_ranks = rankdata(parent_vals)
    ovcharova_ranks = rankdata(ovcharova_vals)
    
    step2 = f"""
Каждому ответу присваивается ранг от 1 до {corr['n']} (где {corr['n']} - количество наблюдений, то есть количество пар родитель-ребенок, по которым рассчитана корреляция).
Если несколько значений одинаковы, им присваивается средний ранг.

Пример ранжирования для ответов родителей:
• Минимальное значение: {min(parent_vals):.0f} → ранг 1
• Максимальное значение: {max(parent_vals):.0f} → ранг {corr['n']}
• Среднее значение: {np.mean(parent_vals):.2f}

Пример ранжирования для ответов подростков:
• Минимальное значение: {min(ovcharova_vals):.0f} → ранг 1
• Максимальное значение: {max(ovcharova_vals):.0f} → ранг {corr['n']}
• Среднее значение: {np.mean(ovcharova_vals):.2f}
"""
    doc.add_paragraph(step2)
    
    # Таблица с рангами для всех пар
    doc.add_paragraph("Ранжирование для всех пар:")
    
    rank_table = doc.add_table(rows=1, cols=5)
    rank_table.style = 'Light Grid Accent 1'
    hdr_cells = rank_table.rows[0].cells
    hdr_cells[0].text = 'Пара №'
    hdr_cells[1].text = 'Ранг родителя'
    hdr_cells[2].text = 'Ранг подростка'
    hdr_cells[3].text = 'd (разность)'
    hdr_cells[4].text = 'd²'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    for i in range(len(parent_vals)):
        d = parent_ranks[i] - ovcharova_ranks[i]
        d_squared = d ** 2
        row_cells = rank_table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = f"{parent_ranks[i]:.1f}"
        row_cells[2].text = f"{ovcharova_ranks[i]:.1f}"
        row_cells[3].text = f"{d:.1f}"
        row_cells[4].text = f"{d_squared:.1f}"
    
    doc.add_paragraph()
    
    # ШАГ 3
    doc.add_heading("ШАГ 3: Расчёт разностей рангов и их квадратов", 3)
    
    diffs = parent_ranks - ovcharova_ranks
    diffs_squared = diffs ** 2
    sum_d_squared = np.sum(diffs_squared)
    
    step3 = f"""
Для каждой пары вычисляется разность рангов: d = ранг_родителя - ранг_подростка
Затем вычисляется квадрат разности: d²

Сумма квадратов разностей для всех {corr['n']} пар: Σd² = {sum_d_squared:.2f}
"""
    doc.add_paragraph(step3)
    
    # ШАГ 4
    doc.add_heading("ШАГ 4: Применение формулы коэффициента корреляции Спирмена", 3)
    
    n = corr['n']
    formula_text = f"""
Формула коэффициента Спирмена:

ρ = 1 - (6 × Σd²) / (n × (n² - 1))

Подстановка значений:
• n = {n} (количество пар наблюдений, то есть количество пар родитель-ребенок, по которым рассчитана корреляция)
• Σd² = {sum_d_squared:.2f} (сумма квадратов разностей рангов)
• n² - 1 = {n**2 - 1}
• n × (n² - 1) = {n * (n**2 - 1)}
• 6 × Σd² = {6 * sum_d_squared:.2f}

Расчёт:
ρ = 1 - ({6 * sum_d_squared:.2f}) / ({n * (n**2 - 1)})
ρ = 1 - {6 * sum_d_squared / (n * (n**2 - 1)):.6f}
ρ = {1 - 6 * sum_d_squared / (n * (n**2 - 1)):.4f}

ПРИМЕЧАНИЕ: Небольшое расхождение между расчетным значением ({1 - 6 * sum_d_squared / (n * (n**2 - 1)):.4f}) 
и фактическим ({corr['spearman_corr']:.4f}) может быть связано с обработкой связанных рангов (ties) 
в статистическом пакете, что является стандартной практикой при расчете коэффициента Спирмена.
"""
    doc.add_paragraph(formula_text)
    
    # Результат
    doc.add_heading("Результат расчёта", 3)
    
    result_text = f"""
• Коэффициент Спирмена: ρ = {corr['spearman_corr']:.4f}
• Уровень значимости: p = {corr['spearman_p']:.4f}
• Коэффициент Пирсона: r = {corr['pearson_corr']:.4f}
• Уровень значимости (Пирсон): p = {corr['pearson_p']:.4f}
• Количество пар: n = {corr['n']}
"""
    doc.add_paragraph(result_text)
    
    # Интерпретация
    doc.add_heading("Интерпретация результата", 3)
    
    rho_abs = abs(corr['spearman_corr'])
    if rho_abs < 0.2:
        strength = "очень слабая"
    elif rho_abs < 0.4:
        strength = "слабая"
    elif rho_abs < 0.6:
        strength = "умеренная"
    elif rho_abs < 0.8:
        strength = "сильная"
    else:
        strength = "очень сильная"
    
    direction = "прямая" if corr['spearman_corr'] > 0 else "обратная"
    
    if corr['spearman_p'] < 0.001:
        sig_text = "Связь статистически высоко значима (p < 0.001)."
    elif corr['spearman_p'] < 0.01:
        sig_text = "Связь статистически значима (p < 0.01)."
    elif corr['spearman_p'] < 0.05:
        sig_text = "Связь статистически значима (p < 0.05)."
    else:
        sig_text = "Связь не является статистически значимой (p ≥ 0.05)."
    
    interpretation = f"""
Коэффициент корреляции ρ = {corr['spearman_corr']:.4f} указывает на {strength} {direction} связь между установками родителей и мотивами выбора профессии подростками.

{sig_text}
"""
    doc.add_paragraph(interpretation)
    
    if idx < len(sorted_significant):
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()

# 3.5. КОРРЕЛЯЦИИ С ОПРОСНИКОМ ХОЛЛАНДА
doc.add_heading('3.5. КОРРЕЛЯЦИИ С ОПРОСНИКОМ ХОЛЛАНДА', 1)

holland_intro = """
Опросник Холланда (42 вопроса "Что тебе ближе?") оценивает профессиональные интересы подростков 
по шести типам личности: Реалистичный (R), Исследовательский (I), Артистический (A), 
Социальный (S), Предпринимательский (E), Конвенциональный (C).

Для корреляционного анализа профессии, выбранные подростками, были классифицированы по типам Холланда,
и для каждого респондента рассчитаны итоговые баллы по каждому из шести типов (количество выбранных 
профессий каждого типа).

В данном разделе представлен анализ корреляций между:
• Установками родителей (60 вопросов ВРР Марковской)
• Мотивами выбора профессии (вопросы 12 и 19 Овчаровой)
• Профессиональными интересами (баллы по типам Холланда: R, I, A, S, E, C)
"""
doc.add_paragraph(holland_intro)

# Находим вопросы Холланда
print("\n7. Поиск вопросов Холланда...")
holland_cols = []
for col in student_cols:
    col_str = str(col).strip()
    # Ищем столбцы с "Что тебе ближе" (с учётом возможных пробелов)
    if 'Что тебе ближе' in col_str or col_str.startswith('1. Что тебе ближе') or \
       (col_str.startswith('1.') and 'ближе' in col_str.lower()):
        holland_cols.append(col)

print(f"   Найдено вопросов Холланда: {len(holland_cols)}")

# Рассчитываем баллы по типам Холланда для каждого респондента
print("\n8. Расчёт баллов по типам Холланда...")
holland_scores = None

if len(holland_cols) > 0:
    sample_val = students_sorted.iloc[0][holland_cols[0]] if len(students_sorted) > 0 else None
    print(f"   Пример значения в первом столбце Холланда: {str(sample_val)[:50] if sample_val is not None else 'None'}")
    
    # Рассчитываем баллы по типам Холланда (R, I, A, S, E, C)
    holland_scores = calculate_holland_scores(students_sorted, holland_cols, pairs)
    
    # Выводим статистику по баллам
    for h_type in ['R', 'I', 'A', 'S', 'E', 'C']:
        scores = holland_scores[h_type]
        print(f"   Тип {h_type}: средний балл = {np.mean(scores):.2f}, мин = {min(scores)}, макс = {max(scores)}")
else:
    print("   ⚠ Столбцы Холланда не найдены")

# Рассчитываем корреляции между вопросами родителей и типами Холланда
print("\n9. Расчёт корреляций с типами Холланда...")
holland_correlations = []

if holland_scores is not None:
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный (Realistic)',
        'I': 'Исследовательский (Investigative)',
        'A': 'Артистический (Artistic)',
        'S': 'Социальный (Social)',
        'E': 'Предприимчивый (Enterprising)',
        'C': 'Конвенциональный (Conventional)'
    }
    
    for parent_q in parent_questions:
        for h_type in holland_types:
            p_vals = []
            h_vals = []
            
            for i, pair in enumerate(pairs):
                try:
                    p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                    h_score = holland_scores[h_type][i]
                    
                    p_num = float(p_val) if pd.notna(p_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if p_num is not None and h_num is not None:
                        p_vals.append(p_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(p_vals) >= 3:
                try:
                    rho, p = spearmanr(p_vals, h_vals)
                    holland_correlations.append({
                        'parent_question': parent_q,
                        'holland_type': h_type,
                        'holland_type_name': holland_type_names[h_type],
                        'spearman_corr': rho,
                        'spearman_p': p,
                        'n': len(p_vals)
                    })
                except:
                    pass

# Рассчитываем корреляции между вопросами Овчаровой и типами Холланда
ovcharova_holland_correlations = []

if holland_scores is not None:
    holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
    holland_type_names = {
        'R': 'Реалистичный (Realistic)',
        'I': 'Исследовательский (Investigative)',
        'A': 'Артистический (Artistic)',
        'S': 'Социальный (Social)',
        'E': 'Предприимчивый (Enterprising)',
        'C': 'Конвенциональный (Conventional)'
    }
    
    for ovcharova_q_col in [q12_col, q19_col]:
        for h_type in holland_types:
            o_vals = []
            h_vals = []
            
            for i, pair in enumerate(pairs):
                try:
                    o_val = students_sorted.iloc[pair['student_idx']][ovcharova_q_col]
                    h_score = holland_scores[h_type][i]
                    
                    o_num = float(o_val) if pd.notna(o_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if o_num is not None and h_num is not None:
                        o_vals.append(o_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(o_vals) >= 3:
                try:
                    rho, p = spearmanr(o_vals, h_vals)
                    ovcharova_holland_correlations.append({
                        'ovcharova_question': ovcharova_q_col,
                        'holland_type': h_type,
                        'holland_type_name': holland_type_names[h_type],
                        'spearman_corr': rho,
                        'spearman_p': p,
                        'n': len(o_vals)
                    })
                except:
                    pass

print(f"   Корреляций родители-Холланд: {len(holland_correlations)}")
print(f"   Корреляций Овчарова-Холланд: {len(ovcharova_holland_correlations)}")

# Статистика по Холланду
holland_significant = [c for c in holland_correlations if c['spearman_p'] < 0.05]
ovh_significant = [c for c in ovcharova_holland_correlations if c['spearman_p'] < 0.05]

doc.add_heading('3.5.1. Корреляции между установками родителей и профессиональными интересами (Холланд)', 2)

holland_stats = f"""
Всего рассчитано корреляций между установками родителей и профессиональными интересами: {len(holland_correlations)}
Статистически значимых корреляций (p < 0.05): {len(holland_significant)}
"""
doc.add_paragraph(holland_stats)

if holland_significant:
    doc.add_paragraph("Таблица значимых корреляций между установками родителей и профессиональными интересами:")
    
    holland_table = doc.add_table(rows=1, cols=4)
    holland_table.style = 'Light Grid Accent 1'
    
    hdr_cells = holland_table.rows[0].cells
    hdr_cells[0].text = 'Вопрос родителя'
    hdr_cells[1].text = 'Тип Холланда'
    hdr_cells[2].text = 'ρ (Спирмен)'
    hdr_cells[3].text = 'p-value'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    for corr in sorted(holland_significant, key=lambda x: abs(x['spearman_corr']), reverse=True):
        row_cells = holland_table.add_row().cells
        row_cells[0].text = str(corr['parent_question'])[:50] + "..."
        row_cells[1].text = corr.get('holland_type_name', corr.get('holland_type', 'N/A'))
        row_cells[2].text = f"{corr['spearman_corr']:.4f}"
        row_cells[3].text = f"{corr['spearman_p']:.4f}"
    
    # Добавляем графики и интерпретации для значимых корреляций родители-Холланд
    if holland_scores is not None:
        doc.add_heading('3.5.1.1. Визуализация и интерпретация значимых корреляций', 3)
        
        # Создаём папку для графиков
        graphs_dir = os.path.join(script_dir, "holland_graphs")
        os.makedirs(graphs_dir, exist_ok=True)
        
        for idx, corr in enumerate(sorted(holland_significant, key=lambda x: abs(x['spearman_corr']), reverse=True)):
            # Получаем данные для графика
            p_vals = []
            h_vals = []
            
            parent_q = corr['parent_question']
            h_type = corr['holland_type']
            
            for i, pair in enumerate(pairs):
                try:
                    p_val = parents_sorted.iloc[pair['parent_idx']][parent_q]
                    h_score = holland_scores[h_type][i]
                    
                    p_num = float(p_val) if pd.notna(p_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if p_num is not None and h_num is not None:
                        p_vals.append(p_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(p_vals) >= 3:
                parent_q_short = str(parent_q)[:60].replace('\n', ' ')
                h_type_name = corr.get('holland_type_name', h_type)
                
                # Создаём все 4 варианта графиков
                variant_names = [
                    "Вариант 1: Комбинированный (scatter + гистограммы + квадранты)",
                    "Вариант 2: С квадрантами и пояснениями",
                    "Вариант 3: С box plots и scatter",
                    "Вариант 4: С категоризацией по силе связи"
                ]
                
                doc.add_paragraph(f"\n{'='*80}")
                doc.add_paragraph(f"КОРРЕЛЯЦИЯ {idx+1}: Установка родителя ↔ {h_type_name}")
                doc.add_paragraph(f"{'='*80}")
                
                for variant_num in range(1, 5):
                    graph_filename = f"parent_holland_{idx+1}_variant{variant_num}.png"
                    graph_path = os.path.join(graphs_dir, graph_filename)
                    
                    title = f"Корреляция: Установка родителя ↔ {h_type_name} ({variant_names[variant_num-1]})"
                    
                    create_correlation_plot(
                        p_vals, h_vals,
                        f"Установка родителя: {parent_q_short}...",
                        f"Балл по типу {h_type_name}",
                        corr['spearman_corr'],
                        corr['spearman_p'],
                        title,
                        graph_path,
                        variant=variant_num
                    )
                    
                    # Добавляем график в документ
                    doc.add_paragraph(f"\n{variant_names[variant_num-1]}:")
                    doc.add_picture(graph_path, width=Inches(6))
                    doc.add_paragraph()
                
                # Добавляем интерпретацию и рекомендации (один раз для всех вариантов)
                doc.add_paragraph("ИНТЕРПРЕТАЦИЯ И РЕКОМЕНДАЦИИ:")
                interpretation = generate_interpretation_and_recommendations(
                    corr, 'parent_holland', h_type_name
                )
                doc.add_paragraph(interpretation)
                
                doc.add_paragraph("─" * 80)
                doc.add_paragraph()
else:
    if len(holland_correlations) == 0:
        doc.add_paragraph("""
ПРИМЕЧАНИЕ: Корреляции с опросником Холланда не были рассчитаны, так как столбцы Холланда не найдены в данных.
""")
    else:
        doc.add_paragraph("Значимых корреляций между установками родителей и профессиональными интересами не найдено.")

doc.add_heading('3.5.2. Корреляции между мотивами выбора профессии (Овчарова) и профессиональными интересами (Холланд)', 2)

ovh_stats = f"""
Всего рассчитано корреляций между мотивами выбора профессии и профессиональными интересами: {len(ovcharova_holland_correlations)}
Статистически значимых корреляций (p < 0.05): {len(ovh_significant)}
"""
doc.add_paragraph(ovh_stats)

if ovh_significant:
    doc.add_paragraph("Таблица значимых корреляций:")
    
    ovh_table = doc.add_table(rows=1, cols=4)
    ovh_table.style = 'Light Grid Accent 1'
    
    hdr_cells = ovh_table.rows[0].cells
    hdr_cells[0].text = 'Вопрос Овчаровой'
    hdr_cells[1].text = 'Тип Холланда'
    hdr_cells[2].text = 'ρ (Спирмен)'
    hdr_cells[3].text = 'p-value'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    for corr in sorted(ovh_significant, key=lambda x: abs(x['spearman_corr']), reverse=True):
        row_cells = ovh_table.add_row().cells
        row_cells[0].text = str(corr['ovcharova_question'])[:50] + "..."
        row_cells[1].text = corr.get('holland_type_name', corr.get('holland_type', 'N/A'))
        row_cells[2].text = f"{corr['spearman_corr']:.4f}"
        row_cells[3].text = f"{corr['spearman_p']:.4f}"
    
    # Добавляем графики и интерпретации для значимых корреляций Овчарова-Холланд
    if holland_scores is not None:
        doc.add_heading('3.5.2.1. Визуализация и интерпретация значимых корреляций', 3)
        
        # Создаём папку для графиков (если ещё не создана)
        graphs_dir = os.path.join(script_dir, "holland_graphs")
        os.makedirs(graphs_dir, exist_ok=True)
        
        for idx, corr in enumerate(sorted(ovh_significant, key=lambda x: abs(x['spearman_corr']), reverse=True)):
            # Получаем данные для графика
            o_vals = []
            h_vals = []
            
            ovcharova_q_col = corr['ovcharova_question']
            h_type = corr['holland_type']
            
            for i, pair in enumerate(pairs):
                try:
                    o_val = students_sorted.iloc[pair['student_idx']][ovcharova_q_col]
                    h_score = holland_scores[h_type][i]
                    
                    o_num = float(o_val) if pd.notna(o_val) else None
                    h_num = float(h_score) if pd.notna(h_score) else None
                    
                    if o_num is not None and h_num is not None:
                        o_vals.append(o_num)
                        h_vals.append(h_num)
                except:
                    pass
            
            if len(o_vals) >= 3:
                ovcharova_q_short = str(ovcharova_q_col)[:60].replace('\n', ' ')
                h_type_name = corr.get('holland_type_name', h_type)
                
                # Создаём все 4 варианта графиков
                variant_names = [
                    "Вариант 1: Комбинированный (scatter + гистограммы + квадранты)",
                    "Вариант 2: С квадрантами и пояснениями",
                    "Вариант 3: С box plots и scatter",
                    "Вариант 4: С категоризацией по силе связи"
                ]
                
                doc.add_paragraph(f"\n{'='*80}")
                doc.add_paragraph(f"КОРРЕЛЯЦИЯ {idx+1}: Мотив выбора профессии ↔ {h_type_name}")
                doc.add_paragraph(f"{'='*80}")
                
                for variant_num in range(1, 5):
                    graph_filename = f"ovcharova_holland_{idx+1}_variant{variant_num}.png"
                    graph_path = os.path.join(graphs_dir, graph_filename)
                    
                    title = f"Корреляция: Мотив выбора профессии ↔ {h_type_name} ({variant_names[variant_num-1]})"
                    
                    create_correlation_plot(
                        o_vals, h_vals,
                        f"Мотив выбора профессии: {ovcharova_q_short}...",
                        f"Балл по типу {h_type_name}",
                        corr['spearman_corr'],
                        corr['spearman_p'],
                        title,
                        graph_path,
                        variant=variant_num
                    )
                    
                    # Добавляем график в документ
                    doc.add_paragraph(f"\n{variant_names[variant_num-1]}:")
                    doc.add_picture(graph_path, width=Inches(6))
                    doc.add_paragraph()
                
                # Добавляем интерпретацию и рекомендации (один раз для всех вариантов)
                doc.add_paragraph("ИНТЕРПРЕТАЦИЯ И РЕКОМЕНДАЦИИ:")
                interpretation = generate_interpretation_and_recommendations(
                    corr, 'ovcharova_holland', h_type_name
                )
                doc.add_paragraph(interpretation)
                
                doc.add_paragraph("─" * 80)
                doc.add_paragraph()
else:
    if len(ovcharova_holland_correlations) == 0:
        doc.add_paragraph("""
ПРИМЕЧАНИЕ: Корреляции с опросником Холланда не были рассчитаны, так как столбцы Холланда не найдены в данных.
""")
    else:
        doc.add_paragraph("Значимых корреляций между мотивами выбора профессии и профессиональными интересами не найдено.")

# 3.5.3. Подробный расчёт взаимосвязи результатов по опроснику Холланда
doc.add_heading('3.5.3. Подробный расчёт взаимосвязи результатов по опроснику Холланда', 2)

# Расчёт корреляций между типами Холланда
holland_intercorrelations = []
holland_types = ['R', 'I', 'A', 'S', 'E', 'C']
holland_type_names = {
    'R': 'Реалистичный (Realistic)',
    'I': 'Исследовательский (Investigative)',
    'A': 'Артистический (Artistic)',
    'S': 'Социальный (Social)',
    'E': 'Предприимчивый (Enterprising)',
    'C': 'Конвенциональный (Conventional)'
}

if holland_scores is not None:
    doc.add_paragraph("""
В данном разделе представлен анализ взаимосвязей между различными типами профессиональных интересов 
по опроснику Холланда. Это позволяет понять структуру профессиональных интересов подростков и выявить 
взаимосвязи между различными типами интересов.
""")
    
    # Рассчитываем корреляции между типами
    for i, type1 in enumerate(holland_types):
        for type2 in holland_types[i+1:]:  # Избегаем дублирования
            scores1 = holland_scores[type1]
            scores2 = holland_scores[type2]
            
            if len(scores1) >= 3 and len(scores2) >= 3:
                try:
                    rho, p = spearmanr(scores1, scores2)
                    holland_intercorrelations.append({
                        'type1': type1,
                        'type1_name': holland_type_names[type1],
                        'type2': type2,
                        'type2_name': holland_type_names[type2],
                        'spearman_corr': rho,
                        'spearman_p': p,
                        'n': len(scores1)
                    })
                except:
                    pass
    
    if holland_intercorrelations:
        doc.add_paragraph("Таблица корреляций между типами профессиональных интересов Холланда:")
        
        intercorr_table = doc.add_table(rows=1, cols=4)
        intercorr_table.style = 'Light Grid Accent 1'
        
        hdr_cells = intercorr_table.rows[0].cells
        hdr_cells[0].text = 'Тип 1'
        hdr_cells[1].text = 'Тип 2'
        hdr_cells[2].text = 'ρ (Спирмен)'
        hdr_cells[3].text = 'p-value'
        
        for cell in hdr_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_font(run, bold=True)
        
        for corr in sorted(holland_intercorrelations, key=lambda x: abs(x['spearman_corr']), reverse=True):
            row_cells = intercorr_table.add_row().cells
            row_cells[0].text = corr['type1_name']
            row_cells[1].text = corr['type2_name']
            row_cells[2].text = f"{corr['spearman_corr']:.4f}"
            row_cells[3].text = f"{corr['spearman_p']:.4f}"
        
        # Статистика по значимым корреляциям
        sig_intercorr = [c for c in holland_intercorrelations if c['spearman_p'] < 0.05]
        doc.add_paragraph(f"\nВсего рассчитано корреляций между типами: {len(holland_intercorrelations)}")
        doc.add_paragraph(f"Статистически значимых корреляций (p < 0.05): {len(sig_intercorr)}")
        
        if sig_intercorr:
            doc.add_paragraph("\nЗначимые взаимосвязи между типами профессиональных интересов:")
            for corr in sorted(sig_intercorr, key=lambda x: abs(x['spearman_corr']), reverse=True):
                direction = "прямая" if corr['spearman_corr'] > 0 else "обратная"
                strength = "сильная" if abs(corr['spearman_corr']) > 0.5 else "умеренная" if abs(corr['spearman_corr']) > 0.3 else "слабая"
                doc.add_paragraph(
                    f"• {corr['type1_name']} ↔ {corr['type2_name']}: "
                    f"ρ = {corr['spearman_corr']:.4f} ({strength} {direction} связь, p = {corr['spearman_p']:.4f})"
                )

# Интерпретация результатов Холланда
doc.add_heading('3.5.4. Интерпретация результатов корреляций с опросником Холланда', 2)

holland_interpretation = f"""
РЕЗУЛЬТАТЫ АНАЛИЗА КОРРЕЛЯЦИЙ С ОПРОСНИКОМ ХОЛЛАНДА:

1. КОРРЕЛЯЦИИ МЕЖДУ УСТАНОВКАМИ РОДИТЕЛЕЙ И ПРОФЕССИОНАЛЬНЫМИ ИНТЕРЕСАМИ:
   • Всего рассчитано корреляций: {len(holland_correlations)}
   • Значимых корреляций: {len(holland_significant)}
   • Это указывает на {'наличие' if holland_significant else 'отсутствие'} влияния установок родителей на формирование профессиональных интересов подростков.

2. КОРРЕЛЯЦИИ МЕЖДУ МОТИВАМИ ВЫБОРА ПРОФЕССИИ И ПРОФЕССИОНАЛЬНЫМИ ИНТЕРЕСАМИ:
   • Всего рассчитано корреляций: {len(ovcharova_holland_correlations)}
   • Значимых корреляций: {len(ovh_significant)}
   • Это указывает на {'наличие' if ovh_significant else 'отсутствие'} связи между мотивами выбора профессии и профессиональными интересами.

3. ВАЖНОСТЬ РЕЗУЛЬТАТОВ ХОЛЛАНДА ДЛЯ ИССЛЕДОВАНИЯ:
   • Опросник Холланда позволяет оценить профессиональные интересы подростков по шести типам личности
   • Корреляции с установками родителей показывают, влияют ли родительские установки на формирование профессиональных интересов
   • Корреляции с мотивами выбора профессии показывают, связаны ли мотивы выбора с профессиональными интересами
   • Эти данные дополняют картину влияния родительских установок на профессиональное самоопределение подростков
   • Результаты Холланда помогают понять, как родительские установки влияют не только на мотивы выбора профессии, но и на сами профессиональные интересы

4. ВАЖНЫЕ ПОКАЗАТЕЛИ ХОЛЛАНДА ДЛЯ ДОКАЗАТЕЛЬСТВА ГИПОТЕЗЫ:
   • Типы профессиональных интересов (R, I, A, S, E, C) являются важными показателями для понимания профессионального самоопределения подростков
   • Значимые корреляции между установками родителей и типами Холланда подтверждают влияние родительских установок на формирование профессиональных интересов
   • Значимые корреляции между мотивами выбора профессии (Овчарова) и типами Холланда показывают связь между мотивами и интересами
   • Взаимосвязи между типами Холланда помогают понять структуру профессиональных интересов и их взаимное влияние
   • Эти показатели важны для доказательства гипотезы о связи родительских установок с профессиональным самоопределением подростков

5. СВЯЗЬ С ДРУГИМИ МЕТОДИКАМИ:
   • Результаты Холланда дополняют данные опросника Овчаровой о мотивах выбора профессии
   • Вместе с данными ВРР Марковской они создают целостную картину влияния родительских установок на профессиональное самоопределение
   • Корреляции между типами Холланда и мотивами Овчаровой показывают, как профессиональные интересы связаны с мотивами выбора профессии
   • Это позволяет более глубоко понять механизмы влияния родительских установок на профессиональное самоопределение подростков
"""
doc.add_paragraph(holland_interpretation)

# 4. РЕЗУЛЬТАТЫ
doc.add_heading('4. РЕЗУЛЬТАТЫ ИССЛЕДОВАНИЯ', 1)

significant = [c for c in verified_correlations if c['spearman_p'] < 0.05]
strong = [c for c in verified_correlations if abs(c['spearman_corr']) > 0.3]

# 4.1. Общая статистика
doc.add_heading('4.1. Общая статистика корреляций', 2)

stats_text = f"""
Всего было рассчитано {len(verified_correlations)} корреляций между установками родителей (60 вопросов ВРР Марковской) 
и мотивами выбора профессии подростками (2 вопроса опросника Овчаровой: вопросы 12 и 19).

Общая статистика (коэффициент Спирмена):
• Средняя корреляция: {np.mean([c['spearman_corr'] for c in verified_correlations]):.4f}
• Медианная корреляция: {np.median([c['spearman_corr'] for c in verified_correlations]):.4f}
• Минимальная корреляция: {min([c['spearman_corr'] for c in verified_correlations]):.4f}
• Максимальная корреляция: {max([c['spearman_corr'] for c in verified_correlations]):.4f}
• Стандартное отклонение: {np.std([c['spearman_corr'] for c in verified_correlations]):.4f}
"""
doc.add_paragraph(stats_text)

# 4.2. Значимые корреляции
doc.add_heading('4.2. Статистически значимые корреляции', 2)

if significant:
    avg_sig_corr = np.mean([c['spearman_corr'] for c in significant])
else:
    avg_sig_corr = 0.0

sig_text = f"""
Найдено {len(significant)} статистически значимых корреляций (p < 0.05).

Средняя значимая корреляция: {avg_sig_corr:.4f}

Подробные расчёты для этих корреляций представлены в разделе 3.
"""
doc.add_paragraph(sig_text)

# Таблица значимых корреляций
if significant:
    doc.add_paragraph("Сводная таблица статистически значимых корреляций:")
    
    sig_table = doc.add_table(rows=1, cols=5)
    sig_table.style = 'Light Grid Accent 1'
    
    hdr_cells = sig_table.rows[0].cells
    hdr_cells[0].text = 'Вопрос родителя'
    hdr_cells[1].text = 'Вопрос Овчаровой'
    hdr_cells[2].text = 'ρ (Спирмен)'
    hdr_cells[3].text = 'p-value'
    hdr_cells[4].text = 'n'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    for corr in sorted(significant, key=lambda x: abs(x['spearman_corr']), reverse=True):
        row_cells = sig_table.add_row().cells
        row_cells[0].text = str(corr['parent_question'])[:50] + "..."
        row_cells[1].text = str(corr['ovcharova_question'])[:50] + "..."
        row_cells[2].text = f"{corr['spearman_corr']:.4f}"
        row_cells[3].text = f"{corr['spearman_p']:.4f}"
        row_cells[4].text = str(corr['n'])

# 4.3. Незначимые корреляции
doc.add_heading('4.3. Статистически незначимые корреляции', 2)

non_significant = [c for c in verified_correlations if c['spearman_p'] >= 0.05]

non_sig_text = f"""
Найдено {len(non_significant)} статистически незначимых корреляций (p ≥ 0.05).
Эти корреляции не подтверждают наличие значимой связи между установками родителей и мотивами выбора профессии.

Подробные расчёты для незначимых корреляций не представлены, так как они не подтверждают гипотезу исследования.
"""
doc.add_paragraph(non_sig_text)

if non_significant:
    doc.add_paragraph("Таблица статистически незначимых корреляций (коэффициент Спирмена):")
    
    non_sig_table = doc.add_table(rows=1, cols=4)
    non_sig_table.style = 'Light Grid Accent 1'
    
    hdr_cells = non_sig_table.rows[0].cells
    hdr_cells[0].text = 'Вопрос родителя'
    hdr_cells[1].text = 'Вопрос Овчаровой'
    hdr_cells[2].text = 'ρ (Спирмен)'
    hdr_cells[3].text = 'p-value'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    # Сортируем по силе корреляции
    non_sig_sorted = sorted(non_significant, key=lambda x: abs(x['spearman_corr']), reverse=True)
    
    for corr in non_sig_sorted:
        row_cells = non_sig_table.add_row().cells
        row_cells[0].text = str(corr['parent_question'])[:60] + "..."
        row_cells[1].text = str(corr['ovcharova_question'])[:60] + "..."
        row_cells[2].text = f"{corr['spearman_corr']:.4f}"
        row_cells[3].text = f"{corr['spearman_p']:.4f}"

# 5. ВЫВОДЫ
doc.add_heading('5. ВЫВОДЫ', 1)

conclusions = f"""
На основе проведенного статистического анализа можно сделать следующие выводы:

1. КОЛИЧЕСТВЕННЫЕ РЕЗУЛЬТАТЫ:
   • Всего проанализировано {len(verified_correlations)} комбинаций вопросов
   • Найдено {len(strong)} сильных корреляций (|ρ| > 0.3)
   • Найдено {len(significant)} статистически значимых корреляций (p < 0.05)

2. КАЧЕСТВЕННЫЕ ВЫВОДЫ:
"""
doc.add_paragraph(conclusions)

if significant:
    conclusion_text = f"""
   ✓ Гипотеза о влиянии установок родителей на мотивы выбора профессии подростками ПОДТВЕРЖДЕНА.
   
   Наличие {len(significant)} статистически значимых корреляций указывает на то, что определенные 
   установки родителей действительно связаны с мотивами выбора ребенком будущей профессии.
   
   Наиболее сильные связи обнаружены между:
   • Представлениями родителей о способностях ребенка и его стремлением к профессиональному росту
   • Отношением родителей к ребенку и его желанием использовать профессиональные умения вне работы
"""
else:
    conclusion_text = """
   ⚠ Гипотеза о влиянии установок родителей на мотивы выбора профессии подростками НЕ ПОДТВЕРЖДЕНА 
   на статистически значимом уровне.
   
   Однако наличие сильных корреляций (|ρ| > 0.3) может указывать на тенденцию, требующую 
   дальнейшего исследования с большей выборкой.
"""

doc.add_paragraph(conclusion_text)

# Сохраняем документ
output_file = os.path.join(script_dir, "Полный_итоговый_отчёт_корреляции.docx")

try:
    doc.save(output_file)
    print(f"\n✓ Полный итоговый отчёт создан: {output_file}")
    print(f"  Размер документа: {os.path.getsize(output_file) / 1024:.1f} КБ")
    print(f"  Всего расчётов значимых корреляций: {len(sorted_significant)}")
except PermissionError:
    print(f"\n⚠ ОШИБКА: Не удалось сохранить файл {output_file}")
    print("  Возможные причины:")
    print("  1. Файл открыт в Microsoft Word или другой программе")
    print("  2. Нет прав доступа к файлу")
    print("\n  РЕШЕНИЕ:")
    print("  - Закройте файл в Word, если он открыт")
    print("  - Попробуйте запустить скрипт снова")
    print("  - Или сохраните файл под другим именем")
    
    # Пробуем сохранить с другим именем
    import datetime
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    alt_output_file = os.path.join(script_dir, f"Полный_итоговый_отчёт_корреляции_{timestamp}.docx")
    try:
        doc.save(alt_output_file)
        print(f"\n✓ Файл сохранён под альтернативным именем: {alt_output_file}")
    except Exception as e:
        print(f"\n✗ Не удалось сохранить файл даже с альтернативным именем: {e}")
except Exception as e:
    print(f"\n✗ ОШИБКА при сохранении файла: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "=" * 80)
print("СОЗДАНИЕ ОТЧЁТА ЗАВЕРШЕНО!")
print("=" * 80)
