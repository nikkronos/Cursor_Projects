"""Извлечение примеров расчётов из существующего отчёта"""
import os
from docx import Document

script_dir = os.path.dirname(os.path.abspath(__file__))
parent_dir = os.path.dirname(script_dir)
newtest_dir = os.path.join(parent_dir, "newtest")

report_file = os.path.join(newtest_dir, "Отчет_Анализ_Корреляций.docx")

print("Чтение документа отчёта...")
doc = Document(report_file)

print(f"Всего параграфов: {len(doc.paragraphs)}")
print(f"Всего таблиц: {len(doc.tables)}")

# Ищем примеры расчётов (страницы 7-12)
print("\n" + "=" * 80)
print("ПРИМЕРЫ РАСЧЁТОВ (страницы 7-12)")
print("=" * 80)

# Ищем раздел с примерами
in_examples_section = False
example_count = 0

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    # Ищем начало раздела с примерами
    if "ПОДРОБНЫЙ РАСЧЕТ" in text.upper() or "ПРИМЕР РАСЧЕТА" in text.upper() or \
       "ШАГ 1" in text.upper() or "ШАГ 2" in text.upper():
        in_examples_section = True
        example_count += 1
        print(f"\n--- Пример {example_count} ---")
    
    if in_examples_section and text:
        # Ограничиваем вывод первыми 200 символами
        if len(text) > 200:
            print(f"{text[:200]}...")
        else:
            print(text)
        
        # Если встретили следующий большой раздел, останавливаемся
        if text.startswith("4.") or text.startswith("РЕЗУЛЬТАТЫ") or text.startswith("ВЫВОДЫ"):
            break

# Ищем таблицы в примерах
print("\n" + "=" * 80)
print("ТАБЛИЦЫ В ПРИМЕРАХ")
print("=" * 80)

for i, table in enumerate(doc.tables):
    print(f"\nТаблица {i+1}:")
    print(f"  Строк: {len(table.rows)}, Столбцов: {len(table.columns) if table.rows else 0}")
    
    if table.rows:
        # Первые 3 строки
        for row_idx in range(min(3, len(table.rows))):
            row_data = [cell.text.strip()[:50] for cell in table.rows[row_idx].cells]
            print(f"  Строка {row_idx+1}: {row_data}")

print("\nГотово!")
