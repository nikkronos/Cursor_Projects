"""
Генерация детальных расчётов корреляций в формате примеров
Создаёт полные пошаговые расчёты для всех корреляций
"""
import os
import pickle
import pandas as pd
import numpy as np
from scipy.stats import rankdata
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

script_dir = os.path.dirname(os.path.abspath(__file__))

# Загружаем детальные данные
detailed_file = os.path.join(script_dir, "detailed_correlation_data.pkl")
if not os.path.exists(detailed_file):
    print(f"ОШИБКА: Файл {detailed_file} не найден. Сначала запустите comprehensive_analysis.py")
    exit(1)

print("Загрузка детальных данных...")
with open(detailed_file, 'rb') as f:
    correlation_data = pickle.load(f)

print(f"Загружено корреляций: {len(correlation_data)}")

# Создаём документ Word
doc = Document()

# Настройка стилей
def set_font(run, font_name='Times New Roman', size=12, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# Заголовок
title = doc.add_heading('ПОЛНЫЕ РАСЧЁТЫ КОРРЕЛЯЦИЙ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Пустая строка

# Введение
intro = doc.add_paragraph()
intro.add_run("В данном разделе представлены полные пошаговые расчёты корреляций между установками родителей (вопросы ВРР Марковской) и мотивами выбора профессии подростками (вопросы опросника Овчаровой).")
set_font(intro.runs[0], size=12)

doc.add_paragraph()

# Сортируем по силе корреляции
sorted_data = sorted(correlation_data, key=lambda x: abs(x['spearman_corr']), reverse=True)

# Генерируем расчёты для всех корреляций
print("\nГенерация детальных расчётов...")

for idx, corr in enumerate(sorted_data, 1):
    # Заголовок расчёта
    heading_text = f"Расчёт {idx}. Корреляция между вопросом родителя и вопросом Овчаровой"
    doc.add_heading(heading_text, 1)
    
    # Описание
    desc_para = doc.add_paragraph()
    desc_para.add_run(f"Вопрос родителя (ВРР Марковской): ").bold = True
    desc_para.add_run(f"{corr['parent_question']}")
    
    desc_para2 = doc.add_paragraph()
    desc_para2.add_run(f"Вопрос Овчаровой (мотив выбора профессии): ").bold = True
    desc_para2.add_run(f"{corr['ovcharova_question']}")
    
    doc.add_paragraph()
    
    # ШАГ 1: Сбор данных
    doc.add_heading("ШАГ 1: Сбор данных по всем парам", 2)
    
    step1_text = f"""
Для данной комбинации вопросов были собраны ответы всех {corr['n']} пар родитель-ребенок.
Каждая пара представляет собой ответы одного родителя и его ребенка на соответствующие вопросы.
"""
    doc.add_paragraph(step1_text)
    
    # Таблица с примерами данных (первые 10 пар)
    doc.add_paragraph("Пример первых 10 пар данных:")
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    
    # Заголовки
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Пара №'
    hdr_cells[1].text = 'Ответ родителя'
    hdr_cells[2].text = 'Ответ подростка'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    # Данные
    parent_vals = corr['parent_values']
    ovcharova_vals = corr['ovcharova_values']
    
    for i in range(min(10, len(parent_vals))):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = f"{parent_vals[i]:.0f}"
        row_cells[2].text = f"{ovcharova_vals[i]:.0f}"
    
    doc.add_paragraph()
    
    # ШАГ 2: Преобразование в ранги
    doc.add_heading("ШАГ 2: Преобразование в ранги", 2)
    
    parent_ranks = rankdata(parent_vals)
    ovcharova_ranks = rankdata(ovcharova_vals)
    
    step2_text = f"""
Каждому ответу присваивается ранг от 1 до {corr['n']} (где {corr['n']} - количество наблюдений).
Если несколько значений одинаковы, им присваивается средний ранг.

Пример ранжирования для ответов родителей:
• Минимальное значение: {min(parent_vals):.0f} → ранг 1
• Максимальное значение: {max(parent_vals):.0f} → ранг {corr['n']}
• Среднее значение: {np.mean(parent_vals):.2f}

Пример ранжирования для ответов подростков:
• Минимальное значение: {min(ovcharova_vals):.0f} → ранг 1
• Максимальное значение: {max(ovcharova_vals):.0f} → ранг {corr['n']}
• Среднее значение: {np.mean(ovcharova_vals):.2f}
"""
    doc.add_paragraph(step2_text)
    
    # Таблица с рангами (первые 10 пар)
    doc.add_paragraph("Пример ранжирования для первых 10 пар:")
    
    rank_table = doc.add_table(rows=1, cols=5)
    rank_table.style = 'Light Grid Accent 1'
    
    hdr_cells = rank_table.rows[0].cells
    hdr_cells[0].text = 'Пара №'
    hdr_cells[1].text = 'Ранг родителя'
    hdr_cells[2].text = 'Ранг подростка'
    hdr_cells[3].text = 'd (разность)'
    hdr_cells[4].text = 'd²'
    
    for cell in hdr_cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_font(run, bold=True)
    
    for i in range(min(10, len(parent_vals))):
        d = parent_ranks[i] - ovcharova_ranks[i]
        d_squared = d ** 2
        
        row_cells = rank_table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = f"{parent_ranks[i]:.1f}"
        row_cells[2].text = f"{ovcharova_ranks[i]:.1f}"
        row_cells[3].text = f"{d:.1f}"
        row_cells[4].text = f"{d_squared:.1f}"
    
    doc.add_paragraph()
    
    # ШАГ 3: Расчёт разностей рангов
    doc.add_heading("ШАГ 3: Расчёт разностей рангов и их квадратов", 2)
    
    diffs = parent_ranks - ovcharova_ranks
    diffs_squared = diffs ** 2
    sum_d_squared = np.sum(diffs_squared)
    
    step3_text = f"""
Для каждой пары вычисляется разность рангов: d = ранг_родителя - ранг_подростка
Затем вычисляется квадрат разности: d²

Сумма квадратов разностей: Σd² = {sum_d_squared:.2f}
"""
    doc.add_paragraph(step3_text)
    
    # ШАГ 4: Применение формулы Спирмена
    doc.add_heading("ШАГ 4: Применение формулы коэффициента корреляции Спирмена", 2)
    
    n = corr['n']
    formula_text = f"""
Формула коэффициента Спирмена:

ρ = 1 - (6 × Σd²) / (n × (n² - 1))

где:
• n = {n} (количество пар наблюдений)
• Σd² = {sum_d_squared:.2f} (сумма квадратов разностей рангов)
• n² - 1 = {n**2 - 1}
• n × (n² - 1) = {n * (n**2 - 1)}
• 6 × Σd² = {6 * sum_d_squared:.2f}

Подстановка значений:

ρ = 1 - ({6 * sum_d_squared:.2f}) / ({n * (n**2 - 1)})
ρ = 1 - {6 * sum_d_squared / (n * (n**2 - 1)):.6f}
ρ = {1 - 6 * sum_d_squared / (n * (n**2 - 1)):.4f}
"""
    doc.add_paragraph(formula_text)
    
    # Результат
    doc.add_heading("Результат расчёта", 2)
    
    result_text = f"""
• Коэффициент Спирмена: ρ = {corr['spearman_corr']:.4f}
• Уровень значимости: p = {corr['spearman_p']:.4f}
• Коэффициент Пирсона: r = {corr['pearson_corr']:.4f}
• Уровень значимости (Пирсон): p = {corr['pearson_p']:.4f}
• Количество пар: n = {corr['n']}
"""
    doc.add_paragraph(result_text)
    
    # Интерпретация
    doc.add_heading("Интерпретация результата", 2)
    
    rho_abs = abs(corr['spearman_corr'])
    if rho_abs < 0.2:
        strength = "очень слабая"
    elif rho_abs < 0.4:
        strength = "слабая"
    elif rho_abs < 0.6:
        strength = "умеренная"
    elif rho_abs < 0.8:
        strength = "сильная"
    else:
        strength = "очень сильная"
    
    if corr['spearman_corr'] > 0:
        direction = "прямая"
    else:
        direction = "обратная"
    
    significance = ""
    if corr['spearman_p'] < 0.001:
        significance = "Статистически высоко значима (p < 0.001)."
    elif corr['spearman_p'] < 0.01:
        significance = "Статистически значима (p < 0.01)."
    elif corr['spearman_p'] < 0.05:
        significance = "Статистически значима (p < 0.05)."
    else:
        significance = "Не является статистически значимой (p ≥ 0.05)."
    
    interpretation_text = f"""
Коэффициент корреляции ρ = {corr['spearman_corr']:.4f} указывает на {strength} {direction} связь между установками родителей и мотивами выбора профессии подростками.

{significance}

"""
    doc.add_paragraph(interpretation_text)
    
    # Разделитель
    if idx < len(sorted_data):
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()

# Сохраняем документ
output_file = os.path.join(script_dir, "Полные_расчёты_корреляций.docx")
doc.save(output_file)

print(f"\nДокумент с полными расчётами создан: {output_file}")
print(f"Всего расчётов: {len(sorted_data)}")
