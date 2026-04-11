# -*- coding: utf-8 -*-
"""
Упрощённый скрипт для исправления документа v20 -> v30:
1. Извлечение списка литературы из документа в отдельный файл (для проверки)
2. Сортировка списка литературы по алфавиту (русские → английские)
3. Обновление всех ссылок в тексте документа
4. Замена списка литературы в документе на отсортированный

Работает по принципу старых версий: сначала список в отдельном файле, затем обновление документа.
"""
from pathlib import Path
from docx import Document
import re
from typing import Dict, List, Tuple

def extract_author_surname(entry: str) -> str:
    """Извлекает фамилию первого автора для сортировки."""
    entry = re.sub(r'^\d+\.\s*', '', entry.strip())
    
    # Русские фамилии
    match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
    if match:
        return match.group(1).upper()
    
    # Английские фамилии
    match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z]\.?)*)', entry)
    if match:
        return match.group(1).split()[0].upper()
    
    # Если не нашлось, пробуем первую заглавную букву
    match = re.search(r'([А-ЯЁA-Z])', entry)
    if match:
        return match.group(1)
    
    return 'Я'

def extract_bibliography_to_file(doc: Document, output_path: Path) -> Dict[int, str]:
    """
    Извлекает список литературы из документа и сохраняет в файл.
    Возвращает словарь {номер: запись}.
    """
    bibliography = {}
    in_bibliography = False
    current_entry = None
    current_number = None
    
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        
        # Ищем начало списка литературы
        if not in_bibliography:
            if 'СПИСОК ЛИТЕРАТУРЫ' in text.upper() or 'ЛИТЕРАТУРА' in text.upper():
                in_bibliography = True
                continue
        
        if not in_bibliography:
            continue
        
        # Если мы в разделе библиографии, ищем записи
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            entry_text = match.group(2).strip()
            
            # Пропускаем записи, которые выглядят как текст, а не библиография
            # (начинаются с глаголов или общих слов типа "Исследование")
            forbidden_words = ['исследование', 'методология', 'вывод', 'заключение', 
                             'результат', 'задача', 'цель', 'разработать', 'изучить',
                             'провести', 'проанализировать', 'применение', 'использование']
            
            first_words = re.findall(r'\b\w+\b', entry_text.lower())
            if first_words and first_words[0] in forbidden_words:
                # Если встретили текст, а не библиографию - конец списка
                break
            
            # Сохраняем предыдущую запись
            if current_number is not None and current_entry:
                bibliography[current_number] = current_entry.strip()
            
            # Начинаем новую запись
            current_number = num
            current_entry = entry_text
        elif text and current_number is not None:
            # Продолжение многострочной записи
            if not re.match(r'^\d+\.\s+', text):
                current_entry += " " + text.strip()
            else:
                bibliography[current_number] = current_entry.strip()
                current_number = None
                current_entry = None
        elif not text:
            # Пустая строка - сохраняем текущую запись
            if current_number is not None and current_entry:
                bibliography[current_number] = current_entry.strip()
                current_number = None
                current_entry = None
        elif text and current_number is None:
            # Проверяем, не конец ли списка (например, "ПРИЛОЖЕНИЯ")
            if re.match(r'^ПРИЛОЖЕНИЯ?\s*$', text.upper().strip()):
                break
    
    # Сохраняем последнюю запись
    if current_number is not None and current_entry:
        bibliography[current_number] = current_entry.strip()
    
    # Фильтруем записи, которые выглядят как текст, а не библиография
    filtered_bibliography = {}
    for num, entry in bibliography.items():
        first_words = re.findall(r'\b\w+\b', entry.lower())
        if first_words and first_words[0] in forbidden_words:
            continue
        if len(entry) > 300:  # Слишком длинные записи - это текст, а не библиография
            continue
        if not re.match(r'^[А-ЯЁA-Z]', entry):  # Должно начинаться с заглавной буквы
            continue
        filtered_bibliography[num] = entry
    
    # Сохраняем в файл
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ\n\n")
        for num in sorted(filtered_bibliography.keys()):
            f.write(f"{num}. {filtered_bibliography[num]}\n\n")
    
    return filtered_bibliography

def parse_literature_list(lit_path: Path) -> Dict[int, str]:
    """Парсит список литературы из файла."""
    literature = {}
    
    if not lit_path.exists():
        print(f"ERROR: Файл списка литературы не найден: {lit_path}")
        return literature
    
    with open(lit_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    current_number = None
    current_entry = []
    
    for line in lines:
        line = line.strip()
        
        if not line or line.startswith('#'):
            continue
        
        match = re.match(r'^(\d+)\.\s*(.+)', line)
        if match:
            if current_number is not None:
                literature[current_number] = ' '.join(current_entry).strip()
            current_number = int(match.group(1))
            current_entry = [match.group(2)]
        elif current_number is not None:
            current_entry.append(line)
    
    if current_number is not None:
        literature[current_number] = ' '.join(current_entry).strip()
    
    return literature

def sort_literature_by_alphabet(literature: Dict[int, str]) -> List[Tuple[int, str]]:
    """Сортирует источники по алфавиту (русские → английские)."""
    sorted_items = []
    
    for old_num, entry in literature.items():
        surname = extract_author_surname(entry)
        sorted_items.append((surname, old_num, entry))
    
    def sort_key(x):
        surname = x[0]
        if surname and surname[0] in 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ':
            return (0, surname)
        elif surname and surname[0].isalpha():
            return (1, surname)
        else:
            return (2, surname)
    
    sorted_items.sort(key=sort_key)
    return [(old_num, entry) for _, old_num, entry in sorted_items]

def create_number_mapping(sorted_items: List[Tuple[int, str]]) -> Dict[int, int]:
    """Создаёт маппинг старый_номер -> новый_номер."""
    mapping = {}
    for new_num, (old_num, _) in enumerate(sorted_items, start=1):
        mapping[old_num] = new_num
    return mapping

def find_all_citations(text: str) -> List[Tuple[int, int, str]]:
    """Находит все ссылки [N] или [N, M] в тексте."""
    citations = []
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    for match in re.finditer(pattern, text):
        citations.append((match.start(), match.end(), match.group(0)))
    return citations

def replace_citation_numbers(citation_text: str, mapping: Dict[int, int]) -> str:
    """Заменяет номера в ссылке согласно маппингу."""
    numbers = re.findall(r'\d+', citation_text)
    if not numbers:
        return citation_text
    
    new_numbers = []
    for num_str in numbers:
        old_num = int(num_str)
        if old_num in mapping:
            new_numbers.append(str(mapping[old_num]))
        else:
            new_numbers.append(num_str)
    
    content = citation_text[1:-1]
    if ',' in content:
        sep = ', '
    elif ';' in content:
        sep = '; '
    elif '-' in content:
        sep = '-'
    else:
        return '[' + new_numbers[0] + ']'
    
    return '[' + sep.join(new_numbers) + ']'

def update_document_citations(doc: Document, mapping: Dict[int, int]) -> int:
    """Обновляет все ссылки в документе."""
    replaced_count = 0
    
    for para in doc.paragraphs:
        text = para.text
        citations = find_all_citations(text)
        if citations:
            for start, end, citation_text in reversed(citations):
                new_citation = replace_citation_numbers(citation_text, mapping)
                if new_citation != citation_text:
                    text = text[:start] + new_citation + text[end:]
                    replaced_count += 1
            if text != para.text:
                para.text = text
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                citations = find_all_citations(text)
                if citations:
                    for start, end, citation_text in reversed(citations):
                        new_citation = replace_citation_numbers(citation_text, mapping)
                        if new_citation != citation_text:
                            text = text[:start] + new_citation + text[end:]
                            replaced_count += 1
                    if text != cell.text:
                        cell.text = text
    
    return replaced_count

def update_bibliography_in_doc(doc: Document, sorted_items: List[Tuple[int, str]], literature: Dict[int, str]):
    """Заменяет список литературы в документе на отсортированный."""
    # Находим начало и конец списка литературы
    bib_start_idx = None
    bib_end_idx = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'СПИСОК ЛИТЕРАТУРЫ' in text or 'ЛИТЕРАТУРА' in text:
            bib_start_idx = i
        elif bib_start_idx is not None and ('ПРИЛОЖЕНИЯ' in text or 'ПРИЛОЖЕНИЕ' in text):
            bib_end_idx = i
            break
    
    if bib_start_idx is None:
        print("  ⚠ Не найден заголовок списка литературы")
        return
    
    if bib_end_idx is None:
        bib_end_idx = len(doc.paragraphs)
    
    # Находим все пронумерованные записи между началом и концом
    entries_to_remove = []
    for i in range(bib_start_idx + 1, bib_end_idx):
        text = doc.paragraphs[i].text.strip()
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            # Проверяем, что это действительно библиографическая запись
            entry_text = match.group(2).strip()
            if num in literature and entry_text and len(entry_text) > 20:
                entries_to_remove.append(i)
    
    # Удаляем старые записи (с конца, чтобы не сбить индексы)
    for i in reversed(entries_to_remove):
        doc.paragraphs[i].clear()
        # Удаляем параграф полностью
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    # Вставляем новые отсортированные записи после заголовка
    insert_idx = bib_start_idx + 1
    for new_num, (old_num, entry) in enumerate(sorted_items, start=1):
        para = doc.paragraphs[insert_idx].insert_paragraph_before()
        para.text = f"{new_num}. {entry}"
        insert_idx += 1

def fix_document():
    version10 = Path(__file__).parent
    version6 = version10.parent / "Version6"
    
    # Пути к файлам
    v20_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
    if not v20_path.exists():
        # Копируем из Version6
        v20_source = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
        if v20_source.exists():
            import shutil
            shutil.copy2(v20_source, v20_path)
            print(f"✓ Скопирован документ из Version6: {v20_path.name}")
        else:
            print(f"ERROR: Документ v20 не найден ни в Version10, ни в Version6")
            return
    
    v30_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v30.docx"
    lit_path = version10 / "СПИСОК_ЛИТЕРАТУРЫ_v20_extracted.md"
    
    print("=" * 70)
    print("УПРОЩЁННОЕ ИСПРАВЛЕНИЕ ДОКУМЕНТА v20 -> v30")
    print("=" * 70)
    print("1. Извлечение списка литературы в отдельный файл")
    print("2. Сортировка списка литературы по алфавиту")
    print("3. Обновление ссылок в тексте документа")
    print("4. Замена списка литературы в документе")
    print("=" * 70)
    print()
    
    # Загружаем документ
    print("📄 Загрузка документа...")
    doc = Document(str(v20_path))
    print(f"✓ Документ загружен: {v20_path.name}\n")
    
    # Извлекаем список литературы
    print("📚 Извлечение списка литературы...")
    bibliography = extract_bibliography_to_file(doc, lit_path)
    print(f"✓ Найдено источников: {len(bibliography)}")
    print(f"✓ Список сохранён в файл: {lit_path.name}\n")
    
    if len(bibliography) == 0:
        print("ERROR: Список литературы пуст. Проверьте документ вручную.")
        return
    
    # Парсим список литературы (для работы с ним)
    print("📖 Парсинг списка литературы...")
    literature = parse_literature_list(lit_path)
    print(f"✓ Распознано источников: {len(literature)}\n")
    
    # Сортируем по алфавиту
    print("🔤 Сортировка списка литературы по алфавиту...")
    sorted_items = sort_literature_by_alphabet(literature)
    print(f"✓ Отсортировано {len(sorted_items)} источников\n")
    
    # Создаём маппинг
    print("🔄 Создание маппинга старых номеров на новые...")
    mapping = create_number_mapping(sorted_items)
    changes_count = sum(1 for old_num, new_num in mapping.items() if old_num != new_num)
    print(f"✓ Изменений номеров: {changes_count}\n")
    
    # Обновляем ссылки в документе
    print("📝 Обновление ссылок в тексте документа...")
    replaced_count = update_document_citations(doc, mapping)
    print(f"✓ Заменено ссылок: {replaced_count}\n")
    
    # Обновляем список литературы в документе
    print("📋 Обновление списка литературы в документе...")
    update_bibliography_in_doc(doc, sorted_items, literature)
    print("✓ Список литературы обновлён\n")
    
    # Сохраняем отсортированный список в файл
    sorted_lit_path = version10 / "СПИСОК_ЛИТЕРАТУРЫ_v30_sorted.md"
    with open(sorted_lit_path, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ (отсортированный)\n\n")
        for new_num, (old_num, entry) in enumerate(sorted_items, start=1):
            f.write(f"{new_num}. {entry}\n\n")
    print(f"✓ Отсортированный список сохранён: {sorted_lit_path.name}\n")
    
    # Сохраняем документ
    print("💾 Сохранение документа...")
    doc.save(str(v30_path))
    print(f"✓ Документ сохранён: {v30_path.name}\n")
    
    print("=" * 70)
    print("РЕЗУЛЬТАТЫ")
    print("=" * 70)
    print(f"✓ Документ сохранён: {v30_path.name}")
    print(f"✓ Отсортированный список: {sorted_lit_path.name}")
    print(f"✓ Всего источников: {len(sorted_items)}")
    print(f"✓ Заменено ссылок в тексте: {replaced_count}")
    print(f"✓ Изменений номеров: {changes_count}")
    print("=" * 70)
    print("\nВАЖНО: Проверьте извлечённый список литературы в файле:")
    print(f"  {lit_path.name}")
    print("\nЕсли там есть не-библиографические записи, удалите их вручную,")
    print("затем переименуйте файл и запустите скрипт заново.\n")

if __name__ == '__main__':
    fix_document()
