# -*- coding: utf-8 -*-
"""
Скрипт для точного определения границ списка литературы
"""
from pathlib import Path
from docx import Document
import re

def find_exact_bibliography():
    version10 = Path(__file__).parent
    v20_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
    
    if not v20_path.exists():
        print(f"ERROR: Документ v20 не найден: {v20_path}")
        return
    
    doc = Document(str(v20_path))
    print("=" * 70)
    print("ТОЧНЫЙ ПОИСК ГРАНИЦ СПИСКА ЛИТЕРАТУРЫ")
    print("=" * 70)
    print(f"Всего параграфов в документе: {len(doc.paragraphs)}\n")
    
    # Находим заголовок списка литературы
    bib_header_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip().upper()
        if any(keyword in text for keyword in ['СПИСОК ЛИТЕРАТУРЫ', 'ЛИТЕРАТУРА', 'БИБЛИОГРАФИЧЕСКИЙ СПИСОК']):
            bib_header_idx = i
            print(f"✓ Заголовок списка литературы на позиции {i}")
            print(f"  Текст: '{doc.paragraphs[i].text[:80]}'\n")
            break
    
    if bib_header_idx is None:
        print("ERROR: Заголовок списка литературы не найден!")
        return
    
    # Ищем все записи вида "N. Автор..." во всём документе после заголовка
    print("=" * 70)
    print("ПОИСК ВСЕХ ЗАПИСЕЙ ВИДА 'N. Автор...' ПОСЛЕ ЗАГОЛОВКА")
    print("=" * 70)
    
    all_entries = []
    for i in range(bib_header_idx + 1, len(doc.paragraphs)):
        text = (doc.paragraphs[i].text or "").strip()
        
        if not text:
            continue
        
        # Игнорируем ссылки в квадратных скобках
        if re.match(r'^\[.*\]\s*$', text):
            continue
        
        # Ищем запись вида "N. Автор Название..."
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            entry_start = match.group(2).strip()[:80]
            # Проверяем, что это похоже на запись списка литературы (начинается с фамилии автора)
            if re.match(r'^[А-ЯЁA-Z]', entry_start):
                all_entries.append((i, num, entry_start))
    
    print(f"\nНайдено всего записей вида 'N. Автор...': {len(all_entries)}\n")
    
    if all_entries:
        print("Первые 20 записей:")
        print("-" * 70)
        for idx, (para_idx, num, entry) in enumerate(all_entries[:20]):
            print(f"[{para_idx:4d}] №{num:3d}: {entry}...")
        
        if len(all_entries) > 20:
            print(f"\n... (пропущено {len(all_entries) - 40} записей) ...\n")
            print("Последние 20 записей:")
            print("-" * 70)
            for idx, (para_idx, num, entry) in enumerate(all_entries[-20:]):
                print(f"[{para_idx:4d}] №{num:3d}: {entry}...")
        
        # Анализируем распределение записей
        print("\n" + "=" * 70)
        print("АНАЛИЗ РАСПРЕДЕЛЕНИЯ ЗАПИСЕЙ")
        print("=" * 70)
        
        # Ищем блоки последовательных записей
        consecutive_blocks = []
        current_block = []
        last_num = None
        
        for para_idx, num, entry in all_entries:
            if last_num is None or num == last_num + 1 or (current_block and num > last_num):
                current_block.append((para_idx, num, entry))
                last_num = num
            else:
                if len(current_block) >= 5:  # Блок из минимум 5 записей
                    consecutive_blocks.append(current_block)
                current_block = [(para_idx, num, entry)]
                last_num = num
        
        # Добавляем последний блок
        if len(current_block) >= 5:
            consecutive_blocks.append(current_block)
        
        if consecutive_blocks:
            print(f"\nНайдено блоков последовательных записей (>= 5 записей): {len(consecutive_blocks)}")
            for idx, block in enumerate(consecutive_blocks, 1):
                print(f"\nБлок {idx}: {len(block)} записей")
                print(f"  Диапазон параграфов: {block[0][0]} - {block[-1][0]}")
                print(f"  Диапазон номеров: {block[0][1]} - {block[-1][1]}")
                print(f"  Первая запись: №{block[0][1]} - {block[0][2]}...")
                print(f"  Последняя запись: №{block[-1][1]} - {block[-1][2]}...")
            
            # Самый большой блок - вероятно, это список литературы
            biggest_block = max(consecutive_blocks, key=len)
            print(f"\n✓ САМЫЙ БОЛЬШОЙ БЛОК (вероятно, список литературы):")
            print(f"  Параграфы: {biggest_block[0][0]} - {biggest_block[-1][0]}")
            print(f"  Номера: {biggest_block[0][1]} - {biggest_block[-1][1]}")
            print(f"  Всего записей: {len(biggest_block)}")
        else:
            print("\n⚠ Не найдено больших блоков последовательных записей")
        
        # Показываем параграфы вокруг предполагаемого списка литературы
        # Используем известные записи из списка (например, "Абульханова-Славская")
        print("\n" + "=" * 70)
        print("ПОИСК КОНКРЕТНЫХ ЗАПИСЕЙ ИЗ СПИСКА ЛИТЕРАТУРЫ")
        print("=" * 70)
        
        search_terms = ["Абульханова-Славская", "Адлер А.", "Ананьев Б.Г.", "Франкл В."]
        found_positions = []
        
        for term in search_terms:
            for para_idx, num, entry in all_entries:
                if term.lower() in entry.lower():
                    found_positions.append((para_idx, num, term))
                    print(f"✓ Найдено '{term}' на позиции {para_idx}, номер {num}")
                    break
        
        if found_positions:
            min_pos = min(p[0] for p in found_positions)
            max_pos = max(p[0] for p in found_positions)
            print(f"\n✓ Предполагаемый диапазон списка литературы: позиции {min_pos} - {max_pos}")

if __name__ == '__main__':
    find_exact_bibliography()
