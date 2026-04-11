# -*- coding: utf-8 -*-
"""
Скрипт для исправления документа v20 -> v30:
1. Извлечение списка литературы из документа
2. Сортировка списка литературы по алфавиту (русские → английские)
3. Создание маппинга старых номеров на новые
4. Обновление всех ссылок в тексте документа
5. Обновление списка литературы в документе
6. Сохранение результата как v30
"""
from pathlib import Path
from docx import Document
import re
import shutil
from typing import Dict, List, Tuple, Optional

def extract_author_surname(entry: str) -> str:
    """
    Извлекает фамилию первого автора из записи библиографии.
    Возвращает фамилию для сортировки (в верхнем регистре).
    """
    # Убираем номер в начале (например, "1. " или "10. ")
    entry = re.sub(r'^\d+\.\s*', '', entry.strip())
    
    # Ищем первую фамилию
    # Паттерн для русских фамилий: начинается с заглавной буквы, содержит буквы
    match = re.match(r'^([А-ЯЁ][а-яё]+(?:\-[А-ЯЁ][а-яё]+)?)', entry)
    if match:
        surname = match.group(1)
        return surname.upper()
    
    # Для английских фамилий
    match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Z]\.?)*)', entry)
    if match:
        surname = match.group(1).split()[0]
        return surname.upper()
    
    # Если не нашлось, пробуем найти любую заглавную букву
    match = re.search(r'([А-ЯЁA-Z])', entry)
    if match:
        return match.group(1)
    
    return 'Я'  # Если не удалось определить, ставим в конец

def parse_bibliography_from_doc(doc: Document) -> Dict[int, str]:
    """
    Извлекает список литературы из документа Word.
    Ищет заголовок "СПИСОК ЛИТЕРАТУРЫ" и собирает все записи вида "N. Автор..." до конца.
    Обрабатывает многострочные записи.
    """
    bibliography: Dict[int, str] = {}
    in_bibliography = False
    current_entry = None
    current_number = None
    
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        
        # Ищем начало списка литературы
        if not in_bibliography:
            text_upper = text.upper()
            if any(keyword in text_upper for keyword in ['СПИСОК ЛИТЕРАТУРЫ', 'ЛИТЕРАТУРА', 'БИБЛИОГРАФИЧЕСКИЙ СПИСОК']):
                in_bibliography = True
                print(f"  ✓ Найден заголовок списка литературы: '{text[:60]}'")
                continue
        
        if not in_bibliography:
            continue
        
        # Если мы в разделе библиографии, ищем записи
        # Ищем запись с номером: "N. Автор Название..."
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            current_number_candidate = int(match.group(1))
            entry_text_candidate = match.group(2).strip()
            
            # КРИТИЧЕСКАЯ ПРОВЕРКА: запись должна быть библиографической
            # (начинаться с фамилии автора, а не с общих слов типа "Исследование", "Вывод" и т.д.)
            
            # Список запрещённых первых слов (текст из выводов/заключения, а не библиография)
            non_bibliography_start_words = [
                # Глаголы в повелительном наклонении
                'разработать', 'изучить', 'провести', 'проанализировать', 'рассмотреть', 
                'определить', 'выявить', 'установить', 'создать', 'подготовить', 'составить',
                # Общие слова начала предложений
                'исследование', 'методология', 'методологический', 'методологическая', 'методологические',
                'вывод', 'выводы', 'заключение', 'результат', 'результаты', 'задача', 'цель', 'задание',
                'рекомендация', 'совет', 'правило', 'особенность', 'особенности',
                'взаимосвязь', 'влияние', 'взаимодействие', 'применение', 'использование',
                'выборка', 'методика', 'метод', 'методы', 'подход', 'подходы',
                # Другие общие слова
                'приложение', 'таблица', 'рисунок', 'график', 'диаграмма',
                # Ещё общие слова из выводов
                'изучение', 'анализ', 'проверка', 'подтверждение', 'наличие',
            ]
            
            # Извлекаем первое слово записи (в нижнем регистре)
            entry_words = re.findall(r'\b\w+\b', entry_text_candidate.lower())
            if not entry_words:
                continue
            
            entry_first_word = entry_words[0]
            
            # СТРОГАЯ ПРОВЕРКА: пропускаем записи, которые начинаются с запрещённых слов
            if entry_first_word in non_bibliography_start_words:
                # Если встречаем запись с запрещённым словом - это может быть конец списка литературы
                # или начало другого раздела (выводы, заключение)
                print(f"  ⚠ Пропущена не-библиографическая запись №{current_number_candidate}: '{entry_text_candidate[:70]}...'")
                # Сохраняем предыдущую запись, если есть
                if current_number is not None and current_entry:
                    bibliography[current_number] = current_entry.strip()
                # Останавливаем извлечение, если это выглядит как конец списка литературы
                # (например, после записи с большим номером может идти текст из выводов)
                # Также останавливаем, если запись содержит много текста (это явно не библиография)
                if current_number_candidate > 50 or len(entry_text_candidate) > 200:
                    print(f"  ℹ Обнаружена запись с большим номером/текстом и запрещённым словом, возможно, конец списка литературы")
                    break
                current_number = None
                current_entry = None
                continue
            
            # Дополнительная проверка: запись должна начинаться с фамилии автора (заглавная буква)
            # и быть достаточно длинной (библиографические записи обычно длиннее 30 символов)
            # НО не должна быть слишком длинной (если больше 300 символов - это скорее всего текст, а не библиография)
            if not re.match(r'^[А-ЯЁA-Z]', entry_text_candidate):
                print(f"  ⚠ Пропущена запись №{current_number_candidate}: не начинается с заглавной буквы")
                continue
            
            if len(entry_text_candidate) < 30:
                print(f"  ⚠ Пропущена запись №{current_number_candidate}: слишком короткая для библиографической записи")
                continue
            
            if len(entry_text_candidate) > 300:
                print(f"  ⚠ Пропущена запись №{current_number_candidate}: слишком длинная (скорее всего это текст, а не библиография): '{entry_text_candidate[:70]}...'")
                # Сохраняем предыдущую запись, если есть
                if current_number is not None and current_entry:
                    bibliography[current_number] = current_entry.strip()
                # Останавливаем извлечение, так как это похоже на начало другого раздела
                print(f"  ℹ Остановка извлечения из-за обнаружения длинного текста (возможно, конец списка литературы)")
                break
            
            # Сохраняем предыдущую запись, если есть
            if current_number is not None and current_entry:
                bibliography[current_number] = current_entry.strip()
            
            # Начинаем новую запись (проверка пройдена)
            current_number = current_number_candidate
            current_entry = entry_text_candidate
        elif text and current_number is not None:
            # Продолжение текущей записи (многострочная запись)
            # Проверяем, что это продолжение, а не новая запись
            # (не начинается с числа и точки)
            if not re.match(r'^\d+\.\s+', text):
                current_entry += " " + text.strip()
            else:
                # Это новая запись, сохраняем предыдущую
                bibliography[current_number] = current_entry.strip()
                current_number = None
                current_entry = None
        elif not text:
            # Пустая строка - сохраняем текущую запись (если есть) и начинаем новую
            if current_number is not None and current_entry:
                bibliography[current_number] = current_entry.strip()
                current_number = None
                current_entry = None
        # Если текст не пустой, не начинается с номера и нет текущей записи - 
        # возможно, мы дошли до конца списка или это другой раздел
        # Проверяем, не раздел ли это ПРИЛОЖЕНИЯ
        elif text and current_number is None:
            text_upper = text.upper().strip()
            # Если это короткий заголовок ПРИЛОЖЕНИЯ без точек - конец списка
            if re.fullmatch(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)(\s+\d+)?\.?\s*$', text_upper):
                print(f"  ✓ Обнаружен раздел '{text[:60]}', конец списка литературы")
                break
    
    # Сохраняем последнюю запись
    if current_number is not None and current_entry:
        bibliography[current_number] = current_entry.strip()
    
    # ФИНАЛЬНАЯ ФИЛЬТРАЦИЯ: удаляем записи, которые начинаются с запрещённых слов
    # Это важно, так как некоторые записи могли пройти проверку на предыдущем этапе
    non_bibliography_start_words = [
        # Глаголы в повелительном наклонении
        'разработать', 'изучить', 'провести', 'проанализировать', 'рассмотреть', 
        'определить', 'выявить', 'установить', 'создать', 'подготовить', 'составить',
        # Общие слова начала предложений
        'исследование', 'методология', 'методологический', 'методологическая', 'методологические',
        'вывод', 'выводы', 'заключение', 'результат', 'результаты', 'задача', 'цель', 'задание',
        'рекомендация', 'совет', 'правило', 'особенность', 'особенности',
        'взаимосвязь', 'влияние', 'взаимодействие', 'применение', 'использование',
        'выборка', 'методика', 'метод', 'методы', 'подход', 'подходы',
        # Другие общие слова
        'приложение', 'таблица', 'рисунок', 'график', 'диаграмма',
        # Ещё общие слова из выводов
        'изучение', 'анализ', 'проверка', 'подтверждение', 'наличие',
    ]
    
    filtered_bibliography = {}
    removed_count = 0
    for num, entry in bibliography.items():
        entry_words = re.findall(r'\b\w+\b', entry.lower())
        if entry_words and entry_words[0] in non_bibliography_start_words:
            print(f"  ⚠ Удалена не-библиографическая запись №{num} при финальной фильтрации: '{entry[:60]}...'")
            removed_count += 1
            continue
        
        # Дополнительная проверка: запись должна начинаться с фамилии автора
        # (заглавная буква, не общее слово)
        if not re.match(r'^[А-ЯЁA-Z]', entry.strip()) or len(entry.strip()) < 30:
            print(f"  ⚠ Удалена запись №{num} при финальной фильтрации: не похожа на библиографическую")
            removed_count += 1
            continue
        
        filtered_bibliography[num] = entry.strip()
    
    if removed_count > 0:
        print(f"  ℹ При финальной фильтрации удалено {removed_count} записей, не соответствующих библиографическому формату")
    
    return filtered_bibliography

def sort_literature_by_alphabet(literature: Dict[int, str]) -> List[Tuple[int, str]]:
    """
    Сортирует источники по алфавиту (по фамилии первого автора).
    Сначала русские источники (А-Я), затем английские.
    Возвращает список кортежей (старый_номер, запись)
    """
    sorted_items = []
    
    for old_num, entry in literature.items():
        surname = extract_author_surname(entry)
        sorted_items.append((surname, old_num, entry))
    
    # Сортируем: русские сначала, затем английские
    def sort_key(x):
        surname = x[0]
        if not surname:
            return (2, '')  # Пустые в конец
        
        first_char = surname[0]
        # Русские буквы
        if first_char in 'АБВГДЕЁЖЗИЙКЛМНОПРСТУФХЦЧШЩЪЫЬЭЮЯ':
            return (0, surname)  # Русские сначала
        # Английские буквы
        elif first_char.isalpha():
            return (1, surname)  # Английские потом
        else:
            return (2, surname)  # Остальное в конце
    
    sorted_items.sort(key=sort_key)
    
    # Возвращаем список (старый_номер, запись) отсортированный по алфавиту
    return [(old_num, entry) for _, old_num, entry in sorted_items]

def create_number_mapping(sorted_items: List[Tuple[int, str]]) -> Dict[int, int]:
    """
    Создаёт маппинг старый_номер -> новый_номер
    """
    mapping = {}
    
    for new_num, (old_num, entry) in enumerate(sorted_items, start=1):
        mapping[old_num] = new_num
    
    return mapping

def find_all_citations(text: str) -> List[Tuple[int, int, str]]:
    """
    Находит все ссылки на источники в тексте.
    Возвращает список (start_pos, end_pos, citation_text)
    Поддерживает форматы: [N], [N, M], [N; M], [N-M], [N, M, K]
    """
    citations = []
    
    # Паттерн для поиска ссылок: [числа, разделённые запятыми, точками с запятой или дефисами]
    pattern = r'\[(\d+(?:\s*[,\-;]\s*\d+)*)\]'
    
    for match in re.finditer(pattern, text):
        start = match.start()
        end = match.end()
        citation_text = match.group(0)
        citations.append((start, end, citation_text))
    
    return citations

def replace_citation_numbers(citation_text: str, mapping: Dict[int, int]) -> str:
    """
    Заменяет номера в ссылке согласно маппингу.
    Сохраняет оригинальные разделители.
    """
    # Извлекаем все числа из ссылки
    numbers = re.findall(r'\d+', citation_text)
    
    if not numbers:
        return citation_text
    
    # Заменяем каждое число
    new_numbers = []
    for num_str in numbers:
        old_num = int(num_str)
        if old_num in mapping:
            new_numbers.append(str(mapping[old_num]))
        else:
            # Если номер не найден в маппинге, оставляем как есть (с предупреждением)
            print(f"    ⚠ Номер {old_num} не найден в маппинге, оставляем без изменений")
            new_numbers.append(num_str)
    
    # Восстанавливаем формат с разделителями
    content = citation_text[1:-1]  # Убираем [ и ]
    
    # Определяем разделитель из оригинальной ссылки
    if ',' in content:
        sep = ', '
        result = '[' + sep.join(sorted(new_numbers, key=int)) + ']'  # Сортируем числа
    elif ';' in content:
        sep = '; '
        result = '[' + sep.join(sorted(new_numbers, key=int)) + ']'
    elif '-' in content:
        # Для диапазонов сохраняем дефис
        if len(new_numbers) == 2:
            result = '[' + '-'.join(sorted(new_numbers, key=int)) + ']'
        else:
            result = '[' + ', '.join(sorted(new_numbers, key=int)) + ']'
    else:
        # Если только одно число
        result = '[' + new_numbers[0] + ']'
    
    return result

def update_citations_in_paragraph(para, mapping: Dict[int, int]) -> int:
    """
    Обновляет ссылки в параграфе.
    Возвращает количество заменённых ссылок.
    """
    text = para.text
    citations = find_all_citations(text)
    
    if not citations:
        return 0
    
    replaced_count = 0
    # Заменяем с конца, чтобы не сбить позиции
    for start, end, citation_text in reversed(citations):
        new_citation = replace_citation_numbers(citation_text, mapping)
        if new_citation != citation_text:
            text = text[:start] + new_citation + text[end:]
            replaced_count += 1
    
    if text != para.text:
        para.clear()
        para.add_run(text)
    
    return replaced_count

def update_document_citations(doc: Document, mapping: Dict[int, int]) -> int:
    """
    Обновляет все ссылки на источники в документе.
    Возвращает количество заменённых ссылок.
    """
    replaced_count = 0
    
    # Обрабатываем параграфы
    for para in doc.paragraphs:
        count = update_citations_in_paragraph(para, mapping)
        replaced_count += count
    
    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Обновляем ссылки в каждой ячейке
                for para in cell.paragraphs:
                    count = update_citations_in_paragraph(para, mapping)
                    replaced_count += count
    
    return replaced_count

def update_bibliography_in_doc(doc: Document, sorted_items: List[Tuple[int, str]], mapping: Dict[int, int], original_bibliography: Dict[int, str]):
    """
    Обновляет список литературы в документе с новой нумерацией.
    Полностью заменяет раздел списка литературы новым отсортированным списком.
    """
    # Находим начало списка литературы
    bib_start_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip().upper()
        if any(keyword in text for keyword in ['СПИСОК ЛИТЕРАТУРЫ', 'ЛИТЕРАТУРА', 'БИБЛИОГРАФИЧЕСКИЙ СПИСОК']):
            bib_start_idx = i
            print(f"  ✓ Найдено начало списка литературы на позиции {i}")
            break
    
    if bib_start_idx is None:
        print("  ⚠ Не найдено начало списка литературы в документе")
        return
    
    # Находим конец списка литературы - ищем раздел ПРИЛОЖЕНИЯ
    # Расширяем поиск, так как список литературы может находиться далеко от заголовка
    bib_end_idx = None
    max_search_range = min(bib_start_idx + 350, len(doc.paragraphs))  # Увеличиваем диапазон
    
    for i in range(bib_start_idx + 1, max_search_range):
        text = (doc.paragraphs[i].text or "").strip()
        
        # Игнорируем строки с большим количеством точек (колонтитулы/номера страниц)
        if text.count('.') > 10:
            continue
        
        text_upper = text.upper().strip()
        # Проверяем, является ли это настоящим заголовком раздела ПРИЛОЖЕНИЯ
        # Заголовок должен быть коротким и не содержать номеров записей
        if (re.fullmatch(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)(\s+\d+)?\.?\s*$', text_upper) or 
            (re.match(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)', text_upper) and len(text.strip()) < 50 and text.count('.') <= 2 and not re.search(r'\d+\.\s+', text))):
            bib_end_idx = i
            print(f"  ✓ Найден конец списка литературы на позиции {i} (раздел '{text[:50]}')")
            break
    
    # Если не нашли конец через раздел ПРИЛОЖЕНИЯ, ищем последнюю запись из original_bibliography
    if bib_end_idx is None:
        print("  ℹ Раздел ПРИЛОЖЕНИЯ не найден в пределах 150 параграфов, ищем все записи списка литературы...")
        
        # Ищем ВСЕ записи из original_bibliography, чтобы определить их диапазон
        found_entry_indices = []
        max_search_extended = min(bib_start_idx + 350, len(doc.paragraphs))  # Увеличиваем диапазон поиска
        
        for i in range(bib_start_idx + 1, max_search_extended):
            text = (doc.paragraphs[i].text or "").strip()
            if not text:
                continue
            
            # Игнорируем ссылки в квадратных скобках
            if re.match(r'^\[.*\]\s*$', text):
                continue
            
            match = re.match(r'^(\d+)\.\s+(.+)', text)
            if match:
                num = int(match.group(1))
                entry_text = match.group(2).strip()
                
                # Проверяем, что это действительно запись списка литературы (не короткая ссылка)
                if len(entry_text) > 20 and num in original_bibliography:
                    # Дополнительная проверка: начало записи должно совпадать
                    original_entry_start = original_bibliography[num][:50].strip().lower()
                    doc_entry_start = entry_text[:50].strip().lower()
                    if original_entry_start == doc_entry_start:
                        found_entry_indices.append(i)
        
        if found_entry_indices:
            # Находим самый большой диапазон последовательных записей
            # Сортируем индексы
            found_entry_indices.sort()
            
            # Ищем самый длинный блок последовательных записей
            if len(found_entry_indices) >= len(original_bibliography) * 0.5:  # Если найдено хотя бы половина записей
                bib_end_idx = found_entry_indices[-1] + 10  # Оставляем запас после последней записи
                print(f"  ✓ Найдено {len(found_entry_indices)} записей списка литературы")
                print(f"     Диапазон: позиции {found_entry_indices[0]} - {found_entry_indices[-1]}")
                print(f"     Установлен конец списка: позиция {bib_end_idx}")
            else:
                bib_end_idx = min(bib_start_idx + 200, len(doc.paragraphs))
                print(f"  ⚠ Найдено только {len(found_entry_indices)} записей из {len(original_bibliography)}, используем позицию {bib_end_idx}")
        else:
            bib_end_idx = min(bib_start_idx + 200, len(doc.paragraphs))
            print(f"  ⚠ Записи не найдены, используем позицию {bib_end_idx}")
    
    # НОВАЯ СТРАТЕГИЯ: находим блок записей вида "N. Автор...", которые СООТВЕТСТВУЮТ 
    # записям из original_bibliography. Это гарантирует, что мы не захватим записи из заключения/приложений.
    
    print(f"  ℹ Поиск записей списка литературы в диапазоне (до позиции {bib_end_idx}):")
    
    # Создаём словарь для быстрого сравнения: начало записи (первые 50 символов без пробелов) -> номер
    original_entry_signatures = {}
    for old_num, entry in original_bibliography.items():
        # Берём первые 60 символов, убираем лишние пробелы, приводим к нижнему регистру
        signature = re.sub(r'\s+', ' ', entry[:60].strip().lower())
        original_entry_signatures[signature] = old_num
    
    all_candidate_entries = []
    search_range = min(bib_end_idx, len(doc.paragraphs))
    
    for i in range(bib_start_idx + 1, search_range):
        text = (doc.paragraphs[i].text or "").strip()
        
        if not text:
            continue
        
        # Игнорируем ссылки в квадратных скобках (это ссылки в тексте)
        if re.match(r'^\[.*\]\s*$', text):
            continue
        
        # Проверяем, является ли это записью библиографии: "N. Автор..."
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            entry_text = match.group(2).strip()
            
            # Проверяем, что запись достаточно длинная и начинается с фамилии автора
            if len(entry_text) < 30 or not re.match(r'^[А-ЯЁA-Z]', entry_text):
                continue
            
            # КРИТИЧЕСКОЕ ПРОВЕРКА: запись должна соответствовать одной из записей из original_bibliography
            # ИСПРАВЛЕНИЕ: сначала проверяем, что запись НЕ является текстом из выводов/заключения
            
            # СПИСОК СЛОВ, которые НЕ могут быть началом библиографической записи
            non_bibliography_start_words = [
                # Глаголы в повелительном наклонении
                'разработать', 'изучить', 'провести', 'проанализировать', 'рассмотреть', 
                'определить', 'выявить', 'установить', 'создать', 'подготовить', 'составить',
                # Общие слова начала предложений
                'исследование', 'методология', 'методологический', 'методологическая', 'методологические',
                'вывод', 'выводы', 'заключение', 'результат', 'результаты', 'задача', 'цель', 'задание',
                'рекомендация', 'совет', 'правило', 'особенность', 'особенности',
                'взаимосвязь', 'влияние', 'взаимодействие', 'применение', 'использование',
                'выборка', 'методика', 'метод', 'методы', 'подход', 'подходы',
                # Другие общие слова
                'приложение', 'таблица', 'рисунок', 'график', 'диаграмма',
                # Ещё общие слова из выводов
                'изучение', 'анализ', 'проверка', 'подтверждение', 'наличие',
            ]
            
            # Извлекаем первое слово записи (в нижнем регистре)
            entry_words = re.findall(r'\b\w+\b', entry_text.lower())
            if not entry_words:
                continue
            
            entry_first_word = entry_words[0]
            
            # СТРОГАЯ ПРОВЕРКА: если первое слово в списке запрещённых - пропускаем запись
            # Даже если она есть в original_bibliography (это означает ошибку при извлечении)
            if entry_first_word in non_bibliography_start_words:
                print(f"  ⚠ Пропущена запись №{num} при поиске: начинается с запрещённого слова '{entry_first_word}': '{entry_text[:60]}...'")
                continue
            
            # Дополнительная проверка: запись не должна быть слишком длинной (это текст, а не библиография)
            if len(entry_text) > 300:
                print(f"  ⚠ Пропущена запись №{num} при поиске: слишком длинная (скорее всего это текст): '{entry_text[:60]}...'")
                continue
            
            # Проверяем, что запись соответствует одной из записей из original_bibliography
            # Сравниваем начало записи (фамилия автора)
            matches_original = False
            
            # Извлекаем первые 2-3 слова из записи (обычно это фамилия автора)
            entry_words = re.findall(r'\b\w+\b', entry_text[:60].lower())
            if len(entry_words) < 1:
                continue
            
            # Сравниваем с записями из original_bibliography
            for orig_num, orig_entry in original_bibliography.items():
                # Извлекаем первые слова из оригинальной записи (фамилия автора)
                orig_words = re.findall(r'\b\w+\b', orig_entry[:60].lower())
                if len(orig_words) < 1:
                    continue
                
                # СТРОГАЯ ПРОВЕРКА: первое слово (фамилия) должно совпадать
                if entry_words[0] == orig_words[0]:
                    # Если есть второе слово, проверяем его (может быть частью фамилии или инициалом)
                    if len(entry_words) >= 2 and len(orig_words) >= 2:
                        # Второе слово должно совпадать (часть фамилии типа "Абульханова-Славская")
                        # или быть достаточно близким (инициалы могут отличаться форматированием)
                        if (entry_words[1] == orig_words[1] or 
                            (len(entry_words) >= 2 and '-' in entry_words[0] and entry_words[1][:3] == orig_words[1][:3])):
                            matches_original = True
                            break
                        # Для составных фамилий сравниваем первые части
                        if '-' in entry_words[0] and '-' in orig_words[0]:
                            entry_first_part = entry_words[0].split('-')[0]
                            orig_first_part = orig_words[0].split('-')[0]
                            if entry_first_part == orig_first_part and len(entry_first_part) > 5:
                                matches_original = True
                                break
                    else:
                        # Если только одно слово - проверяем, что оно достаточно длинное (фамилия, а не короткое слово)
                        if len(entry_words[0]) >= 5:  # Фамилии обычно длиннее 4 букв
                            matches_original = True
                            break
                
                # ДОПОЛНИТЕЛЬНАЯ ПРОВЕРКА: сравниваем начало записи (первые 30 символов без пробелов)
                # Это помогает найти совпадения, даже если форматирование немного отличается
                entry_clean_start = re.sub(r'[^\w]', '', entry_text[:40].lower())[:25]
                orig_clean_start = re.sub(r'[^\w]', '', orig_entry[:40].lower())[:25]
                if entry_clean_start and orig_clean_start and len(entry_clean_start) >= 15:
                    if entry_clean_start[:20] == orig_clean_start[:20]:
                        matches_original = True
                        break
            
            # Добавляем только если запись соответствует original_bibliography
            if matches_original:
                all_candidate_entries.append((i, num, entry_text))
    
    if not all_candidate_entries:
        print("  ⚠ Не найдены записи списка литературы, соответствующие извлечённым")
        print(f"     Проверялся диапазон: позиции {bib_start_idx + 1} - {search_range}")
        print(f"     Ожидалось найти записи из списка литературы (46 источников)")
        
        # Диагностика: показываем все записи вида "N. Автор..." для анализа
        print(f"\n  ℹ Диагностика: все записи вида 'N. Автор...' в диапазоне:")
        diagnostic_entries = []
        for i in range(bib_start_idx + 1, min(bib_start_idx + 200, search_range)):
            text = (doc.paragraphs[i].text or "").strip()
            if text and not re.match(r'^\[.*\]\s*$', text):
                match = re.match(r'^(\d+)\.\s+(.+)', text)
                if match:
                    num = int(match.group(1))
                    entry_text = match.group(2).strip()
                    if len(entry_text) > 20 and re.match(r'^[А-ЯЁA-Z]', entry_text):
                        diagnostic_entries.append((i, num, entry_text[:60]))
        
        if diagnostic_entries:
            print(f"     Найдено {len(diagnostic_entries)} записей (не все соответствуют списку литературы):")
            for para_idx, num, entry in diagnostic_entries[:10]:
                print(f"       [{para_idx:4d}] №{num:3d}: {entry}...")
            if len(diagnostic_entries) > 10:
                print(f"       ... (пропущено {len(diagnostic_entries) - 10} записей)")
        
        return
    
    print(f"     Найдено записей, соответствующих списку литературы: {len(all_candidate_entries)}")
    print(f"     Диапазон позиций: {all_candidate_entries[0][0]} - {all_candidate_entries[-1][0]}")
    print(f"     Первая запись: №{all_candidate_entries[0][1]} - {all_candidate_entries[0][2][:60]}...")
    print(f"     Последняя запись: №{all_candidate_entries[-1][1]} - {all_candidate_entries[-1][2][:60]}...")
    
    # Находим самый большой непрерывный блок записей
    # Записи считаются непрерывными, если они идут подряд (или с небольшими пропусками - пустые строки)
    continuous_blocks = []
    current_block = []
    last_idx = None
    
    for para_idx, num, entry_text in all_candidate_entries:
        if last_idx is None:
            # Первая запись - начинаем новый блок
            current_block = [(para_idx, num, entry_text)]
            last_idx = para_idx
        elif para_idx - last_idx <= 8:  # Допускаем пропуски (до 7 параграфов между записями - учитываем пустые строки)
            # Запись близко к предыдущей - добавляем в текущий блок
            current_block.append((para_idx, num, entry_text))
            last_idx = para_idx
        else:
            # Большой разрыв - сохраняем текущий блок и начинаем новый
            if len(current_block) >= 10:  # Блок должен быть достаточно большим (минимум 10 записей)
                continuous_blocks.append(current_block)
            current_block = [(para_idx, num, entry_text)]
            last_idx = para_idx
    
    # Добавляем последний блок
    if len(current_block) >= 10:
        continuous_blocks.append(current_block)
    
    # Если не нашли блоки с минимумом 10 записей, пробуем с минимумом 5
    if not continuous_blocks and len(all_candidate_entries) >= 20:
        print(f"  ℹ Не найден блок с минимумом 10 записей, пробуем с минимумом 5...")
        continuous_blocks = []
        current_block = []
        last_idx = None
        
        for para_idx, num, entry_text in all_candidate_entries:
            if last_idx is None:
                current_block = [(para_idx, num, entry_text)]
                last_idx = para_idx
            elif para_idx - last_idx <= 8:
                current_block.append((para_idx, num, entry_text))
                last_idx = para_idx
            else:
                if len(current_block) >= 5:
                    continuous_blocks.append(current_block)
                current_block = [(para_idx, num, entry_text)]
                last_idx = para_idx
        
        if len(current_block) >= 5:
            continuous_blocks.append(current_block)
    
    if not continuous_blocks:
        print(f"  ⚠ Не найден достаточно большой непрерывный блок записей")
        print(f"     Найдено только {len(all_candidate_entries)} записей в разных местах")
        # Показываем первые записи для диагностики
        print(f"     Первые 10 записей:")
        for para_idx, num, entry_text in all_candidate_entries[:10]:
            print(f"       [{para_idx:4d}] №{num:3d}: {entry_text[:60]}...")
        return
    
    # Выбираем блок: сначала проверяем все блоки и выбираем тот, который НЕ содержит записей с запрещёнными словами
    non_bib_start_words = [
        # Глаголы в повелительном наклонении
        'разработать', 'изучить', 'провести', 'проанализировать', 'рассмотреть', 
        'определить', 'выявить', 'установить', 'создать', 'подготовить', 'составить',
        # Общие слова начала предложений
        'исследование', 'методология', 'методологический', 'вывод', 'выводы', 
        'заключение', 'результат', 'результаты', 'задача', 'цель', 'задание',
        'рекомендация', 'совет', 'правило', 'особенность', 'особенности',
        'взаимосвязь', 'влияние', 'взаимодействие', 'применение', 'использование',
        'выборка', 'методика', 'метод', 'методы', 'подход', 'подходы',
        # Другие общие слова
        'приложение', 'таблица', 'рисунок', 'график', 'диаграмма',
    ]
    
    # Фильтруем блоки: исключаем блоки, которые содержат записи с запрещёнными словами
    valid_blocks = []
    for block in continuous_blocks:
        has_forbidden_entry = False
        for para_idx, num, entry_text in block:
            entry_words = re.findall(r'\b\w+\b', entry_text[:40].lower())
            if entry_words and entry_words[0] in non_bib_start_words:
                has_forbidden_entry = True
                print(f"  ⚠ Блок с {len(block)} записями содержит запись с запрещённым словом: №{num} - {entry_text[:60]}...")
                break
        
        if not has_forbidden_entry:
            valid_blocks.append(block)
    
    # Если есть валидные блоки, выбираем самый большой из них
    if valid_blocks:
        biggest_block = max(valid_blocks, key=len)
        print(f"  ✓ Найден валидный блок без записей с запрещёнными словами: {len(biggest_block)} записей")
    else:
        # Если все блоки содержат запрещённые записи, используем самый большой блок
        # но предупреждаем об этом
        biggest_block = max(continuous_blocks, key=len)
        print(f"  ⚠ ВНИМАНИЕ: Все блоки содержат записи с запрещёнными словами!")
        print(f"     Используется самый большой блок: {len(biggest_block)} записей")
        print(f"     Будет выполнена дополнительная проверка первой и последней записи...")
    
    # Финальная проверка: первая и последняя записи блока должны быть библиографическими
    first_entry = biggest_block[0][2]
    last_entry = biggest_block[-1][2]
    
    # Проверяем, что первая запись начинается с фамилии автора (а не с запрещённого слова)
    first_words = re.findall(r'\b\w+\b', first_entry[:40].lower())
    first_is_bib = False
    if first_words and first_words[0] not in non_bib_start_words:
        # Проверяем, что первое слово похоже на фамилию (начинается с заглавной буквы в оригинале)
        if re.match(r'^[А-ЯЁA-Z]', first_entry.strip()):
            first_is_bib = True
    
    # Проверяем, что последняя запись тоже библиографическая
    last_words = re.findall(r'\b\w+\b', last_entry[:40].lower())
    last_is_bib = False
    if last_words and last_words[0] not in non_bib_start_words:
        if re.match(r'^[А-ЯЁA-Z]', last_entry.strip()):
            last_is_bib = True
    
    if not first_is_bib or not last_is_bib:
        print(f"  ✗ Блок не соответствует списку литературы!")
        print(f"     Первая запись блока: №{biggest_block[0][1]} - {first_entry[:70]}...")
        print(f"     Последняя запись блока: №{biggest_block[-1][1]} - {last_entry[:70]}...")
        print(f"     Первая запись похожа на библиографическую: {first_is_bib}")
        print(f"     Последняя запись похожа на библиографическую: {last_is_bib}")
        print(f"     ⚠ ВНИМАНИЕ: Блок содержит записи из заключения/приложений, а не из списка литературы!")
        print(f"     Скрипт прерван, чтобы избежать порчи документа.")
        return
    
    print(f"  ✓ Найден блок записей списка литературы: {len(biggest_block)} записей")
    print(f"     Диапазон позиций: {biggest_block[0][0]} - {biggest_block[-1][0]}")
    print(f"     Первая запись: №{biggest_block[0][1]} - {biggest_block[0][2][:60]}...")
    print(f"     Последняя запись: №{biggest_block[-1][1]} - {biggest_block[-1][2][:60]}...")
    
    # Определяем границы для замены: от первой до последней записи блока
    first_bib_para_idx = biggest_block[0][0]
    last_bib_para_idx = biggest_block[-1][0]
    
    # Расширяем диапазон, чтобы включить пустые строки между записями
    # Но ограничиваемся следующим параграфом после последней записи
    paras_to_remove = []
    for para_idx in range(first_bib_para_idx, last_bib_para_idx + 1):
        if para_idx < len(doc.paragraphs):
            text = (doc.paragraphs[para_idx].text or "").strip()
            # Включаем параграф, если это запись списка или пустая строка
            if not text or re.match(r'^(\d+)\.\s+', text):
                paras_to_remove.append(para_idx)
    
    print(f"  ✓ Найдено параграфов для замены: {len(paras_to_remove)}")
    print(f"     Диапазон параграфов: {first_bib_para_idx} - {last_bib_para_idx}")
    
    # Запоминаем параграф ПЕРЕД первой записью блока для вставки после удаления
    # После удаления параграфов из блока этот параграф останется на позиции first_bib_para_idx - 1
    # (так как мы удаляем только параграфы ВНУТРИ блока, а не перед ним)
    insert_after_idx = max(0, first_bib_para_idx - 1)  # Параграф перед блоком
    
    # Удаляем старые записи в обратном порядке (от конца к началу, чтобы индексы не сдвигались)
    removed_count = 0
    for para_idx in reversed(sorted(paras_to_remove)):
        if para_idx < len(doc.paragraphs):
            para = doc.paragraphs[para_idx]
            p_element = para._element
            p_element.getparent().remove(p_element)
            removed_count += 1
    
    print(f"  ✓ Удалено старых записей: {removed_count}")
    
    # После удаления параграфов из блока позиция insert_after_idx не изменится
    # (так как мы удаляли только параграфы внутри блока, а не перед ним)
    # Но если first_bib_para_idx был 0, то insert_after_idx будет -1, что неправильно
    # Используем параграф на позиции insert_after_idx + 1 для вставки ПЕРЕД ним
    if insert_after_idx < 0:
        insert_after_idx = bib_start_idx
        insert_at_idx = bib_start_idx + 1
    else:
        # Вставляем ПОСЛЕ параграфа на позиции insert_after_idx
        # Для этого нужно вставить перед параграфом на позиции insert_after_idx + 1
        insert_at_idx = min(insert_after_idx + 1, len(doc.paragraphs))
    
    # Вставляем новые записи
    inserted_count = 0
    for new_num, (old_num, entry) in enumerate(sorted_items, start=1):
        # Вставляем новый параграф перед параграфом на позиции insert_at_idx
        if insert_at_idx < len(doc.paragraphs):
            new_para = doc.paragraphs[insert_at_idx].insert_paragraph_before()
        else:
            # Если дошли до конца документа, добавляем в конец
            new_para = doc.add_paragraph()
        new_para.add_run(f"{new_num}. {entry}")
        inserted_count += 1
        # После вставки индекс следующей позиции увеличивается
        insert_at_idx += 1
        
        # Добавляем пустую строку между записями (кроме последней)
        if new_num < len(sorted_items):
            if insert_at_idx < len(doc.paragraphs):
                empty_para = doc.paragraphs[insert_at_idx].insert_paragraph_before()
            else:
                empty_para = doc.add_paragraph()
            insert_at_idx += 1
    
    print(f"  ✓ Вставлено новых отсортированных записей: {inserted_count}")

def fix_document():
    """
    Основная функция для исправления документа v20 -> v30
    """
    version10 = Path(__file__).parent
    version6 = version10.parent / "Version6"
    
    # Проверяем наличие исходного документа
    v20_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
    v30_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v30.docx"
    
    # Если v30 уже существует, предупреждаем пользователя
    if v30_path.exists():
        print("⚠ ВНИМАНИЕ: Документ v30 уже существует и будет перезаписан!")
        print(f"   Файл: {v30_path.name}\n")
    
    if not v20_path.exists():
        # Пробуем скопировать из Version6
        source_path = version6 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
        if source_path.exists():
            import shutil
            print(f"📋 Копирование документа v20 из Version6...")
            shutil.copy2(source_path, v20_path)
            print(f"  ✓ Документ скопирован: {v20_path.name}\n")
        else:
            print(f"ERROR: Документ v20 не найден: {v20_path}")
            print(f"  Исходный файл также не найден: {source_path}")
            return
    else:
        print(f"✓ Используется существующий документ v20: {v20_path.name}\n")
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ДОКУМЕНТА v20 -> v30")
    print("=" * 70)
    print("1. Извлечение списка литературы")
    print("2. Сортировка по алфавиту (русские → английские)")
    print("3. Создание маппинга номеров")
    print("4. Обновление ссылок в тексте")
    print("5. Обновление списка литературы в документе")
    print("=" * 70)
    print()
    
    # Загружаем документ
    print("📄 Загрузка документа...")
    doc = Document(str(v20_path))
    print(f"  ✓ Документ загружен: {v20_path.name}")
    print(f"  ✓ Всего параграфов: {len(doc.paragraphs)}\n")
    
    # Извлекаем список литературы
    print("📚 Извлечение списка литературы из документа...")
    bibliography = parse_bibliography_from_doc(doc)
    print(f"  ✓ Найдено источников: {len(bibliography)}")
    
    if not bibliography:
        print("  ❌ ERROR: Не удалось извлечь список литературы!")
        print("  Проверьте, что в документе есть раздел 'СПИСОК ЛИТЕРАТУРЫ'")
        print("  Запустите run_debug.bat для диагностики структуры документа")
        return
    
    # Дополнительная проверка - показываем найденные номера
    found_numbers = sorted(bibliography.keys())
    print(f"  Диапазон найденных номеров: {min(found_numbers)} - {max(found_numbers)}")
    
    # Проверяем на пропуски
    expected_numbers = set(range(1, max(found_numbers) + 1))
    missing_numbers = expected_numbers - set(found_numbers)
    if missing_numbers:
        print(f"  ⚠ ВНИМАНИЕ: Пропущенные номера в списке: {sorted(missing_numbers)}")
    else:
        print(f"  ✓ Все номера от 1 до {max(found_numbers)} присутствуют")
    
    # Показываем первые и последние записи для проверки
    sorted_keys = sorted(bibliography.keys())
    print(f"  Диапазон номеров: {min(sorted_keys)} - {max(sorted_keys)}")
    print(f"  Первая запись (№{sorted_keys[0]}): {bibliography[sorted_keys[0]][:60]}...")
    print(f"  Последняя запись (№{sorted_keys[-1]}): {bibliography[sorted_keys[-1]][:60]}...\n")
    
    # Сортируем по алфавиту
    print("🔤 Сортировка списка литературы по алфавиту...")
    sorted_items = sort_literature_by_alphabet(bibliography)
    print(f"  ✓ Отсортировано {len(sorted_items)} источников")
    
    # Показываем примеры изменений
    print("\n  Примеры новой сортировки (первые 5):")
    for new_num, (old_num, entry) in enumerate(sorted_items[:5], start=1):
        surname = extract_author_surname(entry)
        print(f"    {new_num:3d} (было {old_num:3d}) [{surname}]: {entry[:50]}...")
    print()
    
    # Создаём маппинг
    print("🔄 Создание маппинга старых номеров на новые...")
    mapping = create_number_mapping(sorted_items)
    
    # Показываем примеры изменений номеров
    changed_count = sum(1 for old, new in mapping.items() if old != new)
    print(f"  ✓ Создан маппинг для {len(mapping)} источников")
    print(f"  ✓ Изменений номеров: {changed_count}")
    
    if changed_count > 0:
        print("\n  Примеры изменений номеров (первые 10):")
        count = 0
        for old_num in sorted(mapping.keys()):
            new_num = mapping[old_num]
            if old_num != new_num and count < 10:
                entry = bibliography[old_num][:50]
                print(f"    {old_num:3d} -> {new_num:3d}: {entry}...")
                count += 1
        if changed_count > 10:
            print(f"    ... и ещё {changed_count - 10} изменений")
    print()
    
    # Обновляем ссылки в документе
    print("📝 Обновление ссылок в тексте документа...")
    replaced_count = update_document_citations(doc, mapping)
    print(f"  ✓ Заменено ссылок: {replaced_count}\n")
    
    # Обновляем список литературы в документе
    # ВАЖНО: используем только те записи, которые были извлечены функцией parse_bibliography_from_doc
    print("📚 Обновление списка литературы в документе...")
    update_bibliography_in_doc(doc, sorted_items, mapping, bibliography)
    print(f"  ✓ Список литературы обновлён\n")
    
    # Сохраняем документ
    print("💾 Сохранение документа...")
    doc.save(str(v30_path))
    print(f"  ✓ Документ сохранён: {v30_path.name}\n")
    
    # Сохраняем отсортированный список литературы в отдельный файл для проверки
    lit_file_path = version10 / "СПИСОК_ЛИТЕРАТУРЫ_v30_sorted.md"
    with open(lit_file_path, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ\n\n")
        f.write("*Отсортировано по алфавиту (русские → английские)*\n\n")
        f.write(f"*Всего источников: {len(sorted_items)}*\n\n")
        
        for new_num, (old_num, entry) in enumerate(sorted_items, start=1):
            f.write(f"{new_num}. {entry}\n\n")
    
    print(f"  ✓ Отсортированный список сохранён: {lit_file_path.name}\n")
    
    print("=" * 70)
    print("РЕЗУЛЬТАТЫ")
    print("=" * 70)
    print(f"✓ Документ сохранён: {v30_path.name}")
    print(f"✓ Отсортированный список: {lit_file_path.name}")
    print(f"✓ Всего источников: {len(sorted_items)}")
    print(f"✓ Заменено ссылок в тексте: {replaced_count}")
    print(f"✓ Изменений номеров: {changed_count}")
    print("=" * 70)
    print("\n⚠ ВАЖНО: Проверьте документ v30 вручную перед использованием!")
    print("   Особенно проверьте:")
    print("   1. Правильность сортировки списка литературы")
    print("   2. Корректность всех ссылок в тексте")
    print("   3. Форматирование списка литературы в документе")

if __name__ == '__main__':
    try:
        fix_document()
    except Exception as e:
        print(f"\n❌ ERROR: Произошла ошибка: {e}")
        import traceback
        traceback.print_exc()
        input("\nНажмите Enter для выхода...")
