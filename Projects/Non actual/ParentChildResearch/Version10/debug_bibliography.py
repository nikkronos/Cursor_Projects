# -*- coding: utf-8 -*-
"""
Диагностический скрипт для проверки структуры списка литературы
"""
from pathlib import Path
from docx import Document
import re

def debug_bibliography():
    version10 = Path(__file__).parent
    v20_path = version10 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v20.docx"
    
    if not v20_path.exists():
        print(f"ERROR: Документ v20 не найден: {v20_path}")
        return
    
    doc = Document(str(v20_path))
    print("=" * 70)
    print("ДИАГНОСТИКА СПИСКА ЛИТЕРАТУРЫ")
    print("=" * 70)
    print(f"Всего параграфов в документе: {len(doc.paragraphs)}\n")
    
    # Ищем начало списка литературы
    in_bibliography = False
    bib_start_idx = None
    bib_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        
        if not in_bibliography:
            text_upper = text.upper()
            if any(keyword in text_upper for keyword in ['СПИСОК ЛИТЕРАТУРЫ', 'ЛИТЕРАТУРА', 'БИБЛИОГРАФИЧЕСКИЙ СПИСОК']):
                in_bibliography = True
                bib_start_idx = i
                print(f"✓ Найден заголовок списка литературы на позиции {i}:")
                print(f"  '{text[:100]}'\n")
                continue
        
        if in_bibliography:
            # Выход при разделе ПРИЛОЖЕНИЯ (только если это заголовок раздела)
            text_upper = text.upper().strip()
            
            # Игнорируем строки с большим количеством точек (это обычно колонтитулы/номера страниц)
            if text.count('.') > 5:
                # Продолжаем поиск, это не заголовок раздела
                pass
            # Проверяем, является ли это настоящим заголовком раздела
            elif (re.fullmatch(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)(\s+\d+)?\.?\s*$', text_upper) or 
                  (re.match(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)', text_upper) and 
                   len(text.strip()) < 50 and 
                   not re.search(r'\d+\.\s+', text) and
                   text.count('.') <= 2)):  # Заголовок не должен содержать много точек
                print(f"✓ Обнаружен раздел '{text[:60]}' на позиции {i}, конец списка литературы\n")
                break
            
            # Сохраняем все параграфы из библиографии (включая пустые для понимания структуры)
            bib_paragraphs.append((i, text))
    
    # Дополнительно: проверяем параграфы между заголовком и ПРИЛОЖЕНИЯМИ
    print("=" * 70)
    print("ПРОВЕРКА ВСЕХ ПАРАГРАФОВ МЕЖДУ ЗАГОЛОВКОМ И ПРИЛОЖЕНИЯМИ")
    print("=" * 70)
    
    if bib_start_idx is not None:
        # Ищем, где начинается ПРИЛОЖЕНИЯ
        app_start_idx = None
        for i in range(bib_start_idx + 1, len(doc.paragraphs)):
            text = (doc.paragraphs[i].text or "").strip().upper()
            if re.match(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)', text):
                app_start_idx = i
                break
        
        if app_start_idx is not None:
            print(f"\nЗаголовок библиографии: позиция {bib_start_idx}")
            print(f"Раздел ПРИЛОЖЕНИЯ: позиция {app_start_idx}")
            print(f"Всего параграфов между ними: {app_start_idx - bib_start_idx - 1}\n")
            
            # Ищем настоящий раздел ПРИЛОЖЕНИЯ (не просто текст со словом)
            real_app_idx = None
            for i in range(bib_start_idx + 1, min(bib_start_idx + 150, len(doc.paragraphs))):
                para = doc.paragraphs[i]
                para_text = (para.text or "").strip().upper()
                
                # Игнорируем строки с большим количеством точек (колонтитулы/номера страниц)
                if para.text.count('.') > 5:
                    continue
                
                # Заголовок раздела - короткий, только "ПРИЛОЖЕНИЯ" или "ПРИЛОЖЕНИЯ 1" и т.д.
                if re.fullmatch(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)(\s+\d+)?\.?\s*$', para_text) or \
                   (re.match(r'^(ПРИЛОЖЕНИЯ|ПРИЛОЖЕНИЕ)', para_text) and 
                    len(para.text.strip()) < 50 and 
                    para.text.count('.') <= 2):
                    real_app_idx = i
                    break
            
            if real_app_idx is None:
                real_app_idx = app_start_idx
            
            print(f"Настоящий раздел ПРИЛОЖЕНИЯ (заголовок): позиция {real_app_idx}")
            print(f"Всего параграфов для анализа: {real_app_idx - bib_start_idx - 1}\n")
            
            print("Все параграфы между заголовком и ПРИЛОЖЕНИЯМИ:")
            print("-" * 70)
            entries_found = 0
            for i in range(bib_start_idx + 1, real_app_idx):
                para = doc.paragraphs[i]
                text = (para.text or "").strip()
                if text:
                    # Проверяем, является ли это записью
                    match = re.match(r'^(\d+)\.\s+(.+)', text)
                    if match:
                        num = match.group(1)
                        entry = match.group(2)[:80]
                        print(f"[{i:4d}] №{num:3s}: {entry}...")
                        entries_found += 1
                    else:
                        print(f"[{i:4d}] (не запись): '{text[:100]}'")
                else:
                    print(f"[{i:4d}] (пустая строка)")
            
            print(f"\n✓ Найдено потенциальных записей: {entries_found}")
            
            # Проверяем таблицы в этом диапазоне
            print("\n" + "=" * 70)
            print("ПРОВЕРКА ТАБЛИЦ В ДОКУМЕНТЕ")
            print("=" * 70)
            print(f"Всего таблиц в документе: {len(doc.tables)}")
            
            # Пробуем найти таблицы, которые могут содержать список литературы
            # (хотя обычно список в параграфах, но проверим)
            for table_idx, table in enumerate(doc.tables):
                print(f"\nТаблица {table_idx + 1}: {len(table.rows)} строк, {len(table.columns) if table.rows else 0} столбцов")
                # Показываем первую строку для примера
                if table.rows:
                    first_row_text = " | ".join([cell.text[:30].strip() for cell in table.rows[0].cells])
                    print(f"  Первая строка: {first_row_text}")
    
    print(f"Всего параграфов в разделе библиографии (проверено циклом): {len(bib_paragraphs)}\n")
    
    # Если записей нет, более детально проверим структуру
    non_empty = [t for _, t in bib_paragraphs if t]
    if len(bib_paragraphs) == 0 or len(non_empty) == 0:
        print("⚠ ВНИМАНИЕ: Не найдено параграфов в разделе библиографии!")
        print("Возможно, список литературы находится в другом формате или структуре.\n")
    
    print("=" * 70)
    print("АНАЛИЗ ПАРАГРАФОВ СПИСКА ЛИТЕРАТУРЫ")
    print("=" * 70)
    
    # Анализируем первые 20 и последние 20 параграфов
    print("\nПервые 20 параграфов:")
    print("-" * 70)
    for idx, (para_idx, text) in enumerate(bib_paragraphs[:20]):
        if text:  # Показываем только непустые
            # Проверяем формат
            match_num = re.match(r'^(\d+)\.\s+(.+)', text)
            if match_num:
                num = match_num.group(1)
                entry_preview = match_num.group(2)[:60]
                print(f"[{para_idx:4d}] №{num:3s}: {entry_preview}...")
            else:
                # Не начинается с номера
                print(f"[{para_idx:4d}] (без номера): {text[:80]}")
    
    if len(bib_paragraphs) > 20:
        print(f"\n... (пропущено {len(bib_paragraphs) - 40} параграфов) ...\n")
        print("Последние 20 параграфов:")
        print("-" * 70)
        for idx, (para_idx, text) in enumerate(bib_paragraphs[-20:]):
            if text:
                match_num = re.match(r'^(\d+)\.\s+(.+)', text)
                if match_num:
                    num = match_num.group(1)
                    entry_preview = match_num.group(2)[:60]
                    print(f"[{para_idx:4d}] №{num:3s}: {entry_preview}...")
                else:
                    print(f"[{para_idx:4d}] (без номера): {text[:80]}")
    
    print("\n" + "=" * 70)
    print("ПОИСК ЗАПИСЕЙ С НОМЕРАМИ")
    print("=" * 70)
    
    # Ищем все записи с номерами
    numbered_entries = []
    for para_idx, text in bib_paragraphs:
        if not text:
            continue
        
        match = re.match(r'^(\d+)\.\s+(.+)', text)
        if match:
            num = int(match.group(1))
            entry = match.group(2).strip()
            numbered_entries.append((para_idx, num, entry))
    
    print(f"\n✓ Найдено записей с номерами: {len(numbered_entries)}")
    
    if numbered_entries:
        print(f"Диапазон номеров: {min(e[1] for e in numbered_entries)} - {max(e[1] for e in numbered_entries)}")
        print(f"\nПервые 5 записей:")
        for para_idx, num, entry in numbered_entries[:5]:
            print(f"  №{num:3d} [позиция {para_idx:4d}]: {entry[:70]}...")
        
        if len(numbered_entries) > 5:
            print(f"\nПоследние 5 записей:")
            for para_idx, num, entry in numbered_entries[-5:]:
                print(f"  №{num:3d} [позиция {para_idx:4d}]: {entry[:70]}...")
        
        # Проверяем, есть ли пропуски в номерах
        nums = sorted(set(e[1] for e in numbered_entries))
        expected = set(range(1, max(nums) + 1))
        missing = expected - set(nums)
        if missing:
            print(f"\n⚠ ВНИМАНИЕ: Пропущенные номера: {sorted(missing)}")
        else:
            print(f"\n✓ Номера идут без пропусков от 1 до {max(nums)}")
    
    # Проверяем пустые строки
    empty_count = sum(1 for _, text in bib_paragraphs if not text.strip())
    print(f"\nПустых строк в библиографии: {empty_count}")
    
    # Проверяем записи без номеров (многострочные)
    non_numbered = [text for _, text in bib_paragraphs if text and not re.match(r'^\d+\.\s+', text)]
    if non_numbered:
        print(f"\n⚠ Параграфов без номеров (возможно, продолжения записей): {len(non_numbered)}")
        print("Примеры:")
        for text in non_numbered[:3]:
            print(f"  '{text[:80]}'")

if __name__ == '__main__':
    debug_bibliography()
