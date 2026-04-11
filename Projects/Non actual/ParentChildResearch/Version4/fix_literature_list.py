# -*- coding: utf-8 -*-
"""
Исправление списка литературы:
1. Убрать источник 65 (Ядов) - он заменён на [13; 15]
2. Проверить корреляцию между списком и ссылками
"""
from pathlib import Path
import re

def load_literature_list():
    """Загружает список литературы из файла"""
    v2 = Path(__file__).parent.parent / "Version2"
    lit_file = v2 / "СПИСОК_ЛИТЕРАТУРЫ_50_ИСТОЧНИКОВ_ГОСТ.md"
    
    if not lit_file.exists():
        return {}
    
    literature = {}
    with open(lit_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    current_num = None
    current_text = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        match = re.match(r'^(\d+)\.\s*(.+)', line)
        if match:
            if current_num is not None:
                literature[current_num] = ' '.join(current_text).strip()
            
            current_num = int(match.group(1))
            current_text = [match.group(2)]
        elif current_num is not None:
            current_text.append(line)
    
    if current_num is not None:
        literature[current_num] = ' '.join(current_text).strip()
    
    return literature

def extract_all_references(text):
    """Извлекает все ссылки на источники из текста"""
    pattern = r'\[(\d+(?:[;\s,]*\d+)*)\]'
    matches = re.finditer(pattern, text)
    
    all_refs = set()
    for match in matches:
        ref_str = match.group(1)
        numbers = re.findall(r'\d+', ref_str)
        all_refs.update([int(n) for n in numbers])
    
    return sorted(all_refs)

def create_final_literature_list():
    """Создаёт финальный список литературы без источника 65"""
    v4 = Path(__file__).parent
    
    # Читаем текст документа v7 (актуальная версия)
    # Сначала пытаемся извлечь из v7, если нет - используем v6_text.txt
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    if v7_file.exists():
        try:
            from docx import Document
            doc = Document(str(v7_file))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            text = "\n".join(text_parts)
            print("✓ Текст извлечён из ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx")
        except Exception as e:
            print(f"⚠ Не удалось извлечь из v7, используем v6_text.txt: {e}")
            text_file = v4 / "v6_text.txt"
            if not text_file.exists():
                print(f"ERROR: {text_file} not found")
                return False
            with open(text_file, 'r', encoding='utf-8') as f:
                text = f.read()
    else:
        # Fallback на v6_text.txt
        text_file = v4 / "v6_text.txt"
        if not text_file.exists():
            print(f"ERROR: {text_file} not found")
            return False
        with open(text_file, 'r', encoding='utf-8') as f:
            text = f.read()
    
    print("=" * 70)
    print("ФОРМИРОВАНИЕ ФИНАЛЬНОГО СПИСКА ЛИТЕРАТУРЫ")
    print("=" * 70)
    
    # Извлекаем все ссылки
    used_refs = extract_all_references(text)
    
    print(f"\nНайдено уникальных номеров источников в тексте: {len(used_refs)}")
    print(f"Диапазон: {min(used_refs) if used_refs else 0} - {max(used_refs) if used_refs else 0}")
    
    # Убираем источник 65 (он заменён на [13; 15])
    if 65 in used_refs:
        used_refs.remove(65)
        print(f"\n⚠ Убран источник 65 (Ядов) - он заменён на [13; 15]")
    
    # Загружаем список литературы
    literature = load_literature_list()
    
    # Формируем финальный список (только используемые источники, без 65)
    final_literature = {}
    missing_refs = []
    
    for ref_num in used_refs:
        if ref_num in literature:
            final_literature[ref_num] = literature[ref_num]
        else:
            missing_refs.append(ref_num)
    
    # Сортируем по номерам
    sorted_refs = sorted(final_literature.keys())
    
    print(f"\n✓ Используется в тексте: {len(used_refs)} источников")
    print(f"✓ Найдено в списке литературы: {len(final_literature)} источников")
    
    if missing_refs:
        print(f"\n⚠ Источники, используемые в тексте, но отсутствующие в списке: {missing_refs}")
    
    # Сохраняем финальный список
    output_file = v4 / "СПИСОК_ЛИТЕРАТУРЫ_ФИНАЛЬНЫЙ.md"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("# СПИСОК ЛИТЕРАТУРЫ\n\n")
        f.write(f"*Сформирован на основе реального использования в тексте*\n\n")
        f.write(f"*Всего источников: {len(sorted_refs)}*\n")
        f.write(f"*Источник 65 (Ядов) исключён - заменён на [13; 15]*\n\n")
        
        for idx, ref_num in enumerate(sorted_refs, 1):
            f.write(f"{idx}. {final_literature[ref_num]}\n\n")
        
        if missing_refs:
            f.write("\n\n---\n\n")
            f.write("## ИСТОЧНИКИ, ТРЕБУЮЩИЕ ДОБАВЛЕНИЯ:\n\n")
            for ref_num in missing_refs:
                f.write(f"[{ref_num}] - требуется добавить в список литературы\n")
    
    print(f"\n✓ Финальный список литературы сохранён: {output_file.name}")
    print(f"  Всего источников: {len(sorted_refs)}")
    
    # Проверяем, укладываемся ли в 50
    total = len(sorted_refs)
    print(f"\n{'=' * 70}")
    print(f"ИТОГО ИСТОЧНИКОВ: {total}")
    if total <= 50:
        print(f"✓ Укладываемся в 50 источников (осталось {50 - total})")
    else:
        print(f"⚠ Превышено на {total - 50} источников")
        print(f"  Требуется сократить список до 50 источников")
    print(f"{'=' * 70}")
    
    return True

if __name__ == '__main__':
    create_final_literature_list()

