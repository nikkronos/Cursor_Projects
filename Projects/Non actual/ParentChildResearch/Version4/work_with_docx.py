# -*- coding: utf-8 -*-
"""
Скрипт для извлечения текста из docx документов
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def extract_full_text(docx_path):
    """Извлекает весь текст из docx, включая таблицы"""
    doc = Document(str(docx_path))
    text_parts = []
    
    # Параграфы
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)
    
    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n".join(text_parts)

def main():
    base = Path(__file__).parent.parent
    version4 = Path(__file__).parent
    
    # Копируем файлы если нужно
    v6_source = base / "Version3" / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    source_doc = base / "Version2" / "A1" / "ВКР Бакалавриат_Ворошилова.docx"
    
    v6_local = version4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    source_local = version4 / "ВКР_ИСХОДНИК.docx"
    
    import shutil
    
    if v6_source.exists() and not v6_local.exists():
        shutil.copy2(v6_source, v6_local)
        print(f"✓ Скопирован v6")
    
    if source_doc.exists() and not source_local.exists():
        shutil.copy2(source_doc, source_local)
        print(f"✓ Скопирован исходник")
    
    # Извлекаем текст из v6
    if v6_local.exists():
        print(f"\nИзвлекаю текст из: {v6_local.name}")
        text_v6 = extract_full_text(v6_local)
        output_v6 = version4 / "v6_text.txt"
        with open(output_v6, 'w', encoding='utf-8') as f:
            f.write(text_v6)
        print(f"✓ Сохранено {len(text_v6.split(chr(10)))} строк в {output_v6.name}")
    
    # Извлекаем текст из исходника
    if source_local.exists():
        print(f"\nИзвлекаю текст из: {source_local.name}")
        text_source = extract_full_text(source_local)
        output_source = version4 / "source_text.txt"
        with open(output_source, 'w', encoding='utf-8') as f:
            f.write(text_source)
        print(f"✓ Сохранено {len(text_source.split(chr(10)))} строк в {output_source.name}")

if __name__ == '__main__':
    main()

