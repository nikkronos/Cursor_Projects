# -*- coding: utf-8 -*-
"""
Финальное исправление проблем Ж и П с подробным выводом
"""
from pathlib import Path
from docx import Document

AUTHORS_MAP = {
    1: "Абульханова-Славская",
    2: "Адлер",
    7: "Байярд",
    9: "Марковская",
    18: "Гинзбург",
    22: "Дружинин",
    27: "Кабардова",
    30: "Климов",
    34: "Леви",
    35: "Леонтьев А.Н.",
    41: "Матейчек",
    48: "Поваренков",
    54: "Сатир",
    63: "Шнейдер",
    64: "Эйдемиллер, Юстицкис"
}

def format_citation(ref_numbers):
    """Форматирует цитату с указанием авторов и номеров"""
    authors_list = []
    for num in ref_numbers:
        if num in AUTHORS_MAP:
            authors_list.append(f"{AUTHORS_MAP[num]}, {num}")
        else:
            authors_list.append(str(num))
    return "[" + "; ".join(authors_list) + "]"

def fix_Ж_П_final():
    """Финальное исправление проблем Ж и П"""
    v4 = Path(__file__).parent
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_file.exists():
        print(f"ERROR: {v7_file} not found")
        return False
    
    print("=" * 70)
    print("ФИНАЛЬНОЕ ИСПРАВЛЕНИЕ ПРОБЛЕМ Ж И П")
    print("=" * 70)
    
    doc = Document(str(v7_file))
    changes = []
    found_patterns = []
    
    # Сначала ищем все проблемные места
    print("\nПоиск проблемных мест...")
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        
        # Ж. [22; 35]
        if "[22; 35]" in text:
            found_patterns.append({
                "para": para_idx + 1,
                "type": "Ж",
                "pattern": "[22; 35]",
                "text": text[:200],
                "fixed": "Дружинин, 22" in text or "Леонтьев А.Н., 35" in text
            })
        
        # Ж. [63; 64]
        if "[63; 64]" in text:
            found_patterns.append({
                "para": para_idx + 1,
                "type": "Ж",
                "pattern": "[63; 64]",
                "text": text[:200],
                "fixed": "Шнейдер, 63" in text or "Эйдемиллер, Юстицкис, 64" in text
            })
        
        # П. [1; 2; 7; 9; 34; 41; 54]
        if "[1; 2; 7; 9; 34; 41; 54]" in text:
            found_patterns.append({
                "para": para_idx + 1,
                "type": "П",
                "pattern": "[1; 2; 7; 9; 34; 41; 54]",
                "text": text[:200],
                "fixed": "Абульханова-Славская, 1" in text
            })
        
        # П. [18; 27; 30; 48]
        if "[18; 27; 30; 48]" in text:
            found_patterns.append({
                "para": para_idx + 1,
                "type": "П",
                "pattern": "[18; 27; 30; 48]",
                "text": text[:200],
                "fixed": "Гинзбург, 18" in text
            })
    
    print(f"\nНайдено мест с паттернами: {len(found_patterns)}")
    for item in found_patterns:
        status = "✓ Уже исправлено" if item["fixed"] else "✗ Требует исправления"
        print(f"  Параграф {item['para']}: {item['type']} - {item['pattern']} - {status}")
        if not item["fixed"]:
            print(f"    Текст: {item['text'][:100]}...")
    
    # Исправляем
    print("\n" + "=" * 70)
    print("ВЫПОЛНЕНИЕ ИСПРАВЛЕНИЙ")
    print("=" * 70)
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        original_text = text
        
        # Ж. [22; 35]
        if "[22; 35]" in text and "Дружинин, 22" not in text:
            text = text.replace("[22; 35]", format_citation([22, 35]))
            changes.append(f"Ж: Параграф {para_idx+1} - [22; 35]")
        
        # Ж. [63; 64]
        if "[63; 64]" in text and "Шнейдер, 63" not in text:
            text = text.replace("[63; 64]", format_citation([63, 64]))
            changes.append(f"Ж: Параграф {para_idx+1} - [63; 64]")
        
        # П. [1; 2; 7; 9; 34; 41; 54]
        if "[1; 2; 7; 9; 34; 41; 54]" in text and "Абульханова-Славская, 1" not in text:
            text = text.replace("[1; 2; 7; 9; 34; 41; 54]", format_citation([1, 2, 7, 9, 34, 41, 54]))
            changes.append(f"П: Параграф {para_idx+1} - [1; 2; 7; 9; 34; 41; 54]")
        
        # П. [18; 27; 30; 48]
        if "[18; 27; 30; 48]" in text and "Гинзбург, 18" not in text:
            text = text.replace("[18; 27; 30; 48]", format_citation([18, 27, 30, 48]))
            changes.append(f"П: Параграф {para_idx+1} - [18; 27; 30; 48]")
        
        if text != original_text:
            para.text = text
    
    print(f"\nВнесено изменений: {len(changes)}")
    for change in changes:
        print(f"  ✓ {change}")
    
    if len(changes) == 0:
        print("\n⚠ Изменения не внесены.")
        print("Возможные причины:")
        print("  1. Уже исправлено ранее")
        print("  2. Паттерны не найдены в документе")
        print("  3. Текст находится в таблице")
    
    # Сохраняем
    doc.save(str(v7_file))
    print(f"\n✓ Документ сохранён: {v7_file.name}")
    
    return True

if __name__ == '__main__':
    fix_Ж_П_final()





