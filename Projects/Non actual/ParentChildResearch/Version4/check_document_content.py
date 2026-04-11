# -*- coding: utf-8 -*-
"""
Проверка содержимого документа для отладки
"""
from pathlib import Path
from docx import Document

def check_document():
    """Проверяет содержимое документа и ищет проблемные места"""
    v4 = Path(__file__).parent
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_file.exists():
        print(f"ERROR: {v7_file} not found")
        return False
    
    print("=" * 70)
    print("ПРОВЕРКА СОДЕРЖИМОГО ДОКУМЕНТА v7")
    print("=" * 70)
    
    doc = Document(str(v7_file))
    
    # Ищем проблемные места
    search_patterns = [
        ("[22; 35]", "Ж - должно быть с цитированием"),
        ("[63; 64]", "Ж - должно быть с цитированием"),
        ("[1; 2; 7; 9; 34; 41; 54]", "П - должно быть с цитированием"),
        ("[18; 27; 30; 48]", "П - должно быть с цитированием"),
    ]
    
    found_issues = []
    
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text
        for pattern, description in search_patterns:
            if pattern in text:
                # Проверяем, исправлено ли уже
                if "Дружинин, 22" in text and pattern == "[22; 35]":
                    continue
                if "Шнейдер, 63" in text and pattern == "[63; 64]":
                    continue
                if "Абульханова-Славская, 1" in text and pattern == "[1; 2; 7; 9; 34; 41; 54]":
                    continue
                if "Гинзбург, 18" in text and pattern == "[18; 27; 30; 48]":
                    continue
                
                # Найдена проблема
                found_issues.append({
                    "para": para_idx + 1,
                    "pattern": pattern,
                    "description": description,
                    "text": text[:150] + "..." if len(text) > 150 else text
                })
    
    print(f"\nНайдено проблемных мест: {len(found_issues)}")
    
    if found_issues:
        print("\nПроблемные места:")
        for issue in found_issues:
            print(f"\n  Параграф {issue['para']}: {issue['description']}")
            print(f"  Паттерн: {issue['pattern']}")
            print(f"  Текст: {issue['text']}")
    else:
        print("\n✓ Проблемных мест не найдено - возможно, уже исправлено")
    
    # Также проверяем таблицы
    print(f"\nПроверка таблиц: найдено {len(doc.tables)} таблиц")
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                cell_text = cell.text
                for pattern, description in search_patterns:
                    if pattern in cell_text:
                        print(f"\n  Таблица {table_idx+1}, строка {row_idx+1}, ячейка {cell_idx+1}:")
                        print(f"    {description}")
                        print(f"    Текст: {cell_text[:100]}...")
    
    return True

if __name__ == '__main__':
    check_document()





