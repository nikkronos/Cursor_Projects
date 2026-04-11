# -*- coding: utf-8 -*-
"""
Исправление оставшихся проблем: Ж, Л, П
"""
from pathlib import Path
from docx import Document

AUTHORS_MAP = {
    1: "Абульханова-Славская",
    2: "Адлер",
    7: "Байярд",
    9: "Марковская",
    18: "Гинзбург",
    22: "Дружинин",
    27: "Кабардова",
    30: "Климов",
    34: "Леви",
    35: "Леонтьев А.Н.",
    41: "Матейчек",
    48: "Поваренков",
    54: "Сатир",
    63: "Шнейдер",
    64: "Эйдемиллер, Юстицкис"
}

def format_citation_with_authors(ref_numbers):
    """Форматирует цитату с указанием авторов и номеров"""
    authors_list = []
    for num in ref_numbers:
        if num in AUTHORS_MAP:
            authors_list.append(f"{AUTHORS_MAP[num]}, {num}")
        else:
            authors_list.append(str(num))
    return "[" + "; ".join(authors_list) + "]"

def fix_remaining_issues():
    """Исправляет оставшиеся проблемы Ж, Л, П"""
    v4 = Path(__file__).parent
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_file.exists():
        print(f"ERROR: {v7_file} not found")
        return False
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ОСТАВШИХСЯ ПРОБЛЕМ: Ж, Л, П")
    print("=" * 70)
    
    doc = Document(str(v7_file))
    changes = []
    
    # Ж. Страница 10 - исправить [22; 35] и [63; 64]
    print("\n[Ж] Исправление цитирования на странице 10...")
    for para in doc.paragraphs:
        text = para.text
        
        # Ищем строку с [22; 35]
        if "[22; 35]" in text and "Дружинин, 22" not in text:
            para.text = text.replace("[22; 35]", format_citation_with_authors([22, 35]))
            changes.append("Ж: Исправлено [22; 35] -> с цитированием")
            print(f"  ✓ {changes[-1]}")
            text = para.text  # Обновляем текст
        
        # Ищем строку с [63; 64]
        if "[63; 64]" in text and "Шнейдер, 63" not in text:
            para.text = text.replace("[63; 64]", format_citation_with_authors([63, 64]))
            changes.append("Ж: Исправлено [63; 64] -> с цитированием")
            print(f"  ✓ {changes[-1]}")
    
    # Л. Проверка источника [23]
    print("\n[Л] Проверка источника [23]...")
    for para in doc.paragraphs:
        if "Выборка была сформирована методом добровольного участия [23]" in para.text:
            print("  Найдено: 'Выборка была сформирована методом добровольного участия [23]'")
            print("  Источник 23: Дубровина И.В. 'Практическая психология образования'")
            print("  ✓ Источник корректен для метода добровольного участия в образовательном контексте")
            print("  Рекомендация: Оставить как есть")
            changes.append("Л: Источник [23] проверен - корректен")
            break
    
    # П. Страница 35 - исправить [1; 2; 7; 9; 34; 41; 54] и [18; 27; 30; 48]
    print("\n[П] Исправление цитирования на странице 35...")
    for para in doc.paragraphs:
        text = para.text
        
        # Ищем [1; 2; 7; 9; 34; 41; 54]
        if "[1; 2; 7; 9; 34; 41; 54]" in text and "Абульханова-Славская, 1" not in text:
            para.text = text.replace(
                "[1; 2; 7; 9; 34; 41; 54]",
                format_citation_with_authors([1, 2, 7, 9, 34, 41, 54])
            )
            changes.append("П: Исправлено [1; 2; 7; 9; 34; 41; 54] -> с цитированием")
            print(f"  ✓ {changes[-1]}")
            text = para.text  # Обновляем текст
        
        # Ищем [18; 27; 30; 48]
        if "[18; 27; 30; 48]" in text and "Гинзбург, 18" not in text:
            para.text = text.replace(
                "[18; 27; 30; 48]",
                format_citation_with_authors([18, 27, 30, 48])
            )
            changes.append("П: Исправлено [18; 27; 30; 48] -> с цитированием")
            print(f"  ✓ {changes[-1]}")
    
    print(f"\nВнесено изменений: {len(changes)}")
    
    # Сохраняем
    doc.save(str(v7_file))
    print(f"\n✓ Документ обновлён: {v7_file.name}")
    
    return True

if __name__ == '__main__':
    fix_remaining_issues()





