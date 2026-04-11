# -*- coding: utf-8 -*-
"""
Скрипт для задачи В: Добавление фамилий авторов к первым трём подходам
"""
from pathlib import Path
from docx import Document

# Соответствие номеров источников и авторов
AUTHORS = {
    2: "Адлер",
    5: "Андреева",
    18: "Гинзбург",
    30: "Климов",
    49: "Прихожан, Толстых",
    64: "Эйдемиллер, Юстицкис"
}

def fix_task_V():
    """Добавляет фамилии авторов к первым трём подходам"""
    v4 = Path(__file__).parent
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_file.exists():
        print(f"ERROR: {v7_file} not found")
        return False
    
    print("=" * 70)
    print("ЗАДАЧА В: Добавление фамилий авторов к подходам")
    print("=" * 70)
    
    doc = Document(str(v7_file))
    changes = []
    
    # 1. Системный семейный подход [5; 64]
    for para in doc.paragraphs:
        if "Системный семейный подход [5; 64]" in para.text:
            old = "Системный семейный подход [5; 64]"
            # Формат как в задаче Й: [Автор, номер; Автор, номер]
            new = f"Системный семейный подход [{AUTHORS[5]}, 5; {AUTHORS[64]}, 64]"
            para.text = para.text.replace(old, new)
            changes.append("В1: Добавлены фамилии к системному семейному подходу")
            print(f"  ✓ {changes[-1]}")
            break
    
    # 2. Психологические теории развития и сценариев [2]
    for para in doc.paragraphs:
        if "Психологические теории развития и сценариев [2]" in para.text:
            old = "Психологические теории развития и сценариев [2]"
            new = f"Психологические теории развития и сценариев [{AUTHORS[2]}; 2]"
            para.text = para.text.replace(old, new)
            changes.append("В2: Добавлена фамилия к теориям развития и сценариев")
            print(f"  ✓ {changes[-1]}")
            break
    
    # 3. Теории профессионального самоопределения [18; 30; 49]
    for para in doc.paragraphs:
        if "Теории профессионального самоопределения [18; 30; 49]" in para.text:
            old = "Теории профессионального самоопределения [18; 30; 49]"
            # Формат: [Автор, номер; Автор, номер; Автор, номер]
            new = f"Теории профессионального самоопределения [{AUTHORS[18]}, 18; {AUTHORS[30]}, 30; {AUTHORS[49]}, 49]"
            para.text = para.text.replace(old, new)
            changes.append("В3: Добавлены фамилии к теориям профессионального самоопределения")
            print(f"  ✓ {changes[-1]}")
            break
    
    print(f"\nВнесено изменений: {len(changes)}")
    
    # Сохраняем
    doc.save(str(v7_file))
    print(f"\n✓ Документ обновлён: {v7_file.name}")
    
    return True

if __name__ == '__main__':
    fix_task_V()

