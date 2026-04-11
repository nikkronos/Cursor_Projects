#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для извлечения текста из docx и анализа документа
"""
import sys
import os
from pathlib import Path

# Добавляем родительскую директорию в путь для импорта
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Устанавливаю python-docx...")
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def extract_text_from_docx(docx_path):
    """Извлекает весь текст из docx файла"""
    doc = Document(docx_path)
    full_text = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            full_text.append(text)
    
    # Также извлекаем текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    full_text.append(cell_text)
    
    return "\n".join(full_text)

def save_text_to_file(text, output_path):
    """Сохраняет текст в файл"""
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(text)
    print(f"✓ Текст сохранён в: {output_path}")

if __name__ == "__main__":
    # Определяем пути
    script_dir = Path(__file__).parent
    v6_path = script_dir / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    source_path = script_dir / "ВКР_ИСХОДНИК.docx"
    
    # Копируем файлы если их нет
    if not v6_path.exists():
        v3_path = script_dir.parent / "Version3" / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
        if v3_path.exists():
            import shutil
            shutil.copy2(v3_path, v6_path)
            print(f"✓ Скопирован v6 из Version3")
    
    if not source_path.exists():
        source_orig = script_dir.parent / "Version2" / "A1" / "ВКР Бакалавриат_Ворошилова.docx"
        if source_orig.exists():
            import shutil
            shutil.copy2(source_orig, source_path)
            print(f"✓ Скопирован исходник из Version2/A1")
    
    # Извлекаем текст
    if v6_path.exists():
        print(f"\nИзвлекаю текст из: {v6_path.name}")
        text_v6 = extract_text_from_docx(v6_path)
        save_text_to_file(text_v6, script_dir / "ВКР_v6_текст.txt")
        print(f"Всего строк: {len(text_v6.split(chr(10)))}")
    
    if source_path.exists():
        print(f"\nИзвлекаю текст из: {source_path.name}")
        text_source = extract_text_from_docx(source_path)
        save_text_to_file(text_source, script_dir / "ВКР_ИСХОДНИК_текст.txt")
        print(f"Всего строк: {len(text_source.split(chr(10)))}")





