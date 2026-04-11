# -*- coding: utf-8 -*-
"""
Скрипт для поиска фамилий авторов из исходника для задачи В
"""
from pathlib import Path
from docx import Document

def extract_source_text():
    """Извлекает текст из исходника"""
    v4 = Path(__file__).parent
    source_file = v4 / "ВКР_ИСХОДНИК.docx"
    
    if not source_file.exists():
        print(f"ERROR: {source_file} not found")
        return None
    
    print("Извлекаю текст из исходника...")
    doc = Document(str(source_file))
    
    text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)
    
    # Также извлекаем из таблиц
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    full_text = "\n".join(text_parts)
    
    # Сохраняем для анализа
    output_file = v4 / "source_text_extracted.txt"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write(full_text)
    
    print(f"✓ Текст сохранён: {output_file.name}")
    print(f"  Размер: {len(full_text)} символов")
    
    return full_text

def find_approaches_in_source(text):
    """Ищет информацию о первых трёх подходах в исходнике"""
    print("\n" + "=" * 70)
    print("ПОИСК ИНФОРМАЦИИ О ПОДХОДАХ")
    print("=" * 70)
    
    # Ищем упоминания подходов
    approaches = {
        "Системный семейный подход": {
            "keywords": ["системный", "семейный", "подход", "эмоциональная система", "Эйдемиллер", "Юстицкис", "Андреева"],
            "sources": [5, 64]
        },
        "Психологические теории развития и сценариев": {
            "keywords": ["теории развития", "сценариев", "Адлер", "бессознательные", "жизненные сценарии"],
            "sources": [2]
        },
        "Теории профессионального самоопределения": {
            "keywords": ["профессиональное самоопределение", "Гинзбург", "Климов", "Поваренков", "классификации профессий"],
            "sources": [18, 30, 49]
        }
    }
    
    results = {}
    
    for approach_name, info in approaches.items():
        print(f"\n[{approach_name}]")
        print(f"Источники в тексте: {info['sources']}")
        print("Ищем упоминания...")
        
        # Ищем упоминания ключевых слов
        found_lines = []
        lines = text.split('\n')
        
        for i, line in enumerate(lines):
            line_lower = line.lower()
            for keyword in info['keywords']:
                if keyword.lower() in line_lower:
                    # Проверяем, есть ли в этой строке фамилии авторов
                    found_lines.append((i+1, line[:200]))
                    break
        
        if found_lines:
            print(f"  Найдено упоминаний: {len(found_lines)}")
            for line_num, line_text in found_lines[:3]:  # Показываем первые 3
                print(f"    Строка {line_num}: {line_text}...")
        else:
            print("  Упоминания не найдены")
        
        results[approach_name] = found_lines
    
    return results

def main():
    print("=" * 70)
    print("ПОИСК ФАМИЛИЙ АВТОРОВ ИЗ ИСХОДНИКА (Задача В)")
    print("=" * 70)
    
    # Извлекаем текст
    text = extract_source_text()
    if not text:
        return
    
    # Ищем информацию о подходах
    results = find_approaches_in_source(text)
    
    # Сохраняем результаты
    v4 = Path(__file__).parent
    results_file = v4 / "РЕЗУЛЬТАТЫ_ПОИСКА_АВТОРОВ.md"
    
    with open(results_file, 'w', encoding='utf-8') as f:
        f.write("# Результаты поиска авторов из исходника\n\n")
        f.write("## Задача В: Найти фамилии для первых трёх подходов\n\n")
        
        for approach_name, found_lines in results.items():
            f.write(f"### {approach_name}\n\n")
            if found_lines:
                f.write(f"Найдено упоминаний: {len(found_lines)}\n\n")
                for line_num, line_text in found_lines[:5]:
                    f.write(f"**Строка {line_num}:**\n{line_text}\n\n")
            else:
                f.write("Упоминания не найдены в исходнике.\n\n")
    
    print(f"\n✓ Результаты сохранены: {results_file.name}")
    print("\n" + "=" * 70)
    print("СЛЕДУЮЩИЙ ШАГ:")
    print("Откройте source_text_extracted.txt и РЕЗУЛЬТАТЫ_ПОИСКА_АВТОРОВ.md")
    print("Найдите фамилии авторов для каждого подхода")
    print("=" * 70)

if __name__ == '__main__':
    main()





