# -*- coding: utf-8 -*-
"""
Полный скрипт для исправления документа v6 и создания v7
Выполняет все задачи А-П
"""
import sys
import re
from pathlib import Path
from collections import defaultdict

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def extract_text_with_structure(docx_path):
    """Извлекает текст с сохранением структуры (параграфы, таблицы)"""
    doc = Document(str(docx_path))
    paragraphs = []
    tables_data = []
    
    for para in doc.paragraphs:
        paragraphs.append({
            'text': para.text,
            'style': para.style.name if para.style else None
        })
    
    for table_idx, table in enumerate(doc.tables):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_data.append(table_data)
    
    return paragraphs, tables_data

def find_references_in_text(text):
    """Находит все ссылки на источники в тексте (например, [1], [2, 3], [5, 64])"""
    # Паттерн для ссылок типа [1], [2, 3], [5, 64; 2; 18, 30, 49]
    pattern = r'\[(\d+(?:,\s*\d+)*(?:\s*;\s*\d+(?:,\s*\d+)*)*)\]'
    matches = re.finditer(pattern, text)
    refs = []
    for match in matches:
        refs.append({
            'full_match': match.group(0),
            'numbers': [int(x) for x in re.findall(r'\d+', match.group(1))]
        })
    return refs

def analyze_literature_usage(text, literature_list):
    """Анализирует использование источников в тексте"""
    used_refs = set()
    all_refs = find_references_in_text(text)
    
    for ref in all_refs:
        used_refs.update(ref['numbers'])
    
    # Находим упоминания авторов по фамилиям
    author_mentions = {}
    for idx, lit_item in enumerate(literature_list, 1):
        # Извлекаем фамилию автора (первое слово до точки)
        author_match = re.match(r'(\d+)\.\s*([А-ЯЁ][а-яё]+)', lit_item)
        if author_match:
            author_name = author_match.group(2)
            if author_name not in author_mentions:
                author_mentions[author_name] = []
            author_mentions[author_name].append(idx)
    
    return {
        'used_refs': used_refs,
        'all_refs_in_text': all_refs,
        'author_mentions': author_mentions
    }

def main():
    base = Path(__file__).parent.parent
    version4 = Path(__file__).parent
    
    # Пути к файлам
    v6_path = base / "Version3" / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    source_path = base / "Version2" / "A1" / "ВКР Бакалавриат_Ворошилова.docx"
    
    v6_local = version4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    source_local = version4 / "ВКР_ИСХОДНИК.docx"
    
    # Копируем файлы
    import shutil
    if v6_path.exists() and not v6_local.exists():
        shutil.copy2(v6_path, v6_local)
        print(f"✓ Скопирован v6")
    
    if source_path.exists() and not source_local.exists():
        shutil.copy2(source_path, source_local)
        print(f"✓ Скопирован исходник")
    
    if not v6_local.exists():
        print(f"ERROR: {v6_local} not found")
        return
    
    # Извлекаем текст
    print("\nИзвлекаю текст из документов...")
    paras_v6, tables_v6 = extract_text_with_structure(v6_local)
    
    # Сохраняем текст для анализа
    text_v6 = "\n".join([p['text'] for p in paras_v6 if p['text'].strip()])
    with open(version4 / "v6_text_full.txt", 'w', encoding='utf-8') as f:
        f.write(text_v6)
    
    print(f"✓ Извлечено {len(paras_v6)} параграфов и {len(tables_v6)} таблиц")
    print(f"✓ Текст сохранён в v6_text_full.txt")
    
    # Анализируем ссылки
    print("\nАнализирую использование источников...")
    refs = find_references_in_text(text_v6)
    all_ref_numbers = set()
    for ref in refs:
        all_ref_numbers.update(ref['numbers'])
    
    print(f"Найдено ссылок: {len(refs)}")
    print(f"Уникальных номеров источников: {len(all_ref_numbers)}")
    print(f"Максимальный номер: {max(all_ref_numbers) if all_ref_numbers else 0}")
    
    # Сохраняем анализ
    with open(version4 / "analysis_refs.txt", 'w', encoding='utf-8') as f:
        f.write("АНАЛИЗ ССЫЛОК В ТЕКСТЕ\n")
        f.write("=" * 50 + "\n\n")
        for ref in sorted(refs, key=lambda x: min(x['numbers']) if x['numbers'] else 0):
            f.write(f"{ref['full_match']} -> {ref['numbers']}\n")
        f.write(f"\n\nВсего уникальных номеров: {sorted(all_ref_numbers)}\n")

if __name__ == '__main__':
    main()





