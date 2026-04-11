# -*- coding: utf-8 -*-
"""
Исправление проблем Ж и П - более надёжная версия
"""
from pathlib import Path
from docx import Document
import re

AUTHORS_MAP = {
    1: "Абульханова-Славская",
    2: "Адлер",
    7: "Байярд",
    9: "Марковская",
    18: "Гинзбург",
    22: "Дружинин",
    27: "Кабардова",
    30: "Климов",
    34: "Леви",
    35: "Леонтьев А.Н.",
    41: "Матейчек",
    48: "Поваренков",
    54: "Сатир",
    63: "Шнейдер",
    64: "Эйдемиллер, Юстицкис"
}

def format_citation_with_authors(ref_numbers):
    """Форматирует цитату с указанием авторов и номеров"""
    authors_list = []
    for num in ref_numbers:
        if num in AUTHORS_MAP:
            authors_list.append(f"{AUTHORS_MAP[num]}, {num}")
        else:
            authors_list.append(str(num))
    return "[" + "; ".join(authors_list) + "]"

def fix_issues_Ж_П():
    """Исправляет проблемы Ж и П"""
    v4 = Path(__file__).parent
    v7_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v7.docx"
    
    if not v7_file.exists():
        print(f"ERROR: {v7_file} not found")
        return False
    
    print("=" * 70)
    print("ИСПРАВЛЕНИЕ ПРОБЛЕМ Ж И П")
    print("=" * 70)
    
    doc = Document(str(v7_file))
    changes = []
    
    # Ж. Страница 10 - исправить [22; 35] и [63; 64]
    print("\n[Ж] Исправление цитирования на странице 10...")
    for para in doc.paragraphs:
        text = para.text
        
        # Ищем [22; 35] - более гибкий поиск
        pattern_22_35 = r'\[22\s*;\s*35\]'
        if re.search(pattern_22_35, text):
            # Проверяем, не исправлено ли уже
            if "Дружинин, 22" not in text and "Леонтьев А.Н., 35" not in text:
                new_text = format_citation_with_authors([22, 35])
                para.text = re.sub(pattern_22_35, new_text, text)
                changes.append("Ж: Исправлено [22; 35]")
                print(f"  ✓ {changes[-1]}")
                text = para.text
        
        # Ищем [63; 64]
        pattern_63_64 = r'\[63\s*;\s*64\]'
        if re.search(pattern_63_64, text):
            if "Шнейдер, 63" not in text and "Эйдемиллер, Юстицкис, 64" not in text:
                new_text = format_citation_with_authors([63, 64])
                para.text = re.sub(pattern_63_64, new_text, text)
                changes.append("Ж: Исправлено [63; 64]")
                print(f"  ✓ {changes[-1]}")
    
    # П. Страница 35 - исправить [1; 2; 7; 9; 34; 41; 54] и [18; 27; 30; 48]
    print("\n[П] Исправление цитирования на странице 35...")
    for para in doc.paragraphs:
        text = para.text
        
        # Ищем [1; 2; 7; 9; 34; 41; 54]
        pattern_1_2_7 = r'\[1\s*;\s*2\s*;\s*7\s*;\s*9\s*;\s*34\s*;\s*41\s*;\s*54\]'
        if re.search(pattern_1_2_7, text):
            if "Абульханова-Славская, 1" not in text:
                new_text = format_citation_with_authors([1, 2, 7, 9, 34, 41, 54])
                para.text = re.sub(pattern_1_2_7, new_text, text)
                changes.append("П: Исправлено [1; 2; 7; 9; 34; 41; 54]")
                print(f"  ✓ {changes[-1]}")
                text = para.text
        
        # Ищем [18; 27; 30; 48]
        pattern_18_27 = r'\[18\s*;\s*27\s*;\s*30\s*;\s*48\]'
        if re.search(pattern_18_27, text):
            if "Гинзбург, 18" not in text:
                new_text = format_citation_with_authors([18, 27, 30, 48])
                para.text = re.sub(pattern_18_27, new_text, text)
                changes.append("П: Исправлено [18; 27; 30; 48]")
                print(f"  ✓ {changes[-1]}")
    
    print(f"\nВнесено изменений: {len(changes)}")
    
    # Сохраняем
    doc.save(str(v7_file))
    print(f"\n✓ Документ обновлён: {v7_file.name}")
    
    return True

if __name__ == '__main__':
    fix_issues_Ж_П()





