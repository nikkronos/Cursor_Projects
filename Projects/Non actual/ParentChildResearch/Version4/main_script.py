# -*- coding: utf-8 -*-
"""
Главный скрипт для всех задач: копирование, извлечение, анализ и правки
"""
import sys
from pathlib import Path
import shutil

def main():
    base = Path(__file__).parent.parent
    v4 = Path(__file__).parent
    
    print("=" * 60)
    print("СКРИПТ ДЛЯ РАБОТЫ С ВКР v6 -> v7")
    print("=" * 60)
    
    # Шаг 1: Копирование файлов
    print("\n[1/4] Копирование файлов...")
    v6_src = base / "Version3" / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    src_doc = base / "Version2" / "A1" / "ВКР Бакалавриат_Ворошилова.docx"
    
    if v6_src.exists():
        shutil.copy2(v6_src, v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx")
        print(f"  ✓ Скопирован: {v6_src.name}")
    else:
        print(f"  ✗ Не найден: {v6_src}")
        return False
    
    if src_doc.exists():
        shutil.copy2(src_doc, v4 / "ВКР_ИСХОДНИК.docx")
        print(f"  ✓ Скопирован: {src_doc.name}")
    else:
        print(f"  ✗ Не найден: {src_doc}")
    
    # Шаг 2: Извлечение текста
    print("\n[2/4] Извлечение текста...")
    try:
        from docx import Document
        
        v6_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
        if v6_file.exists():
            doc = Document(str(v6_file))
            text_parts = []
            
            # Параграфы
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            output_file = v4 / "v6_text.txt"
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(text)
            
            print(f"  ✓ Извлечено: {len(text_parts)} параграфов")
            print(f"  ✓ Сохранено в: {output_file.name}")
            print(f"  ✓ Размер текста: {len(text)} символов")
            
            return True
        else:
            print(f"  ✗ Файл не найден: {v6_file}")
            return False
            
    except ImportError:
        print("  ✗ ОШИБКА: python-docx не установлен")
        print("  Установите: pip install python-docx")
        return False
    except Exception as e:
        print(f"  ✗ ОШИБКА: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = main()
    if success:
        print("\n" + "=" * 60)
        print("✓ ИЗВЛЕЧЕНИЕ ЗАВЕРШЕНО УСПЕШНО!")
        print("=" * 60)
        print("\nСледующие шаги:")
        print("1. Откройте v6_text.txt для анализа")
        print("2. Запустите скрипт для правок (будет создан)")
    else:
        print("\n" + "=" * 60)
        print("✗ ОШИБКА ПРИ ВЫПОЛНЕНИИ")
        print("=" * 60)
        sys.exit(1)





