# -*- coding: utf-8 -*-
"""Выполняет извлечение текста из docx"""
import sys
import os
from pathlib import Path

# Добавляем текущую директорию в путь
sys.path.insert(0, str(Path(__file__).parent))

# Пытаемся импортировать и выполнить quick_extract
try:
    import quick_extract
    # Выполняем main логику
    base = Path(__file__).parent.parent
    v4 = Path(__file__).parent
    
    import shutil
    
    # Копируем файлы
    v6_src = base / "Version3" / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
    src_doc = base / "Version2" / "A1" / "ВКР Бакалавриат_Ворошилова.docx"
    
    if v6_src.exists():
        shutil.copy2(v6_src, v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx")
        print("OK: v6 copied")
    
    if src_doc.exists():
        shutil.copy2(src_doc, v4 / "ВКР_ИСХОДНИК.docx")
        print("OK: source copied")
    
    # Извлекаем текст
    try:
        from docx import Document
        
        v6_file = v4 / "ВКР_ПОЛНЫЙ_ДОКУМЕНТ_v6.docx"
        if v6_file.exists():
            doc = Document(str(v6_file))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Также извлекаем из таблиц
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            text = "\n".join(text_parts)
            with open(v4 / "v6_text.txt", "w", encoding="utf-8") as f:
                f.write(text)
            print(f"OK: extracted {len(text)} chars, {len(text_parts)} paragraphs")
            
    except ImportError:
        print("ERROR: Need to install python-docx: pip install python-docx")
        sys.exit(1)
        
except Exception as e:
    print(f"ERROR: {e}")
    import traceback
    traceback.print_exc()
    sys.exit(1)





