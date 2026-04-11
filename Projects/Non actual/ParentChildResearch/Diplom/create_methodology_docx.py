# -*- coding: utf-8 -*-
"""Создание документа с объяснением методики расчёта корреляций."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Заголовок
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Как рассчитывалась корреляция в исследовании')
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = subtitle.add_run('Пошаговое объяснение методики расчёта')
run2.font.size = Pt(12)
run2.font.name = 'Times New Roman'
doc.add_paragraph()

# Раздел 1
h1 = doc.add_paragraph()
h1.add_run('1. Почему мы считали корреляции по отдельным вопросам (установкам), а не по шкалам опросника Марковской?').bold = True
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Короткий ответ: ').bold = True
p.add_run('потому что нас интересовали конкретные родительские установки, а не обобщённые шкалы.')
doc.add_paragraph()

doc.add_paragraph('Подробнее:')
doc.add_paragraph(
    'Опросник ВРР Марковской устроен так: 60 вопросов группируются в 10 шкал '
    '(требовательность, контроль, эмоциональная близость и т.д.). Обычно исследователи '
    'суммируют баллы по шкалам и работают с этими суммами.'
)
doc.add_paragraph(
    'Но в данной работе мы хотели понять, какие именно установки (конкретные утверждения родителей) '
    'связаны с профессиональным выбором. Это более тонкий анализ. Например:'
)
doc.add_paragraph(
    '• Шкала «контроль» — это 10 вопросов. Но внутри неё есть разные оттенки: '
    '«контролирую все действия» — это одно, а «знаю, с кем дружит» — немного другое.'
)
doc.add_paragraph(
    '• Если считать только по шкалам, мы бы потеряли эту детализацию.'
)
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Обоснование:').bold = True
doc.add_paragraph('1. Исследовательский характер работы — мы искали, где именно есть связи.')
doc.add_paragraph('2. Небольшая выборка (50 пар) — при малых выборках иногда информативнее смотреть на отдельные пункты, чем на агрегированные шкалы.')
doc.add_paragraph('3. Практическая ценность — для рекомендаций родителям полезнее знать «вот эта конкретная установка связана с тем-то», чем абстрактное «контроль влияет».')
doc.add_paragraph()

# Раздел 2
h2 = doc.add_paragraph()
h2.add_run('2. Как технически это считалось: пошаговая инструкция').bold = True
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 1. Подготовка данных (Excel или Google Таблицы)').bold = True
doc.add_paragraph()

doc.add_paragraph('Представь, что у тебя есть Excel-файл. В нём:')
doc.add_paragraph()

doc.add_paragraph('Лист 1 — ответы родителей (ВРР):')
doc.add_paragraph('Таблица: ID пары | Вопрос 1 | Вопрос 2 | ... | Вопрос 60')
doc.add_paragraph('Каждая строка — один родитель (всего 50 строк). Каждый столбец — ответ на вопрос (от 1 до 5 баллов).')
doc.add_paragraph()

doc.add_paragraph('Лист 2 — ответы подростков (Овчарова, мотивы):')
doc.add_paragraph('Таблица: ID пары | Мотив 1 | Мотив 2 | ... | Мотив 20')
doc.add_paragraph('Каждая строка — один подросток (всего 50 строк).')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Важно: ').bold = True
p.add_run('ID пары должен совпадать! Строка 1 у родителей — это родитель подростка из строки 1.')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 2. Выбор программы').bold = True
doc.add_paragraph()

doc.add_paragraph('Можно использовать:')
doc.add_paragraph('• SPSS (классика для психологов)')
doc.add_paragraph('• JASP (бесплатный, похож на SPSS, скачивается с jasp-stats.org)')
doc.add_paragraph('• Excel (с надстройкой «Анализ данных» или вручную через ранги)')
doc.add_paragraph('• Python (библиотека scipy)')
doc.add_paragraph('• Jamovi (бесплатный, очень простой интерфейс)')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 3. Что такое корреляция Спирмена?').bold = True
doc.add_paragraph()

doc.add_paragraph('Корреляция Спирмена — это когда мы:')
doc.add_paragraph('1. Берём два столбца данных (например, «Вопрос 19 родителя» и «Мотив 12 подростка»).')
doc.add_paragraph('2. Переводим значения в ранги (кто на каком месте по возрастанию).')
doc.add_paragraph('3. Смотрим, насколько ранги совпадают.')
doc.add_paragraph()

doc.add_paragraph('Пример на пальцах:')
doc.add_paragraph('Допустим, у нас 5 пар «родитель-подросток»:')
doc.add_paragraph('ID 1: Вопрос 19 = 2, Мотив 12 = 5')
doc.add_paragraph('ID 2: Вопрос 19 = 5, Мотив 12 = 2')
doc.add_paragraph('ID 3: Вопрос 19 = 3, Мотив 12 = 4')
doc.add_paragraph('ID 4: Вопрос 19 = 4, Мотив 12 = 3')
doc.add_paragraph('ID 5: Вопрос 19 = 1, Мотив 12 = 5')
doc.add_paragraph()

doc.add_paragraph('Переводим в ранги:')
doc.add_paragraph('Вопрос 19: значения 2, 5, 3, 4, 1 → ранги: 2, 5, 3, 4, 1 (кто меньше — тот ранг ниже)')
doc.add_paragraph('Мотив 12: значения 5, 2, 4, 3, 5 → ранги: 4.5, 1, 3, 2, 4.5 (два значения «5» делят 4-е и 5-е места → среднее = 4.5)')
doc.add_paragraph()

doc.add_paragraph('Формула Спирмена:')
doc.add_paragraph('ρ = 1 - (6 × Σd²) / (n × (n² - 1))')
doc.add_paragraph('где d — разница рангов между парами, n — количество пар.')
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run('Но вручную это не считают! ').bold = True
p.add_run('Программа делает это за секунду.')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 4. Как это сделать в JASP (пошагово)').bold = True
doc.add_paragraph()

doc.add_paragraph('1. Открываешь JASP, загружаешь CSV или Excel файл с данными.')
doc.add_paragraph('   В файле должны быть все столбцы: и вопросы родителей, и мотивы/типы подростков.')
doc.add_paragraph('2. Идёшь в меню: Regression → Correlation Matrix')
doc.add_paragraph('3. Выбираешь переменные:')
doc.add_paragraph('   В левом списке выделяешь нужные столбцы (например, «Вопрос_19_родитель» и «Мотив_12_подросток»).')
doc.add_paragraph('   Перетаскиваешь их в правое окно.')
doc.add_paragraph('4. В настройках ставишь галочку: Spearman (вместо Pearson).')
doc.add_paragraph('5. Нажимаешь — готово! Программа выдаёт таблицу:')
doc.add_paragraph('   Вопрос_19 — Мотив_12: ρ = -0.304, p = 0.032')
doc.add_paragraph()
doc.add_paragraph('Если p < 0.05 — связь статистически значима.')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 5. Как это сделать в Excel (если нет JASP)').bold = True
doc.add_paragraph()

doc.add_paragraph('В Excel нет встроенной функции Спирмена, но можно посчитать через ранги:')
doc.add_paragraph()
doc.add_paragraph('1. Создаёшь два новых столбца для рангов.')
doc.add_paragraph('   В ячейке рядом с данными пишешь формулу:')
doc.add_paragraph('   =РАНГ.СР(B2;$B$2:$B$51)')
doc.add_paragraph('   Это присвоит ранг значению из B2 среди всех значений B2:B51.')
doc.add_paragraph('   Протягиваешь формулу вниз на все 50 строк.')
doc.add_paragraph('   То же самое делаешь для второго столбца.')
doc.add_paragraph()
doc.add_paragraph('2. Считаешь корреляцию Пирсона между рангами:')
doc.add_paragraph('   =КОРРЕЛ(столбец_рангов_1; столбец_рангов_2)')
doc.add_paragraph('   Корреляция Пирсона между рангами = корреляция Спирмена.')
doc.add_paragraph()
doc.add_paragraph('3. Для p-value — это сложнее, нужна отдельная формула или таблица критических значений.')
doc.add_paragraph('   Проще использовать JASP или онлайн-калькулятор.')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 6. Что конкретно считалось в этой работе').bold = True
doc.add_paragraph()

doc.add_paragraph('Связи ВРР и Овчаровой:')
doc.add_paragraph('• 60 вопросов родителей × 2 выбранных мотива (12 и 19) = 120 корреляций')
doc.add_paragraph('• Из них 5 оказались значимыми (p < 0.05)')
doc.add_paragraph()

doc.add_paragraph('Связи ВРР и Голланда:')
doc.add_paragraph('• 60 вопросов родителей × 6 типов Голланда = 360 корреляций')
doc.add_paragraph('• Из них 11 оказались значимыми')
doc.add_paragraph()

doc.add_paragraph('Итого: 492 корреляции, 17 значимых (3.5%).')
doc.add_paragraph()

h3 = doc.add_paragraph()
h3.add_run('Шаг 7. Почему выбрали только мотивы 12 и 19 (а не все 20)?').bold = True
doc.add_paragraph()

doc.add_paragraph('Это объясняется в работе:')
doc.add_paragraph('• Мотив 12 («Даёт возможности для роста профессионального мастерства») — отражает ориентацию на развитие.')
doc.add_paragraph('• Мотив 19 («Позволяет использовать профессиональные умения вне работы») — отражает интеграцию профессии в жизнь.')
doc.add_paragraph()
doc.add_paragraph(
    'Оба связаны с внутренней мотивацией и самореализацией — именно это, по гипотезе, '
    'должно быть чувствительно к родительским установкам. Остальные мотивы (зарплата, престиж, '
    'мнение друзей) — более внешние, менее связаны с семейным контекстом.'
)
doc.add_paragraph()
doc.add_paragraph('Это фокусировка исследования — не «проверим всё подряд», а «проверим то, что теоретически обосновано».')
doc.add_paragraph()

# Раздел 3
h2 = doc.add_paragraph()
h2.add_run('3. Резюме: алгоритм работы').bold = True
doc.add_paragraph()

doc.add_paragraph('1. Собрал данные в таблицу (50 строк × много столбцов).')
doc.add_paragraph('2. Убедился, что ID пар совпадают (родитель 1 → подросток 1).')
doc.add_paragraph('3. Загрузил в JASP (или SPSS, или Python).')
doc.add_paragraph('4. Для каждой пары «вопрос родителя — показатель подростка» посчитал корреляцию Спирмена.')
doc.add_paragraph('5. Отобрал те, где p < 0.05 (статистически значимые).')
doc.add_paragraph('6. Интерпретировал:')
doc.add_paragraph('   • Положительная корреляция = «чем больше одно, тем больше другое»')
doc.add_paragraph('   • Отрицательная корреляция = «чем больше одно, тем меньше другое»')

doc.save('Методика_расчёта_корреляций.docx')
print('Создан файл: Методика_расчёта_корреляций.docx')
