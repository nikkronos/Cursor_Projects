"""
Скрипт для конвертации Markdown файла в Word документ (.docx)
Обрабатывает заголовки, списки, таблицы, форматирование текста
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os

def add_table_from_markdown(doc, table_lines):
    """Добавляет таблицу из Markdown строк"""
    if len(table_lines) < 2:
        return
    
    # Определяем количество столбцов
    header = table_lines[0]
    columns = [col.strip() for col in header.split('|') if col.strip()]
    num_cols = len(columns)
    
    if num_cols == 0:
        return
    
    # Создаём таблицу
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Light Grid Accent 1'
    
    # Заголовок
    header_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        header_cells[i].text = col
        # Делаем заголовок жирным
        for paragraph in header_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Данные (пропускаем строку с разделителями)
    for row_line in table_lines[2:]:
        if not row_line.strip() or '|' not in row_line:
            continue
        
        cells = [cell.strip() for cell in row_line.split('|') if cell.strip() or len(cell) > 0]
        # Иногда первая и последняя ячейка могут быть пустыми
        while len(cells) > 0 and not cells[0]:
            cells.pop(0)
        while len(cells) > num_cols:
            cells.pop()
        
        if len(cells) < num_cols:
            cells.extend([''] * (num_cols - len(cells)))
        
        row_cells = table.add_row().cells
        for i in range(min(num_cols, len(cells))):
            row_cells[i].text = cells[i]

def markdown_to_docx(md_file, docx_file):
    """Конвертирует Markdown файл в Word документ"""
    
    # Читаем Markdown файл
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Создаём новый документ Word
    doc = Document()
    
    # Настраиваем стили
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    in_table = False
    table_lines = []
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Заголовок уровня 1
        if stripped.startswith('# ') and not stripped.startswith('##'):
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            text = stripped[2:].strip()
            doc.add_heading(text, level=1)
            i += 1
            continue
        
        # Заголовок уровня 2
        elif stripped.startswith('## ') and not stripped.startswith('###'):
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            text = stripped[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        # Заголовок уровня 3
        elif stripped.startswith('### '):
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            text = stripped[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Таблица (строка содержит |)
        elif '|' in stripped and not stripped.startswith('```'):
            in_table = True
            # Пропускаем строку с разделителями (содержит только |---|)
            if not re.match(r'^[\s\|\-:]+$', stripped):
                table_lines.append(line)
            else:
                table_lines.append(line)  # Добавляем разделитель тоже
            i += 1
            continue
        
        # Конец таблицы
        elif in_table and '|' not in stripped:
            if table_lines:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
            in_table = False
            # Не увеличиваем i, обработаем эту строку как обычную
        
        # Горизонтальная линия
        elif stripped == '---':
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            paragraph = doc.add_paragraph()
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            i += 1
            continue
        
        # Нумерованный список
        elif re.match(r'^\d+\.\s', stripped):
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            text = re.sub(r'^\d+\.\s+', '', stripped)
            para = doc.add_paragraph(text, style='List Number')
            i += 1
            continue
        
        # Маркированный список
        elif stripped.startswith('- ') or stripped.startswith('* '):
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            text = stripped[2:].strip()
            
            # Обрабатываем форматирование
            para = doc.add_paragraph(style='List Bullet')
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', text)
            for part in parts:
                if not part:
                    continue
                if part.startswith('**') and part.endswith('**'):
                    run = para.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('`') and part.endswith('`'):
                    run = para.add_run(part[1:-1])
                    run.font.name = 'Courier New'
                else:
                    if part:
                        para.add_run(part)
            i += 1
            continue
        
        # Обычный текст
        elif stripped:
            if in_table:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            
            para = doc.add_paragraph()
            
            # Разбиваем строку на части с форматированием
            parts = re.split(r'(\*\*.*?\*\*|`.*?`)', stripped)
            
            for part in parts:
                if not part:
                    continue
                
                # Жирный текст
                if part.startswith('**') and part.endswith('**'):
                    text = part[2:-2]
                    run = para.add_run(text)
                    run.bold = True
                
                # Код (моноширинный)
                elif part.startswith('`') and part.endswith('`'):
                    text = part[1:-1]
                    run = para.add_run(text)
                    run.font.name = 'Courier New'
                
                # Обычный текст
                else:
                    para.add_run(part)
            
            i += 1
            continue
        
        # Пустая строка
        else:
            if in_table and table_lines:
                add_table_from_markdown(doc, table_lines)
                table_lines = []
                in_table = False
            i += 1
            continue
    
    # Если осталась незавершённая таблица
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)
    
    # Сохраняем документ
    doc.save(docx_file)
    print(f"Файл успешно сохранён: {docx_file}")

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    md_file = os.path.join(script_dir, "Теоретические_вопросы_корреляция.md")
    docx_file = os.path.join(script_dir, "Теоретические_вопросы_корреляция.docx")
    
    if not os.path.exists(md_file):
        print(f"ОШИБКА: Файл не найден: {md_file}")
        return
    
    print("Конвертация Markdown в Word...")
    print(f"Исходный файл: {md_file}")
    print(f"Результирующий файл: {docx_file}")
    print()
    
    try:
        markdown_to_docx(md_file, docx_file)
        print()
        print("Готово! Файл успешно создан.")
    except Exception as e:
        print(f"ОШИБКА: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
